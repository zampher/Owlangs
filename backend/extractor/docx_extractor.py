# SPDX-FileCopyrightText: 2026 Zampher
# SPDX-License-Identifier: MPL-2.0

from __future__ import annotations

from io import BytesIO
from typing import List, Optional, Dict, Any, Set, Tuple
from .base import Extractor, ExtractResult
from utils.docx_utils import get_run_formatting_key, is_image_run, paragraph_has_toc_field
from utils.table_utils import (
    get_all_merged_regions_docx,
    is_merged_cell_start_at_position_docx,
    is_cell_in_merged_region_docx,
)
from logger import unified_logger as logger
from logger.logger import LogModule


class DocxExtractor(Extractor):
    def __init__(self, file_bytes: bytes, chunk_size: int = 3000):
        self.file_bytes = file_bytes
        self.chunk_size = chunk_size

    def _is_wps_file(self) -> bool:
        """
        Detect if this is a WPS-generated DOCX file based on filename and content analysis.
        """
        # Check filename patterns that indicate WPS origin
        # Note: filename is not available here, but we can check content patterns

        # WPS files often have different structure or missing standard OOXML components
        # Check for common WPS indicators in the first few bytes
        if len(self.file_bytes) < 4:
            return False

        # Check if it starts with PK (standard ZIP signature) but fails ZIP validation
        # This indicates a malformed ZIP that might be WPS format
        if self.file_bytes[:2] == b'PK':
            try:
                import zipfile
                zipfile.ZipFile(BytesIO(self.file_bytes))
                return False  # Valid ZIP, not WPS issue
            except zipfile.BadZipFile:
                # Starts with PK but not valid ZIP - likely WPS or corrupted
                return True

        return False

    def _extract_wps_content(self) -> ExtractResult:
        """
        Attempt to extract content from WPS-generated DOCX files using fallback methods.
        """
        logger.info(
            LogModule.EXTRACT,
            "[DOCX_EXTRACTOR] Attempting WPS content extraction fallback"
        )

        # Method 1: Try to decode as UTF-8 text (some WPS files contain readable text)
        try:
            text = self.file_bytes.decode('utf-8')
            # Look for XML-like content that might contain document text
            if '<w:p' in text or '<w:r' in text or 'word/document.xml' in text:
                # This looks like it might contain Word XML content
                # Try to extract text between XML tags
                import re
                # Extract text content from Word XML
                text_matches = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', text)
                if text_matches:
                    combined_text = ' '.join(text_matches)
                    segments = [s.strip() for s in combined_text.split('.') if s.strip()]
                    if segments:
                        logger.info(
                            LogModule.EXTRACT,
                            f"[DOCX_EXTRACTOR] Successfully extracted {len(segments)} segments from WPS file using XML parsing"
                        )
                        return ExtractResult(segments=segments)
        except UnicodeDecodeError:
            pass

        # Method 2: Try different encodings
        for encoding in ['utf-16', 'gb2312', 'gbk', 'big5']:
            try:
                text = self.file_bytes.decode(encoding)
                if len(text) > 100:  # Has substantial content
                    segments = [p for p in text.split('\n\n') if p.strip()]
                    if segments:
                        logger.info(
                            LogModule.EXTRACT,
                            f"[DOCX_EXTRACTOR] Successfully extracted content from WPS file using {encoding} encoding"
                        )
                        return ExtractResult(segments=segments)
            except (UnicodeDecodeError, LookupError):
                continue

        # Method 3: Last resort - extract any printable text
        try:
            # Extract printable ASCII characters
            printable_chars = []
            for byte in self.file_bytes:
                if 32 <= byte <= 126 or byte in [9, 10, 13]:  # printable ASCII + tab, LF, CR
                    printable_chars.append(chr(byte))

            text = ''.join(printable_chars)
            segments = [p for p in text.split('\n\n') if p.strip()]
            if segments:
                logger.warning(
                    LogModule.EXTRACT,
                    "[DOCX_EXTRACTOR] Used best-effort text extraction for WPS file - results may be incomplete"
                )
                return ExtractResult(segments=segments)
        except Exception as e:
            logger.warning(
                LogModule.EXTRACT,
                f"[DOCX_EXTRACTOR] Best-effort extraction failed: {e}"
            )

        # If all methods fail, raise an informative error
        raise ValueError(
            "Failed to extract segments from file. The file may be corrupted, encrypted by a third-party system, "
            "in an unsupported format (e.g., WPS format), or incompatible with the extraction process. "
            "Please try converting the file to a standard format (e.g., standard DOCX) and try again. "
            "If the problem persists, please contact the developer."
        )

    def extract(self) -> ExtractResult:
        """
        Extract text segments from DOCX with format-aware splitting.
        Splits segments based on formatting changes (font, bold, italic, underline, size, color).
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            # Fallback: decode as UTF-8 to avoid empty preview
            try:
                text = self.file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text = self.file_bytes.decode('utf-8', errors='replace')
            # Best-effort split by double newlines
            segments = [p for p in text.split('\n\n') if p.strip()]
            return ExtractResult(segments=segments)

        # Check if this is a WPS file before attempting standard DOCX processing
        if self._is_wps_file():
            logger.warning(
                LogModule.EXTRACT,
                "[DOCX_EXTRACTOR] Detected WPS file format, attempting fallback extraction methods"
            )
            try:
                return self._extract_wps_content()
            except ValueError as wps_error:
                # Re-raise WPS-specific errors
                raise wps_error
            except Exception as wps_error:
                logger.error(
                    LogModule.EXTRACT,
                    f"[DOCX_EXTRACTOR] WPS fallback extraction failed: {wps_error}"
                )
                raise ValueError(
                    "Failed to extract segments from file. The file may be corrupted, encrypted by a third-party system, "
                    "in an unsupported format (e.g., WPS format), or incompatible with the extraction process. "
                    "Please try converting the file to a standard format (e.g., standard DOCX) and try again. "
                    "If the problem persists, please contact the developer."
                ) from wps_error

        # Validate file format before attempting to open
        # DOCX files are ZIP archives, so check if it's a valid ZIP first
        try:
            import zipfile
            zipfile.ZipFile(BytesIO(self.file_bytes))
        except zipfile.BadZipFile as e:
            logger.error(
                LogModule.EXTRACT,
                f"[DOCX_EXTRACTOR] File is not a valid ZIP archive (may be WPS format or corrupted): {e}. "
                f"File size: {len(self.file_bytes)} bytes, first 16 bytes (hex): {self.file_bytes[:16].hex()}"
            )
            raise ValueError(
                "Failed to extract segments from file. The file may be corrupted, encrypted by a third-party system, "
                "in an unsupported format (e.g., WPS format), or incompatible with the extraction process. "
                "Please try converting the file to a standard format (e.g., standard DOCX) and try again. "
                "If the problem persists, please contact the developer."
            ) from e
        except Exception as e:
            logger.warning(
                LogModule.EXTRACT,
                f"[DOCX_EXTRACTOR] ZIP validation check failed (non-fatal): {e}"
            )
            # Continue anyway - python-docx will handle it

        try:
            doc = DocxDocument(BytesIO(self.file_bytes))
        except Exception as e:
            error_msg = str(e)
            if "not a zip file" in error_msg.lower() or "badzipfile" in error_msg.lower():
                logger.error(
                    LogModule.EXTRACT,
                    f"[DOCX_EXTRACTOR] Failed to open DOCX file as ZIP archive. "
                    f"This may be a WPS format file (.wps.docx) which uses a different format. "
                    f"File size: {len(self.file_bytes)} bytes, first 16 bytes (hex): {self.file_bytes[:16].hex()}. "
                    f"Error: {e}"
                )
                raise ValueError(
                    "Failed to extract segments from file. The file may be corrupted, encrypted by a third-party system, "
                    "in an unsupported format (e.g., WPS format), or incompatible with the extraction process. "
                    "Please try converting the file to a standard format (e.g., standard DOCX) and try again. "
                    "If the problem persists, please contact the developer."
                ) from e
            else:
                logger.error(
                    LogModule.EXTRACT,
                    f"[DOCX_EXTRACTOR] Failed to open DOCX file: {e}. "
                    f"File size: {len(self.file_bytes)} bytes"
                )
                raise
        segments: List[str] = []
        segment_info: List[Dict[str, Any]] = []

        para_index = 0

        def _get_run_format(run) -> Dict[str, Any]:
            """Extract run-level formatting from a single run."""
            fmt: Dict[str, Any] = {}
            try:
                if run.font.name:
                    fmt['font_name'] = run.font.name
            except Exception:
                pass
            try:
                if run.font.bold is not None:
                    fmt['bold'] = run.font.bold
            except Exception:
                pass
            try:
                if run.font.italic is not None:
                    fmt['italic'] = run.font.italic
            except Exception:
                pass
            try:
                if run.font.underline is not None:
                    fmt['underline'] = run.font.underline
            except Exception:
                pass
            try:
                if run.font.size is not None:
                    fmt['font_size_pt'] = round(run.font.size.pt, 2)
            except Exception:
                pass
            try:
                if run.font.color and run.font.color.rgb:
                    fmt['font_color'] = str(run.font.color.rgb)
            except Exception:
                pass
            return fmt

        def _get_paragraph_format(para) -> Dict[str, Any]:
            """Extract paragraph-level formatting (alignment, indentation, spacing)."""
            fmt: Dict[str, Any] = {}
            try:
                if para.alignment is not None:
                    fmt['alignment'] = str(para.alignment)
            except Exception:
                pass
            try:
                pf = para.paragraph_format
                for key, attr in [
                    ('left_indent_pt', 'left_indent'),
                    ('right_indent_pt', 'right_indent'),
                    ('first_line_indent_pt', 'first_line_indent'),
                    ('space_before_pt', 'space_before'),
                    ('space_after_pt', 'space_after'),
                ]:
                    try:
                        val = getattr(pf, attr)
                        if val is not None:
                            fmt[key] = round(val.pt, 2)
                    except Exception:
                        pass
                try:
                    if pf.line_spacing is not None:
                        fmt['line_spacing'] = round(float(pf.line_spacing), 2)
                except Exception:
                    pass
            except Exception:
                pass
            return fmt

        def process_paragraph(
            para,
            is_table_cell: bool = False,
            table_idx: Optional[int] = None,
            row_idx: Optional[int] = None,
            cell_idx: Optional[int] = None,
            merge_range: Optional[Tuple[int, int, int, int]] = None,
            cell_local_idx: Optional[int] = None,
        ):
            """Process a paragraph and extract segments with format detection."""
            nonlocal para_index

            # Skip TOC paragraphs (same logic as DocxTranslator)
            if paragraph_has_toc_field(para):
                # Don't increment para_index for TOC paragraphs to match DocxTranslator behavior
                # TOC paragraphs are skipped but don't affect the index
                return

            # Extract paragraph-level formatting once (shared by all segments in this paragraph)
            para_format = _get_paragraph_format(para)

            # Handle paragraphs with no runs (empty paragraphs or paragraphs with only formatting)
            # If paragraph has text but no runs, extract the text directly
            if not para.runs:
                para_text = para.text.strip()
                if para_text:
                    seg_info = {
                        'para_index': para_index,
                        'run_start_index': 0,
                        'run_end_index': 0,
                        'is_table_cell': is_table_cell,
                        'table_index': table_idx,
                        'row_index': row_idx,
                        'cell_index': cell_idx,
                        # For table cells, this is the paragraph index local to the cell
                        # For non-table paragraphs this remains None
                        'cell_local_idx': cell_local_idx,
                    }
                    seg_info.update(para_format)
                    if merge_range is not None:
                        seg_info['is_merged_cell'] = True
                        seg_info['merge_range'] = merge_range
                    segments.append(para_text)
                    segment_info.append(seg_info)
                para_index += 1
                return

            current_text_segment = ""
            current_run_start = 0
            previous_run_formatting = None
            # Capture run format of the first run in the current segment
            current_segment_first_run_format: Dict[str, Any] = {}

            for run_idx, run in enumerate(para.runs):
                if is_image_run(run):
                    # Encounter image, treat previously accumulated text as a segment
                    if current_text_segment.strip():
                        seg_info = {
                            'para_index': para_index,
                            'run_start_index': current_run_start,
                            'run_end_index': run_idx,
                            'is_table_cell': is_table_cell,
                            'table_index': table_idx,
                            'row_index': row_idx,
                            'cell_index': cell_idx,
                            'cell_local_idx': cell_local_idx,
                        }
                        seg_info.update(para_format)
                        seg_info.update(current_segment_first_run_format)
                        # Add merge information if this is a merged cell
                        if merge_range is not None:
                            seg_info['is_merged_cell'] = True
                            seg_info['merge_range'] = merge_range
                        segments.append(current_text_segment)
                        segment_info.append(seg_info)
                    # Reset accumulator
                    current_text_segment = ""
                    current_run_start = run_idx + 1
                    previous_run_formatting = None
                    current_segment_first_run_format = {}
                else:
                    # Capture formatting of the first run in this segment
                    if current_text_segment == "":
                        current_segment_first_run_format = _get_run_format(run)

                    # Check if formatting changed
                    current_run_formatting = get_run_formatting_key(run)

                    # If formatting changed and we have accumulated text, start a new segment
                    if (previous_run_formatting is not None and
                        current_run_formatting != previous_run_formatting and
                        current_text_segment.strip()):
                        # Save current segment
                        seg_info = {
                            'para_index': para_index,
                            'run_start_index': current_run_start,
                            'run_end_index': run_idx,
                            'is_table_cell': is_table_cell,
                            'table_index': table_idx,
                            'row_index': row_idx,
                            'cell_index': cell_idx,
                            'cell_local_idx': cell_local_idx,
                        }
                        seg_info.update(para_format)
                        seg_info.update(current_segment_first_run_format)
                        # Add merge information if this is a merged cell
                        if merge_range is not None:
                            seg_info['is_merged_cell'] = True
                            seg_info['merge_range'] = merge_range
                        segments.append(current_text_segment)
                        segment_info.append(seg_info)
                        # Start new segment
                        current_text_segment = ""
                        current_run_start = run_idx
                        # Capture format of the first run of the NEW segment
                        current_segment_first_run_format = _get_run_format(run)

                    # Accumulate text run
                    current_text_segment += run.text
                    previous_run_formatting = current_run_formatting

            # Process the last text block at the end of the paragraph
            if current_text_segment.strip():
                seg_info = {
                    'para_index': para_index,
                    'run_start_index': current_run_start,
                    'run_end_index': len(para.runs),
                    'is_table_cell': is_table_cell,
                    'table_index': table_idx,
                    'row_index': row_idx,
                    'cell_index': cell_idx,
                    'cell_local_idx': cell_local_idx,
                }
                seg_info.update(para_format)
                seg_info.update(current_segment_first_run_format)
                # Add merge information if this is a merged cell
                if merge_range is not None:
                    seg_info['is_merged_cell'] = True
                    seg_info['merge_range'] = merge_range
                segments.append(current_text_segment)
                segment_info.append(seg_info)

            para_index += 1

        # Process paragraphs in document body (skip TOC paragraphs to match DocxTranslator)
        for para in doc.paragraphs:
            if not paragraph_has_toc_field(para):
                # For non-table paragraphs, cell_local_idx is always None
                process_paragraph(para, is_table_cell=False, cell_local_idx=None)
            # Note: para_index is only incremented inside process_paragraph for non-TOC paragraphs

        # Process paragraphs in tables
        # First, get all merged regions for each table to avoid processing merged cells multiple times
        table_merged_regions: Dict[int, List[Tuple[int, int, int, int]]] = {}
        logger.info(
            LogModule.EXTRACT,
            f"[DOCX_TABLE] Found {len(doc.tables)} tables in document"
        )
        
        for table_idx, table in enumerate(doc.tables):
            merged_regions = get_all_merged_regions_docx(table)
            table_merged_regions[table_idx] = merged_regions
            
            # Debug: log table information
            row_count = len(table.rows)
            col_count = len(table.rows[0].cells) if table.rows else 0
            total_cells = sum(len(row.cells) for row in table.rows)
            logger.info(
                LogModule.EXTRACT,
                f"[DOCX_TABLE] Table {table_idx}: "
                f"rows={row_count}, cols={col_count}, total_cells={total_cells}, "
                f"merged_regions={len(merged_regions)}"
            )
            if merged_regions:
                logger.info(
                    LogModule.EXTRACT,
                    f"[DOCX_TABLE] Table {table_idx}: Merged regions: {merged_regions}"
                )
        
        for table_idx, table in enumerate(doc.tables):
            merged_regions = table_merged_regions.get(table_idx, [])
            
            # Track processed cells to avoid duplicates
            processed_cells: Set[Tuple[int, int]] = set()
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Skip if already processed (part of a merged region)
                    if (row_idx, cell_idx) in processed_cells:
                        logger.debug(
                            LogModule.EXTRACT,
                            f"[DOCX_MERGED_CELL] Table {table_idx}, Row {row_idx}, Cell {cell_idx}: "
                            f"Skipping (already processed as part of merged region)"
                        )
                        continue
                    
                    # Check if this cell is part of a merged region
                    is_in_merged, merge_range = is_cell_in_merged_region_docx(
                        table, row_idx, cell_idx, merged_regions
                    )
                    
                    if is_in_merged and merge_range is not None:
                        # Mark all cells in this merged region as processed
                        start_row, start_col, end_row, end_col = merge_range
                        for r in range(start_row, end_row + 1):
                            for c in range(start_col, end_col + 1):
                                processed_cells.add((r, c))
                        
                        # Only process if this is the start of the merged region
                        is_start = is_merged_cell_start_at_position_docx(
                            table, row_idx, cell_idx, merged_regions
                        )
                        if not is_start:
                            # Skip continuation cells (they are part of a merged region but not the start)
                            logger.debug(
                                LogModule.EXTRACT,
                                f"[DOCX_MERGED_CELL] Table {table_idx}, Row {row_idx}, Cell {cell_idx}: "
                                f"Skipping continuation cell (merge_range={merge_range})"
                            )
                            continue
                        else:
                            # This is the start of a merged region, process it
                            cell_text = cell.text[:50] if cell.text else "(empty)"
                            logger.info(
                                LogModule.EXTRACT,
                                f"[DOCX_MERGED_CELL] Table {table_idx}, Row {row_idx}, Cell {cell_idx}: "
                                f"Processing merged cell start (range={merge_range}, text='{cell_text}...')"
                            )
                    
                    # Process paragraphs in this cell
                    cell_text_parts = []
                    cell_para_local_idx = 0
                    for para in cell.paragraphs:
                        if not paragraph_has_toc_field(para):
                            para_text = para.text.strip()
                            if para_text:
                                cell_text_parts.append(para_text)
                            process_paragraph(
                                para, 
                                is_table_cell=True,
                                table_idx=table_idx,
                                row_idx=row_idx,
                                cell_idx=cell_idx,
                                merge_range=merge_range if is_in_merged else None,
                                cell_local_idx=cell_para_local_idx,
                            )
                            cell_para_local_idx += 1
                    
                    # Process nested tables in this cell (if any)
                    # Note: python-docx supports nested tables via cell.tables
                    # Recursively process nested tables to extract all content
                    def process_nested_table(nested_table_obj, parent_table_idx: int, parent_row_idx: int, parent_cell_idx: int, depth: int = 0):
                        """Recursively process nested tables within a cell."""
                        if depth > 10:  # Prevent infinite recursion
                            logger.warning(
                                LogModule.EXTRACT,
                                f"[DOCX_TABLE] Nested table depth exceeds 10, stopping recursion"
                            )
                            return
                        
                        nested_merged_regions = get_all_merged_regions_docx(nested_table_obj)
                        nested_processed_cells: Set[Tuple[int, int]] = set()
                        
                        for nested_row_idx, nested_row in enumerate(nested_table_obj.rows):
                            for nested_cell_idx, nested_cell in enumerate(nested_row.cells):
                                # Skip if already processed (part of a merged region in nested table)
                                if (nested_row_idx, nested_cell_idx) in nested_processed_cells:
                                    continue
                                
                                # Check if this nested cell is part of a merged region
                                nested_is_in_merged, nested_merge_range = is_cell_in_merged_region_docx(
                                    nested_table_obj, nested_row_idx, nested_cell_idx, nested_merged_regions
                                )
                                
                                if nested_is_in_merged and nested_merge_range is not None:
                                    # Mark all cells in this merged region as processed
                                    nested_start_row, nested_start_col, nested_end_row, nested_end_col = nested_merge_range
                                    for r in range(nested_start_row, nested_end_row + 1):
                                        for c in range(nested_start_col, nested_end_col + 1):
                                            nested_processed_cells.add((r, c))
                                    
                                    # Only process if this is the start of the merged region
                                    nested_is_start = is_merged_cell_start_at_position_docx(
                                        nested_table_obj, nested_row_idx, nested_cell_idx, nested_merged_regions
                                    )
                                    if not nested_is_start:
                                        continue
                                
                                # Process paragraphs in nested cell
                                for nested_para in nested_cell.paragraphs:
                                    if not paragraph_has_toc_field(nested_para):
                                        process_paragraph(
                                            nested_para,
                                            is_table_cell=True,
                                            table_idx=parent_table_idx,  # Keep parent table index
                                            row_idx=parent_row_idx,
                                            cell_idx=parent_cell_idx,
                                            merge_range=merge_range if is_in_merged else None,
                                            # Nested tables are treated as part of the same logical cell,
                                            # but we do not assign a deterministic cell_local_idx here.
                                            # Callers should fall back to existing para_index-based or
                                            # text-based matching for these rare cases.
                                            cell_local_idx=None,
                                        )
                                
                                # Recursively process any nested tables within this nested cell
                                if hasattr(nested_cell, 'tables') and nested_cell.tables:
                                    for deeper_nested_table in nested_cell.tables:
                                        process_nested_table(
                                            deeper_nested_table,
                                            parent_table_idx,
                                            parent_row_idx,
                                            parent_cell_idx,
                                            depth + 1
                                        )
                    
                    if hasattr(cell, 'tables') and cell.tables:
                        logger.info(
                            LogModule.EXTRACT,
                            f"[DOCX_TABLE] Table {table_idx}, Row {row_idx}, Cell {cell_idx}: "
                            f"Found {len(cell.tables)} nested table(s) in cell"
                        )
                        for nested_table in cell.tables:
                            process_nested_table(nested_table, table_idx, row_idx, cell_idx, depth=0)
                    
                    # Debug: log table cell information
                    if cell_text_parts:
                        cell_text = " | ".join(cell_text_parts)
                        cell_text_preview = cell_text[:100] + "..." if len(cell_text) > 100 else cell_text
                        merge_info = f", merge_range={merge_range}" if is_in_merged and merge_range else ""
                        logger.debug(
                            LogModule.EXTRACT,
                            f"[DOCX_TABLE] Table {table_idx}, Row {row_idx}, Cell {cell_idx}: "
                            f"text='{cell_text_preview}'{merge_info}"
                        )

        return ExtractResult(segments=segments, segment_info=segment_info)
