# SPDX-FileCopyrightText: 2026 Zampher
# SPDX-License-Identifier: MPL-2.0
import ast
import asyncio
import json
import sys
from dataclasses import dataclass
from io import BytesIO
from typing import Self, Literal, List, Dict, Any, Tuple, Optional

import docx
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from agents.segments_agent import SegmentsTranslateAgentConfig, SegmentsTranslateAgent
from ir.document import Document
from translator.ai_translator.base import AiTranslatorConfig, AiTranslator
from utils.docx_utils import get_run_formatting_key, is_image_run, paragraph_has_toc_field
from logger.logger import LogModule


def _normalize_merge_range(merge_range: Any) -> Optional[Tuple[int, int, int, int]]:
    """
    Normalize merge_range to (start_row, start_col, end_row, end_col).
    Accepts tuple, list, or str (e.g. '[1, 4, 10, 4]' from JSON / serialization).
    Returns None if parsing fails.
    """
    if merge_range is None:
        return None
    if isinstance(merge_range, (list, tuple)):
        if len(merge_range) >= 4:
            try:
                return (
                    int(merge_range[0]), int(merge_range[1]),
                    int(merge_range[2]), int(merge_range[3]),
                )
            except (ValueError, TypeError):
                return None
        return None
    if isinstance(merge_range, str):
        s = merge_range.strip()
        if not s:
            return None
        try:
            # JSON array: "[1, 4, 10, 4]"
            parsed = json.loads(s)
            if isinstance(parsed, (list, tuple)) and len(parsed) >= 4:
                return (
                    int(parsed[0]), int(parsed[1]),
                    int(parsed[2]), int(parsed[3]),
                )
        except (json.JSONDecodeError, ValueError, TypeError):
            pass
        try:
            # Python literal: "(1, 4, 10, 4)" or "[1, 4, 10, 4]"
            parsed = ast.literal_eval(s)
            if isinstance(parsed, (list, tuple)) and len(parsed) >= 4:
                return (
                    int(parsed[0]), int(parsed[1]),
                    int(parsed[2]), int(parsed[3]),
                )
        except (ValueError, TypeError, SyntaxError):
            pass
        # Comma-separated: "1,4,10,4"
        parts = [p.strip() for p in s.strip("[]()").split(",")]
        if len(parts) >= 4:
            try:
                return (
                    int(parts[0]), int(parts[1]),
                    int(parts[2]), int(parts[3]),
                )
            except (ValueError, TypeError):
                pass
    return None


def has_page_break(run: Run) -> bool:
    """Check if a run contains a page break."""
    return '<w:br' in run._element.xml and 'w:type="page"' in run._element.xml


def preserve_page_breaks_in_run(run: Run, new_text: str) -> None:
    """Set run text while preserving page breaks."""
    # Check if run has page breaks
    breaks = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
    page_breaks = [br for br in breaks if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page']

    if page_breaks:
        # Clear existing text elements but keep breaks
        text_elements = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        for text_elem in text_elements:
            text_elem.text = new_text
            # Only keep the first text element, remove others
            break
        else:
            # No text elements found, create one
            from docx.oxml import OxmlElement
            t_elem = OxmlElement('w:t')
            t_elem.text = new_text
            run._element.append(t_elem)
    else:
        # No page breaks, safe to set text directly
        run.text = new_text


def safe_get_highlight_color(font) -> str | None:
    """
    Safely get highlight_color from font object with WPS compatibility.

    Some third-party editors (like early WPS) may write unsupported values
    like w:val="none", which causes ValueError: WD_COLOR_INDEX has no XML mapping for 'none'
    when accessing font.highlight_color.

    Args:
        font: Font object from python-docx

    Returns:
        String representation of highlight color, or None if not available or invalid
    """
    try:
        # Direct access to highlight_color property without using hasattr
        # hasattr may not catch all exceptions during property access
        hc = font.highlight_color
        if hc:
            return str(hc)
    except ValueError as e:
        # Handle known WPS compatibility issues
        msg = str(e)
        if "no XML mapping for 'none'" in msg or "has no XML mapping for 'none'" in msg:
            # This is a known compatibility issue with WPS and other third-party editors
            # Log warning but don't fail the translation task
            try:
                import logging
                logging.getLogger(__name__).warning(
                    "Ignored unsupported highlight_color 'none' in DOCX run: %s", msg
                )
            except Exception:
                # Logging failure shouldn't affect main processing
                pass
            return None
        else:
            # Re-raise unknown ValueError to avoid masking real issues
            raise
    except (AttributeError, TypeError):
        # Font object doesn't have highlight_color property or other access issues
        return None

    return None


def get_font_for_language(target_language: str) -> str:
    """
    Return appropriate font based on target language and operating system.
    Accepts both language names (e.g. "Chinese", "English") and codes (e.g. "zh", "zh-CN", "en").

    Args:
        target_language: Target language name or code (e.g. "Chinese", "zh", "zh-CN")

    Returns:
        Recommended font name
    """
    if not target_language:
        return "Calibri"
    lang_lower = target_language.strip().lower()
    
    # 根据操作系统选择字体
    is_macos = sys.platform == "darwin"
    is_linux = sys.platform.startswith("linux")
    is_windows = sys.platform == "win32"
    
    # Language code to font mapping (payload/task_state often store codes like "zh", "en")
    # Aligned with frontend languageMap: zh, en, ja, ko, fr, de, es, ru, it, pt, ar, th, vi, he, hi,
    # pl, nl, da, nb, sv, fi, el, lt, ro, uk, ca, cs, hr, tr, ur, bn, ms, sl, mk, km, fil
    # Windows fonts
    code_font_map_windows = {
        "zh": "Microsoft YaHei",
        "zh-cn": "Microsoft YaHei",
        "zh-hans": "Microsoft YaHei",
        "zh-tw": "Microsoft JhengHei",
        "zh-hant": "Microsoft JhengHei",
        "en": "Calibri",
        "en-us": "Calibri",
        "en-gb": "Calibri",
        "ja": "Yu Gothic",
        "ja-jp": "Yu Gothic",
        "ko": "Malgun Gothic",
        "ko-kr": "Malgun Gothic",
        "ru": "Times New Roman",
        "ar": "Arial Unicode MS",
        "vi": "Arial Unicode MS",
        "he": "Arial Unicode MS",
        "th": "Arial Unicode MS",
        "hi": "Arial Unicode MS",
        "es": "Calibri",
        "fr": "Calibri",
        "de": "Calibri",
        "pt": "Calibri",
        "it": "Calibri",
        "pl": "Calibri",
        "nl": "Calibri",
        "da": "Calibri",
        "nb": "Calibri",
        "sv": "Calibri",
        "fi": "Calibri",
        "el": "Calibri",
        "lt": "Calibri",
        "ro": "Calibri",
        "uk": "Times New Roman",
        "ca": "Calibri",
        "cs": "Calibri",
        "hr": "Calibri",
        "tr": "Calibri",
        "ur": "Arial Unicode MS",
        "bn": "Arial Unicode MS",
        "ms": "Calibri",
        "sl": "Calibri",
        "mk": "Times New Roman",
        "km": "Arial Unicode MS",
        "fil": "Calibri",
    }
    
    # macOS fonts
    code_font_map_macos = {
        "zh": "PingFang SC",
        "zh-cn": "PingFang SC",
        "zh-hans": "PingFang SC",
        "zh-tw": "PingFang TC",
        "zh-hant": "PingFang TC",
        "en": "Helvetica Neue",
        "en-us": "Helvetica Neue",
        "en-gb": "Helvetica Neue",
        "ja": "Hiragino Sans",
        "ja-jp": "Hiragino Sans",
        "ko": "AppleGothic",
        "ko-kr": "AppleGothic",
        "ru": "Times New Roman",
        "ar": "Arial Unicode MS",
        "vi": "Arial Unicode MS",
        "he": "Arial Unicode MS",
        "th": "Arial Unicode MS",
        "hi": "Arial Unicode MS",
        "es": "Helvetica Neue",
        "fr": "Helvetica Neue",
        "de": "Helvetica Neue",
        "pt": "Helvetica Neue",
        "it": "Helvetica Neue",
        "pl": "Helvetica Neue",
        "nl": "Helvetica Neue",
        "da": "Helvetica Neue",
        "nb": "Helvetica Neue",
        "sv": "Helvetica Neue",
        "fi": "Helvetica Neue",
        "el": "Helvetica Neue",
        "lt": "Helvetica Neue",
        "ro": "Helvetica Neue",
        "uk": "Times New Roman",
        "ca": "Helvetica Neue",
        "cs": "Helvetica Neue",
        "hr": "Helvetica Neue",
        "tr": "Helvetica Neue",
        "ur": "Arial Unicode MS",
        "bn": "Arial Unicode MS",
        "ms": "Helvetica Neue",
        "sl": "Helvetica Neue",
        "mk": "Times New Roman",
        "km": "Arial Unicode MS",
        "fil": "Helvetica Neue",
    }
    
    # Linux fonts
    code_font_map_linux = {
        "zh": "Noto Sans CJK SC",
        "zh-cn": "Noto Sans CJK SC",
        "zh-hans": "Noto Sans CJK SC",
        "zh-tw": "Noto Sans CJK TC",
        "zh-hant": "Noto Sans CJK TC",
        "en": "DejaVu Sans",
        "en-us": "DejaVu Sans",
        "en-gb": "DejaVu Sans",
        "ja": "Noto Sans CJK JP",
        "ja-jp": "Noto Sans CJK JP",
        "ko": "Noto Sans CJK KR",
        "ko-kr": "Noto Sans CJK KR",
        "ru": "DejaVu Sans",
        "ar": "DejaVu Sans",
        "vi": "DejaVu Sans",
        "he": "DejaVu Sans",
        "th": "DejaVu Sans",
        "hi": "DejaVu Sans",
        "es": "DejaVu Sans",
        "fr": "DejaVu Sans",
        "de": "DejaVu Sans",
        "pt": "DejaVu Sans",
        "it": "DejaVu Sans",
        "pl": "DejaVu Sans",
        "nl": "DejaVu Sans",
        "da": "DejaVu Sans",
        "nb": "DejaVu Sans",
        "sv": "DejaVu Sans",
        "fi": "DejaVu Sans",
        "el": "DejaVu Sans",
        "lt": "DejaVu Sans",
        "ro": "DejaVu Sans",
        "uk": "DejaVu Sans",
        "ca": "DejaVu Sans",
        "cs": "DejaVu Sans",
        "hr": "DejaVu Sans",
        "tr": "DejaVu Sans",
        "ur": "DejaVu Sans",
        "bn": "DejaVu Sans",
        "ms": "DejaVu Sans",
        "sl": "DejaVu Sans",
        "mk": "DejaVu Sans",
        "km": "DejaVu Sans",
        "fil": "DejaVu Sans",
    }
    
    # 根据操作系统选择字体映射
    if is_macos:
        code_font_map = code_font_map_macos
    elif is_linux:
        code_font_map = code_font_map_linux
    else:
        code_font_map = code_font_map_windows
    
    if lang_lower in code_font_map:
        return code_font_map[lang_lower]
    
    # 如果语言代码没有找到，返回默认字体
    # 根据操作系统选择默认字体
    if is_macos:
        return "Helvetica Neue"
    elif is_linux:
        return "DejaVu Sans"
    else:
        return "Calibri"


@dataclass
class DocxTranslatorConfig(AiTranslatorConfig):
    """
    Configuration class for DocxTranslator.
    """
    insert_mode: Literal["replace", "append", "prepend"] = "replace"
    separator: str = "\n"


class DocxTranslator(AiTranslator):
    """
    Translator for .docx files.
    This version is optimized to handle mixed text and image paragraphs without losing images.
    """

    def __init__(self, config: DocxTranslatorConfig):
        super().__init__(config=config)
        self.chunk_size = config.chunk_size
        self.translate_agent = None
        if not self.skip_translate:
            agent_config = SegmentsTranslateAgentConfig(
                custom_prompt=config.custom_prompt,
                to_lang=config.to_lang,
                base_url=config.base_url,
                api_key=config.api_key,
                model_id=config.model_id,
                api_type=getattr(config, 'api_type', None) or getattr(config, 'api_protocol', None) or 'openai',
                temperature=config.temperature,
                thinking=config.thinking,
                concurrent=config.concurrent,
                connect_timeout=getattr(config, 'connect_timeout', 15),
                timeout=config.timeout,
                logger=self.logger,
                glossary_dict=config.glossary_dict,
                retry=config.retry,
                max_tokens=getattr(config, 'max_tokens', None),  # Get max_tokens from platform config
                use_seg_tags=True,  # Use SEG-tag format for DOCX segments
            )
            self.translate_agent = SegmentsTranslateAgent(agent_config)
        self.insert_mode = config.insert_mode
        self.separator = config.separator

    def _pre_translate(self, document: Document) -> Tuple[DocumentObject, List[Dict[str, Any]], List[str]]:
        """
        [Refactored] Preprocess .docx file, extract text at Run level to avoid breaking images.
        Uses format metadata from Extract phase if available, otherwise performs format detection.
        :param document: Document object containing .docx file content.
        :return: A tuple containing:
                 - docx.Document object
                 - A list containing text block information (each element represents a group of consecutive text runs)
                 - A list containing all original texts to be translated
        """
        doc = docx.Document(BytesIO(document.content))
        elements_to_translate = []
        original_texts = []
        
        # Try to load format metadata from task_state (from Extract phase)
        task_id = getattr(self, '_task_id', None)
        format_metadata = None
        if task_id:
            try:
                from backend.app.services.task import task_manager
                task_state = task_manager.get_task(task_id)
                if task_state:
                    segments_metadata = task_state.get("segments_metadata", {})
                    format_metadata = segments_metadata.get("segment_info")
            except Exception as e:
                try:
                    if getattr(self, "logger", None):
                        self.logger.warning(LogModule.TRANS, f"[DOCX_FORMATTING] Failed to load format metadata: {e}")
                except Exception:
                    pass
        
        # Use format metadata from Extract phase (should always be available after refactoring)
        if format_metadata:
            result = self._pre_translate_with_metadata(doc, format_metadata)
            if result[0] is not None:  # Successfully used metadata
                return result
            # If metadata failed, log warning and fall back to format detection
            # This should rarely happen, but we keep it for robustness
            try:
                if getattr(self, "logger", None):
                    self.logger.warning(
                        LogModule.TRANS,
                        "[DOCX_FORMATTING] Failed to use format metadata, falling back to format detection. "
                        "This should not happen in normal flow after refactoring.",
                    )
            except Exception:
                pass
            # Fall through to format detection
        
        # Format detection fallback (should rarely be used after refactoring)
        # Keep for backward compatibility and robustness
        try:
            if getattr(self, "logger", None):
                self.logger.warning(
                    LogModule.TRANS,
                    "[DOCX_FORMATTING] Format metadata not available, using format detection fallback. "
                    "This should not happen in normal flow after refactoring.",
                )
        except Exception:
            pass
        
        skipped_toc = 0

        def process_paragraph(para: Paragraph):
            nonlocal elements_to_translate, original_texts, skipped_toc
            
            # Skip paragraphs that contain TOC fields
            if paragraph_has_toc_field(para):
                skipped_toc += 1
                try:
                    snippet = (para.text or "")[:120]
                except Exception:
                    snippet = ""
                try:
                    if getattr(self, "logger", None):
                        self.logger.info(LogModule.TRANS, f"[TOC] Skipping TOC paragraph: '{snippet}'")
                except Exception:
                    pass
                return
            
            current_text_segment = ""
            current_runs = []
            previous_run_formatting = None  # Track previous run's formatting

            for run in para.runs:
                if is_image_run(run):
                    # Encounter image, treat previously accumulated text as a translation unit
                    if current_text_segment.strip():
                        elements_to_translate.append({"type": "text_runs", "runs": current_runs})
                        original_texts.append(current_text_segment)
                    # Reset accumulator
                    current_text_segment = ""
                    current_runs = []
                    previous_run_formatting = None
                else:
                    # Check if formatting changed
                    current_run_formatting = get_run_formatting_key(run)
                    
                    # If formatting changed and we have accumulated text, start a new segment
                    if (previous_run_formatting is not None and 
                        current_run_formatting != previous_run_formatting and 
                        current_text_segment.strip()):
                        # Save current segment and start a new one
                        elements_to_translate.append({"type": "text_runs", "runs": current_runs})
                        original_texts.append(current_text_segment)
                        current_text_segment = ""
                        current_runs = []
                    
                    # Accumulate text run
                    current_runs.append(run)
                    current_text_segment += run.text
                    previous_run_formatting = current_run_formatting

            # Process the last text block at the end of the paragraph
            if current_text_segment.strip():
                elements_to_translate.append({"type": "text_runs", "runs": current_runs})
                original_texts.append(current_text_segment)
            

        # Traverse all paragraphs (skip TOC paragraphs)
        for para in doc.paragraphs:
            if not paragraph_has_toc_field(para):
                process_paragraph(para)

        # Traverse all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not paragraph_has_toc_field(para):
                            process_paragraph(para)

        # Log summary for TOC skipping
        try:
            if getattr(self, "logger", None):
                self.logger.info(LogModule.TRANS, f"[TOC] Skipped paragraphs counted as TOC: {skipped_toc}")
        except Exception:
            pass
        
        # Log summary for format detection

        return doc, elements_to_translate, original_texts

    def _pre_translate_with_metadata(self, doc: DocumentObject, format_metadata: List[Dict[str, Any]]) -> Tuple[DocumentObject, List[Dict[str, Any]], List[str]]:
        """
        Preprocess DOCX using format metadata from Extract phase.
        Locates runs based on metadata instead of re-detecting format changes.
        """
        elements_to_translate = []
        original_texts = []
        
        # Get segments text from task_state
        task_id = getattr(self, '_task_id', None)
        segments_text = None
        if task_id:
            try:
                from backend.app.services.task import task_manager
                task_state = task_manager.get_task(task_id)
                if task_state:
                    source_chunks_cache = task_state.get("source_chunks_cache", {})
                    segments_text = source_chunks_cache.get("segments")
            except Exception:
                pass
        
        if not segments_text:
            # Segments text should always be available after refactoring
            # If not available, this is an error condition
            try:
                if getattr(self, "logger", None):
                    self.logger.error(LogModule.TRANS, "[DOCX_FORMATTING] segments text not found in task_state. This should not happen after refactoring.")
            except Exception:
                pass
            raise ValueError("Segments text not found in task_state. Extract phase should always generate segments.")
        
        # Build paragraph and table index maps (skip TOC paragraphs to match Extract phase)
        para_list = list(doc.paragraphs)
        table_list = list(doc.tables)
        
        # Track paragraph indices for document body and table cells separately
        # Only include non-TOC paragraphs to match Extract phase indexing
        para_index_map = {}  # Maps (is_table_cell, table_idx, row_idx, cell_idx, para_local_idx) -> paragraph
        para_local_idx = 0
        for para in para_list:
            if not paragraph_has_toc_field(para):
                para_index_map[(False, None, None, None, para_local_idx)] = para
                para_local_idx += 1
        
        for table_idx, table in enumerate(table_list):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_para_local_idx = 0
                    for para in cell.paragraphs:
                        if not paragraph_has_toc_field(para):
                            para_index_map[(True, table_idx, row_idx, cell_idx, cell_para_local_idx)] = para
                            cell_para_local_idx += 1
        
        # Process each segment using format metadata
        for seg_idx, seg_info in enumerate(format_metadata):
            if seg_idx >= len(segments_text):
                break
                
            para_index = seg_info.get('para_index')
            run_start = seg_info.get('run_start_index', 0)
            run_end = seg_info.get('run_end_index')
            is_table_cell = seg_info.get('is_table_cell', False)
            table_idx = seg_info.get('table_index')
            row_idx = seg_info.get('row_index')
            cell_idx = seg_info.get('cell_index')
            cell_local_idx = seg_info.get('cell_local_idx')
            
            # Locate the paragraph using the index map
            # 优先使用 Extract 阶段提供的 cell_local_idx（单一真相），避免在这里重新计数
            para = None
            if is_table_cell and table_idx is not None and row_idx is not None and cell_idx is not None:
                if cell_local_idx is not None:
                    # 新路径：直接用 cell-local index 查表（与 Extract 完全一致）
                    para_key = (True, table_idx, row_idx, cell_idx, cell_local_idx)
                    para = para_index_map.get(para_key)
                else:
                    # 兼容旧数据或嵌套表格：退回到原来的 para_index 计数逻辑
                    if table_idx < len(table_list):
                        table = table_list[table_idx]
                        if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[cell_idx]
                            # Count non-TOC paragraphs in document body before this table
                            body_para_count = 0
                            for p in doc.paragraphs:
                                if not paragraph_has_toc_field(p):
                                    body_para_count += 1
                            
                            # Count non-TOC paragraphs in tables before current table
                            for t_idx in range(table_idx):
                                for r in table_list[t_idx].rows:
                                    for c in r.cells:
                                        for p in c.paragraphs:
                                            if not paragraph_has_toc_field(p):
                                                body_para_count += 1
                            
                            # Count non-TOC paragraphs in current table before current cell
                            for r_idx in range(row_idx):
                                for c_idx in range(len(table.rows[r_idx].cells)):
                                    for p in table.rows[r_idx].cells[c_idx].paragraphs:
                                        if not paragraph_has_toc_field(p):
                                            body_para_count += 1
                            
                            # Count non-TOC paragraphs in current row before current cell
                            for c_idx in range(cell_idx):
                                for p in table.rows[row_idx].cells[c_idx].paragraphs:
                                    if not paragraph_has_toc_field(p):
                                        body_para_count += 1
                            
                            # Now find the cell-local para_index that matches the global para_index
                            cell_local_idx_fallback = 0
                            for p in cell.paragraphs:
                                if not paragraph_has_toc_field(p):
                                    if body_para_count == para_index:
                                        # Found the matching paragraph
                                        para_key = (is_table_cell, table_idx, row_idx, cell_idx, cell_local_idx_fallback)
                                        para = para_index_map.get(para_key)
                                        break
                                    body_para_count += 1
                                    cell_local_idx_fallback += 1
            else:
                # For non-table paragraphs, para_index is already the correct index
                para_key = (is_table_cell, table_idx, row_idx, cell_idx, para_index)
                para = para_index_map.get(para_key)
            
            # Check if this is a merged cell and log it
            is_merged_cell_in_seg = seg_info.get('is_merged_cell', False)
            if is_merged_cell_in_seg and para is None:
                try:
                    if getattr(self, "logger", None):
                        self.logger.warning(
                            LogModule.TRANS,
                            f"[DOCX_FORMATTING] Paragraph not found for merged cell segment {seg_idx}: "
                            f"para_index={para_index}, is_table_cell={is_table_cell}, "
                            f"table_idx={table_idx}, row_idx={row_idx}, cell_idx={cell_idx}, "
                            f"merge_range={seg_info.get('merge_range')}"
                        )
                except Exception:
                    pass
            
            if para is None:
                try:
                    if getattr(self, "logger", None) and seg_idx < 5:
                        self.logger.warning(
                            LogModule.TRANS,
                            f"[DOCX_FORMATTING] Paragraph not found for segment {seg_idx}: "
                            f"para_index={para_index}, is_table_cell={is_table_cell}, "
                            f"table_idx={table_idx}, row_idx={row_idx}, cell_idx={cell_idx}"
                        )
                except Exception:
                    pass
                # Use segment text from cache as fallback
                segment_text = segments_text[seg_idx] if seg_idx < len(segments_text) else ""
                if segment_text.strip():
                    original_texts.append(segment_text)
                    element_info = {"type": "text_runs", "runs": []}  # Empty runs, will be handled in _after_translate
                    # Add merge information if this is a merged cell (even if para is None)
                    if is_merged_cell_in_seg:
                        element_info['is_merged_cell'] = True
                        element_info['merge_range'] = seg_info.get('merge_range')
                        element_info['table_index'] = table_idx
                        element_info['row_index'] = row_idx
                        element_info['cell_index'] = cell_idx
                    elements_to_translate.append(element_info)
                continue
            
            # Extract runs for this segment
            runs = []
            segment_text = ""
            para_runs = list(para.runs)
            
            if run_end is None:
                run_end = len(para_runs)
            
            for run_idx in range(run_start, min(run_end, len(para_runs))):
                run = para_runs[run_idx]
                if not is_image_run(run):
                    runs.append(run)
                    segment_text += run.text
            
            # Verify segment text matches (with some tolerance for whitespace)
            expected_text = segments_text[seg_idx] if seg_idx < len(segments_text) else ""
            if expected_text and segment_text.strip() != expected_text.strip():
                try:
                    if getattr(self, "logger", None) and seg_idx < 5:
                        self.logger.warning(
                            LogModule.TRANS,
                            f"[DOCX_FORMATTING] Segment text mismatch for segment {seg_idx}: "
                            f"expected length={len(expected_text)}, actual length={len(segment_text)}"
                        )
                except Exception:
                    pass
                # Use expected text from cache
                segment_text = expected_text
            
            if segment_text.strip() and runs:
                element_info = {"type": "text_runs", "runs": runs}
                # Add merge information if this is a merged cell
                if seg_info.get('is_merged_cell', False):
                    element_info['is_merged_cell'] = True
                    element_info['merge_range'] = seg_info.get('merge_range')
                    element_info['table_index'] = table_idx
                    element_info['row_index'] = row_idx
                    element_info['cell_index'] = cell_idx
                elements_to_translate.append(element_info)
                original_texts.append(segment_text)
            elif segment_text.strip():
                # Has text but no runs (shouldn't happen, but handle gracefully)
                element_info = {"type": "text_runs", "runs": []}
                # Add merge information if this is a merged cell
                if seg_info.get('is_merged_cell', False):
                    element_info['is_merged_cell'] = True
                    element_info['merge_range'] = seg_info.get('merge_range')
                    element_info['table_index'] = table_idx
                    element_info['row_index'] = row_idx
                    element_info['cell_index'] = cell_idx
                elements_to_translate.append(element_info)
                original_texts.append(segment_text)
        
        
        return doc, elements_to_translate, original_texts

    def _after_translate(self, doc: DocumentObject, elements_to_translate: List[Dict[str, Any]],
                         translated_texts: List[str], original_texts: List[str]) -> bytes:
        """
        [Refactored] Write translated text back to corresponding text runs, preserving images and styles.
        """
        
        translation_map = dict(zip(original_texts, translated_texts))

        # Track merged cell text accumulation: (table_idx, row_idx, cell_idx) -> list of (final_text, preserved_formatting)
        merged_cell_data: Dict[Tuple[int, int, int], List[Tuple[str, Dict]]] = {}

        for i, element_info in enumerate(elements_to_translate):
            runs = element_info["runs"]
            original_text = original_texts[i]
            translated_text = translated_texts[i]

            # Determine final text based on insert mode
            if self.insert_mode == "replace":
                final_text = translated_text
            elif self.insert_mode == "append":
                final_text = original_text + self.separator + translated_text
            elif self.insert_mode == "prepend":
                final_text = translated_text + self.separator + original_text
            else:
                self.logger.error(LogModule.TRANS, "Invalid DocxTranslatorConfig parameter")
                final_text = translated_text

            # Check if this is a merged cell
            is_merged_cell = element_info.get('is_merged_cell', False)
            merge_range = element_info.get('merge_range')
            table_idx = element_info.get('table_index')
            row_idx = element_info.get('row_index')
            cell_idx = element_info.get('cell_index')
            
            
            if is_merged_cell and merge_range is not None and table_idx is not None:
                # This is a merged cell, accumulate text for multiple paragraphs in the same cell
                # Key: (table_idx, row_idx, cell_idx) identifies the cell
                cell_key = (table_idx, row_idx, cell_idx)
                
                # Get the first cell to preserve its formatting (only on first encounter)
                if cell_key not in merged_cell_data:
                    merged_cell_data[cell_key] = []
                    try:
                        if table_idx < len(doc.tables):
                            table = doc.tables[table_idx]
                            first_cell = table.rows[row_idx].cells[cell_idx]
                            first_para = first_cell.paragraphs[0] if first_cell.paragraphs else None
                            
                            # Preserve formatting from the first cell's first run (if exists) with WPS compatibility
                            preserved_formatting = None
                            if first_para and first_para.runs:
                                first_run = first_para.runs[0]
                                if first_run.font:
                                    preserved_formatting = {
                                        'bold': None,
                                        'italic': None,
                                        'underline': None,
                                        'size': None,
                                        'color': None,
                                    }

                                    try:
                                        preserved_formatting['bold'] = first_run.font.bold
                                    except Exception:
                                        # WPS compatibility: some font properties may not be accessible
                                        pass

                                    try:
                                        preserved_formatting['italic'] = first_run.font.italic
                                    except Exception:
                                        # WPS compatibility: some font properties may not be accessible
                                        pass

                                    try:
                                        preserved_formatting['underline'] = first_run.font.underline
                                    except Exception:
                                        # WPS compatibility: some font properties may not be accessible
                                        pass

                                    try:
                                        preserved_formatting['size'] = first_run.font.size
                                    except Exception:
                                        # WPS compatibility: some font properties may not be accessible
                                        pass

                                    try:
                                        if first_run.font.color and hasattr(first_run.font.color, 'rgb') and first_run.font.color.rgb:
                                            preserved_formatting['color'] = first_run.font.color.rgb
                                    except Exception:
                                        # WPS compatibility: color access may fail
                                        pass
                            
                            # Store preserved formatting with the first text
                            merged_cell_data[cell_key].append((final_text, preserved_formatting, merge_range))
                        else:
                            merged_cell_data[cell_key].append((final_text, None, merge_range))
                    except Exception as e:
                        pass
                        merged_cell_data[cell_key].append((final_text, None, merge_range))
                else:
                    # Additional paragraph in the same cell, just accumulate text
                    merged_cell_data[cell_key].append((final_text, None, merge_range))
                
                # Skip immediate write, will write all accumulated text at the end
                continue

            # Process non-merged cells and normal runs
            if not runs:
                continue


            # --- Core modification section ---
            # Distribute translated text across runs proportionally to preserve formatting
            if len(runs) == 1:
                # Single run: write all text to it
                first_run = runs[0]
                preserve_page_breaks_in_run(first_run, final_text)
                
                # Set appropriate font based on target language
                if first_run.font:
                    target_font = get_font_for_language(self.config.to_lang)
                    # Preserve original formatting (bold, italic, etc.) with WPS compatibility
                    original_bold = None
                    original_italic = None
                    original_underline = None
                    original_size = None
                    original_color = None

                    try:
                        original_bold = first_run.font.bold
                    except Exception:
                        # WPS compatibility: some font properties may not be accessible
                        pass

                    try:
                        original_italic = first_run.font.italic
                    except Exception:
                        # WPS compatibility: some font properties may not be accessible
                        pass

                    try:
                        original_underline = first_run.font.underline
                    except Exception:
                        # WPS compatibility: some font properties may not be accessible
                        pass

                    try:
                        original_size = first_run.font.size
                    except Exception:
                        # WPS compatibility: some font properties may not be accessible
                        pass

                    try:
                        if first_run.font.color and hasattr(first_run.font.color, 'rgb') and first_run.font.color.rgb:
                            original_color = first_run.font.color.rgb
                    except Exception:
                        # WPS compatibility: color access may fail
                        pass
                    
                    first_run.font.name = target_font
                    
                    # Restore original formatting
                    if original_bold is not None:
                        first_run.font.bold = original_bold
                    if original_italic is not None:
                        first_run.font.italic = original_italic
                    if original_underline is not None:
                        first_run.font.underline = original_underline
                    if original_size is not None:
                        first_run.font.size = original_size
                    if original_color is not None:
                        try:
                            from docx.shared import RGBColor
                            if isinstance(original_color, str):
                                # Handle hex color strings
                                if original_color.startswith('#'):
                                    original_color = original_color[1:]
                                original_color = RGBColor.from_string(original_color)
                            first_run.font.color.rgb = original_color
                        except Exception:
                            pass
                    
                    # If primary font is not available, try fallback fonts
                    if not first_run.font.name:
                        # Select fallback fonts based on language type
                        if any(char in self.config.to_lang for char in ['Chinese', 'Chinese', 'Simplified', 'Traditional']):
                            fallback_fonts = ['SimSun', 'SimHei', 'Arial Unicode MS', 'Times New Roman']
                        elif any(char in self.config.to_lang for char in ['Japanese', 'Japanese', 'Japanese']):
                            fallback_fonts = ['MS Gothic', 'Arial Unicode MS', 'Times New Roman']
                        elif any(char in self.config.to_lang for char in ['Korean', 'Korean', '한국어']):
                            fallback_fonts = ['Gulim', 'Arial Unicode MS', 'Times New Roman']
                        elif any(char in self.config.to_lang for char in ['Russian', 'Russian', 'Русский']):
                            fallback_fonts = ['Times New Roman', 'Arial', 'Calibri']
                        elif any(char in self.config.to_lang for char in ['Arabic', 'Arabic', 'العَرَبِيَّة']):
                            fallback_fonts = ['Arial Unicode MS', 'Times New Roman', 'Arial']
                        else:
                            fallback_fonts = ['Calibri', 'Times New Roman', 'Arial']
                        
                        # Try fallback fonts
                        for fallback_font in fallback_fonts:
                            first_run.font.name = fallback_font
                            if first_run.font.name:
                                break
            else:
                # Multiple runs: distribute text proportionally based on original text length
                # Calculate original text length for each run
                original_lengths = [len(run.text) for run in runs]
                total_original_length = sum(original_lengths)
                
                if total_original_length == 0:
                    # All runs are empty, write all text to first run
                    preserve_page_breaks_in_run(runs[0], final_text)
                else:
                    # Distribute translated text proportionally
                    current_pos = 0
                    for idx, run in enumerate(runs):
                        # Calculate proportion of this run
                        proportion = original_lengths[idx] / total_original_length
                        # Calculate text length for this run
                        run_text_length = int(len(final_text) * proportion)
                        
                        # Last run gets remaining text
                        if idx == len(runs) - 1:
                            run_text = final_text[current_pos:]
                        else:
                            run_text = final_text[current_pos:current_pos + run_text_length]
                            current_pos += run_text_length
                        
                        # Write text to run, preserving formatting
                        preserve_page_breaks_in_run(run, run_text)
                        
                        # Set appropriate font based on target language (preserve original formatting)
                        if run.font:
                            target_font = get_font_for_language(self.config.to_lang)
                            # Preserve original formatting (bold, italic, etc.) with WPS compatibility
                            original_bold = None
                            original_italic = None
                            original_underline = None
                            original_size = None
                            original_color = None

                            try:
                                original_bold = run.font.bold
                            except Exception:
                                # WPS compatibility: some font properties may not be accessible
                                pass

                            try:
                                original_italic = run.font.italic
                            except Exception:
                                # WPS compatibility: some font properties may not be accessible
                                pass

                            try:
                                original_underline = run.font.underline
                            except Exception:
                                # WPS compatibility: some font properties may not be accessible
                                pass

                            try:
                                original_size = run.font.size
                            except Exception:
                                # WPS compatibility: some font properties may not be accessible
                                pass

                            try:
                                if run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                                    original_color = run.font.color.rgb
                            except Exception:
                                # WPS compatibility: color access may fail
                                pass
                            
                            run.font.name = target_font
                            
                            # Restore original formatting
                            if original_bold is not None:
                                run.font.bold = original_bold
                            if original_italic is not None:
                                run.font.italic = original_italic
                            if original_underline is not None:
                                run.font.underline = original_underline
                            if original_size is not None:
                                run.font.size = original_size
                            if original_color is not None:
                                try:
                                    from docx.shared import RGBColor
                                    if isinstance(original_color, str):
                                        # Handle hex color strings
                                        if original_color.startswith('#'):
                                            original_color = original_color[1:]
                                        original_color = RGBColor.from_string(original_color)
                                    run.font.color.rgb = original_color
                                except Exception:
                                    pass
                            
                            # If primary font is not available, try fallback fonts
                            if not run.font.name:
                                # Select fallback fonts based on language type
                                if any(char in self.config.to_lang for char in ['Chinese', 'Chinese', 'Simplified', 'Traditional']):
                                    fallback_fonts = ['SimSun', 'SimHei', 'Arial Unicode MS', 'Times New Roman']
                                elif any(char in self.config.to_lang for char in ['Japanese', 'Japanese', 'Japanese']):
                                    fallback_fonts = ['MS Gothic', 'Arial Unicode MS', 'Times New Roman']
                                elif any(char in self.config.to_lang for char in ['Korean', 'Korean', '한국어']):
                                    fallback_fonts = ['Gulim', 'Arial Unicode MS', 'Times New Roman']
                                elif any(char in self.config.to_lang for char in ['Russian', 'Russian', 'Русский']):
                                    fallback_fonts = ['Times New Roman', 'Arial', 'Calibri']
                                elif any(char in self.config.to_lang for char in ['Arabic', 'Arabic', 'العَرَبِيَّة']):
                                    fallback_fonts = ['Arial Unicode MS', 'Times New Roman', 'Arial']
                                else:
                                    fallback_fonts = ['Calibri', 'Times New Roman', 'Arial']
                                
                                # Try fallback fonts
                                for fallback_font in fallback_fonts:
                                    run.font.name = fallback_font
                                    if run.font.name:
                                        break
            # --- End of modification ---

        # Write accumulated merged cell texts (for cells with multiple paragraphs)
        if merged_cell_data:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from utils.table_utils import is_merged_cell_continuation_docx
            
            for cell_key, text_data_list in merged_cell_data.items():
                table_idx, row_idx, cell_idx = cell_key
                if table_idx >= len(doc.tables):
                    continue
                
                table = doc.tables[table_idx]
                if row_idx >= len(table.rows):
                    continue
                
                # Get merge_range from first entry (expected: (start_row, start_col, end_row, end_col))
                merge_range_raw = text_data_list[0][2] if text_data_list else None
                if not merge_range_raw:
                    continue
                # Normalize merge_range: may be tuple, list, or str (e.g. from JSON/serialization)
                parsed = _normalize_merge_range(merge_range_raw)
                if parsed is None:
                    self.logger.warning(
                        LogModule.TRANS,
                        f"[DOCX_AFTER_TRANSLATE] merge_range could not be parsed: "
                        f"type={type(merge_range_raw).__name__}, value={merge_range_raw!r}"
                    )
                    continue
                start_row, start_col, end_row, end_col = parsed
                
                # Combine all paragraph texts (preserve line breaks between paragraphs)
                combined_texts = []
                preserved_formatting = None
                for final_text, fmt, _ in text_data_list:
                    if fmt is not None:
                        preserved_formatting = fmt
                    combined_texts.append(final_text)
                
                # Join paragraphs with line breaks (matching how Extract phase reads)
                combined_text = '\n'.join(combined_texts)
                
                # Write to the start cell only (for both horizontal and vertical merges)
                if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                    try:
                        cell = table.rows[row_idx].cells[cell_idx]
                        
                        # Remove all existing paragraphs
                        paras = list(cell.paragraphs)
                        for para in paras:
                            p = para._element
                            p.getparent().remove(p)
                        
                        # Split combined text by line breaks and create paragraphs
                        text_lines = combined_text.split('\n')
                        for line_idx, line_text in enumerate(text_lines):
                            if line_idx == 0:
                                # First paragraph
                                new_para = cell.add_paragraph(line_text)
                            else:
                                # Additional paragraphs
                                new_para = cell.add_paragraph(line_text)
                            
                            # Apply formatting to the first run of each paragraph
                            if new_para.runs:
                                new_run = new_para.runs[0]
                                target_font = get_font_for_language(self.config.to_lang)
                                if new_run.font:
                                    new_run.font.name = target_font
                                    
                                    # Restore preserved formatting (from first paragraph)
                                    if preserved_formatting:
                                        if preserved_formatting.get('bold') is not None:
                                            new_run.font.bold = preserved_formatting['bold']
                                        if preserved_formatting.get('italic') is not None:
                                            new_run.font.italic = preserved_formatting['italic']
                                        if preserved_formatting.get('underline') is not None:
                                            new_run.font.underline = preserved_formatting['underline']
                                        if preserved_formatting.get('size') is not None:
                                            new_run.font.size = preserved_formatting['size']
                                        if preserved_formatting.get('color') is not None:
                                            try:
                                                from docx.shared import RGBColor
                                                color = preserved_formatting['color']
                                                if isinstance(color, str):
                                                    if color.startswith('#'):
                                                        color = color[1:]
                                                    color = RGBColor.from_string(color)
                                                new_run.font.color.rgb = color
                                            except Exception:
                                                pass
                        
                        # Clear continuation cells for vertical merges
                        is_vertical_merge = (end_row - start_row) > 0
                        if is_vertical_merge:
                            for r in range(start_row + 1, end_row + 1):
                                if r < len(table.rows):
                                    # Check if this cell is a continuation
                                    row_xml = table.rows[r]._tr
                                    cell_elements = row_xml.findall(qn('w:tc'))
                                    actual_col = 0
                                    for tc_elem in cell_elements:
                                        tcPr = tc_elem.find(qn('w:tcPr'))
                                        gridSpan = None
                                        if tcPr is not None:
                                            gridSpan_elem = tcPr.find(qn('w:gridSpan'))
                                            if gridSpan_elem is not None:
                                                gridSpan_val = gridSpan_elem.get(qn('w:val'))
                                                if gridSpan_val:
                                                    try:
                                                        gridSpan = int(gridSpan_val)
                                                    except (ValueError, TypeError):
                                                        pass
                                        
                                        if actual_col == cell_idx:
                                            # Check if this is a continuation cell
                                            try:
                                                # Get python-docx cell object to check
                                                if actual_col < len(table.rows[r].cells):
                                                    continuation_cell = table.rows[r].cells[actual_col]
                                                    if is_merged_cell_continuation_docx(continuation_cell):
                                                        # Clear paragraphs but ensure one empty paragraph exists
                                                        for para in list(continuation_cell.paragraphs):
                                                            continuation_cell._element.remove(para._element)
                                                        # Create empty paragraph
                                                        continuation_cell.add_paragraph()
                                            except Exception:
                                                pass
                                            break
                                        
                                        actual_col += (gridSpan if gridSpan else 1)
                        
                    except Exception as e:
                        try:
                            if getattr(self, "logger", None):
                                self.logger.warning(
                                    LogModule.TRANS,
                                    f"[DOCX_TABLE_EXPORT] Failed to write accumulated merged cell text "
                                    f"for {cell_key}: {e}"
                                )
                        except Exception:
                            pass

        # Save the modified document to BytesIO stream
        doc_output_stream = BytesIO()
        doc.save(doc_output_stream)
        return doc_output_stream.getvalue()

    def translate(self, document: Document) -> Self:
        """
        Synchronously translate .docx file.
        """
        doc, elements_to_translate, original_texts = self._pre_translate(document)
        if not original_texts:
            # Use i18n logger for translation messages
            from logger.logger import i18n_logger
            i18n_logger.info("backend.translation.task.no_text_found")
            output_stream = BytesIO()
            doc.save(output_stream)
            document.content = output_stream.getvalue()
            return self

        if self.glossary_agent:
            self.glossary_dict_gen = self.glossary_agent.send_segments(original_texts, self.chunk_size)
            if self.translate_agent:
                self.translate_agent.update_glossary_dict(self.glossary_dict_gen)

        # Call translation agent
        task_id = getattr(self, '_task_id', None)
        excluded_set = self._get_excluded_segments(task_id)
        if excluded_set:
            self.logger.info(LogModule.TRANS, f"[DOCX_TRANSLATOR] Skipping translation for excluded segments: {sorted(excluded_set)}")

        translate_indices = [i for i in range(len(original_texts)) if i not in excluded_set]
        texts_for_translation = [original_texts[i] for i in translate_indices]

        # Use generic chunk translation helper to save segments to cache and translate with chunk merging
        chunk_to_segment_map = None
        if self.translate_agent and texts_for_translation:
            try:
                from utils.chunk_translation_helper import translate_segments_with_agent
                from backend.app.services.task import task_manager
                task_state = task_manager.get_task(task_id) if task_id else None
                
                translated_segments, metadata = translate_segments_with_agent(
                    segments=texts_for_translation,
                    chunk_size=self.chunk_size,
                    translate_agent=self.translate_agent,
                    task_id=task_id,
                    task_state=task_state,
                    original_filename=getattr(self, '_original_filename', None),
                    file_contents=document.content,
                    segment_indices=translate_indices,  # CRITICAL: Pass original segment indices to preserve index mapping
                )
                chunk_to_segment_map = metadata.get("chunk_to_segment_map")
            except Exception as e:
                # Fallback to direct translation if helper fails
                self.logger.warning(LogModule.TRANS, f"[DOCX_TRANSLATOR] Failed to use chunk translation helper: {e}, falling back to direct translation")
                translated_segments = self.translate_agent.send_segments(texts_for_translation, self.chunk_size, segment_indices=translate_indices)
        else:
            translated_segments = texts_for_translation

        # Validate translated_segments length matches translate_indices
        if len(translated_segments) != len(translate_indices):
            self.logger.error(
                LogModule.TRANS,
                f"[DOCX_TRANSLATOR] Length mismatch: translated_segments={len(translated_segments)}, "
                f"translate_indices={len(translate_indices)}. This will cause incorrect segment mapping!"
            )
            # Pad or truncate to match length
            if len(translated_segments) < len(translate_indices):
                self.logger.warning(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] translated_segments is shorter, padding with original text"
                )
                translated_segments = list(translated_segments) + [
                    original_texts[translate_indices[i]] 
                    for i in range(len(translated_segments), len(translate_indices))
                ]
            else:
                self.logger.warning(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] translated_segments is longer, truncating to match translate_indices"
                )
                translated_segments = translated_segments[:len(translate_indices)]
        
        # Log first few mappings for debugging
        if len(translate_indices) > 0 and len(translated_segments) > 0:
            self.logger.debug(
                LogModule.TRANS,
                f"[DOCX_TRANSLATOR] Mapping translated segments to final_translated_texts: "
                f"translate_indices[:5]={translate_indices[:5]}, "
                f"translated_segments preview (first 50 chars): {[str(s)[:50] for s in translated_segments[:5]]}"
            )
        
        final_translated_texts = list(original_texts)
        for idx, segment_idx in enumerate(translate_indices):
            if idx < len(translated_segments):
                final_translated_texts[segment_idx] = translated_segments[idx]
            else:
                self.logger.warning(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] Index {idx} out of range for translated_segments "
                    f"(length: {len(translated_segments)}), keeping original text"
                )

        # Write translation results back to document
        document.content = self._after_translate(doc, elements_to_translate, final_translated_texts, original_texts)
        return self

    async def translate_async(self, document: Document, progress_callback=None) -> Self:
        """
        Asynchronously translate .docx file.
        """
        doc, elements_to_translate, original_texts = await asyncio.to_thread(self._pre_translate, document)
        if not original_texts:
            # Use i18n logger for translation messages
            from logger.logger import i18n_logger
            i18n_logger.info("backend.translation.task.no_text_found")
            # Correctly save and return in async environment
            output_stream = BytesIO()
            doc.save(output_stream)
            document.content = output_stream.getvalue()
            return self

        # CRITICAL: Check for applied glossary in task_state and update agent if available
        # This handles the case where glossary is applied after workflow config is built
        task_id = getattr(self, '_task_id', None)
        self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] translate_async: task_id={task_id}, translate_agent={self.translate_agent is not None}")
        if task_id and self.translate_agent:
            try:
                from backend.app.services.task import task_manager
                task_state = task_manager.get_task(task_id)
                self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: task_state exists: {task_state is not None}")
                if task_state:
                    applied_glossary = task_state.get("applied_glossary")
                    self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: applied_glossary exists: {applied_glossary is not None}, type: {type(applied_glossary)}")
                    if applied_glossary and isinstance(applied_glossary, dict):
                        glossary_dict = applied_glossary.get("glossary_dict", {})
                        self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: glossary_dict exists: {glossary_dict is not None}, size: {len(glossary_dict) if glossary_dict else 0}")
                        if glossary_dict:
                            self.logger.info(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: Updating translate_agent with {len(glossary_dict)} glossary entries from task_state")
                            self.translate_agent.update_glossary_dict(glossary_dict)
                            sample = dict(list(glossary_dict.items())[:3])
                            self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: Sample glossary entries: {sample}")
                            # Verify the update was successful
                            self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: translate_agent.glossary_dict after update: {self.translate_agent.glossary_dict is not None}, size: {len(self.translate_agent.glossary_dict) if self.translate_agent.glossary_dict else 0}")
                        else:
                            self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: glossary_dict is empty or None")
                    else:
                        self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: applied_glossary is not a dict or is None")
                else:
                    self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: task_state not found")
            except Exception as e:
                self.logger.warning(LogModule.TRANS, f"[DOCX_TRANSLATOR] Task {task_id}: Failed to load glossary from task_state: {e}", exc_info=True)
        else:
            if not task_id:
                self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] translate_async: No task_id available")
            if not self.translate_agent:
                self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] translate_async: No translate_agent available")

        if self.glossary_agent:
            self.glossary_dict_gen = await self.glossary_agent.send_segments_async(original_texts, self.chunk_size)
            if self.translate_agent:
                self.translate_agent.update_glossary_dict(self.glossary_dict_gen)

        # Asynchronously call translation agent
        task_id = getattr(self, '_task_id', None)
        excluded_set = self._get_excluded_segments(task_id)
        if excluded_set:
            self.logger.info(LogModule.TRANS, f"[DOCX_TRANSLATOR] Skipping translation for excluded segments: {sorted(excluded_set)}")

        translate_indices = [i for i in range(len(original_texts)) if i not in excluded_set]
        texts_for_translation = [original_texts[i] for i in translate_indices]

        # Use generic chunk translation helper to save segments to cache and translate with chunk merging
        chunk_to_segment_map = None
        if self.translate_agent and texts_for_translation:
            try:
                from utils.chunk_translation_helper import translate_segments_with_agent_async
                from backend.app.services.task import task_manager
                task_state = task_manager.get_task(task_id) if task_id else None
                
                translated_segments, metadata = await translate_segments_with_agent_async(
                    segments=texts_for_translation,
                    chunk_size=self.chunk_size,
                    translate_agent=self.translate_agent,
                    task_id=task_id,
                    task_state=task_state,
                    original_filename=getattr(self, '_original_filename', None),
                    file_contents=document.content,
                    progress_callback=progress_callback,
                    segment_indices=translate_indices,  # CRITICAL: Pass original segment indices to preserve index mapping
                )
                chunk_to_segment_map = metadata.get("chunk_to_segment_map")
            except Exception as e:
                # Fallback to direct translation if helper fails
                self.logger.warning(LogModule.TRANS, f"[DOCX_TRANSLATOR] Failed to use chunk translation helper: {e}, falling back to direct translation")
                # CRITICAL: Set task_state on agent even in fallback path to ensure debug files are saved
                if task_id and self.translate_agent:
                    try:
                        from backend.app.services.task import task_manager
                        task_state = task_manager.get_task(task_id) if task_id else None
                        if task_state:
                            self.translate_agent.task_state = task_state
                            self.translate_agent.task_id = task_id
                            self.logger.debug(LogModule.TRANS, f"[DOCX_TRANSLATOR] Set task_state on agent in fallback path: task_id={task_id}")
                    except Exception as e2:
                        self.logger.warning(LogModule.TRANS, f"[DOCX_TRANSLATOR] Failed to set task_state in fallback path: {e2}")
                translated_segments = await self.translate_agent.send_segments_async(
                    texts_for_translation, self.chunk_size, progress_callback, segment_indices=translate_indices
                )
        else:
            # No translate_agent or no texts_for_translation, use original texts
            translated_segments = texts_for_translation

        # Validate translated_segments length matches translate_indices
        if len(translated_segments) != len(translate_indices):
            self.logger.error(
                LogModule.TRANS,
                f"[DOCX_TRANSLATOR] Length mismatch: translated_segments={len(translated_segments)}, "
                f"translate_indices={len(translate_indices)}. This will cause incorrect segment mapping!"
            )
            # Pad or truncate to match length
            if len(translated_segments) < len(translate_indices):
                self.logger.warning(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] translated_segments is shorter, padding with original text"
                )
                translated_segments = list(translated_segments) + [
                    original_texts[translate_indices[i]] 
                    for i in range(len(translated_segments), len(translate_indices))
                ]
            else:
                self.logger.warning(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] translated_segments is longer, truncating to match translate_indices"
                )
                translated_segments = translated_segments[:len(translate_indices)]
        
        # Log first few mappings for debugging
        if len(translate_indices) > 0 and len(translated_segments) > 0:
            self.logger.debug(
                LogModule.TRANS,
                f"[DOCX_TRANSLATOR] Mapping translated segments to final_translated_texts: "
                f"translate_indices[:5]={translate_indices[:5]}, "
                f"translated_segments preview (first 50 chars): {[str(s)[:50] for s in translated_segments[:5]]}"
            )
        
        final_translated_texts = list(original_texts)
        for idx, segment_idx in enumerate(translate_indices):
            if idx < len(translated_segments):
                final_translated_texts[segment_idx] = translated_segments[idx]
            else:
                self.logger.warning(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] Index {idx} out of range for translated_segments "
                    f"(length: {len(translated_segments)}), keeping original text"
                )
        
        # CRITICAL: Extract translation results from llm_api_output
        # With the new implementation, API returns chunks with original segment indices (may be non-continuous, e.g., 0, 1, 3, 5, 6)
        # So we can directly use the indices from API response without mapping
        if task_id:
            try:
                from backend.app.services.task import task_manager
                task_state = task_manager.get_task(task_id) if task_id else None
                if task_state and 'llm_api_output' in task_state:
                    chunk_data = task_state.get('llm_api_output')
                    if chunk_data and isinstance(chunk_data, list):
                        # Build map from chunk data (each chunk is a dict with original segment indices as keys)
                        api_translations = {}  # original_segment_index -> translated_text
                        for chunk_idx, chunk in enumerate(chunk_data):
                            if isinstance(chunk, dict):
                                for seg_idx_str, translated_text in chunk.items():
                                    try:
                                        original_seg_idx = int(seg_idx_str)
                                        # Check if translation differs from source (valid translation exists)
                                        if original_seg_idx < len(original_texts):
                                            source_text = original_texts[original_seg_idx]
                                            if translated_text and translated_text.strip() and translated_text.strip() != source_text.strip():
                                                api_translations[original_seg_idx] = translated_text
                                                # Log for excluded segments (they might not be in final_translated_texts yet)
                                                if original_seg_idx in excluded_set:
                                                    self.logger.info(
                                                        LogModule.TRANS,
                                                        f"[DOCX_TRANSLATOR] Found translation for excluded segment {original_seg_idx} "
                                                        f"in API output: source='{source_text[:50]}...', target='{translated_text[:50]}...'"
                                                    )
                                    except (ValueError, TypeError):
                                        continue
                        
                        # Update final_translated_texts with API translations
                        # This ensures excluded segments that were sent to API get their translations
                        if api_translations:
                            for original_seg_idx, translated_text in api_translations.items():
                                if original_seg_idx < len(final_translated_texts):
                                    # Only update if current value is source text (not already translated)
                                    current_value = final_translated_texts[original_seg_idx]
                                    source_text = original_texts[original_seg_idx]
                                    if current_value == source_text or original_seg_idx in excluded_set:
                                        final_translated_texts[original_seg_idx] = translated_text
                                        if original_seg_idx in excluded_set:
                                            self.logger.info(
                                                LogModule.TRANS,
                                                f"[DOCX_TRANSLATOR] Updated final_translated_texts[{original_seg_idx}] "
                                                f"with translation from API output for excluded segment: '{translated_text[:50]}...'"
                                            )
            except Exception as e:
                self.logger.warning(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] Failed to extract translations from llm_api_output: {e}",
                    exc_info=True
                )
        
        # Record translation segments if task_id is provided
        if task_id and len(original_texts) == len(final_translated_texts):
            try:
                from utils.translation_segments import record_translation_segments
                from backend.app.services.task import task_manager
                task_state = task_manager.get_task(task_id)
                
                if task_state:
                    # Get platform key from task_state (set during task initialization)
                    platform_key = task_state.get("platform_key")
                    self.logger.info(LogModule.TRANS, f"Recording {len(original_texts)} translation segments for task {task_id}")
                    # Get excluded segments using _get_excluded_segments (includes segments_metadata.excluded_segment_indices)
                    excluded_set_for_recording = self._get_excluded_segments(task_id)
                    excluded_segments_for_recording = sorted(excluded_set_for_recording) if excluded_set_for_recording else None
                    if excluded_segments_for_recording:
                        self.logger.info(LogModule.TRANS, f"[DOCX_TRANSLATOR] Using {len(excluded_segments_for_recording)} excluded_segments for recording: {excluded_segments_for_recording[:10]}...")
                    
                    # CRITICAL: For DOCX workflow, original_texts and final_translated_texts are segments (not chunks)
                    # We need to pass chunk_to_segment_map=None to let record_translation_segments correctly identify them as segments
                    # If we pass chunk_to_segment_map_from_state, it will be treated as chunks, causing incorrect mapping
                    record_translation_segments(
                        task_id=task_id,
                        source_chunks=original_texts,
                        target_chunks=final_translated_texts,
                        original_filename=getattr(self, '_original_filename', None),
                        workflow_type=getattr(self, '_workflow_type', None),
                        source_lang=None,
                        target_lang=self.config.to_lang if hasattr(self.config, 'to_lang') else None,
                        platform_key=platform_key,
                        task_state=task_state,
                        excluded_segments=excluded_segments_for_recording,
                        chunk_to_segment_map=None,  # CRITICAL: Pass None to indicate these are segments, not chunks
                    )
                    self.logger.info(LogModule.TRANS, f"Successfully recorded translation segments for task {task_id}")
                else:
                    self.logger.warning(LogModule.TRANS, f"Task state not found for task {task_id}, cannot record segments")
            except Exception as e:
                # Log error but don't fail translation
                self.logger.warning(LogModule.TRANS, f"Failed to record translation segments for task {task_id}: {e}", exc_info=True)
        else:
            if not task_id:
                self.logger.debug(LogModule.TRANS, "No task_id provided, skipping segment recording")
            elif len(original_texts) != len(final_translated_texts):
                self.logger.warning(
                    LogModule.TRANS,
                    f"Source chunks ({len(original_texts)}) and target chunks ({len(final_translated_texts)}) "
                    f"count mismatch, skipping segment recording"
                )
        
        # Write translation results back to document
        document.content = await asyncio.to_thread(self._after_translate, doc, elements_to_translate, final_translated_texts,
                                                   original_texts)
        return self

    def _get_excluded_segments(self, task_id: str | None) -> set[int]:
        """
        Get excluded segment indices using ExclusionManager (single source of truth).
        This ensures that manually excluded segments from Extract phase are correctly identified.
        """
        if not task_id:
            return set()
        try:
            from backend.app.services.task import task_manager
        except ImportError:
            return set()
        task_state = task_manager.get_task(task_id)
        if not task_state:
            return set()
        
        # CRITICAL: Use ExclusionManager.get_excluded_segments as the single source of truth
        # This ensures that manually excluded segments from Extract phase are correctly identified
        from exclusion.core import ExclusionManager
        excluded_segments_with_reasons = ExclusionManager.get_excluded_segments(task_state)
        excluded_set = set(excluded_segments_with_reasons.keys())
        
        # CRITICAL: Validate consistency with Extract phase
        segments_metadata = task_state.get("segments_metadata", {})
        excluded_segments_dict = segments_metadata.get("excluded_segments", {})
        if excluded_segments_dict and isinstance(excluded_segments_dict, dict):
            extract_phase_count = len(excluded_segments_dict)
            translate_phase_count = len(excluded_set)
            if extract_phase_count != translate_phase_count:
                self.logger.error(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] Task {task_id}: INCONSISTENCY DETECTED! "
                    f"Extract phase excluded_segments count: {extract_phase_count}, "
                    f"Translate phase retrieved count: {translate_phase_count}. "
                    f"Difference: {extract_phase_count - translate_phase_count} segments missing. "
                    f"This indicates exclusion data was not properly passed from Extract to Translate phase.",
                )
            else:
                self.logger.info(
                    LogModule.TRANS,
                    f"[DOCX_TRANSLATOR] Task {task_id}: Retrieved {len(excluded_set)} excluded_segments from ExclusionManager "
                    f"(consistent with Extract phase: {extract_phase_count}). "
                    f"Excluded indices: {sorted(excluded_set)[:20]}{'...' if len(excluded_set) > 20 else ''}",
                )
        elif excluded_set:
            self.logger.info(
                LogModule.TRANS,
                f"[DOCX_TRANSLATOR] Task {task_id}: Retrieved {len(excluded_set)} excluded_segments from ExclusionManager "
                f"(single source of truth). Excluded indices: {sorted(excluded_set)[:20]}{'...' if len(excluded_set) > 20 else ''}",
            )
        else:
            self.logger.debug(
                LogModule.TRANS,
                f"[DOCX_TRANSLATOR] Task {task_id}: No excluded_segments found from ExclusionManager",
            )
        
        return excluded_set
