#!/usr/bin/env python3
# SPDX-FileCopyrightText: 2026 Zamphersss
# SPDX-License-Identifier: MPL-2.0

"""
Verify merged cell translation export in DOCX.
This script checks if translated text was correctly written to merged cells.
"""

import sys
from pathlib import Path

# Add project root to path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

# Setup logger module before importing table_utils
import logging
from types import ModuleType
logger_module = ModuleType('logger')
verify_logger = logging.getLogger('verify')
logger_module.unified_logger = verify_logger  # For table_utils and any logger import
sys.modules['logger'] = logger_module

import docx
from backend.utils.table_utils import get_all_merged_regions_docx, get_merged_cell_range_docx

def verify_merged_cell_export(docx_path: str):
    """
    Verify that translated text was correctly written to merged cells.
    
    Args:
        docx_path: Path to the translated DOCX file
    """
    print(f"[VERIFY] Loading DOCX file: {docx_path}")
    doc = docx.Document(docx_path)
    
    if not doc.tables:
        print("[VERIFY] ERROR: No tables found in document")
        return False
    
    print(f"[VERIFY] Found {len(doc.tables)} table(s) in document")
    
    all_verified = True
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n[VERIFY] Checking table {table_idx}")
        print(f"[VERIFY] Table has {len(table.rows)} rows")
        
        # Get all merged regions
        merged_regions = get_all_merged_regions_docx(table)
        print(f"[VERIFY] Found {len(merged_regions)} merged regions: {merged_regions}")
        
        # Check each merged region
        for merge_range in merged_regions:
            start_row, start_col, end_row, end_col = merge_range
            print(f"\n[VERIFY] Checking merged region: {merge_range}")
            
            # Get text from the start cell
            try:
                start_cell = table.rows[start_row].cells[start_col]
                start_text = start_cell.text.strip()
                print(f"[VERIFY] Start cell text (first 100 chars): {start_text[:100]}...")
                print(f"[VERIFY] Start cell text length: {len(start_text)}")
                
                if not start_text:
                    print(f"[VERIFY] WARNING: Start cell is empty!")
                    all_verified = False
                    continue
                
                # Check all cells in the merge range
                cells_checked = 0
                cells_with_text = 0
                cells_matching = 0
                
                for r in range(start_row, end_row + 1):
                    if r >= len(table.rows):
                        break
                    
                    # Get row XML to access all cells
                    row_xml = table.rows[r]._tr
                    cell_elements = row_xml.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                    
                    # Calculate actual column index
                    actual_col = 0
                    for c_idx, tc_elem in enumerate(cell_elements):
                        # Check gridSpan
                        tcPr = tc_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
                        gridSpan = None
                        if tcPr is not None:
                            gridSpan_elem = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                            if gridSpan_elem is not None:
                                gridSpan_val = gridSpan_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if gridSpan_val:
                                    try:
                                        gridSpan = int(gridSpan_val)
                                    except (ValueError, TypeError):
                                        pass
                        
                        # Check if this cell is in the merge range
                        cell_start_col = actual_col
                        cell_end_col = actual_col + (gridSpan - 1 if gridSpan else 0)
                        
                        is_in_range = False
                        if gridSpan and gridSpan > 1:
                            if cell_start_col <= start_col <= cell_end_col:
                                is_in_range = True
                        else:
                            if start_col <= cell_end_col and end_col >= cell_start_col:
                                is_in_range = True
                        
                        if is_in_range:
                            cells_checked += 1
                            
                            # Get text from this cell
                            cell_text = ""
                            para_elems = tc_elem.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
                            for para_elem in para_elems:
                                text_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                for text_elem in text_elems:
                                    if text_elem.text:
                                        cell_text += text_elem.text
                            
                            cell_text = cell_text.strip()
                            
                            if cell_text:
                                cells_with_text += 1
                                if cell_text == start_text:
                                    cells_matching += 1
                                else:
                                    print(f"[VERIFY] WARNING: Cell at row {r}, col {actual_col} has different text!")
                                    print(f"[VERIFY]   Expected: {start_text[:50]}...")
                                    print(f"[VERIFY]   Got: {cell_text[:50]}...")
                            
                            # Also try python-docx API as fallback
                            try:
                                if r < len(table.rows) and actual_col < len(table.rows[r].cells):
                                    api_cell = table.rows[r].cells[actual_col]
                                    api_text = api_cell.text.strip()
                                    if api_text and api_text != cell_text:
                                        print(f"[VERIFY] NOTE: python-docx API shows different text for row {r}, col {actual_col}")
                                        print(f"[VERIFY]   XML text: {cell_text[:50]}...")
                                        print(f"[VERIFY]   API text: {api_text[:50]}...")
                            except Exception as e:
                                pass
                        
                        actual_col += (gridSpan if gridSpan else 1)
                        if actual_col > end_col:
                            break
                
                print(f"[VERIFY] Cells checked: {cells_checked}, with text: {cells_with_text}, matching: {cells_matching}")
                
                if cells_with_text == 0:
                    print(f"[VERIFY] ERROR: No cells have text in this merge range!")
                    all_verified = False
                elif cells_matching < cells_checked:
                    print(f"[VERIFY] WARNING: Not all cells have matching text!")
                    all_verified = False
                else:
                    print(f"[VERIFY] OK: All cells in merge range have matching text")
            
            except Exception as e:
                print(f"[VERIFY] ERROR: Failed to check merge range {merge_range}: {e}")
                import traceback
                traceback.print_exc()
                all_verified = False
    
    return all_verified

def main():
    """Main function"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Verify merged cell translation export in DOCX')
    parser.add_argument('docx_path', help='Path to the translated DOCX file')
    
    args = parser.parse_args()
    
    if not Path(args.docx_path).exists():
        print(f"[VERIFY] ERROR: File not found: {args.docx_path}")
        return 1
    
    success = verify_merged_cell_export(args.docx_path)
    
    if success:
        print("\n[VERIFY] Verification completed successfully!")
        return 0
    else:
        print("\n[VERIFY] Verification completed with errors!")
        return 1

if __name__ == '__main__':
    sys.exit(main())


