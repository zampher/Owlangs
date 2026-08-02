# SPDX-FileCopyrightText: 2026 Zampher
# SPDX-License-Identifier: MPL-2.0

"""Format conversion utilities: HTML to DOCX and Markdown to PDF (via Pandoc)."""

import base64
import hashlib
import io
import mimetypes
import os
import re
import shutil
import sys
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document
from docx.oxml.ns import qn
from logger import unified_logger as logger
from logger.logger import LogModule
from utils.latex_repair_payload import extract_latex_error_context
from utils.docx_algorithm_latex_wrap import wrap_bare_latex_for_docx_algorithms
from utils.docx_md_normalize import normalize_docx_markdown_sup_sub
from utils.math_md_normalize import normalize_md_math_for_pandoc_export


def _to_short_path_if_needed(path: Path) -> Path:
    """On Windows, return 8.3 short path when path contains non-ASCII (e.g. CJK).
    TeX/kpathsea tools can fail with path encoding; short path avoids that."""
    if sys.platform != "win32" or not path.exists():
        return path
    s = str(path)
    if all(ord(c) <= 127 for c in s):
        return path
    try:
        import ctypes
        from ctypes import wintypes
        buf_size = 0
        kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
        GetShortPathNameW = kernel32.GetShortPathNameW
        GetShortPathNameW.argtypes = [wintypes.LPCWSTR, wintypes.LPWSTR, wintypes.DWORD]
        GetShortPathNameW.restype = wintypes.DWORD
        needed = GetShortPathNameW(s, None, 0)
        if needed == 0 or needed > 32767:
            return path
        buf = ctypes.create_unicode_buffer(int(needed))
        if GetShortPathNameW(s, buf, needed) == 0:
            return path
        return Path(buf.value)
    except Exception:
        return path


def _ensure_ascii_path_for_tex(tex_root: Path) -> Path:
    """Return an ASCII-only path for the TeX root so Pandoc/XeLaTeX subprocesses get paths they can use.
    Tries 8.3 short path first; if that fails (e.g. 8.3 disabled on volume), creates a directory junction
    in %%LOCALAPPDATA%%\\OwlangsPdflatex so kpathsea and child processes see ASCII paths."""
    if sys.platform != "win32" or not tex_root.exists():
        return tex_root
    s = str(tex_root)
    if all(ord(c) <= 127 for c in s):
        return tex_root
    short = _to_short_path_if_needed(tex_root)
    if short != tex_root and all(ord(c) <= 127 for c in str(short)):
        return short
    # Short path failed (e.g. 8.3 disabled); create junction in ASCII-only dir
    import subprocess
    base = Path(os.environ.get("LOCALAPPDATA", os.environ.get("TEMP", "."))) / "OwlangsPdflatex"
    base.mkdir(parents=True, exist_ok=True)
    key = hashlib.sha256(s.encode("utf-8")).hexdigest()[:8]
    link_path = base / key
    if link_path.exists():
        if link_path.is_dir():
            logger.debug(LogModule.RESTOR, f"[PDF-EXPORT] Using existing junction for TeX root: {link_path}")
            return link_path
        link_path.unlink()
    try:
        subprocess.run(
            ["cmd", "/c", "mklink", "/J", str(link_path), s],
            check=True,
            capture_output=True,
            timeout=10,
        )
        logger.info(
            LogModule.RESTOR,
            f"[PDF-EXPORT] Created junction for CJK TeX path: {link_path} -> {tex_root}",
        )
        return link_path
    except (subprocess.CalledProcessError, FileNotFoundError, OSError) as e:
        logger.warning(
            LogModule.RESTOR,
            f"[PDF-EXPORT] Junction fallback failed: {e}. PDF export may fail in CJK install path.",
        )
        return tex_root


def _get_user_texmfvar_dir() -> Path:
    r"""Return a user-writable texmf-var directory (outside Program Files).
    Uses %LOCALAPPDATA%\Owlangs\texmf-var on Windows, ~/.cache/owlangs/texmf-var elsewhere."""
    if sys.platform == "win32":
        local_appdata = os.environ.get("LOCALAPPDATA")
        if local_appdata:
            return Path(local_appdata) / "Owlangs" / "texmf-var"
    return Path.home() / ".cache" / "owlangs" / "texmf-var"


def _copy_bundled_fmt_if_needed(pdflatex_root: Path, user_texmfvar: Path) -> bool:
    """Copy pre-built xelatex.fmt from bundle to user texmf-var so fmtutil doesn't need
    to regenerate it (avoids long first-run times and potential permission issues)."""
    bundle_fmt = pdflatex_root / "texmf-var" / "web2c" / "xetex" / "xelatex.fmt"
    user_fmt = user_texmfvar / "web2c" / "xetex" / "xelatex.fmt"
    if not bundle_fmt.exists():
        return False
    if user_fmt.exists():
        return True
    try:
        user_fmt.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(bundle_fmt), str(user_fmt))
        logger.info(
            LogModule.RESTOR,
            f"[PDF-EXPORT] Copied bundled xelatex.fmt to {user_fmt}",
        )
        return True
    except (OSError, shutil.Error) as e:
        logger.warning(
            LogModule.RESTOR,
            f"[PDF-EXPORT] Could not copy bundled xelatex.fmt: {e}. Will let fmtutil regenerate.",
        )
        return False


def _ensure_xelatex_fmt(pdflatex_root: Path, env: Dict[str, str]) -> None:
    """Ensure xelatex.fmt exists before Pandoc invokes xelatex.
    If a usable format file is already present (bundled or previously generated), skip fmtutil
    to avoid unnecessary work and potential permission/version issues.
    Only runs fmtutil-sys when the format is actually missing.
    Caller must set TEXMFCNF (incl. web2c) and TEXMFSYSVAR in env."""
    if sys.platform != "win32":
        return
    bin_win = pdflatex_root / "bin" / "windows"
    fmtutil_sys = bin_win / "fmtutil-sys.exe"
    if not fmtutil_sys.exists():
        logger.warning(
            LogModule.RESTOR,
            "[PDF-EXPORT] fmtutil-sys.exe not found under bin/windows. "
            "Ensure 3rdParty pdflatex is fully deployed. PDF export may fail.",
        )
        return
    user_texmfvar = Path(env.get("TEXMFSYSVAR", str(pdflatex_root / "texmf-var")))
    user_fmt = user_texmfvar / "web2c" / "xetex" / "xelatex.fmt"
    # If format already exists and is non-empty, assume it's usable and skip fmtutil
    if user_fmt.exists() and user_fmt.stat().st_size > 0:
        logger.info(LogModule.RESTOR, f"[PDF-EXPORT] xelatex.fmt already exists at {user_fmt}, skipping fmtutil")
        return
    # No format file: try copying from bundle first, then run fmtutil as last resort
    if _copy_bundled_fmt_if_needed(pdflatex_root, user_texmfvar):
        if user_fmt.exists() and user_fmt.stat().st_size > 0:
            logger.info(LogModule.RESTOR, "[PDF-EXPORT] Using bundled xelatex.fmt, skipping fmtutil")
            return
    import subprocess
    try:
        logger.info(LogModule.RESTOR, "[PDF-EXPORT] Building xelatex format (first run or missing fmt); running fmtutil-sys")
        proc = subprocess.run(
            [str(fmtutil_sys), "--byfmt", "xelatex"],
            env=env,
            capture_output=True,
            timeout=120,
            encoding="utf-8",
            errors="replace",
        )
        if proc.returncode != 0:
            logger.warning(
                LogModule.RESTOR,
                f"[PDF-EXPORT] fmtutil exit code {proc.returncode}, stderr: {(proc.stderr or '')[:500]}",
            )
        else:
            logger.info(LogModule.RESTOR, "[PDF-EXPORT] fmtutil completed (xelatex format ready)")
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError) as e:
        logger.warning(
            LogModule.RESTOR,
            f"[PDF-EXPORT] fmtutil failed: {e}. PDF export may fail.",
        )


def _check_latex_packages_macos() -> bool:
    """Check if required LaTeX packages are installed on macOS.
    Returns True if all required packages are available, False otherwise."""
    if sys.platform != "darwin":
        return True
    
    required_packages = ["titlesec", "xecjk", "ctex", "ragged2e", "hyperref", "graphicx", "etoolbox"]
    missing_packages = []
    
    for package in required_packages:
        try:
            import subprocess
            result = subprocess.run(
                ["kpsewhich", f"{package}.sty"],
                capture_output=True,
                timeout=10,
            )
            if result.returncode != 0:
                missing_packages.append(package)
        except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
            missing_packages.append(package)
    
    if missing_packages:
        logger.warning(
            LogModule.RESTOR,
            f"[PDF-EXPORT] Missing LaTeX packages on macOS: {', '.join(missing_packages)}. "
            "Run: cd 3rdParty/macos && ./install_latex_packages.sh"
        )
        return False
    
    return True


def _get_pandoc_path() -> Optional[Path]:
    """Get Pandoc executable file path, similar to Redis path detection.
    
    Checks multiple possible locations:
    1. PyInstaller environment (packaged executable)
    2. Installation directory (production - standard Program Files)
    3. Development directory (3rdParty/windows/pandoc-*/pandoc.exe)
    4. Current working directory
    
    Returns:
        Path to pandoc.exe if found, None otherwise
    """
    if sys.platform == "win32":
        # Windows - Check multiple possible locations
        
        # 1. Check PyInstaller environment (packaged executable)
        if hasattr(sys, '_MEIPASS'):
            # Running from PyInstaller - try to find Pandoc relative to executable
            exe_path = Path(sys.executable)
            # Onedir: EXE and 3rdParty are in the same folder
            onedir_base = exe_path.parent
            pandoc_base = onedir_base / "3rdParty" / "windows"
            if pandoc_base.exists():
                for pandoc_dir in pandoc_base.glob("pandoc-*"):
                    pandoc_exe = pandoc_dir / "pandoc.exe"
                    if pandoc_exe.exists():
                        logger.debug(LogModule.TRANS, f"Found Pandoc in onedir directory: {pandoc_exe}")
                        return pandoc_exe
            # Onedir / onefile fallback: _MEIPASS contains bundled 3rdParty
            meipass_pandoc_base = Path(sys._MEIPASS) / "3rdParty" / "windows"
            if meipass_pandoc_base.exists():
                for pandoc_dir in meipass_pandoc_base.glob("pandoc-*"):
                    pandoc_exe = pandoc_dir / "pandoc.exe"
                    if pandoc_exe.exists():
                        logger.info(LogModule.TRANS, f"Found Pandoc in PyInstaller bundle directory: {pandoc_exe}")
                        return pandoc_exe
            # Legacy: EXE is in a subdir (e.g. bin/), 3rdParty is in parent
            install_base = exe_path.parent.parent
            pandoc_base = install_base / "3rdParty" / "windows"
            if pandoc_base.exists():
                for pandoc_dir in pandoc_base.glob("pandoc-*"):
                    pandoc_exe = pandoc_dir / "pandoc.exe"
                    if pandoc_exe.exists():
                        logger.debug(LogModule.TRANS, f"Found Pandoc in installation directory: {pandoc_exe}")
                        return pandoc_exe
        
        # 2. Check installation directory (production - detect via registry/env/common paths)
        install_dir = _get_owlangs_install_dir()
        if install_dir:
            pandoc_base = install_dir / "3rdParty" / "windows"
            if pandoc_base.exists():
                for pandoc_dir in pandoc_base.glob("pandoc-*"):
                    pandoc_exe = pandoc_dir / "pandoc.exe"
                    if pandoc_exe.exists():
                        logger.debug(LogModule.TRANS, f"Found Pandoc in installation directory: {pandoc_exe}")
                        return pandoc_exe
        # Fallback to legacy hard-coded path for backwards compatibility
        legacy_dir = Path("C:/Program Files/Owlangs")
        if legacy_dir.exists():
            pandoc_base = legacy_dir / "3rdParty" / "windows"
            if pandoc_base.exists():
                for pandoc_dir in pandoc_base.glob("pandoc-*"):
                    pandoc_exe = pandoc_dir / "pandoc.exe"
                    if pandoc_exe.exists():
                        logger.info(LogModule.TRANS, f"Found Pandoc in legacy path: {pandoc_exe}")
                        return pandoc_exe
        
        # 3. Check development directory
        dev_pandoc_base = Path(__file__).parent.parent.parent / "3rdParty" / "windows"
        if dev_pandoc_base.exists():
            for pandoc_dir in dev_pandoc_base.glob("pandoc-*"):
                pandoc_exe = pandoc_dir / "pandoc.exe"
                if pandoc_exe.exists():
                    logger.info(LogModule.TRANS, f"Found Pandoc in development directory: {pandoc_exe}")
                    return pandoc_exe
        
        # 4. Check current working directory
        cwd_pandoc_base = Path.cwd() / "3rdParty" / "windows"
        if cwd_pandoc_base.exists():
            for pandoc_dir in cwd_pandoc_base.glob("pandoc-*"):
                pandoc_exe = pandoc_dir / "pandoc.exe"
                if pandoc_exe.exists():
                    logger.info(LogModule.TRANS, f"Found Pandoc in current directory: {pandoc_exe}")
                    return pandoc_exe

    # macOS / non-Windows: GUI or frozen app often has minimal PATH (no Homebrew).
    # Search explicit locations so packaged app can find user-installed pandoc.
    if sys.platform == "darwin" or (sys.platform != "win32" and getattr(sys, "frozen", False)):
        # 1. Current PATH (works in terminal, often empty in .app)
        found = shutil.which("pandoc")
        if found:
            p = Path(found)
            if p.is_file():
                logger.info(LogModule.TRANS, f"Found Pandoc via PATH: {p}")
                return p
        # 2. macOS Homebrew and common install locations
        for candidate in (
            Path("/opt/homebrew/bin/pandoc"),   # Apple Silicon Homebrew
            Path("/usr/local/bin/pandoc"),      # Intel Homebrew / universal
            Path(os.path.expanduser("~/.local/bin/pandoc")),
        ):
            if candidate.is_file() and os.access(candidate, os.X_OK):
                logger.info(LogModule.TRANS, f"Found Pandoc at: {candidate}")
                return candidate

    # For other non-Windows platforms, pypandoc will use system PATH
    return None


def _get_owlangs_install_dir() -> Optional[Path]:
    """Try to locate the Owlangs installation directory on Windows.
    Checks env var, registry (Uninstall info), and common hard-coded paths."""
    if sys.platform != "win32":
        return None
    # 1. Environment variable override
    env_dir = os.environ.get("OWLANGS_HOME")
    if env_dir:
        p = Path(env_dir)
        if p.exists():
            return p
    # 2. Registry: look for Uninstall entries with DisplayName containing Owlangs
    try:
        import winreg
        for hive, key_path in [
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Microsoft\Windows\CurrentVersion\Uninstall"),
            (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall"),
        ]:
            try:
                with winreg.OpenKey(hive, key_path) as uninstall_key:
                    for i in range(winreg.QueryInfoKey(uninstall_key)[0]):
                        try:
                            subkey_name = winreg.EnumKey(uninstall_key, i)
                            with winreg.OpenKey(uninstall_key, subkey_name) as app_key:
                                display_name, _ = winreg.QueryValueEx(app_key, "DisplayName")
                                if "Owlangs" in display_name:
                                    install_location, _ = winreg.QueryValueEx(app_key, "InstallLocation")
                                    if install_location and Path(install_location).exists():
                                        return Path(install_location)
                        except (OSError, FileNotFoundError):
                            continue
            except (OSError, FileNotFoundError):
                continue
    except ImportError:
        pass
    # 3. Common hard-coded paths (install dir first so user can override by placing
    # 3rdParty there; ProgramData fallback for the new installer layout)
    for candidate in [
        Path("C:/Program Files/Owlangs"),
        Path("C:/Program Files (x86)/Owlangs"),
        Path(os.environ.get("PROGRAMDATA", "")) / "Owlangs",
        Path(os.environ.get("LOCALAPPDATA", "")) / "Owlangs",
        Path("C:/Owlangs"),
        Path("D:/Owlangs"),
    ]:
        if candidate.exists():
            return candidate
    return None


def _is_system_readonly_path(path: Path) -> bool:
    r"""Check whether *path* lives under a Windows system directory that is typically read-only
    for non-elevated users (e.g. C:\Program Files, C:\Windows)."""
    if sys.platform != "win32":
        return False
    try:
        resolved = path.resolve()
    except (OSError, ValueError):
        resolved = path.absolute()
    lower = str(resolved).lower()
    system_prefixes = (
        os.environ.get("PROGRAMFILES", "C:\\Program Files").lower(),
        os.environ.get("PROGRAMFILES(X86)", "C:\\Program Files (x86)").lower(),
        os.environ.get("WINDIR", "C:\\Windows").lower(),
        "C:\\Windows",
        "C:\\Program Files",
        "C:\\Program Files (x86)",
    )
    return any(lower.startswith(p) for p in system_prefixes)


def _mirror_pdflatex_to_local(pdflatex_root: Path) -> Optional[Path]:
    r"""Mirror the bundled pdflatex directory into %%LOCALAPPDATA%%\Owlangs\3rdParty\windows\pdflatex
    so that XeLaTeX can write format files, font caches, etc. even when the original bundle
    resides in a read-only location such as C:\Program Files.

    Returns the path to the mirrored xelatex.exe, or None if mirroring failed."""
    local_appdata = os.environ.get("LOCALAPPDATA")
    if not local_appdata:
        return None
    local_pdflatex = Path(local_appdata) / "Owlangs" / "3rdParty" / "windows" / "pdflatex"
    local_xelatex = local_pdflatex / "bin" / "windows" / "xelatex.exe"
    if local_xelatex.exists():
        return local_xelatex
    try:
        local_pdflatex.parent.mkdir(parents=True, exist_ok=True)
        shutil.copytree(pdflatex_root, local_pdflatex, dirs_exist_ok=True)
        logger.info(
            LogModule.TRANS,
            f"[PDF-EXPORT] Mirrored pdflatex from readonly path '{pdflatex_root}' to '{local_pdflatex}'",
        )
        return local_xelatex if local_xelatex.exists() else None
    except (OSError, shutil.Error) as e:
        logger.warning(
            LogModule.TRANS,
            f"[PDF-EXPORT] Failed to mirror pdflatex to writable path: {e}. "
            "Will attempt to use the original readonly path (TEXMFVAR redirection may still work).",
        )
        return None


def _get_xelatex_path() -> Optional[Path]:
    """Get XeLaTeX executable path for Pandoc PDF engine (e.g. TinyTeX under 3rdParty/windows/pdflatex).

    Returns:
        Path to xelatex.exe if found, None otherwise.
    """
    if sys.platform != "win32":
        return None
    candidates: list[Path] = []
    # 0. ProgramData (preferred for deployed installations - installer moves pdflatex here)
    program_data = Path(os.environ.get("PROGRAMDATA", "")) / "Owlangs" / "3rdParty" / "windows"
    if program_data.exists():
        candidates.append(program_data)
    # 1. PyInstaller environment (packaged executable)
    if hasattr(sys, "_MEIPASS"):
        # Onedir: EXE and 3rdParty are in the same folder
        candidates.append(Path(sys.executable).parent / "3rdParty" / "windows")
        candidates.append(Path(sys._MEIPASS) / "3rdParty" / "windows")
        # Legacy: EXE is in a subdir, 3rdParty is in parent
        candidates.append(Path(sys.executable).parent.parent / "3rdParty" / "windows")
    # 2. Installation directory (production)
    install_dir = _get_owlangs_install_dir()
    if install_dir:
        candidates.append(install_dir / "3rdParty" / "windows")
    # 3. Development directory
    candidates.append(Path(__file__).parent.parent.parent / "3rdParty" / "windows")
    # 4. Current working directory
    candidates.append(Path.cwd() / "3rdParty" / "windows")
    found_path: Optional[Path] = None
    for base in candidates:
        if not base.exists():
            continue
        # pdflatex dir contains bin/windows/xelatex.exe (TinyTeX layout)
        xelatex_exe = base / "pdflatex" / "bin" / "windows" / "xelatex.exe"
        if xelatex_exe.exists():
            found_path = xelatex_exe
            break
    if found_path is None:
        return None
    # If the discovered pdflatex lives in a system/readonly directory (e.g. C:\Program Files),
    # mirror it to %LOCALAPPDATA% so that subprocesses can write without elevation.
    pdflatex_root = found_path.parent.parent.parent
    if _is_system_readonly_path(pdflatex_root):
        mirrored = _mirror_pdflatex_to_local(pdflatex_root)
        if mirrored is not None:
            logger.info(LogModule.TRANS, f"Using mirrored XeLaTeX: {mirrored}")
            return mirrored
    logger.info(LogModule.TRANS, f"Found XeLaTeX: {found_path}")
    return found_path


_BILINGUAL_STYLED_SPAN_RE = re.compile(
    r'<span\s+style="([^"]*)">\s*([\s\S]*?)\s*</span>',
    re.IGNORECASE,
)


def parse_bilingual_styled_spans(source_content: str) -> List[Tuple[str, bool, Optional[str]]]:
    """Parse styled bilingual spans from markdown/HTML; return (plain_text, italic, #RRGGBB) in order."""
    import html as html_module

    parts: List[Tuple[str, bool, Optional[str]]] = []
    if not source_content or "<span" not in source_content.lower():
        return parts
    for match in _BILINGUAL_STYLED_SPAN_RE.finditer(source_content):
        style = match.group(1).lower()
        inner = html_module.unescape(match.group(2).strip())
        inner = inner.replace("<br/>", "\n").replace("<br>", "\n")
        if not inner:
            continue
        italic = "font-style:italic" in style
        color_match = re.search(r"color:(#[0-9a-f]{6})", style, re.IGNORECASE)
        color_hex = color_match.group(1).upper() if color_match else None
        parts.append((inner, italic, color_hex))
    return parts


def _iter_all_docx_runs(doc: Document):
    """Yield runs in document order (body, tables, headers, footers)."""
    for para in doc.paragraphs:
        for run in para.runs:
            yield run
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        yield run
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                yield run
        for para in section.footer.paragraphs:
            for run in para.runs:
                yield run


def _flatten_docx_run_text(doc: Document) -> Tuple[str, List[Tuple[Any, int, int]]]:
    """Build flat text and (run, start, end) index map for styled-span matching."""
    flat_parts: List[str] = []
    index_map: List[Tuple[Any, int, int]] = []
    pos = 0
    for run in _iter_all_docx_runs(doc):
        text = run.text or ""
        if not text:
            continue
        start = pos
        pos += len(text)
        flat_parts.append(text)
        index_map.append((run, start, pos))
    return "".join(flat_parts), index_map


def _runs_covering_range(
    index_map: List[Tuple[Any, int, int]], start: int, end: int
) -> List[Any]:
    """Return runs whose text overlaps [start, end) in flattened document text."""
    selected: List[Any] = []
    seen: set[int] = set()
    for run, run_start, run_end in index_map:
        if run_end <= start or run_start >= end:
            continue
        run_id = id(run)
        if run_id in seen:
            continue
        seen.add(run_id)
        selected.append(run)
    return selected


def apply_bilingual_styled_spans_to_docx(docx_path: str, source_content: str) -> int:
    """
    Post-process a Pandoc-generated DOCX: re-apply italic/color from bilingual <span style="..."> tags.

    Pandoc (MD/HTML -> DOCX) keeps text but drops inline CSS; this walks source spans and matches
    plain text back onto DOCX runs.
    """
    styled_parts = parse_bilingual_styled_spans(source_content)
    if not styled_parts:
        return 0

    try:
        from docx.shared import RGBColor
    except ImportError:
        logger.warning(
            LogModule.RESTOR,
            "[DOCX-EXPORT] apply_bilingual_styled_spans_to_docx: python-docx unavailable, skip",
        )
        return 0

    doc = Document(docx_path)
    flat_text, index_map = _flatten_docx_run_text(doc)
    if not flat_text:
        logger.warning(
            LogModule.RESTOR,
            "[DOCX-EXPORT] apply_bilingual_styled_spans_to_docx: empty DOCX text, skip",
        )
        return 0

    cursor = 0
    applied = 0
    for plain_text, italic, color_hex in styled_parts:
        idx = flat_text.find(plain_text, cursor)
        if idx < 0:
            logger.warning(
                LogModule.RESTOR,
                "[DOCX-EXPORT] Bilingual span text not found in DOCX after Pandoc: "
                f"{plain_text[:80]!r}",
            )
            continue
        end = idx + len(plain_text)
        for run in _runs_covering_range(index_map, idx, end):
            if italic:
                run.italic = True
            if color_hex:
                hex_digits = color_hex.lstrip("#")
                run.font.color.rgb = RGBColor(
                    int(hex_digits[0:2], 16),
                    int(hex_digits[2:4], 16),
                    int(hex_digits[4:6], 16),
                )
        applied += 1
        cursor = end

    if applied:
        doc.save(docx_path)
        logger.info(
            LogModule.RESTOR,
            f"[DOCX-EXPORT] Applied bilingual styled spans to Pandoc DOCX: {applied}/{len(styled_parts)}",
        )
    elif styled_parts:
        logger.warning(
            LogModule.RESTOR,
            "[DOCX-EXPORT] Bilingual spans present in source but none matched in Pandoc DOCX",
        )
    return applied


def _apply_font_to_run(run, font_name: str) -> None:
    """Set run font name and w:eastAsia so Word uses our font for CJK (avoids 等线 default)."""
    run.font.name = font_name
    try:
        r = run._element
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def _apply_font_to_docx_runs(docx_path: str, font_name: str) -> None:
    """Open DOCX, set font (name + eastAsia) on every run, save. Overrides Pandoc/Word default (e.g. 等线)."""
    doc = Document(docx_path)
    for para in doc.paragraphs:
        for run in para.runs:
            _apply_font_to_run(run, font_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _apply_font_to_run(run, font_name)
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                _apply_font_to_run(run, font_name)
        for para in section.footer.paragraphs:
            for run in para.runs:
                _apply_font_to_run(run, font_name)
    doc.save(docx_path)


def convert_html_to_docx(
    html_content: str,
    output_path: str,
    output_dir: Optional[Path] = None,
    to_lang: Optional[str] = None,
) -> None:
    """
    Convert HTML content to DOCX using Pandoc (via pypandoc). Falls back to basic python-docx output if Pandoc fails.
    
    This function adds unified CSS styles (font-size, and font-family when to_lang is given) to the HTML
    before conversion to ensure consistent appearance in the DOCX output.
    
    Args:
        html_content: HTML content string
        output_path: Path to output DOCX file
        output_dir: Optional output directory for saving images (if data URIs need to be converted to files)
        to_lang: Optional target language name (e.g. "Chinese", "English"); when set, font-family is chosen by language.
    """
    if not html_content:
        raise ValueError("HTML content is empty, cannot generate DOCX.")

    try:
        import pypandoc
    except ImportError:
        logger.error(LogModule.RESTOR, "pypandoc is not installed; cannot convert HTML to DOCX.")
        raise

    # Try to find pandoc.exe in 3rdParty directory (Windows only)
    pandoc_path = _get_pandoc_path()
    if pandoc_path:
        # Set environment variable for pypandoc to use our pandoc.exe
        os.environ['PYPANDOC_PANDOC'] = str(pandoc_path)
        logger.info(LogModule.TRANS, f"[DOCX-EXPORT] Using Pandoc from 3rdParty: {pandoc_path}")
    else:
        logger.info(LogModule.TRANS, "[DOCX-EXPORT] Pandoc path not found in 3rdParty, will use system PATH or download")

    try:
        try:
            pypandoc.get_pandoc_version()
        except OSError:
            # If pandoc_path was set but still not found, try downloading
            if not pandoc_path:
                logger.info(LogModule.TRANS, "Pandoc not found, downloading via pypandoc...")
                pypandoc.download_pandoc()
            else:
                # pandoc_path was set but pypandoc still can't find it
                logger.warning(LogModule.RESTOR,f"Pandoc path was set to {pandoc_path} but pypandoc cannot access it")
                raise

        sanitized_html = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", html_content)
        
        # CRITICAL: Pandoc may not support data URIs in images, so convert them to file paths
        # Extract output directory from output_path if not provided
        if output_dir is None:
            output_dir = Path(output_path).parent
        
        # Create images directory
        images_dir = output_dir / "images"
        images_dir.mkdir(parents=True, exist_ok=True)
        
        # Find all data URI images and convert them to files
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(sanitized_html, 'html.parser')
        data_uri_pattern = re.compile(r'data:image/([^;]+);base64,([^"\']+)')
        images_converted = 0
        
        for img in soup.find_all('img'):
            src = img.get('src', '')
            if src.startswith('data:image/'):
                try:
                    # Parse data URI
                    match = data_uri_pattern.search(src)
                    if match:
                        mime_type = match.group(1)
                        base64_data = match.group(2)
                        
                        # Determine file extension
                        extension = mimetypes.guess_extension(f"image/{mime_type}") or ".png"
                        
                        # Generate unique filename from base64 data
                        image_id = hashlib.md5(base64_data.encode()).hexdigest()[:8]
                        image_filename = f"{image_id}{extension}"
                        image_path = images_dir / image_filename
                        
                        # Decode and save image
                        image_bytes = base64.b64decode(base64_data)
                        image_path.write_bytes(image_bytes)
                        
                        # Update img src to relative path (Pandoc can handle relative paths)
                        relative_path = f"images/{image_filename}"
                        img['src'] = relative_path
                        images_converted += 1
                        logger.debug(LogModule.RESTOR,f"Converted data URI image to file: {relative_path}")
                except Exception as img_error:
                    logger.warning(LogModule.RESTOR,f"Failed to convert data URI image to file: {img_error}, keeping data URI")
        
        if images_converted > 0:
            logger.info(LogModule.TRANS, f"[DOCX-EXPORT] Converted {images_converted} data URI images to files for Pandoc conversion")
            sanitized_html = str(soup)
        else:
            logger.debug(LogModule.RESTOR,"[DOCX-EXPORT] No data URI images found, HTML already uses file paths")
        
        # Add unified CSS styles for consistent font size (and font-family when to_lang is set) in DOCX output
        font_family_css = ""
        if to_lang:
            try:
                from translator.ai_translator.docx_translator import get_font_for_language
                font_name = get_font_for_language(to_lang)
                if font_name:
                    font_family_css = f"\n                font-family: '{font_name}';"
                    logger.info(LogModule.TRANS, f"[DOCX-EXPORT] to_lang={to_lang}, font_name={font_name}")
                else:
                    logger.info(LogModule.TRANS, f"[DOCX-EXPORT] to_lang={to_lang}, font_name=(default)")
            except Exception as _e:
                logger.debug(LogModule.RESTOR,f"[DOCX-EXPORT] Could not get font for to_lang={to_lang}: {_e}, using default")
        font_css = f"""
        <style>
            body {{
                font-size: 14pt;
                line-height: 1.5;{font_family_css}
            }}
            p, li, td, th {{
                font-size: 14pt;{font_family_css}
            }}
            h1 {{
                font-size: 24pt;{font_family_css}
            }}
            h2 {{
                font-size: 20pt;{font_family_css}
            }}
            h3 {{
                font-size: 18pt;{font_family_css}
            }}
            h4 {{
                font-size: 16pt;{font_family_css}
            }}
            h5 {{
                font-size: 14pt;{font_family_css}
            }}
            h6 {{
                font-size: 12pt;{font_family_css}
            }}
        </style>
        """
        
        # Insert CSS into HTML head section
        if "<head>" in sanitized_html:
            # Insert CSS after <head> tag
            sanitized_html = sanitized_html.replace("<head>", f"<head>{font_css}", 1)
        elif "<html>" in sanitized_html:
            # If no <head> tag, add one with CSS
            sanitized_html = sanitized_html.replace("<html>", f"<html><head>{font_css}</head>", 1)
        else:
            # If no HTML structure, wrap with HTML and add CSS
            sanitized_html = f"<!DOCTYPE html><html><head>{font_css}</head><body>{sanitized_html}</body></html>"
        
        # Use pypandoc with working directory set to output_dir so relative image paths work
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
            temp_html.write(sanitized_html)
            temp_html_path = temp_html.name
        
        try:
            # Convert HTML to DOCX
            # Set working directory to output_dir so Pandoc can find relative image paths
            import subprocess
            original_cwd = os.getcwd()
            try:
                os.chdir(str(output_dir))
                pypandoc.convert_file(
                    temp_html_path,
                    "docx",
                    format="html",
                    outputfile=output_path,
                )
            finally:
                os.chdir(original_cwd)
            
            output_file = Path(output_path)
            if not output_file.exists() or output_file.stat().st_size == 0:
                raise ValueError("Pandoc produced an empty DOCX file.")
            logger.info(LogModule.TRANS, f"[DOCX-EXPORT] Generated DOCX via pandoc: {output_path}")
            # Post-process: set font on all runs so Word uses our font instead of template default (e.g. 等线)
            if to_lang and font_family_css:
                try:
                    font_name = get_font_for_language(to_lang) if to_lang else None
                    if font_name:
                        _apply_font_to_docx_runs(output_path, font_name)
                        logger.info(LogModule.TRANS, f"[DOCX-EXPORT] Applied font '{font_name}' to all runs in DOCX (post-process)")
                except Exception as font_err:
                    logger.warning(LogModule.RESTOR,f"[DOCX-EXPORT] Post-process font apply failed: {font_err}, DOCX keeps Pandoc default font")
            try:
                apply_bilingual_styled_spans_to_docx(output_path, html_content)
            except Exception as style_err:
                logger.warning(
                    LogModule.RESTOR,
                    f"[DOCX-EXPORT] convert_html_to_docx bilingual span post-process failed: {style_err}",
                )
        finally:
            # Clean up temp HTML file
            try:
                os.unlink(temp_html_path)
            except:
                pass
    except Exception as e:
        logger.error(LogModule.RESTOR,f"[DOCX-EXPORT] Failed to convert HTML to DOCX via pandoc: {e}", exc_info=True)
        doc = Document()
        doc.add_paragraph("Owlangs translation could not be rendered from HTML.")
        doc.save(output_path)


def _html_table_block_to_pipe(html_block: str) -> str:
    """Convert one HTML <table>...</table> block to Pandoc pipe table (markdown).
    Used so Pandoc can render tables as DOCX tables instead of plain text.
    """
    try:
        from bs4 import BeautifulSoup
    except Exception:
        return html_block
    soup = BeautifulSoup(html_block, "html.parser")
    table = soup.find("table")
    if not table:
        return html_block
    header_cells: list[str] = []
    body_rows: list[list[str]] = []
    thead = table.find("thead")
    header_row = thead.find("tr") if thead else None
    if not header_row:
        header_row = table.find("tr")
    if header_row:
        for th in header_row.find_all(["th", "td"]):
            header_cells.append(th.get_text(strip=True))
    tbody = table.find("tbody")
    rows = tbody.find_all("tr") if tbody else table.find_all("tr")[1:]
    for tr in rows:
        row_cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
        if row_cells:
            body_rows.append(row_cells)
    if not header_cells:
        return html_block
    if not body_rows:
        body_rows = [[]]
    col_count = len(header_cells)
    normalized = []
    for row in body_rows:
        if len(row) < col_count:
            row = row + [""] * (col_count - len(row))
        elif len(row) > col_count:
            row = row[:col_count]
        normalized.append(row)

    def _escape(c: str) -> str:
        return c.replace("|", "\\|").replace("\n", " ")

    lines = ["| " + " | ".join(_escape(c) for c in header_cells) + " |"]
    lines.append("| " + " | ".join("---" for _ in range(col_count)) + " |")
    for row in normalized:
        lines.append("| " + " | ".join(_escape(c) for c in row) + " |")
    return "\n".join(lines)


def _html_tables_in_md_to_pipe_tables(md_content: str) -> str:
    """Replace all HTML <table>...</table> blocks in markdown with pipe tables for Pandoc DOCX."""
    if "<table" not in md_content.lower():
        return md_content
    result: list[str] = []
    rest = md_content
    while True:
        idx = rest.lower().find("<table")
        if idx < 0:
            result.append(rest)
            break
        result.append(rest[:idx])
        rest = rest[idx:]
        depth = 0
        start = 0
        i = 0
        while i < len(rest):
            if (i + 6 <= len(rest) and rest[i : i + 6].lower() == "<table"
                    and (i + 6 == len(rest) or rest[i + 6] in "> ")):
                depth += 1
                if depth == 1:
                    start = i
                i += 1
            elif i + 8 <= len(rest) and rest[i : i + 8].lower() == "</table>":
                depth -= 1
                if depth == 0:
                    block = rest[start : i + 8]
                    pipe = _html_table_block_to_pipe(block)
                    result.append(pipe)
                    rest = rest[i + 8 :]
                    break
                i += 8
            else:
                i += 1
        else:
            result.append(rest)
            break
    return "".join(result)


def _is_pipe_table_row(line: str) -> bool:
    """True if line looks like a pipe table row (header, separator, or body)."""
    s = line.strip()
    return len(s) >= 2 and s.startswith("|") and s.endswith("|")


def _normalize_pipe_table_separators(md_content: str) -> str:
    """Normalize pipe table separator lines so Pandoc recognizes them as tables.
    Layout/translation may produce em dash (—) or en dash (–); Pandoc requires ASCII hyphen (-).
    """
    lines = md_content.split("\n")
    out = []
    for line in lines:
        stripped = line.strip()
        # Pipe table separator: starts with |, contains only | spaces and dashes (and optional colons)
        if (len(stripped) >= 3 and stripped.startswith("|")
                and stripped.endswith("|")
                and not re.search(r"[^\s|\-:—–]", stripped)):
            # Replace Unicode dashes with ASCII hyphens (at least 3 per column for Pandoc)
            line = line.replace("\u2014", "---").replace("\u2013", "---")
        out.append(line)
    return "\n".join(out)


def _ensure_blank_line_before_pipe_tables(md_content: str) -> str:
    """Ensure a blank line before each pipe table so Pandoc parses it as a table, not paragraph text."""
    lines = md_content.split("\n")
    out = []
    prev_non_empty_was_table = False
    for line in lines:
        stripped = line.strip()
        is_table_row = _is_pipe_table_row(line)
        if is_table_row and not prev_non_empty_was_table and out:
            # This line starts a table; previous non-empty line was not a table row -> insert blank
            if out[-1].strip() != "":
                out.append("")
        out.append(line)
        if stripped:
            prev_non_empty_was_table = is_table_row
    return "\n".join(out)


def _sanitize_md_for_pdf(md_content: str) -> str:
    r"""Sanitize Markdown before Pandoc PDF conversion.

    Fixes common failure modes:
    1. Unclosed HTML <div> tags cause Pandoc to emit malformed LaTeX.
    2. Unicode blackboard-bold letters (ℝ, ℤ, ℕ, etc.) are converted by Pandoc
       into \symbb{...} which is only valid in math mode.
    3. Raw LaTeX math commands (\mathbb, \mathcal, \symbb, etc.) appearing
       outside $...$ math mode are passed verbatim to LaTeX, causing
       "allowed only in math mode" errors.

    Returns sanitized markdown string.
    """
    import re

    # 1. Strip HTML <div> tags (with or without attributes). They have no useful
    #    effect in PDF output and unclosed ones break Pandoc's LaTeX writer.
    md_content = re.sub(r'<div\b[^>]*>', '', md_content)
    md_content = re.sub(r'</div>', '', md_content)

    # 2. Replace Unicode blackboard-bold letters with plain ASCII equivalents.
    #    Covers both Letterlike Symbols (U+21xx) and Mathematical Alphanumeric
    #    Symbols (U+1D5xx / U+1D7xx) blocks. Pandoc maps these to \symbb{…}
    #    which is only valid in math mode.
    _bb_map = {
        # Letterlike Symbols
        '\u2102': 'C', '\u210D': 'H', '\u2115': 'N', '\u2119': 'P',
        '\u211A': 'Q', '\u211D': 'R', '\u2124': 'Z',
        # Mathematical Double-Struck Capital A–Y (missing C, H, N, P, Q, R, Z above)
        '\U0001d538': 'A', '\U0001d539': 'B', '\U0001d53b': 'D',
        '\U0001d53c': 'E', '\U0001d53d': 'F', '\U0001d53e': 'G',
        '\U0001d540': 'I', '\U0001d541': 'J', '\U0001d542': 'K',
        '\U0001d543': 'L', '\U0001d544': 'M', '\U0001d546': 'O',
        '\U0001d54a': 'S', '\U0001d54b': 'T', '\U0001d54c': 'U',
        '\U0001d54d': 'V', '\U0001d54e': 'W', '\U0001d54f': 'X',
        '\U0001d550': 'Y',
        # Mathematical Double-Struck Small a–z
        '\U0001d552': 'a', '\U0001d553': 'b', '\U0001d554': 'c',
        '\U0001d555': 'd', '\U0001d556': 'e', '\U0001d557': 'f',
        '\U0001d558': 'g', '\U0001d559': 'h', '\U0001d55a': 'i',
        '\U0001d55b': 'j', '\U0001d55c': 'k', '\U0001d55d': 'l',
        '\U0001d55e': 'm', '\U0001d55f': 'n', '\U0001d560': 'o',
        '\U0001d561': 'p', '\U0001d562': 'q', '\U0001d563': 'r',
        '\U0001d564': 's', '\U0001d565': 't', '\U0001d566': 'u',
        '\U0001d567': 'v', '\U0001d568': 'w', '\U0001d569': 'x',
        '\U0001d56a': 'y', '\U0001d56b': 'z',
        # Mathematical Double-Struck Digits 0–9
        '\U0001d7d8': '0', '\U0001d7d9': '1', '\U0001d7da': '2',
        '\U0001d7db': '3', '\U0001d7dc': '4', '\U0001d7dd': '5',
        '\U0001d7de': '6', '\U0001d7df': '7', '\U0001d7e0': '8',
        '\U0001d7e1': '9',
    }
    _bb_pattern = re.compile('[' + ''.join(_bb_map.keys()) + ']')
    md_content = _bb_pattern.sub(lambda m: _bb_map[m.group(0)], md_content)

    # 3. Escape raw LaTeX commands that appear OUTSIDE math mode.
    #    Pandoc's raw_tex extension passes \command verbatim to LaTeX.
    #    When \command is actually plain text (e.g. \htm from a rewrite rule
    #    or \hat outside math), LaTeX fails with "Undefined control sequence".
    #    We double the backslash so Pandoc treats it as literal text.
    _latex_cmd_re = re.compile(r'(?<!\\)\\[a-zA-Z]+')

    # Math-mode detectors: $$...$$, $...$, \(...\), \[...\]
    # $...$ requires at least one \command inside to avoid matching ordinary
    # dollar signs (e.g. price $5, regex end \$) that happen to have another
    # $ somewhere later in the same line.
    #
    # Do NOT require a non-letter before opening $ — algorithm text uses
    # CER$(...), f$(x), etc. Blocking those left \mathbf outside math and
    # doubled it to \\mathbf → XeLaTeX "There's no line here to end".
    # Closing $ must not be followed by a digit (Pandoc currency rule).
    # Optional spaces inside "$ ... $" cover pipe-table cells from OCR/LLM.
    _math_re = re.compile(
        r'(?<!\\)\$\$[\s\S]*?(?<!\\)\$\$'  # $$...$$
        r'|(?<!\\)\$\s*[^$\n]*?(?:\\[a-zA-Z]+)[^$\n]*?\s*(?<!\\)\$(?![0-9])'  # $...$ with \cmd
        r'|\\\([^)]*\\\)'  # \(...\)
        r'|\\\[[^\]]*\\\]'  # \[...\]
    )

    _parts: list[str] = []
    _last_end = 0
    for _m in _math_re.finditer(md_content):
        _before = md_content[_last_end:_m.start()]
        _before = _latex_cmd_re.sub(lambda mm: '\\' + mm.group(0), _before)
        _parts.append(_before)
        _parts.append(_m.group(0))
        _last_end = _m.end()

    _after = md_content[_last_end:]
    _after = _latex_cmd_re.sub(lambda mm: '\\' + mm.group(0), _after)
    _parts.append(_after)

    return ''.join(_parts)


def convert_md_to_docx(
    md_content: str,
    output_path: str,
    output_dir: Optional[Path] = None,
    to_lang: Optional[str] = None,
) -> bool:
    """
    Convert Markdown to DOCX via Pandoc; apply font by target language. Used for pandoc-first DOCX flow.

    Returns:
        True if conversion succeeded, False if Pandoc is missing or conversion failed (caller should fallback).
    """
    if not (md_content and md_content.strip()):
        logger.debug(LogModule.RESTOR, "[DOCX-EXPORT] convert_md_to_docx: empty md_content, skip")
        return False
    md_content = normalize_md_math_for_pandoc_export(md_content)
    md_content = normalize_docx_markdown_sup_sub(md_content)
    md_content = wrap_bare_latex_for_docx_algorithms(md_content)
    try:
        import pypandoc
    except ImportError:
        logger.debug(LogModule.RESTOR, "[DOCX-EXPORT] convert_md_to_docx: pypandoc not installed, skip")
        return False
    pandoc_path = _get_pandoc_path()
    if not pandoc_path:
        logger.debug(LogModule.RESTOR, "[DOCX-EXPORT] convert_md_to_docx: Pandoc not found in 3rdParty, skip")
        return False
    os.environ["PYPANDOC_PANDOC"] = str(pandoc_path)
    out_dir = output_dir or Path(output_path).parent
    out_dir.mkdir(parents=True, exist_ok=True)
    # Replace HTML tables with pipe tables so Pandoc outputs DOCX tables instead of plain text
    md_for_pandoc = _html_tables_in_md_to_pipe_tables(md_content)
    # Normalize separator lines: replace em dash (—) / en dash (–) with ASCII hyphen so Pandoc recognizes tables
    md_for_pandoc = _normalize_pipe_table_separators(md_for_pandoc)
    # Ensure blank line before each pipe table so Pandoc parses tables instead of treating as paragraph
    md_for_pandoc = _ensure_blank_line_before_pipe_tables(md_for_pandoc)
    # Resolve data URI image refs to files under output_dir so Pandoc can embed them in DOCX
    images_dir = out_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    _data_uri_pattern = re.compile(r"!\[([^\]]*)\]\((data:image/[^)]+)\)")

    def _replace_data_uri(match: re.Match) -> str:
        alt_text, data_uri = match.group(1), match.group(2)
        if not data_uri.startswith("data:image/") or "," not in data_uri:
            return match.group(0)
        try:
            header, b64 = data_uri.split(",", 1)
            mime = header.split(";")[0].split(":")[-1] if ":" in header else "image/png"
            ext = mimetypes.guess_extension(mime) or ".png"
            raw = base64.b64decode(b64)
            name = hashlib.md5(raw[:500]).hexdigest()[:8] + ext
            path = images_dir / name
            path.write_bytes(raw)
            rel = f"./images/{name}"
            logger.debug(LogModule.RESTOR, f"[DOCX-EXPORT] Resolved data URI image to file: {rel}")
            return f"![{alt_text}]({rel})"
        except Exception as e:
            logger.warning(LogModule.RESTOR, f"[DOCX-EXPORT] Failed to resolve data URI image: {e}")
            return match.group(0)

    md_for_pandoc = _data_uri_pattern.sub(_replace_data_uri, md_for_pandoc)
    # Debug: write MD input to output/debug for DOCX export debugging (Pandoc path)
    try:
        debug_dir = out_dir / "debug"
        debug_dir.mkdir(parents=True, exist_ok=True)
        (debug_dir / "docx_export_input_pandoc.md").write_text(md_for_pandoc, encoding="utf-8")
        logger.debug(LogModule.RESTOR, f"[DOCX-EXPORT] Wrote Pandoc MD input to {debug_dir / 'docx_export_input_pandoc.md'}")
    except Exception as e:
        logger.warning(LogModule.RESTOR, f"[DOCX-EXPORT] Failed to write debug MD for Pandoc: {e}")
    # Use +hard_line_breaks so single newlines become line breaks (avoid merging references/short segments)
    # +raw_html so remaining <sup>/<sub> etc. pass through to DOCX OMML when not normalized to Unicode
    pandoc_format = "markdown+pipe_tables+hard_line_breaks+raw_html"
    import tempfile
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(md_for_pandoc)
        temp_md = f.name
    try:
        original_cwd = os.getcwd()
        try:
            os.chdir(str(out_dir))
            pypandoc.convert_file(temp_md, "docx", format=pandoc_format, outputfile=output_path)
        finally:
            os.chdir(original_cwd)
        out_file = Path(output_path)
        if not out_file.exists() or out_file.stat().st_size == 0:
            logger.warning(LogModule.RESTOR, "[DOCX-EXPORT] convert_md_to_docx: Pandoc produced empty file")
            return False
        if to_lang:
            try:
                from translator.ai_translator.docx_translator import get_font_for_language
                font_name = get_font_for_language(to_lang)
                if font_name:
                    _apply_font_to_docx_runs(output_path, font_name)
            except Exception as font_err:
                logger.warning(LogModule.RESTOR, f"[DOCX-EXPORT] convert_md_to_docx font post-process failed: {font_err}")
        try:
            apply_bilingual_styled_spans_to_docx(output_path, md_content)
        except Exception as style_err:
            logger.warning(
                LogModule.RESTOR,
                f"[DOCX-EXPORT] convert_md_to_docx bilingual span post-process failed: {style_err}",
            )
        logger.info(LogModule.TRANS, f"[DOCX-EXPORT] convert_md_to_docx succeeded: {output_path}")
        return True
    except Exception as e:
        logger.warning(LogModule.RESTOR, f"[DOCX-EXPORT] convert_md_to_docx failed: {e}", exc_info=False)
        return False
    finally:
        try:
            os.unlink(temp_md)
        except Exception:
            pass


class PdfExportLatexError(RuntimeError):
    """
    Raised when Pandoc+XeLaTeX PDF export fails and we can extract a local LaTeX/Markdown context.
    Intended for callers to surface a segment-level hint to the frontend, instead of auto-repairing.
    """

    def __init__(
        self,
        message: str,
        *,
        stderr: str = "",
        error_type: str = "unknown",
        line_no: Optional[int] = None,
        error_token: str = "",
        tex_snippet: str = "",
        md_snippet: str = "",
        debug_tex_path: Optional[Path] = None,
        debug_md_path: Optional[Path] = None,
    ) -> None:
        super().__init__(message)
        self.stderr = stderr
        self.error_type = error_type
        self.line_no = line_no
        self.error_token = error_token
        self.tex_snippet = tex_snippet
        self.md_snippet = md_snippet
        self.debug_tex_path = debug_tex_path
        self.debug_md_path = debug_md_path


def convert_md_to_pdf(
    md_content: str,
    output_path: str,
    output_dir: Optional[Path] = None,
    to_lang: Optional[str] = None,
    image_block_indices: Optional[List[int]] = None,
    path_to_block_index: Optional[Dict[str, int]] = None,
    layout_document: Optional[Any] = None,
    layout_block_bbox: Optional[Dict[int, Tuple[float, float, float, float]]] = None,
) -> bool:
    """
    Convert Markdown to PDF via Pandoc with XeLaTeX; mainfont is chosen by target language. Used for pandoc-first PDF flow.

    When path_to_block_index and layout_document are provided, only consecutive images on the same row (bbox y overlap) are laid out side-by-side in LaTeX.
    path_to_block_index maps normalized image paths (filename, lowercase) to layout block indices.

    Returns:
        True if conversion succeeded, False if conversion failed (e.g. pandoc exit non-zero).

    Raises:
        RuntimeError: If Pandoc or XeLaTeX is not installed or not found (clear error message and server log for debugging).
    """
    if not (md_content and md_content.strip()):
        logger.debug(LogModule.RESTOR, "[PDF-EXPORT] convert_md_to_pdf: empty md_content, skip")
        return False
    md_content = normalize_md_math_for_pandoc_export(md_content)
    # Keep tagged formulas as $$...\\tag{n}...$$. Do NOT rewrite to
    # \\begin{equation}: convert_md_to_pdf uses markdown without raw_tex, and
    # _sanitize_md_for_pdf only protects $$/$ math — equation envs get
    # backslash-doubled and XeLaTeX fails with Missing $.
    pandoc_path = _get_pandoc_path()
    if not pandoc_path:
        logger.error(
            LogModule.RESTOR,
            "[PDF-EXPORT] Pandoc is required for PDF export but was not found. "
            "Please install Pandoc and ensure it is available in PATH or in 3rdParty/windows (e.g. pandoc-3.x)."
        )
        raise RuntimeError(
            "Pandoc is required for PDF export but was not found. "
            "Install Pandoc and ensure it is in PATH or in 3rdParty/windows. See server logs for details."
        )
    os.environ["PYPANDOC_PANDOC"] = str(pandoc_path)
    xelatex_path_orig = _get_xelatex_path()
    xelatex_path = None
    pdflatex_root_use = None
    if xelatex_path_orig:
        pdflatex_root = xelatex_path_orig.parent.parent.parent
        pdflatex_root_use = _ensure_ascii_path_for_tex(pdflatex_root)
        xelatex_path = pdflatex_root_use / "bin" / "windows" / "xelatex.exe"
        if not xelatex_path.exists():
            xelatex_path = _to_short_path_if_needed(xelatex_path_orig)
            pdflatex_root_use = xelatex_path.parent.parent.parent
    if not xelatex_path:
        xelatex_in_path = shutil.which("xelatex")
        
        # 检查 macOS 上常见的 xelatex 路径
        if not xelatex_in_path and sys.platform == "darwin":
            common_paths = [
                "/Library/TeX/texbin/xelatex",  # MacTeX / BasicTeX
                "/usr/local/texlive/current/bin/universal-darwin/xelatex",  # TeX Live
                "/usr/local/texlive/2026/bin/universal-darwin/xelatex",  # TeX Live 2026
            ]
            for path in common_paths:
                if Path(path).exists():
                    xelatex_in_path = path
                    logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Found XeLaTeX at: {xelatex_in_path}")
                    break
        
        if not xelatex_in_path:
            # 根据操作系统提供不同的安装指导
            if sys.platform == "darwin":
                install_msg = (
                    "XeLaTeX is required for PDF export but was not found. "
                    "Install TeX Live: brew install --cask mactex "
                    "or TinyTeX: brew install --cask tinytex. "
                    "After installation, run: cd 3rdParty/macos && ./install_latex_packages.sh "
                    "to install required LaTeX packages. See server logs for details."
                )
            elif sys.platform == "linux":
                install_msg = (
                    "XeLaTeX is required for PDF export but was not found. "
                    "Install TeX Live: sudo apt-get install texlive-xetex texlive-lang-chinese "
                    "or TinyTeX. See server logs for details."
                )
            else:  # Windows
                install_msg = (
                    "XeLaTeX is required for PDF export but was not found. "
                    "Please install XeLaTeX (e.g. TeX Live, TinyTeX) and ensure the xelatex executable is in PATH or in 3rdParty/windows/pdflatex."
                    "See server logs for details."
                )
            logger.error(
                LogModule.RESTOR,
                f"[PDF-EXPORT] {install_msg}"
            )
            raise RuntimeError(install_msg)
    pdf_engine = str(xelatex_path) if xelatex_path else (xelatex_in_path if xelatex_in_path else "xelatex")
    import tempfile
    import subprocess
    try:
        from translator.ai_translator.docx_translator import get_font_for_language
        mainfont = get_font_for_language(to_lang) if to_lang else ("Helvetica Neue" if sys.platform == "darwin" else "Calibri")
    except Exception:
        mainfont = "Helvetica Neue" if sys.platform == "darwin" else "Calibri"
    # CJK languages need xeCJK for proper line breaking (otherwise translated lines overflow)
    _cjk_codes = ("zh", "chinese", "zh-cn", "zh-tw", "ja", "japanese", "jp", "ko", "korean", "kr")
    to_lang_lower = (to_lang or "").strip().lower()
    use_xecjk = any(to_lang_lower.startswith(c) or to_lang_lower == c for c in _cjk_codes)
    
    # On macOS, check if required LaTeX packages are installed
    if sys.platform == "darwin":
        _check_latex_packages_macos()
    
    out_dir = output_dir or Path(output_path).parent
    out_dir.mkdir(parents=True, exist_ok=True)
    # Run Pandoc/XeLaTeX from the system temp directory so auxiliary files (.aux, .log, .xdv)
    # are written to a guaranteed-writable location instead of the possibly-read-only out_dir.
    temp_work_dir = Path(tempfile.gettempdir())
    output_path = str(Path(output_path).resolve())
    # Copy images into temp_work_dir so xelatex (running with cwd=temp_work_dir) can resolve
    # relative paths like ./images/... used in both markdown and raw LaTeX side-by-side blocks.
    images_dir = out_dir / "images"
    if images_dir.exists():
        temp_images_dir = temp_work_dir / "images"
        if temp_images_dir.exists():
            shutil.rmtree(temp_images_dir)
        shutil.copytree(images_dir, temp_images_dir)
    # Same table preprocessing as DOCX: HTML->pipe, normalize separators, blank line before tables
    md_for_pdf = _html_tables_in_md_to_pipe_tables(md_content)
    md_for_pdf = _normalize_pipe_table_separators(md_for_pdf)
    md_for_pdf = _ensure_blank_line_before_pipe_tables(md_for_pdf)
    # Sanitize for PDF: strip unclosed HTML divs and replace Unicode math symbols that break LaTeX
    md_for_pdf = _sanitize_md_for_pdf(md_for_pdf)

    # Remove alt text from all images to prevent pandoc from generating "Figure n: ..." captions
    # Pattern: ![any alt text](path) -> ![](path)
    import re
    import base64
    import hashlib
    import mimetypes
    image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

    # Convert data URI image refs to file paths so we can read dimensions and avoid LaTeX default width (stretch)
    # Rebuilt MD can contain ![](data:image/...;base64,...) when segment content had data URIs and image_data_map
    # has no entry for that URI; those images were not replaced earlier and would render at default width (stretched)
    images_dir = out_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    data_uri_pattern = re.compile(r'!\[([^\]]*)\]\((data:image/[^)]+)\)')

    def _replace_data_uri_with_file(match):
        alt_text, data_uri = match.group(1), match.group(2)
        if not data_uri.startswith("data:image/") or "," not in data_uri:
            return match.group(0)
        try:
            header, b64 = data_uri.split(",", 1)
            mime = header.split(";")[0].split(":")[-1] if ":" in header else "image/png"
            ext = mimetypes.guess_extension(mime) or ".png"
            raw = base64.b64decode(b64)
            name = hashlib.md5(raw[:500]).hexdigest()[:8] + ext
            path = images_dir / name
            path.write_bytes(raw)
            rel = f"./images/{name}"
            logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Resolved data URI image to file: {rel}")
            return f"![{alt_text}]({rel})"
        except Exception as e:
            logger.warning(LogModule.RESTOR, f"[PDF-EXPORT] Failed to resolve data URI image: {e}")
            return match.group(0)

    md_for_pdf = data_uri_pattern.sub(_replace_data_uri_with_file, md_for_pdf)

    # Process images: remove alt text and add size attributes to preserve original dimensions
    image_refs = image_pattern.findall(md_for_pdf)
    logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Found {len(image_refs)} image references in markdown before PDF conversion")
    
    # Build a mapping of image paths to their dimensions (in cm for LaTeX)
    # All images from Extract phase use 96 DPI, so use 96 DPI for all images
    image_sizes = {}
    # Paths classified as formula/table: never merge side-by-side (used when alt is missing after transforms)
    formula_or_table_paths = set()
    max_width_cm = 16.0   # A4 usable width
    max_height_cm = 22.0  # A4 usable height (~29.7 - margins) so image does not overflow vertically
    dpi_default = 96.0    # All images use 96 DPI (from Extract phase)

    for alt_text, img_path in image_refs:
        if img_path in image_sizes:
            continue  # Already processed
        # Try to resolve image path (avoid resolve() to prevent long paths)
        img_file_path = None
        if img_path.startswith('./') or img_path.startswith('../'):
            img_file_path = out_dir / img_path.lstrip('./').lstrip('../')
        elif not os.path.isabs(img_path):
            img_file_path = out_dir / img_path
        else:
            img_file_path = Path(img_path)

        # Determine image type for logging and for side-by-side exclusion (formula/table never merged)
        is_formula_or_table = alt_text and (
            "equation" in alt_text.lower() or "formula" in alt_text.lower() or "table" in alt_text.lower()
        )
        if is_formula_or_table:
            formula_or_table_paths.add(img_path)
        image_type = "formula/table" if is_formula_or_table else "original"

        if img_file_path and img_file_path.exists():
            try:
                from PIL import Image
                pil_image = Image.open(str(img_file_path))
                img_width_px, img_height_px = pil_image.size

                # Log original dimensions first (before any processing)
                logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Image original size: path={img_path}, type={image_type}, {img_width_px}x{img_height_px}px")

                # Prefer image metadata DPI when present; otherwise use 96 DPI (from Extract phase)
                dpi_x = pil_image.info.get('dpi', (None, None))[0] if 'dpi' in pil_image.info else None
                dpi_y = pil_image.info.get('dpi', (None, None))[1] if 'dpi' in pil_image.info else None
                if dpi_x and dpi_x > 0 and dpi_y and dpi_y > 0:
                    dpi = (dpi_x + dpi_y) / 2.0
                    logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Image DPI from metadata: path={img_path}, DPI={dpi:.1f}")
                else:
                    dpi = dpi_default
                    logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Image using default DPI: path={img_path}, DPI={dpi:.1f}")

                # Convert pixels to cm: 1 inch = 2.54 cm, 1 inch = dpi px
                width_cm = (img_width_px * 2.54) / dpi
                height_cm = (img_height_px * 2.54) / dpi

                # Scale to fit within page (both width and height) keeping aspect ratio
                original_width_cm, original_height_cm = width_cm, height_cm
                scale_w = max_width_cm / width_cm if width_cm > max_width_cm else 1.0
                scale_h = max_height_cm / height_cm if height_cm > max_height_cm else 1.0
                scale = min(scale_w, scale_h, 1.0)
                if scale < 1.0:
                    width_cm = width_cm * scale
                    height_cm = height_cm * scale
                    logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Image scaled to fit page: path={img_path}, type={image_type}, orig_cm={original_width_cm:.2f}x{original_height_cm:.2f}, final={width_cm:.2f}x{height_cm:.2f}cm")

                image_sizes[img_path] = (width_cm, height_cm, img_width_px, img_height_px, dpi)
                logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Image final size: path={img_path}, type={image_type}, {img_width_px}x{img_height_px}px @ DPI={dpi:.1f}, size={width_cm:.2f}x{height_cm:.2f}cm")
            except Exception as img_err:
                logger.warning(LogModule.RESTOR, f"[PDF-EXPORT] Failed to get image dimensions for {img_path}, type={image_type}: {img_err}")
        else:
            # Log when image file is not found (path may be data URI if not resolved earlier)
            path_preview = img_path[:80] + "..." if isinstance(img_path, str) and len(img_path) > 80 else img_path
            logger.warning(LogModule.RESTOR, f"[PDF-EXPORT] Image file not found: path={path_preview}, type={image_type}, resolved_path={img_file_path}")
    
    # Group consecutive image refs (only whitespace between them) for side-by-side layout
    # When layout_document + path_to_block_index are provided, only merge runs that are on the same row (bbox y overlap).
    has_layout = layout_document is not None and path_to_block_index is not None and len(path_to_block_index) >= 2
    logger.info(
        LogModule.RESTOR,
        f"[PDF-EXPORT] Layout for side-by-side: layout_document={layout_document is not None}, "
        f"path_to_block_index_len={len(path_to_block_index) if path_to_block_index else 0}, has_layout={has_layout}"
    )
    # Prefer bbox from Layout extraction phase (layout_block_bbox) so we do not iterate layout_document at export.
    # Normalize keys to int and bbox to tuple of float (task_state may have been JSON round-trip with string keys/values).
    block_index_to_bbox: Optional[Dict[int, Tuple[float, float, float, float]]] = None
    if layout_block_bbox:
        try:
            block_index_to_bbox = {}
            for k, v in layout_block_bbox.items():
                if v is None or len(v) < 4:
                    continue
                bidx = int(k) if not isinstance(k, int) else k
                block_index_to_bbox[bidx] = (float(v[0]), float(v[1]), float(v[2]), float(v[3]))
        except (TypeError, ValueError, IndexError) as e:
            logger.debug(LogModule.RESTOR, f"[PDF-EXPORT] Normalize layout_block_bbox failed: {e}, will build from layout_document")
            block_index_to_bbox = None
    if has_layout and not block_index_to_bbox:
        try:
            from layout.base import LayoutDocument as _LD
            if isinstance(layout_document, _LD):
                block_index_to_bbox = {}
                for block in layout_document.iter_blocks():
                    if block.index is not None and hasattr(block, "bbox") and block.bbox:
                        block_index_to_bbox[block.index] = block.bbox
                logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Built block_index_to_bbox with {len(block_index_to_bbox)} blocks")
            else:
                logger.info(LogModule.RESTOR, f"[PDF-EXPORT] layout_document is not LayoutDocument instance, type={type(layout_document).__name__}")
        except Exception as e:
            logger.warning(LogModule.RESTOR, f"[PDF-EXPORT] Failed to build block_index_to_bbox: {e}")
            block_index_to_bbox = None

    all_matches = list(image_pattern.finditer(md_for_pdf))
    runs = []
    i = 0
    while i < len(all_matches):
        run = [all_matches[i]]
        while i + 1 < len(all_matches):
            gap = md_for_pdf[run[-1].end() : all_matches[i + 1].start()]
            if not gap.strip():  # only whitespace between
                run.append(all_matches[i + 1])
                i += 1
            else:
                break
        runs.append(run)
        i += 1

    # Total row width for side-by-side (must fit in text block)
    # Increased from 12.0 to 14.0 to allow wider side-by-side images
    row_width_cm = 14.0
    # For runs of 2+ that we do NOT merge, still cap each image width so they don't overflow when laid out
    constrained_width_cm: Dict[str, float] = {}

    # Replace runs of 2+ images with one LaTeX figure (side-by-side) only when same row if layout provided
    # Formula images are never merged side-by-side (they are always rendered on separate lines)
    parts = []
    pos = 0
    image_index = 0  # Track position in image_block_indices list
    for run in runs:
        parts.append(md_for_pdf[pos : run[0].start()])
        
        # Separate formula/table images from regular images in this run
        # Formula/table images must always be rendered individually, never side-by-side
        # Use alt text first; fall back to path in formula_or_table_paths (set when building image_sizes)
        # so we still treat as formula when alt is missing after data-URI or other transforms
        formula_images = []
        regular_images = []
        for m in run:
            alt_text = m.group(1) or ""
            img_path = m.group(2)
            by_alt = alt_text and (
                "equation" in alt_text.lower() or "formula" in alt_text.lower() or "table" in alt_text.lower()
            )
            by_path = img_path in formula_or_table_paths
            is_formula_or_table = by_alt or by_path
            if by_path and not by_alt:
                logger.debug(
                    LogModule.RESTOR,
                    f"[PDF-EXPORT] Formula/table image identified by path (alt missing or empty): path={img_path[:60]}..."
                )
            if is_formula_or_table:
                formula_images.append(m)
            else:
                regular_images.append(m)
        
        # Process formula/table images first - each on a separate line, never side-by-side
        if formula_images:
            logger.info(
                LogModule.RESTOR,
                f"[PDF-EXPORT] Found {len(formula_images)} formula/table image(s) in run, rendering individually (not side-by-side)"
            )
            for j, m in enumerate(formula_images):
                if j > 0:
                    parts.append("\n\n")  # Force each formula/table image on its own paragraph
                parts.append(m.group(0))
                image_index += 1
        
        # Process regular images - can be merged side-by-side if conditions are met
        if not regular_images:
            # Only formula images in this run, continue to next run
            pos = run[-1].end()
            continue
        
        # For regular images: enable side-by-side merge when same row (formula/table images handled separately above)
        merge_run = len(regular_images) >= 2
        if merge_run and block_index_to_bbox is not None and image_block_indices is not None:
            run_block_indices = []
            # Try path-based matching first, fallback to order-based matching
            # Only process regular images (formula images already handled above)
            for m in regular_images:
                img_path = m.group(2)
                block_idx = None
                
                # Method 1: Try path-based matching
                if path_to_block_index:
                    normalized = img_path.replace("./", "").replace("images/", "").replace("images\\", "")
                    filename = os.path.basename(normalized).lower()
                    block_idx = path_to_block_index.get(filename)
                
                # Method 2: Fallback to order-based matching (use image_index to get block index from image_block_indices)
                if block_idx is None and image_index < len(image_block_indices):
                    block_idx = image_block_indices[image_index]
                    logger.info(
                        LogModule.RESTOR,
                        f"[PDF-EXPORT] Using order-based mapping: img_path={img_path}, image_index={image_index}, block_index={block_idx}"
                    )
                
                if block_idx is not None:
                    run_block_indices.append(block_idx)
                else:
                    logger.info(
                        LogModule.RESTOR,
                        f"[PDF-EXPORT] Image path not found in mapping: img_path={img_path}, image_index={image_index}, "
                        f"path_mapping_keys={list(path_to_block_index.keys())[:5] if path_to_block_index else []}, "
                        f"image_block_indices_len={len(image_block_indices) if image_block_indices else 0}"
                    )
                image_index += 1
            same_row = len(run_block_indices) == len(regular_images)
            if same_row:
                for j in range(len(run_block_indices) - 1):
                    bbox_a = block_index_to_bbox.get(run_block_indices[j])
                    bbox_b = block_index_to_bbox.get(run_block_indices[j + 1])
                    if bbox_a is None or bbox_b is None or not _bbox_y_overlap(bbox_a, bbox_b):
                        same_row = False
                        logger.info(
                            LogModule.RESTOR,
                            f"[PDF-EXPORT] Consecutive images NOT same row (no merge): blocks {run_block_indices[j]}, {run_block_indices[j + 1]}"
                        )
                        break
            merge_run = merge_run and same_row
            logger.info(
                LogModule.RESTOR,
                f"[PDF-EXPORT] Run of {len(regular_images)} regular images: run_block_indices={run_block_indices}, same_row={same_row}, merge_run={merge_run}"
            )
        else:
            # Update image_index even when not merging (only for regular images, formula images already processed)
            image_index += len(regular_images)

        n_run = len(regular_images)
        # Same-row merge: LaTeX uses w_frac (e.g. 0.48\\textwidth each) so images are scaled to fit one row; do not split by natural total width
        width_cm_per = row_width_cm / n_run if n_run >= 2 else None
        # Constrain width for consecutive regular images that are NOT merged (to prevent Pandoc from auto-laying them side-by-side and overflowing)
        # Even if layout says different rows, Pandoc may still place them side-by-side, so we constrain width to be safe
        # Formula images are already handled separately above and never constrained
        if n_run >= 2 and not merge_run:
            for m in regular_images:
                constrained_width_cm[m.group(2)] = width_cm_per
            if block_index_to_bbox is None or path_to_block_index is None:
                logger.info(LogModule.RESTOR, f"[PDF-EXPORT] No layout: constrained {n_run} consecutive regular images to {width_cm_per:.2f}cm each")
            else:
                logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Layout says different rows but constraining {n_run} consecutive regular images to {width_cm_per:.2f}cm each (prevent Pandoc auto-side-by-side overflow)")

        if merge_run:
            n = n_run
            paths = [m.group(2) for m in regular_images]
            paths_safe = [p.replace("\\", "/").replace("_", "\\_") for p in paths]
            # Use fraction of \textwidth so total fits regardless of template; leave small gap between images
            # Increased from 0.46 to 0.48 to allow wider side-by-side images
            w_frac = min(0.48, 0.95 / n)
            incl = "".join(
                f"\\includegraphics[width={w_frac:.2f}\\textwidth,keepaspectratio]{{{p}}}"
                + ("\\hfill\n  " if j < n - 1 else "")
                for j, p in enumerate(paths_safe)
            )
            latex_block = f"\\begin{{figure}}[h]\n\\centering\n  {incl}\n\\end{{figure}}"
            parts.append(latex_block)
            logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Side-by-side figure: {n} regular images (same row), {w_frac:.2f}\\textwidth each")
        else:
            # Add regular images individually, each in its own paragraph so Pandoc does not lay them side-by-side
            for j, m in enumerate(regular_images):
                if j > 0:
                    parts.append("\n\n")
                parts.append(m.group(0))
        pos = run[-1].end()
    parts.append(md_for_pdf[pos:])
    md_for_pdf = "".join(parts)

    # Replace single images: remove alt text and add size attribute. Use constrained width for consecutive (non-merged) runs.
    def _replace_image_with_size(match):
        alt_text = match.group(1)
        img_path = match.group(2)
        if img_path in constrained_width_cm:
            w = constrained_width_cm[img_path]
            return f"![]({img_path}){{width={w:.2f}cm}}"
        if img_path in image_sizes:
            width_cm, height_cm, width_px, height_px, dpi = image_sizes[img_path]
            return f"![]({img_path}){{width={width_cm:.2f}cm}}"
        return f"![]({img_path})"

    md_for_pdf = image_pattern.sub(_replace_image_with_size, md_for_pdf)
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(md_for_pdf)
        temp_md = f.name
    env = os.environ.copy()
    if xelatex_path and pdflatex_root_use is not None:
        env["PATH"] = str(xelatex_path.parent) + os.pathsep + env.get("PATH", "")
        # TEXMFCNF: dirs for texmf.cnf; include web2c so base config (TEXMFSYSVAR etc.) is read
        env["TEXMFCNF"] = str(pdflatex_root_use) + os.pathsep + str(pdflatex_root_use / "texmf-dist" / "web2c")
        env["TEXMFROOT"] = str(pdflatex_root_use)
        # Use a user-writable texmf-var (outside Program Files) to avoid permission denied
        user_texmfvar = _get_user_texmfvar_dir()
        user_texmfvar.mkdir(parents=True, exist_ok=True)
        texmfvar = str(user_texmfvar)
        env["TEXMFVAR"] = texmfvar
        env["TEXMFSYSVAR"] = texmfvar  # mktexfmt/fmtutil use this; must match bundle texmf-var
        env["TEXMFOUTPUT"] = str(temp_work_dir)  # fallback output dir for TeX security mode
        # Fontconfig: point XeLaTeX to bundled fonts.conf and user-writable cache dir
        # so it does not fail with "Cannot load default config file" or permission errors.
        fontconfig_file = pdflatex_root_use / "texmf-var" / "fonts" / "conf" / "fonts.conf"
        if fontconfig_file.exists():
            env["FONTCONFIG_FILE"] = str(fontconfig_file)
            fc_cache_dir = user_texmfvar / "fontconfig-cache"
            fc_cache_dir.mkdir(parents=True, exist_ok=True)
            env["FC_CACHEDIR"] = str(fc_cache_dir)
        _ensure_xelatex_fmt(pdflatex_root_use, env)
    # Geometry: constrain text to page and avoid overflow (default template may use small margins)
    geometry_opts = "margin=2.5cm"
    # For CJK (e.g. Chinese): xeCJK enables line breaking so translated text does not overflow
    cjk_preamble = ""
    if use_xecjk:
        # Escape font name for LaTeX (braces in template; escape literal { } in name)
        cjk_font = mainfont.replace("\\", "").replace("}", "\\}").replace("{", "\\{")
        cjk_preamble = f"\\usepackage{{xeCJK}}\\setCJKmainfont{{{cjk_font}}}"
    # ragged2e: better paragraph alignment and line breaks; titlesec: larger title/section fonts
    # Use etoolbox \AtEndPreamble so section formatting runs last and is not overridden by template
    # graphicx: control image sizing (pandoc width/height attributes will override defaults)
    header_includes = (
        cjk_preamble +
        "\\usepackage{ragged2e}\\AtBeginDocument{\\RaggedRight}"
        "\\PassOptionsToPackage{hyphens}{url}\\usepackage{hyperref}\\hypersetup{breaklinks=true}"
        "\\usepackage{titlesec}"
        "\\usepackage{graphicx}"
        "\\usepackage{etoolbox}"
        "\\makeatletter"
        "\\renewcommand{\\@maketitle}{\\begin{center}\\LARGE\\bfseries\\@title\\par\\vskip 0.5em\\large\\@author\\par\\vskip 0.3em\\normalsize\\@date\\end{center}\\par\\vskip 1em}"
        "\\makeatother"
        "\\AtEndPreamble{"
        "\\titleformat*{\\section}{\\LARGE\\bfseries}"
        "\\titleformat*{\\subsection}{\\Large\\bfseries}"
        "\\titleformat*{\\subsubsection}{\\large\\bfseries}"
        "}"
        "\\AtBeginDocument{"
        "\\sloppy\\setlength{\\emergencystretch}{5em}"
        "}"
    )
    # Input format: +hard_line_breaks so single newlines in markdown become line breaks in PDF
    # +link_attributes enables image size attributes like {width=XXcm}
    pandoc_from = "markdown+pipe_tables+hard_line_breaks+link_attributes"
    header_without_cjk = (
        "\\usepackage{ragged2e}\\AtBeginDocument{\\RaggedRight}"
        "\\PassOptionsToPackage{hyphens}{url}\\usepackage{hyperref}\\hypersetup{breaklinks=true}"
        "\\usepackage{titlesec}"
        "\\usepackage{graphicx}"
        "\\usepackage{etoolbox}"
        "\\makeatletter"
        "\\renewcommand{\\@maketitle}{\\begin{center}\\LARGE\\bfseries\\@title\\par\\vskip 0.5em\\large\\@author\\par\\vskip 0.3em\\normalsize\\@date\\end{center}\\par\\vskip 1em}"
        "\\makeatother"
        "\\AtEndPreamble{"
        "\\titleformat*{\\section}{\\LARGE\\bfseries}"
        "\\titleformat*{\\subsection}{\\Large\\bfseries}"
        "\\titleformat*{\\subsubsection}{\\large\\bfseries}"
        "}"
        "\\AtBeginDocument{"
        "\\sloppy\\setlength{\\emergencystretch}{5em}"
        "}"
    )
    attempts = [(header_includes, "with xeCJK" if use_xecjk else "default")]
    if use_xecjk:
        attempts.append((header_without_cjk, "without xeCJK (fallback)"))

    # DEBUG: write intermediate Markdown to a debug directory under out_dir so we
    # can inspect the exact inputs Pandoc sees when PDF export fails.
    # Note: generating debug LaTeX requires an extra pandoc run; we only do it on failure.
    debug_md_path: Optional[Path] = None
    debug_tex_path: Optional[Path] = None

    try:
        debug_dir = out_dir / "debug"
        debug_dir.mkdir(parents=True, exist_ok=True)
        debug_md_path = debug_dir / (Path(output_path).stem + ".md")
        try:
            with open(debug_md_path, "w", encoding="utf-8") as f_md:
                f_md.write(md_for_pdf)
            logger.info(
                LogModule.RESTOR,
                "[PDF-EXPORT] Wrote debug Markdown to {path}",
                path=str(debug_md_path),
            )
        except Exception as md_err:  # noqa: BLE001
            logger.debug(
                LogModule.RESTOR,
                f"[PDF-EXPORT] Writing debug Markdown failed: {md_err}",
            )
    except Exception as tex_err:
        logger.debug(
            LogModule.RESTOR,
            f"[PDF-EXPORT] Writing debug artifacts failed: {tex_err}",
        )

    try:
        for current_header, attempt_name in attempts:
            cmd = [
                str(pandoc_path),
                "-f", pandoc_from,
                temp_md,
                "-o", output_path,
                "--pdf-engine=" + pdf_engine,
                "-V", f"mainfont={mainfont}",
                "-V", f"geometry={geometry_opts}",
                "-V", f"papersize=a4",
                "-V", f"header-includes={current_header}",
                "--resource-path", str(out_dir),
            ]
            proc = subprocess.run(
                cmd,
                cwd=str(temp_work_dir),
                env=env,
                capture_output=True,
                timeout=300,
                encoding="utf-8",
                errors="replace",
            )
            if proc.returncode == 0:
                out_file = Path(output_path)
                if out_file.exists() and out_file.stat().st_size > 0:
                    if attempt_name.startswith("without"):
                        logger.info(LogModule.RESTOR, "[PDF-EXPORT] convert_md_to_pdf succeeded without xeCJK (xeCJK.sty not installed); install xecjk for better CJK line breaking")
                    logger.info(LogModule.TRANS, f"[PDF-EXPORT] convert_md_to_pdf succeeded: {output_path}")
                    return True
            stderr = proc.stderr or ""
            stdout = proc.stdout or ""
            if "xeCJK" in stderr and attempt_name.startswith("with"):
                logger.info(LogModule.RESTOR, "[PDF-EXPORT] convert_md_to_pdf failed (xeCJK.sty not found), retrying without xeCJK...")
                continue
            logger.warning(
                LogModule.RESTOR,
                f"[PDF-EXPORT] convert_md_to_pdf pandoc exit code {proc.returncode}, "
                f"stderr: {stderr[:800]}, stdout: {stdout[:800]}",
            )
            # When PDF export fails, try to extract a small LaTeX context window
            # around the first reported error line so that downstream components
            # (or an LLM-based repair) can focus on a local snippet instead of
            # the entire document.
            # On failure, generate debug LaTeX (extra pandoc run) if not yet generated.
            if debug_tex_path is None:
                try:
                    debug_tex_path = (out_dir / "debug") / (Path(output_path).stem + ".tex")
                    tex_cmd = [
                        str(pandoc_path),
                        "-f",
                        pandoc_from,
                        temp_md,
                        "-t",
                        "latex",
                        "-o",
                        str(debug_tex_path),
                        "-V",
                        f"mainfont={mainfont}",
                        "-V",
                        f"geometry={geometry_opts}",
                        "-V",
                        "papersize=a4",
                        "-V",
                        f"header-includes={current_header}",
                        "--resource-path", str(out_dir),
                    ]
                    tex_proc = subprocess.run(  # noqa: S603
                        tex_cmd,
                        cwd=str(temp_work_dir),
                        env=env,
                        capture_output=True,
                        timeout=120,
                        encoding="utf-8",
                        errors="replace",
                    )
                    if tex_proc.returncode == 0:
                        logger.info(
                            LogModule.RESTOR,
                            "[PDF-EXPORT] Wrote debug LaTeX to {path}",
                            path=str(debug_tex_path),
                        )
                    else:
                        logger.warning(
                            LogModule.RESTOR,
                            "[PDF-EXPORT] Failed to write debug LaTeX (exit={code}), stderr={stderr}",
                            code=tex_proc.returncode,
                            stderr=(tex_proc.stderr or "")[:300],
                        )
                except Exception as tex_err:  # noqa: BLE001
                    logger.debug(
                        LogModule.RESTOR,
                        f"[PDF-EXPORT] Writing debug LaTeX failed: {tex_err}",
                    )

            if debug_tex_path is not None and debug_tex_path.exists():
                ctx = extract_latex_error_context(stderr, debug_tex_path, debug_md_path)
                if ctx is not None and (ctx.md_snippet or ctx.tex_snippet):
                    logger.warning(
                        LogModule.RESTOR,
                        "[PDF-EXPORT] LaTeX compile context (type={etype}, line={line}, token={tok}, debug_tex={tex_path}, debug_md={md_path})",
                        etype=ctx.error_type,
                        line=ctx.line_no,
                        tok=ctx.error_token or "<none>",
                        tex_path=str(debug_tex_path),
                        md_path=str(debug_md_path) if debug_md_path else "<none>",
                    )
                    raise PdfExportLatexError(
                        "Pandoc+XeLaTeX failed to compile the generated LaTeX. See error context for details.",
                        stderr=stderr,
                        error_type=ctx.error_type or "unknown",
                        line_no=ctx.line_no,
                        error_token=ctx.error_token or "",
                        tex_snippet=ctx.tex_snippet or "",
                        md_snippet=ctx.md_snippet or "",
                        debug_tex_path=debug_tex_path,
                        debug_md_path=debug_md_path,
                    )
        return False
    except subprocess.TimeoutExpired:
        logger.warning(LogModule.RESTOR, "[PDF-EXPORT] convert_md_to_pdf: pandoc timed out")
        return False
    except Exception as e:
        # Let callers handle LaTeX compilation failures with extracted context.
        # They can use it to locate the bad segment and guide users to repair.
        if isinstance(e, PdfExportLatexError):
            raise
        logger.warning(LogModule.RESTOR, f"[PDF-EXPORT] convert_md_to_pdf failed: {e}", exc_info=False)
        return False
    finally:
        try:
            os.unlink(temp_md)
        except Exception:
            pass


def get_layout_block_bbox(layout_document: Any) -> Dict[int, Tuple[float, float, float, float]]:
    """
    Build mapping from layout block index to bbox (x0, y0, x1, y1) from a LayoutDocument.
    Intended to be called at Layout extraction / segment recording so export can use cached bbox
    without iterating layout_document again.
    """
    out: Dict[int, Tuple[float, float, float, float]] = {}
    try:
        from layout.base import LayoutDocument as _LD
        if not isinstance(layout_document, _LD):
            return out
        for block in layout_document.iter_blocks():
            if block.index is not None and getattr(block, "bbox", None):
                out[block.index] = tuple(float(x) for x in block.bbox)
    except Exception as e:
        logger.debug(LogModule.RESTOR, f"get_layout_block_bbox failed: {e}")
    return out


def normalize_layout_block_bbox_map(
    raw_map: Any,
) -> Dict[int, Tuple[float, float, float, float]]:
    """Normalize block-index -> bbox map (int keys, float 4-tuple values)."""
    out: Dict[int, Tuple[float, float, float, float]] = {}
    if not isinstance(raw_map, dict):
        return out
    for k, v in raw_map.items():
        if v is None:
            continue
        try:
            bidx = int(k)
            if not hasattr(v, "__iter__") or isinstance(v, (str, bytes)):
                continue
            coords = list(v)
            if len(coords) < 4:
                continue
            out[bidx] = (
                float(coords[0]),
                float(coords[1]),
                float(coords[2]),
                float(coords[3]),
            )
        except (TypeError, ValueError, IndexError):
            continue
    return out


def build_layout_block_page_number_map(layout_document: Any) -> Dict[int, int]:
    """Map layout block index to 1-based PDF page number."""
    result: Dict[int, int] = {}
    if layout_document is None:
        return result
    try:
        for page in layout_document.pages:
            page_num = int(page.page_index) + 1
            for block in page.blocks:
                if block.index is not None:
                    result[int(block.index)] = page_num
    except Exception:
        pass
    return result


def page_numbers_for_layout_block_indices(
    block_indices: Any,
    *,
    layout_document: Any = None,
    page_number_map: Optional[Dict[int, int]] = None,
) -> List[int]:
    """Return 1-based PDF page numbers aligned with layout_block_indices."""
    if not block_indices:
        return []
    lookup = page_number_map or build_layout_block_page_number_map(layout_document)
    pages: List[int] = []
    for raw_idx in block_indices:
        try:
            bidx = int(raw_idx)
        except (TypeError, ValueError):
            pages.append(0)
            continue
        pages.append(int(lookup.get(bidx, 0)))
    return pages


def sync_segment_layout_block_page_numbers(
    segment: Dict[str, Any],
    layout_document: Any,
    *,
    page_number_map: Optional[Dict[int, int]] = None,
) -> bool:
    """Attach layout_block_page_numbers aligned with layout_block_indices."""
    if not isinstance(segment, dict):
        return False
    indices_raw = segment.get("layout_block_indices") or []
    if not indices_raw:
        segment.pop("layout_block_page_numbers", None)
        return False
    pages = page_numbers_for_layout_block_indices(
        indices_raw,
        layout_document=layout_document,
        page_number_map=page_number_map,
    )
    if not pages:
        segment.pop("layout_block_page_numbers", None)
        return False
    if segment.get("layout_block_page_numbers") != pages:
        segment["layout_block_page_numbers"] = pages
        return True
    return False


def bboxes_for_layout_block_indices(
    block_indices: Any,
    bbox_map: Optional[Dict[Any, Any]] = None,
    *,
    layout_document: Any = None,
    return_miss_detail: bool = False,
) -> Union[List[List[float]], Dict[str, Any]]:
    """Resolve per-segment layout_block_bbox from block indices (JSON-safe lists)."""
    if not block_indices:
        if return_miss_detail:
            return {"bboxes": [], "missed": [], "map_keys_sample": []}
        return []
    lookup = normalize_layout_block_bbox_map(bbox_map)
    if not lookup and layout_document is not None:
        lookup = get_layout_block_bbox(layout_document)

    doc_by_index: Dict[int, Tuple[float, float, float, float]] = {}
    if layout_document is not None:
        try:
            from layout.base import LayoutDocument as _LD

            if isinstance(layout_document, _LD):
                for block in layout_document.iter_blocks():
                    if block.index is not None and getattr(block, "bbox", None):
                        doc_by_index[int(block.index)] = tuple(
                            float(x) for x in block.bbox[:4]
                        )
        except Exception:
            pass

    result: List[List[float]] = []
    missed: List[int] = []
    for raw_idx in block_indices:
        try:
            bidx = int(raw_idx)
        except (TypeError, ValueError):
            continue
        bbox = lookup.get(bidx) or doc_by_index.get(bidx)
        if bbox is not None:
            result.append(
                [float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])]
            )
        else:
            missed.append(bidx)

    if return_miss_detail:
        map_keys = sorted(set(lookup.keys()) | set(doc_by_index.keys()))
        return {
            "bboxes": result,
            "missed": missed,
            "map_keys_sample": map_keys[:12],
            "map_block_count": len(map_keys),
        }
    return result


def expand_segment_layout_group_bboxes(
    segment: Dict[str, Any],
    layout_doc: Any,
    bbox_map: Optional[Dict[Any, Any]] = None,
) -> bool:
    """Ensure segment lists layout group companion indices and bboxes for preview."""
    if layout_doc is None or not isinstance(segment, dict):
        return False

    from layout.layout_group_pair_utils import (
        cross_page_pairs_from_raw,
        resolve_layout_group_pairs_for_block,
        sort_layout_block_indices_reading_order,
    )

    indices_raw = segment.get("layout_block_indices") or []
    indices: List[int] = []
    for raw_idx in indices_raw:
        try:
            indices.append(int(raw_idx))
        except (TypeError, ValueError):
            continue
    if not indices:
        return False

    block_by_index: Dict[int, Any] = {}
    for block in layout_doc.iter_blocks():
        if block.index is not None:
            block_by_index[int(block.index)] = block

    expanded_indices = list(indices)
    changed = False
    for idx in list(indices):
        block = block_by_index.get(idx)
        if block is None:
            continue
        block_raw = getattr(block, "raw", None) or {}
        for pair in cross_page_pairs_from_raw(block_raw):
            companion_index = pair.get("index")
            if companion_index is None:
                continue
            try:
                companion_int = int(companion_index)
            except (TypeError, ValueError):
                continue
            if companion_int not in expanded_indices:
                expanded_indices.append(companion_int)
                changed = True
        for pair in resolve_layout_group_pairs_for_block(block, layout_doc):
            companion_index = pair.get("index")
            if companion_index is None:
                continue
            try:
                companion_int = int(companion_index)
            except (TypeError, ValueError):
                continue
            if companion_int not in expanded_indices:
                expanded_indices.append(companion_int)
                changed = True

    if len(expanded_indices) > 1:
        primary_hint = expanded_indices[0]
        sorted_indices = sort_layout_block_indices_reading_order(
            expanded_indices,
            layout_doc,
            primary_hint,
        )
        if sorted_indices != expanded_indices:
            expanded_indices = sorted_indices
            changed = True

    existing_bboxes = segment.get("layout_block_bbox")
    needs_bbox_refresh = (
        changed
        or not isinstance(existing_bboxes, list)
        or len(existing_bboxes) < len(expanded_indices)
    )
    if not needs_bbox_refresh:
        page_changed = sync_segment_layout_block_page_numbers(segment, layout_doc)
        if (
            expanded_indices != list(segment.get("layout_block_indices") or [])
            and len(expanded_indices) > len(indices)
        ):
            segment["layout_block_indices"] = expanded_indices
            return True
        return page_changed

    segment["layout_block_indices"] = expanded_indices

    seg_bboxes = bboxes_for_layout_block_indices(
        expanded_indices,
        bbox_map,
        layout_document=layout_doc,
    )
    if not seg_bboxes:
        sync_segment_layout_block_page_numbers(segment, layout_doc)
        return True

    if seg_bboxes != existing_bboxes:
        segment["layout_block_bbox"] = seg_bboxes
        changed = True
    else:
        changed = bool(changed)

    page_changed = sync_segment_layout_block_page_numbers(
        segment,
        layout_doc,
    )
    return changed or page_changed or needs_bbox_refresh


def segment_needs_layout_block_bbox(segment: Dict[str, Any]) -> bool:
    """True when segment has block indices but no usable layout_block_bbox."""
    if not segment.get("layout_block_indices"):
        return False
    existing = segment.get("layout_block_bbox")
    if existing is None:
        return True
    return isinstance(existing, list) and len(existing) == 0


def _log_latex_error_context(stderr: str, tex_path: Path, md_path: Optional[Path]) -> None:
    """
    Deprecated: use extract_latex_error_context() for context extraction and logging.

    Kept as a no-op shim for backward compatibility in case external callers import it.
    """
    _ = (stderr, tex_path, md_path)
    return None


def _bbox_y_overlap(bbox1: Tuple[float, float, float, float], bbox2: Tuple[float, float, float, float], tolerance: float = 2.0) -> bool:
    """
    Check if two layout bboxes (x0, y0, x1, y1) overlap in y (same row).
    Uses tolerance in same units as bbox to handle rounding.
    Coerces values to float so bbox from JSON (e.g. task_state) with string numbers still works.
    """
    y0_1, y1_1 = float(bbox1[1]), float(bbox1[3])
    y0_2, y1_2 = float(bbox2[1]), float(bbox2[3])
    return not (y1_1 <= y0_2 - tolerance or y1_2 <= y0_1 - tolerance)


def get_image_block_indices_from_layout(
    segments: List[Dict[str, Any]],
    layout_document: Any,
    equation_format: Optional[str] = None,
    table_body_format: Optional[str] = None,
    chart_body_format: Optional[str] = None,
) -> Tuple[List[int], Dict[str, int]]:
    """
    Build list of layout block indices for each image segment in document order, and a mapping from image paths to block indices.
    Used to decide which consecutive images are on the same row (bbox y overlap).
    Only includes blocks that will actually emit an image ref in the current export (e.g. when equation_format=text,
    equation blocks are excluded so the list matches the order of image refs in the rebuilt markdown).

    Args:
        segments: List of translation segments (in order)
        layout_document: LayoutDocument with iter_blocks() and blocks with .index, .bbox, .has_image(), .type
        equation_format: Optional "text" | "latex" | "image" – when not "image", equation blocks are excluded
        table_body_format: Optional "html" | "image" – when not "image", table body image blocks are excluded
        chart_body_format: Optional "html" | "image" – when not "image", chart body image blocks are excluded

    Returns:
        Tuple of (list of block indices in document order, dict mapping normalized image path to block index).
        The dict uses normalized paths (filename only, lowercase) for matching.
    """
    try:
        from layout.base import LayoutDocument as _LD
        if not isinstance(layout_document, _LD):
            logger.debug(LogModule.RESTOR, "[PDF-EXPORT] get_image_block_indices_from_layout: layout_document is not LayoutDocument")
            return ([], {})
    except Exception as e:
        logger.debug(LogModule.RESTOR, f"[PDF-EXPORT] get_image_block_indices_from_layout: {e}")
        return ([], {})
    block_index_to_block: Dict[int, Any] = {}
    block_index_to_type: Dict[int, str] = {}
    for block in layout_document.iter_blocks():
        if block.index is not None:
            block_index_to_block[block.index] = block
            block_index_to_type[block.index] = getattr(block, "type", "") or ""
    image_blocks = [idx for idx, b in block_index_to_block.items() if getattr(b, "has_image", None) and b.has_image()]
    out: List[int] = []
    path_to_block_index: Dict[str, int] = {}
    segs_with_layout = 0
    eq_fmt = (equation_format or "text").strip().lower()
    tbl_fmt = (table_body_format or "html").strip().lower()
    chart_fmt = (chart_body_format or "image").strip().lower()
    from utils.translation_segments import _is_image_segment

    for seg in segments:
        source_text = seg.get("source_text") or seg.get("text") or ""
        # Captions after images share layout block indices; only map true image segments
        if not _is_image_segment(source_text):
            continue
        bidxs = seg.get("layout_block_indices", [])
        if bidxs:
            segs_with_layout += 1
        for bidx in bidxs:
            block = block_index_to_block.get(bidx)
            if not block or not getattr(block, "has_image", None) or not block.has_image():
                continue
            btype = block_index_to_type.get(bidx, "")
            # Only include block if current export will emit an image ref for it
            if btype == "image":
                include = True
            elif btype in ("interline_equation", "formula", "equation"):
                include = eq_fmt == "image"
            elif btype == "table":
                include = tbl_fmt == "image"
            elif btype == "chart":
                include = chart_fmt == "image"
            else:
                include = True
            if not include:
                continue
            out.append(bidx)
            # Try to extract image path from segment for mapping
            # Check multiple possible fields: image_path, placeholder_id (may be filename), source_text
            image_path = seg.get("image_path") or seg.get("source_text", "")
            placeholder_id = seg.get("placeholder_id", "")

            # If we have placeholder_id, it might be the filename (e.g., "e077a90c.jpg")
            # Check if placeholder_id looks like a filename (has extension)
            if not image_path and placeholder_id:
                if "." in placeholder_id and any(placeholder_id.lower().endswith(ext) for ext in [".jpg", ".jpeg", ".png", ".gif", ".webp"]):
                    image_path = placeholder_id

            # For formula/table/chart blocks: extract image_path from the layout block itself
            # when the segment does not carry image_path (e.g. equation_image_path from MinerU spans)
            if not image_path and btype in ("interline_equation", "formula", "equation", "table", "chart"):
                block_img_path = getattr(block, "image_path", None)
                if block_img_path:
                    image_path = str(block_img_path)
                # Also check inside raw spans (e.g. interline_equation span.image_path)
                if not image_path:
                    raw_block = getattr(block, "raw", None) or {}
                    for line in raw_block.get("lines", []):
                        for span in line.get("spans", []):
                            span_img = span.get("image_path")
                            if span_img:
                                image_path = str(span_img)
                                break
                        if image_path:
                            break

            if image_path:
                # Normalize path: extract filename, lowercase, remove path separators
                # Handle markdown image syntax: ![alt](path) -> extract path
                md_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', image_path)
                if md_match:
                    image_path = md_match.group(2)
                # Handle placeholder format: <ph-xxx> -> try to extract filename from placeholder_id
                if image_path.startswith("<ph-") and placeholder_id:
                    # placeholder_id might be the filename
                    if "." in placeholder_id:
                        image_path = placeholder_id
                # Extract filename and normalize (handle both ./images/file.jpg and images/file.jpg and file.jpg)
                # Remove leading ./ and images/ prefix if present
                normalized = image_path.replace("./", "").replace("images/", "").replace("images\\", "")
                filename = os.path.basename(normalized).lower()
                if filename:
                    path_to_block_index[filename] = bidx
                    logger.debug(
                        LogModule.RESTOR,
                        f"[PDF-EXPORT] Mapped image: original={image_path}, placeholder_id={placeholder_id}, "
                        f"normalized_filename={filename}, block_index={bidx}",
                    )
            break
    if not out and segments:
        logger.debug(
            LogModule.RESTOR,
            f"[PDF-EXPORT] get_image_block_indices_from_layout: 0 image indices (segments={len(segments)}, "
            f"segments_with_layout_block_indices={segs_with_layout}, layout_image_blocks={len(image_blocks)})"
        )
    logger.info(
        LogModule.RESTOR,
        f"[PDF-EXPORT] get_image_block_indices_from_layout: {len(out)} image indices, "
        f"{len(path_to_block_index)} path mappings "
        f"(segments={len(segments) if segments else 0})",
    )
    return (out, path_to_block_index)


def group_consecutive_images_for_markdown(
    md_content: str,
    image_block_indices: Optional[List[int]] = None,
    layout_document: Optional[Any] = None,
    layout_block_bbox: Optional[Dict[int, Tuple[float, float, float, float]]] = None,
) -> str:
    """
    Group consecutive image references (separated only by whitespace) into side-by-side HTML layout.
    
    When image_block_indices and layout_document (or layout_block_bbox) are provided, only groups a run of images
    if their layout bboxes overlap in y (same row); otherwise keeps them stacked.
    Prefer layout_block_bbox from Layout extraction phase so layout_document is not iterated at export.

    Args:
        md_content: Markdown content with image references
        image_block_indices: Optional list of layout block index per image in document order
        layout_document: Optional LayoutDocument to resolve bbox for same-row check (used when layout_block_bbox not provided)
        layout_block_bbox: Optional precomputed block index -> bbox from Layout extraction phase

    Returns:
        Modified markdown with consecutive images grouped in HTML divs only when same row (or always if no layout).
    """
    import re
    # Pattern to match markdown image syntax: ![alt](path)
    image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
    
    all_matches = list(image_pattern.finditer(md_content))
    if len(all_matches) < 2:
        return md_content  # No consecutive images possible
    
    # Prefer bbox from Layout extraction phase; normalize keys to int and bbox to tuple of float (JSON round-trip safe)
    block_index_to_bbox: Optional[Dict[int, Tuple[float, float, float, float]]] = None
    if layout_block_bbox:
        try:
            block_index_to_bbox = {}
            for k, v in layout_block_bbox.items():
                if v is None or len(v) < 4:
                    continue
                bidx = int(k) if not isinstance(k, int) else k
                block_index_to_bbox[bidx] = (float(v[0]), float(v[1]), float(v[2]), float(v[3]))
        except (TypeError, ValueError, IndexError):
            block_index_to_bbox = None
    if block_index_to_bbox is None and layout_document and image_block_indices is not None and len(image_block_indices) >= 2:
        try:
            from layout.base import LayoutDocument as _LD
            if isinstance(layout_document, _LD):
                block_index_to_bbox = {}
                for block in layout_document.iter_blocks():
                    if block.index is not None and hasattr(block, "bbox") and block.bbox:
                        block_index_to_bbox[block.index] = block.bbox
        except Exception:
            block_index_to_bbox = None
    
    runs = []
    i = 0
    while i < len(all_matches):
        run = [all_matches[i]]
        while i + 1 < len(all_matches):
            gap = md_content[run[-1].end() : all_matches[i + 1].start()]
            if not gap.strip():  # only whitespace between
                run.append(all_matches[i + 1])
                i += 1
            else:
                break
        runs.append(run)
        i += 1
    
    # Build new markdown with HTML divs for runs of 2+ images (only when same row if layout provided)
    parts = []
    pos = 0
    image_index = 0  # global index of image in document order
    for run in runs:
        parts.append(md_content[pos : run[0].start()])
        merge_run = len(run) >= 2
        if merge_run and block_index_to_bbox is not None and image_block_indices is not None:
            # Only merge if all images in run have same-row bboxes (pairwise y overlap)
            run_block_indices = []
            for j in range(len(run)):
                idx = image_index + j
                if idx < len(image_block_indices):
                    run_block_indices.append(image_block_indices[idx])
            same_row = True
            if len(run_block_indices) == len(run):
                for j in range(len(run_block_indices) - 1):
                    bbox_a = block_index_to_bbox.get(run_block_indices[j])
                    bbox_b = block_index_to_bbox.get(run_block_indices[j + 1])
                    if bbox_a is None or bbox_b is None or not _bbox_y_overlap(bbox_a, bbox_b):
                        same_row = False
                        logger.debug(
                            LogModule.RESTOR,
                            f"[MD-EXPORT] Consecutive images not merged (different row): blocks {run_block_indices[j]}, {run_block_indices[j + 1]}"
                        )
                        break
            else:
                same_row = False
            merge_run = merge_run and same_row
        image_index += len(run)
        
        if merge_run:
            # Create HTML div with inline-block images for side-by-side layout
            n = len(run)
            img_tags = []
            for m in run:
                alt_text = m.group(1) or ""
                img_path = m.group(2)
                # Escape HTML special chars in alt_text and path
                alt_escaped = alt_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
                # Escape path for HTML attribute (quote and escape quotes)
                path_escaped = img_path.replace("&", "&amp;").replace('"', "&quot;")
                # Use HTML img tag with inline-block for side-by-side
                # Each image gets (100/n)% width minus margins for spacing
                width_pct = (100 / n) - 2  # Subtract 2% for margins
                img_tags.append(f'<img src="{path_escaped}" alt="{alt_escaped}" style="display: inline-block; max-width: {width_pct:.1f}%; margin: 0 1%; vertical-align: top;" />')
            
            html_block = f'<div style="text-align: center; margin: 1em 0;">\n' + "\n".join(img_tags) + "\n</div>"
            parts.append(html_block)
            logger.debug(LogModule.RESTOR, f"[MD-EXPORT] Side-by-side images: {n} images grouped in HTML div (same row)")
        else:
            for m in run:
                parts.append(m.group(0))
        pos = run[-1].end()
    parts.append(md_content[pos:])
    return "".join(parts)


#
# NOTE: Legacy HTML->PDF backend has been removed.
# If HTML-to-PDF is needed (e.g. MOBI/EPUB workflows), use Pandoc → XeLaTeX.

_CJK_LANG_CODES = (
    "zh",
    "chinese",
    "zh-cn",
    "zh-tw",
    "ja",
    "japanese",
    "jp",
    "ko",
    "korean",
    "kr",
)
_CJK_CHAR_RE = re.compile(
    r"[\u3000-\u303f\u3040-\u30ff\u31f0-\u31ff\u4e00-\u9fff\uac00-\ud7af\uf900-\ufaff]"
)


def _content_has_cjk(text: str) -> bool:
    return bool(text and _CJK_CHAR_RE.search(text))


def _should_use_xecjk_for_pdf(to_lang: Optional[str], content: Optional[str] = None) -> bool:
    to_lang_lower = (to_lang or "").strip().lower()
    if any(to_lang_lower.startswith(code) or to_lang_lower == code for code in _CJK_LANG_CODES):
        return True
    return _content_has_cjk(content or "")


def _cjk_mainfont_fallback(lang_code: str) -> str:
    lang = (lang_code or "zh").strip().lower()
    if sys.platform == "darwin":
        if lang.startswith("ja"):
            return "Hiragino Sans"
        if lang.startswith("ko"):
            return "Apple SD Gothic Neo"
        return "PingFang SC"
    if sys.platform == "win32":
        if lang.startswith("ja"):
            return "Yu Gothic"
        if lang.startswith("ko"):
            return "Malgun Gothic"
        if lang.startswith("zh") and ("tw" in lang or "hant" in lang):
            return "Microsoft JhengHei"
        return "Microsoft YaHei"
    if lang.startswith("ja"):
        return "Noto Sans CJK JP"
    if lang.startswith("ko"):
        return "Noto Sans CJK KR"
    return "Noto Sans CJK SC"


def _resolve_mainfont_for_pdf(to_lang: Optional[str], content: Optional[str] = None) -> str:
    lang_code = (to_lang or "").strip().lower()
    if not lang_code and content and _content_has_cjk(content):
        if re.search(r"[\u3040-\u30ff\u31f0-\u31ff]", content):
            lang_code = "ja"
        elif re.search(r"[\uac00-\ud7af]", content):
            lang_code = "ko"
        else:
            lang_code = "zh"
    try:
        from translator.ai_translator.docx_translator import get_font_for_language

        if lang_code:
            return get_font_for_language(lang_code)
    except Exception:
        if lang_code and any(
            lang_code.startswith(code) or lang_code == code for code in _CJK_LANG_CODES
        ):
            return _cjk_mainfont_fallback(lang_code)
        if not lang_code and content and _content_has_cjk(content):
            return _cjk_mainfont_fallback(lang_code or "zh")
    return "Helvetica Neue" if sys.platform == "darwin" else "Calibri"


def _resolve_pandoc_lang_for_pdf(to_lang: Optional[str], content: Optional[str] = None) -> str:
    """IETF lang tag for Pandoc PDF metadata.

    HTML workflow wraps exports with ``lang=\"en\"``, which makes Pandoc load babel
    ``english``/``american`` on bundled TeX and fail. Override with a supported tag,
    or return empty to skip babel language setup.
    """
    lang_code = (to_lang or "").strip().lower()
    if not lang_code and content and _content_has_cjk(content):
        if re.search(r"[\u3040-\u30ff\u31f0-\u31ff]", content):
            lang_code = "ja"
        elif re.search(r"[\uac00-\ud7af]", content):
            lang_code = "ko"
        else:
            lang_code = "zh"
    if lang_code.startswith("zh"):
        if "tw" in lang_code or "hant" in lang_code:
            return "zh-TW"
        return "zh-CN"
    if lang_code.startswith("ja"):
        return "ja-JP"
    if lang_code.startswith("ko"):
        return "ko-KR"
    if lang_code.startswith("fr"):
        return "fr-FR"
    if lang_code.startswith("de"):
        return "de-DE"
    if lang_code.startswith("es"):
        return "es-ES"
    # Avoid babel english/american on bundled TeX (also overrides HTML lang="en").
    return ""


def _resolve_cjk_mainfont_for_pdf(to_lang: Optional[str], content: Optional[str] = None) -> str:
    """CJK font for xeCJK; independent of Latin mainfont (e.g. Calibri when to_lang=en)."""
    lang_code = (to_lang or "").strip().lower()
    if content:
        if re.search(r"[\u3040-\u30ff\u31f0-\u31ff]", content):
            lang_code = "ja"
        elif re.search(r"[\uac00-\ud7af]", content):
            lang_code = "ko"
        elif _content_has_cjk(content) and not any(
            lang_code.startswith(code) or lang_code == code for code in _CJK_LANG_CODES
        ):
            lang_code = "zh"
    if "tw" in lang_code or "hant" in lang_code:
        return _cjk_mainfont_fallback("zh-tw")
    if lang_code.startswith("ja"):
        return _cjk_mainfont_fallback("ja")
    if lang_code.startswith("ko"):
        return _cjk_mainfont_fallback("ko")
    return _cjk_mainfont_fallback("zh")


def _pandoc_pdf_header_includes(
    mainfont: str,
    use_xecjk: bool,
    cjk_mainfont: Optional[str] = None,
) -> Tuple[str, str]:
    cjk_preamble = ""
    if use_xecjk:
        cjk_font_name = (cjk_mainfont or mainfont).replace("\\", "").replace("}", "\\}").replace("{", "\\{")
        cjk_preamble = f"\\usepackage{{xeCJK}}\\setCJKmainfont{{{cjk_font_name}}}"
    shared = (
        "\\usepackage{ragged2e}\\AtBeginDocument{\\RaggedRight}"
        "\\PassOptionsToPackage{hyphens}{url}\\usepackage{hyperref}\\hypersetup{breaklinks=true}"
        "\\usepackage{titlesec}"
        "\\usepackage{graphicx}"
        "\\usepackage{etoolbox}"
        "\\makeatletter"
        "\\renewcommand{\\@maketitle}{\\begin{center}\\LARGE\\bfseries\\@title\\par\\vskip 0.5em\\large\\@author\\par\\vskip 0.3em\\normalsize\\@date\\end{center}\\par\\vskip 1em}"
        "\\makeatother"
        "\\AtEndPreamble{"
        "\\titleformat*{\\section}{\\LARGE\\bfseries}"
        "\\titleformat*{\\subsection}{\\Large\\bfseries}"
        "\\titleformat*{\\subsubsection}{\\large\\bfseries}"
        "}"
        "\\AtBeginDocument{"
        "\\sloppy\\setlength{\\emergencystretch}{5em}"
        "}"
    )
    return cjk_preamble + shared, shared


def _xelatex_subprocess_env(
    xelatex_path: Optional[Path],
    pdflatex_root_use: Optional[Path],
    temp_work_dir: Path,
) -> Dict[str, str]:
    env = os.environ.copy()
    if not xelatex_path or pdflatex_root_use is None:
        return env
    env["PATH"] = str(xelatex_path.parent) + os.pathsep + env.get("PATH", "")
    env["TEXMFCNF"] = str(pdflatex_root_use) + os.pathsep + str(
        pdflatex_root_use / "texmf-dist" / "web2c"
    )
    env["TEXMFROOT"] = str(pdflatex_root_use)
    user_texmfvar = _get_user_texmfvar_dir()
    user_texmfvar.mkdir(parents=True, exist_ok=True)
    texmfvar = str(user_texmfvar)
    env["TEXMFVAR"] = texmfvar
    env["TEXMFSYSVAR"] = texmfvar
    env["TEXMFOUTPUT"] = str(temp_work_dir)
    fontconfig_file = pdflatex_root_use / "texmf-var" / "fonts" / "conf" / "fonts.conf"
    if fontconfig_file.exists():
        env["FONTCONFIG_FILE"] = str(fontconfig_file)
        fc_cache_dir = user_texmfvar / "fontconfig-cache"
        fc_cache_dir.mkdir(parents=True, exist_ok=True)
        env["FC_CACHEDIR"] = str(fc_cache_dir)
    _ensure_xelatex_fmt(pdflatex_root_use, env)
    return env


def _ensure_html_utf8_meta(html_content: str) -> str:
    """Ensure Pandoc/XeLaTeX sees UTF-8 when converting HTML to PDF."""
    lower = html_content.lower()
    if "charset" in lower or "encoding=" in lower:
        return html_content
    if re.search(r"<head\b", html_content, re.IGNORECASE):
        return re.sub(
            r"(<head\b[^>]*>)",
            r'\1<meta charset="utf-8" />',
            html_content,
            count=1,
            flags=re.IGNORECASE,
        )
    return f'<!DOCTYPE html><html><head><meta charset="utf-8" /></head><body>{html_content}</body></html>'


async def convert_html_to_pdf(
    html_content: str,
    output_path: str,
    output_dir: Optional[Path] = None,
    to_lang: Optional[str] = None,
) -> None:
    """
    Convert HTML content to PDF via Pandoc with XeLaTeX.

    This is used by workflows that naturally produce HTML (e.g. MOBI/EPUB).
    It intentionally does NOT depend on any browser-based renderer or PyMuPDF.
    """
    if not html_content or not html_content.strip():
        raise ValueError("HTML content is empty, cannot generate PDF.")

    out_dir = output_dir or Path(output_path).parent
    out_dir.mkdir(parents=True, exist_ok=True)
    # Run Pandoc/XeLaTeX from the system temp directory so auxiliary files (.aux, .log, .xdv)
    # are written to a guaranteed-writable location instead of the possibly-read-only out_dir.
    temp_work_dir = Path(tempfile.gettempdir())
    output_path = str(Path(output_path).resolve())
    # Copy images into temp_work_dir so xelatex (running with cwd=temp_work_dir) can resolve
    # relative paths like ./images/... used in both markdown and raw LaTeX side-by-side blocks.
    images_dir = out_dir / "images"
    if images_dir.exists():
        temp_images_dir = temp_work_dir / "images"
        if temp_images_dir.exists():
            shutil.rmtree(temp_images_dir)
        shutil.copytree(images_dir, temp_images_dir)

    pandoc_path = _get_pandoc_path()
    if not pandoc_path:
        raise RuntimeError(
            "Pandoc is required for PDF export but was not found. "
            "Install Pandoc and ensure it is in PATH or in 3rdParty/windows."
        )
    os.environ["PYPANDOC_PANDOC"] = str(pandoc_path)

    xelatex_path_orig = _get_xelatex_path()
    xelatex_path = None
    pdflatex_root_use = None
    if xelatex_path_orig:
        pdflatex_root = xelatex_path_orig.parent.parent.parent
        pdflatex_root_use = _ensure_ascii_path_for_tex(pdflatex_root)
        xelatex_path = pdflatex_root_use / "bin" / "windows" / "xelatex.exe"
        if not xelatex_path.exists():
            xelatex_path = _to_short_path_if_needed(xelatex_path_orig)
            pdflatex_root_use = xelatex_path.parent.parent.parent
    
    # 检查 xelatex 是否在 PATH 中
    xelatex_in_path = shutil.which("xelatex")
    
    # 检查 macOS 上常见的 xelatex 路径
    if not xelatex_in_path and sys.platform == "darwin":
        common_paths = [
            "/Library/TeX/texbin/xelatex",  # MacTeX / BasicTeX
            "/usr/local/texlive/current/bin/universal-darwin/xelatex",  # TeX Live
            "/usr/local/texlive/2026/bin/universal-darwin/xelatex",  # TeX Live 2026
        ]
        for path in common_paths:
            if Path(path).exists():
                xelatex_in_path = path
                logger.info(LogModule.RESTOR, f"[PDF-EXPORT] Found XeLaTeX at: {xelatex_in_path}")
                break
    
    # 如果仍然找不到 xelatex，抛出错误
    if not xelatex_path and not xelatex_in_path:
        if sys.platform == "darwin":
            install_msg = (
                "XeLaTeX is required for PDF export but was not found. "
                "Install TeX Live: brew install --cask mactex "
                "or TinyTeX: brew install --cask tinytex. "
                "After installation, run: cd 3rdParty/macos && ./install_latex_packages.sh "
                "to install required LaTeX packages. See server logs for details."
            )
        elif sys.platform == "linux":
            install_msg = (
                "XeLaTeX is required for PDF export but was not found. "
                "Install TeX Live: sudo apt-get install texlive-xetex texlive-lang-chinese "
                "or TinyTeX. See server logs for details."
            )
        else:  # Windows
            install_msg = (
                "XeLaTeX is required for PDF export but was not found. "
                "Please install XeLaTeX (e.g. TeX Live, TinyTeX) and ensure the xelatex executable is in PATH or in 3rdParty/windows/pdflatex."
                "See server logs for details."
            )
        logger.error(
            LogModule.RESTOR,
            f"[PDF-EXPORT] {install_msg}"
        )
        raise RuntimeError(install_msg)

    pdf_engine = str(xelatex_path) if xelatex_path else (xelatex_in_path if xelatex_in_path else "xelatex")
    if sys.platform == "darwin":
        _check_latex_packages_macos()

    html_content = _ensure_html_utf8_meta(html_content)
    mainfont = _resolve_mainfont_for_pdf(to_lang, html_content)
    use_xecjk = _should_use_xecjk_for_pdf(to_lang, html_content)
    cjk_mainfont = _resolve_cjk_mainfont_for_pdf(to_lang, html_content) if use_xecjk else None
    pandoc_lang = _resolve_pandoc_lang_for_pdf(to_lang, html_content)
    header_with_cjk, header_without_cjk = _pandoc_pdf_header_includes(
        mainfont, use_xecjk, cjk_mainfont=cjk_mainfont
    )
    logger.info(
        LogModule.RESTOR,
        f"[PDF-EXPORT] convert_html_to_pdf fonts: to_lang={to_lang!r}, mainfont={mainfont!r}, "
        f"cjk_mainfont={cjk_mainfont!r}, use_xecjk={use_xecjk}, pandoc_lang={pandoc_lang!r}",
    )
    geometry_opts = "margin=2.5cm"
    env = _xelatex_subprocess_env(xelatex_path, pdflatex_root_use, temp_work_dir)
    import asyncio

    tmp_html = None
    try:
        fd, tmp_html = tempfile.mkstemp(suffix=".html", prefix="owlangs_pdf_")
        os.close(fd)
        Path(tmp_html).write_text(html_content, encoding="utf-8", errors="ignore")

        attempts: List[Tuple[str, str]] = [(header_with_cjk, "with xeCJK" if use_xecjk else "default")]
        if use_xecjk:
            attempts.append((header_without_cjk, "without xeCJK (fallback)"))

        def _run_pandoc() -> None:
            import subprocess

            last_error = ""
            for current_header, attempt_name in attempts:
                cmd = [
                    str(pandoc_path),
                    tmp_html,
                    "-f",
                    "html",
                    "-o",
                    output_path,
                    "--pdf-engine",
                    pdf_engine,
                    "--resource-path",
                    str(out_dir),
                    "-V",
                    f"mainfont={mainfont}",
                    "-V",
                    f"lang={pandoc_lang}",
                    "-V",
                    f"geometry={geometry_opts}",
                    "-V",
                    "papersize=a4",
                    "-V",
                    f"header-includes={current_header}",
                ]
                proc = subprocess.run(
                    cmd,
                    cwd=str(temp_work_dir),
                    capture_output=True,
                    encoding="utf-8",
                    errors="replace",
                    env=env,
                )
                if proc.returncode == 0:
                    if attempt_name.startswith("without xeCJK"):
                        logger.info(
                            LogModule.RESTOR,
                            "[PDF-EXPORT] convert_html_to_pdf succeeded without xeCJK (fallback)",
                        )
                    return
                last_error = (proc.stderr or proc.stdout or "")[:800]
                logger.warning(
                    LogModule.RESTOR,
                    f"[PDF-EXPORT] convert_html_to_pdf pandoc failed ({attempt_name}): {last_error}",
                )
            raise RuntimeError(
                f"Pandoc HTML->PDF failed after {len(attempts)} attempt(s): {last_error}"
            )

        await asyncio.to_thread(_run_pandoc)

        out_file = Path(output_path)
        if not out_file.exists() or out_file.stat().st_size == 0:
            raise RuntimeError("Pandoc produced an empty PDF file.")
    finally:
        try:
            if tmp_html:
                os.unlink(tmp_html)
        except Exception:
            pass
