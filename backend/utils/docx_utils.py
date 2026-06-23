# SPDX-FileCopyrightText: 2026 Zampher
# SPDX-License-Identifier: MPL-2.0

"""
Shared utility functions for DOCX processing.
Used by both DocxExtractor and DocxTranslator to avoid code duplication.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx.text.run import Run
    from docx.text.paragraph import Paragraph


def get_run_formatting_key(run: "Run") -> tuple:
    """
    Get a hashable key representing the formatting of a run.
    Used to detect formatting changes between runs.
    
    Args:
        run: A docx Run object
        
    Returns:
        Tuple of (font_name, bold, italic, underline, size, color_rgb)
    """
    if not run.font:
        return (None, None, None, None, None, None)
    
    try:
        color_rgb = None
        if run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb:
            color_rgb = run.font.color.rgb
    except Exception:
        color_rgb = None
    
    return (
        run.font.name,
        run.font.bold,
        run.font.italic,
        run.font.underline,
        run.font.size,
        color_rgb,
    )


def is_image_run(run: "Run") -> bool:
    """
    Check if a run contains an image.
    
    Args:
        run: A docx Run object
        
    Returns:
        True if the run contains an image, False otherwise
    """
    return '<w:drawing' in run.element.xml or '<w:pict' in run.element.xml


def paragraph_has_toc_field(paragraph: "Paragraph") -> bool:
    """
    Check if a paragraph contains a TOC field.
    
    Args:
        paragraph: A docx Paragraph object
        
    Returns:
        True if the paragraph contains a TOC field, False otherwise
    """
    try:
        p = paragraph._p  # lxml element
        
        # Check for TOC field codes
        fldChars = p.xpath('.//*[local-name()="fldChar"]')
        if not fldChars:
            # quick check for instruction text
            instrs = p.xpath('.//*[local-name()="instrText"]')
            for it in instrs:
                if 'TOC' in (it.text or ''):
                    return True
        else:
            instrs = p.xpath('.//*[local-name()="instrText"]')
            for it in instrs:
                if 'TOC' in (it.text or ''):
                    return True
                    
    except Exception:
        pass
    return False


def count_non_toc_paragraphs_in_nested_table(nested_table, depth: int = 0) -> int:
    """Count non-TOC paragraphs inside a nested table (matches docx_extractor order)."""
    if depth > 10:
        return 0

    from utils.table_utils import (
        get_all_merged_regions_docx,
        is_cell_in_merged_region_docx,
        is_merged_cell_start_at_position_docx,
    )

    nested_merged_regions = get_all_merged_regions_docx(nested_table)
    nested_processed_cells: set = set()
    count = 0

    for nested_row_idx, nested_row in enumerate(nested_table.rows):
        for nested_cell_idx, nested_cell in enumerate(nested_row.cells):
            if (nested_row_idx, nested_cell_idx) in nested_processed_cells:
                continue

            nested_is_in_merged, nested_merge_range = is_cell_in_merged_region_docx(
                nested_table, nested_row_idx, nested_cell_idx, nested_merged_regions
            )
            if nested_is_in_merged and nested_merge_range is not None:
                nested_start_row, nested_start_col, nested_end_row, nested_end_col = nested_merge_range
                for r in range(nested_start_row, nested_end_row + 1):
                    for c in range(nested_start_col, nested_end_col + 1):
                        nested_processed_cells.add((r, c))
                nested_is_start = is_merged_cell_start_at_position_docx(
                    nested_table, nested_row_idx, nested_cell_idx, nested_merged_regions
                )
                if not nested_is_start:
                    continue

            for nested_para in nested_cell.paragraphs:
                if not paragraph_has_toc_field(nested_para):
                    count += 1

            if hasattr(nested_cell, "tables") and nested_cell.tables:
                for deeper_nested_table in nested_cell.tables:
                    count += count_non_toc_paragraphs_in_nested_table(
                        deeper_nested_table, depth + 1
                    )

    return count


def count_non_toc_paragraphs_in_cell(cell, include_nested: bool = True) -> int:
    """Count non-TOC paragraphs in a table cell, optionally including nested tables."""
    count = 0
    for para in cell.paragraphs:
        if not paragraph_has_toc_field(para):
            count += 1
    if include_nested and hasattr(cell, "tables") and cell.tables:
        for nested_table in cell.tables:
            count += count_non_toc_paragraphs_in_nested_table(nested_table)
    return count

