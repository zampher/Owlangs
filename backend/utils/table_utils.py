# SPDX-FileCopyrightText: 2026 Zamphersss
# SPDX-License-Identifier: MPL-2.0

"""
Shared utilities for handling merged cells in DOCX, PPTX, and XLSX.
"""

from __future__ import annotations

from typing import Tuple, Optional, List, Set, Dict
from docx.oxml.ns import qn
from logger import unified_logger as logger
from logger.logger import LogModule

# DOCX namespace
DOCX_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def is_merged_cell_start_docx(cell) -> Tuple[bool, Optional[int], Optional[int]]:
    """
    Check if a DOCX cell is the start of a merged region.
    
    Args:
        cell: DOCX Cell object (from python-docx)
        
    Returns:
        Tuple of (is_merged_start, rowspan, colspan)
        - is_merged_start: True if this is the start of a merged region
        - rowspan: Number of rows merged (None if not merged vertically, or calculated value)
        - colspan: Number of columns merged (None if not merged horizontally)
    """
    tc = cell._tc  # Table cell element
    is_merged_start = False
    rowspan = None
    colspan = None
    
    # Check for vertical merge (vMerge)
    # In DOCX XML specification:
    # - <w:vMerge w:val="restart"/> = start of vertical merge (explicit)
    # - <w:vMerge/> (no val attribute) = continuation of vertical merge
    # However, some Word versions may use <w:vMerge/> for start cells too.
    # We use a more reliable method: check if cell has actual paragraph content.
    # Start cells have paragraph elements with text, continuation cells are usually empty.
    vMerge = tc.find(qn('w:vMerge'))
    if vMerge is not None:
        val = vMerge.get(qn('w:val'))
        # If val is "restart", it's definitely the start
        if val == 'restart':
            is_merged_start = True
            rowspan = 1  # Will be calculated by get_merged_cell_range_docx
        # If val is None, we need to determine if it's start or continuation
        # Check if cell has paragraph elements with actual text content
        elif val is None:
            # Check XML directly for paragraph elements with text
            # Start cells have <w:p> elements with <w:t> text elements
            # Continuation cells may have empty paragraphs or no paragraphs
            has_text_content = False
            # Check paragraphs in XML
            paragraphs = tc.findall(qn('w:p'))
            for para_elem in paragraphs:
                # Check for text elements in this paragraph
                text_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                for text_elem in text_elems:
                    if text_elem.text and text_elem.text.strip():
                        has_text_content = True
                        break
                if has_text_content:
                    break
            # If no text in XML, also check via python-docx API as fallback
            if not has_text_content:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        has_text_content = True
                        break
            # If cell has text content, it's likely the start (heuristic)
            if has_text_content:
                is_merged_start = True
                rowspan = 1
    
    # Additional check: Look for vMerge in tcPr (table cell properties)
    # Some Word versions may store vMerge differently (as seen in the logs: vMerge_in_tcPr=True)
    if not is_merged_start and rowspan is None:
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            vMerge_in_tcPr = tcPr.find(qn('w:vMerge'))
            if vMerge_in_tcPr is not None:
                val = vMerge_in_tcPr.get(qn('w:val'))
                # If val is "restart", it's definitely the start
                if val == 'restart':
                    is_merged_start = True
                    rowspan = 1
                # If val is None, check if cell has text content to determine if it's start
                elif val is None:
                    has_text_content = False
                    # Check paragraphs in XML
                    paragraphs = tc.findall(qn('w:p'))
                    for para_elem in paragraphs:
                        text_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                        for text_elem in text_elems:
                            if text_elem.text and text_elem.text.strip():
                                has_text_content = True
                                break
                        if has_text_content:
                            break
                    # If no text in XML, also check via python-docx API as fallback
                    if not has_text_content:
                        for para in cell.paragraphs:
                            if para.text and para.text.strip():
                                has_text_content = True
                                break
                    # If cell has text content, it's likely the start
                    if has_text_content:
                        is_merged_start = True
                        rowspan = 1
    
    # Check for horizontal merge (gridSpan)
    gridSpan = tc.find(qn('w:gridSpan'))
    if gridSpan is not None:
        colspan_val = gridSpan.get(qn('w:val'))
        if colspan_val:
            try:
                colspan = int(colspan_val)
                if colspan > 1:
                    is_merged_start = True
                    logger.debug(
                        LogModule.EXTRACT,
                        f"[DOCX_MERGED_CELL] Cell has gridSpan={colspan_val}, "
                        f"detected as horizontal merge start"
                    )
            except (ValueError, TypeError):
                colspan = None
        else:
            logger.debug(
                LogModule.EXTRACT,
                f"[DOCX_MERGED_CELL] Cell has gridSpan element but no val attribute"
            )
    
    # Additional check: Look for gridSpan in tcPr (table cell properties)
    # Some Word versions may store gridSpan differently
    if colspan is None:
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            gridSpan_in_tcPr = tcPr.find(qn('w:gridSpan'))
            if gridSpan_in_tcPr is not None:
                colspan_val = gridSpan_in_tcPr.get(qn('w:val'))
                if colspan_val:
                    try:
                        colspan = int(colspan_val)
                        if colspan > 1:
                            is_merged_start = True
                            logger.debug(
                                LogModule.EXTRACT,
                                f"[DOCX_MERGED_CELL] Cell has gridSpan in tcPr={colspan_val}, "
                                f"detected as horizontal merge start"
                            )
                    except (ValueError, TypeError):
                        colspan = None
    
    return (is_merged_start, rowspan, colspan)


def is_merged_cell_continuation_docx(cell) -> bool:
    """
    Check if a DOCX cell is a continuation of a merged region (not the start).
    
    Args:
        cell: DOCX Cell object
        
    Returns:
        True if this cell is part of a merged region but not the start
    """
    tc = cell._tc
    vMerge = tc.find(qn('w:vMerge'))
    if vMerge is not None:
        val = vMerge.get(qn('w:val'))
        # If val is "restart", it's the start, not a continuation
        if val == 'restart':
            return False
        # If val is None, check if cell has text content
        # According to DOCX spec, continuation cells typically have no text
        # (they inherit from the start cell)
        if val is None:
            # Check XML directly for paragraph elements with text
            has_text_content = False
            paragraphs = tc.findall(qn('w:p'))
            for para_elem in paragraphs:
                text_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                for text_elem in text_elems:
                    if text_elem.text and text_elem.text.strip():
                        has_text_content = True
                        break
                if has_text_content:
                    break
            # If no text in XML, also check via python-docx API as fallback
            if not has_text_content:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        has_text_content = True
                        break
            # If no text, it's likely a continuation
            if not has_text_content:
                return True
            # If has text, it might be the start (heuristic used in is_merged_cell_start_docx)
            # So we return False (not a continuation) to be consistent
            return False
    
    # Additional check: Look for vMerge in tcPr (table cell properties)
    # Some Word versions may store vMerge differently
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        vMerge_in_tcPr = tcPr.find(qn('w:vMerge'))
        if vMerge_in_tcPr is not None:
            val = vMerge_in_tcPr.get(qn('w:val'))
            # If val is "restart", it's the start, not a continuation
            if val == 'restart':
                return False
            # If val is None, check if cell has text content
            if val is None:
                has_text_content = False
                paragraphs = tc.findall(qn('w:p'))
                for para_elem in paragraphs:
                    text_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    for text_elem in text_elems:
                        if text_elem.text and text_elem.text.strip():
                            has_text_content = True
                            break
                    if has_text_content:
                        break
                if not has_text_content:
                    for para in cell.paragraphs:
                        if para.text and para.text.strip():
                            has_text_content = True
                            break
                # If no text, it's likely a continuation
                if not has_text_content:
                    return True
                # If has text, it might be the start
                return False
    
    return False


def get_merged_cell_range_docx(table, start_row: int, start_col: int) -> Tuple[int, int, int, int]:
    """
    Calculate the merged cell range for a DOCX table.
    
    Args:
        table: DOCX Table object
        start_row: Starting row index (0-based)
        start_col: Starting column index (0-based)
        
    Returns:
        Tuple of (start_row, start_col, end_row, end_col)
        - end_row and end_col are inclusive
    """
    if start_row >= len(table.rows) or start_col >= len(table.rows[start_row].cells):
        return (start_row, start_col, start_row, start_col)
    
    start_cell = table.rows[start_row].cells[start_col]
    _, rowspan, colspan = is_merged_cell_start_docx(start_cell)
    
    end_row = start_row
    end_col = start_col
    
    # Calculate colspan (horizontal merge)
    if colspan and colspan > 1:
        end_col = start_col + colspan - 1
        logger.debug(
            LogModule.EXTRACT,
            f"[DOCX_MERGED_CELL] Horizontal merge detected: "
            f"start_col={start_col}, colspan={colspan}, end_col={end_col}"
        )
    
    # Calculate rowspan (vertical merge)
    # Check if start cell has vMerge with val="restart" OR val=None with text content
    # Also check tcPr for vMerge (some Word versions store it there)
    tc = start_cell._tc
    vMerge = tc.find(qn('w:vMerge'))
    has_vertical_merge = False
    if vMerge is not None:
        val = vMerge.get(qn('w:val'))
        if val == 'restart':
            has_vertical_merge = True
        elif val is None:
            # Check if this cell has text content (heuristic: if it has text, it might be the start)
            # This handles cases where Word doesn't set val="restart"
            has_text = False
            for para in start_cell.paragraphs:
                if para.text and para.text.strip():
                    has_text = True
                    break
            if has_text:
                has_vertical_merge = True
                logger.debug(
                    LogModule.EXTRACT,
                    f"[DOCX_MERGED_CELL] Vertical merge detected (heuristic): "
                    f"vMerge exists with val=None, but cell has text content"
                )
    
    # Additional check: Look for vMerge in tcPr (table cell properties)
    # Some Word versions may store vMerge differently
    if not has_vertical_merge:
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            vMerge_in_tcPr = tcPr.find(qn('w:vMerge'))
            if vMerge_in_tcPr is not None:
                val = vMerge_in_tcPr.get(qn('w:val'))
                if val == 'restart':
                    has_vertical_merge = True
                    logger.debug(
                        LogModule.EXTRACT,
                        f"[DOCX_MERGED_CELL] Vertical merge detected in tcPr: "
                        f"vMerge in tcPr with val='restart'"
                    )
                elif val is None:
                    # Check if this cell has text content (heuristic: if it has text, it might be the start)
                    has_text = False
                    for para in start_cell.paragraphs:
                        if para.text and para.text.strip():
                            has_text = True
                            break
                    if has_text:
                        has_vertical_merge = True
                        logger.debug(
                            LogModule.EXTRACT,
                            f"[DOCX_MERGED_CELL] Vertical merge detected in tcPr (heuristic): "
                            f"vMerge in tcPr with val=None, but cell has text content"
                        )
    
    if has_vertical_merge:
        # Traverse down to find where the vertical merge ends
        # Special handling: if cells below also have vMerge in tcPr with val='restart',
        # they are also part of the merge (Word marks all cells in a vertical merge as 'restart')
        current_row = start_row + 1
        while current_row < len(table.rows):
            if start_col >= len(table.rows[current_row].cells):
                break
            cell = table.rows[current_row].cells[start_col]
            
            # Check if this is a continuation cell (standard case)
            if is_merged_cell_continuation_docx(cell):
                end_row = current_row
                current_row += 1
                continue
            
            # Special case: if this cell also has vMerge in tcPr with val='restart',
            # we need to check if it's a continuation of the current merge or a new merge start
            cell_tc = cell._tc
            cell_tcPr = cell_tc.find(qn('w:tcPr'))
            if cell_tcPr is not None:
                cell_vMerge_in_tcPr = cell_tcPr.find(qn('w:vMerge'))
                if cell_vMerge_in_tcPr is not None:
                    cell_val = cell_vMerge_in_tcPr.get(qn('w:val'))
                    if cell_val == 'restart':
                        # Check if this cell has text content
                        cell_has_text = False
                        cell_text_preview = ""
                        cell_text_full = ""
                        for para in cell.paragraphs:
                            para_text = para.text.strip() if para.text else ""
                            if para_text:
                                cell_has_text = True
                                cell_text_preview = para_text[:50]
                                cell_text_full = para_text
                                break
                        
                        # Check if the cell above has text content
                        above_cell_has_text = False
                        above_cell_text_preview = ""
                        above_cell_text_full = ""
                        if current_row > 0 and start_col < len(table.rows[current_row - 1].cells):
                            above_cell = table.rows[current_row - 1].cells[start_col]
                            for para in above_cell.paragraphs:
                                para_text = para.text.strip() if para.text else ""
                                if para_text:
                                    above_cell_has_text = True
                                    above_cell_text_preview = para_text[:50]
                                    above_cell_text_full = para_text
                                    break
                        
                        # Log detailed information for debugging
                        logger.debug(
                            LogModule.EXTRACT,
                            f"[DOCX_MERGED_CELL] Checking cell at ({current_row}, {start_col}) for merge continuation: "
                            f"cell_has_text={cell_has_text}, cell_text='{cell_text_preview}', "
                            f"above_cell_has_text={above_cell_has_text}, above_cell_text='{above_cell_text_preview}'"
                        )
                        
                        # Determine if this is a new merge start:
                        # 1. If current cell has text and above cell has no text -> new merge start
                        # 2. If current cell has text and above cell has text but different -> new merge start
                        #    (This handles cases where Word copies text to continuation cells, but new merge has different text)
                        is_new_merge_start = False
                        if cell_has_text:
                            if not above_cell_has_text:
                                # Case 1: Current has text, above has no text -> new merge start
                                is_new_merge_start = True
                            elif cell_text_full != above_cell_text_full:
                                # Case 2: Both have text but different -> new merge start
                                is_new_merge_start = True
                                logger.debug(
                                    LogModule.EXTRACT,
                                    f"[DOCX_MERGED_CELL] Cell at ({current_row}, {start_col}) has different text "
                                    f"('{cell_text_preview}') than above cell ('{above_cell_text_preview}'). "
                                    f"This indicates a new merge start."
                                )
                        
                        if is_new_merge_start:
                            # This is a new merge start, stop the current merge
                            logger.info(
                                LogModule.EXTRACT,
                                f"[DOCX_MERGED_CELL] Cell at ({current_row}, {start_col}) has vMerge='restart' with text='{cell_text_preview}', "
                                f"and cell above ({current_row - 1}, {start_col}) has text='{above_cell_text_preview}' "
                                f"(has_text={above_cell_has_text}). This is a new merge start, "
                                f"stopping current merge at row {current_row - 1}"
                            )
                            break
                        else:
                            # This cell is part of the current merge
                            logger.debug(
                                LogModule.EXTRACT,
                                f"[DOCX_MERGED_CELL] Cell at ({current_row}, {start_col}) is part of current merge: "
                                f"cell_has_text={cell_has_text}, above_cell_has_text={above_cell_has_text}, "
                                f"texts_match={cell_text_full == above_cell_text_full if (cell_has_text and above_cell_has_text) else 'N/A'}"
                            )
                            end_row = current_row
                            current_row += 1
                            continue
            
            # If we get here, this cell is not part of the merge
            break
        
        logger.debug(
            LogModule.EXTRACT,
            f"[DOCX_MERGED_CELL] Vertical merge range: "
            f"start_row={start_row}, end_row={end_row}"
        )
    
    return (start_row, start_col, end_row, end_col)


def _find_merged_region_start_docx(table, row: int, col: int) -> Tuple[int, int]:
    """
    Find the start of a merged region given any cell in the region.
    
    Args:
        table: DOCX Table object
        row: Row index (0-based)
        col: Column index (0-based)
        
    Returns:
        Tuple of (start_row, start_col)
    """
    current_row = row
    current_col = col
    
    # Move up to find vertical merge start
    # Special handling: if current cell has vMerge in tcPr with val='restart',
    # and the cell above also has vMerge in tcPr with val='restart',
    # then continue moving up to find the true start
    while current_row > 0:
        if current_col >= len(table.rows[current_row - 1].cells):
            break
        above_cell = table.rows[current_row - 1].cells[current_col]
        
        # Check if above cell is a continuation cell
        if is_merged_cell_continuation_docx(above_cell):
            current_row -= 1
            continue
        
        # Get vMerge status for current and above cells
        current_tc = table.rows[current_row].cells[current_col]._tc
        current_tcPr = current_tc.find(qn('w:tcPr'))
        current_has_vmerge_restart = False
        if current_tcPr is not None:
            current_vMerge_in_tcPr = current_tcPr.find(qn('w:vMerge'))
            if current_vMerge_in_tcPr is not None:
                current_val = current_vMerge_in_tcPr.get(qn('w:val'))
                if current_val == 'restart':
                    current_has_vmerge_restart = True
        
        above_tc = above_cell._tc
        above_tcPr = above_tc.find(qn('w:tcPr'))
        above_has_vmerge_restart = False
        if above_tcPr is not None:
            above_vMerge_in_tcPr = above_tcPr.find(qn('w:vMerge'))
            if above_vMerge_in_tcPr is not None:
                above_val = above_vMerge_in_tcPr.get(qn('w:val'))
                if above_val == 'restart':
                    above_has_vmerge_restart = True
        
        # If both have val='restart', continue moving up
        if current_has_vmerge_restart and above_has_vmerge_restart:
            logger.debug(
                LogModule.EXTRACT,
                f"[DOCX_MERGED_CELL] _find_merged_region_start_docx: "
                f"Cell ({current_row}, {current_col}) and above cell ({current_row - 1}, {current_col}) "
                f"both have vMerge in tcPr with val='restart'. Continuing to move up."
            )
            current_row -= 1
            continue
        
        # If above cell has val='restart' but current cell doesn't, above cell is the start
        if above_has_vmerge_restart:
            logger.debug(
                LogModule.EXTRACT,
                f"[DOCX_MERGED_CELL] _find_merged_region_start_docx: "
                f"Above cell ({current_row - 1}, {current_col}) has vMerge in tcPr with val='restart', "
                f"but current cell ({current_row}, {current_col}) doesn't. "
                f"Above cell is the start."
            )
            current_row -= 1
            continue
        
        # If we get here, the above cell is not part of the merge
        break
    
    # Move left to find horizontal merge start
    # Check if previous cells have colspan that includes current column
    while current_col > 0:
        if current_row >= len(table.rows):
            break
        prev_cell = table.rows[current_row].cells[current_col - 1]
        _, _, prev_colspan = is_merged_cell_start_docx(prev_cell)
        if prev_colspan and current_col < (current_col - 1) + prev_colspan:
            current_col -= 1
        else:
            break
    
    return (current_row, current_col)


def get_all_merged_regions_docx(table) -> List[Tuple[int, int, int, int]]:
    """
    Get all merged cell regions in a DOCX table.
    
    Args:
        table: DOCX Table object
        
    Returns:
        List of (start_row, start_col, end_row, end_col) tuples
    """
    merged_regions = []
    processed_cells: Set[Tuple[int, int]] = set()
    
    # Debug: log all cells to understand the structure
    logger.trace(LogModule.EXTRACT, f"[DOCX_MERGED_CELL] Scanning table with {len(table.rows)} rows")
    
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if (row_idx, col_idx) in processed_cells:
                continue
            
            # Check if this cell is part of a merged region
            is_start, rowspan, colspan = is_merged_cell_start_docx(cell)
            is_continuation = is_merged_cell_continuation_docx(cell)
            
            # Debug: log cell properties and XML structure
            try:
                tc = cell._tc
                vMerge = tc.find(qn('w:vMerge'))
                gridSpan = tc.find(qn('w:gridSpan'))
                vMerge_val = vMerge.get(qn('w:val')) if vMerge is not None else None
                gridSpan_val = gridSpan.get(qn('w:val')) if gridSpan is not None else None
                cell_text = cell.text[:50] if cell.text else "(empty)"
                
                # Check tcPr for merge attributes (some Word versions store them there)
                tcPr = tc.find(qn('w:tcPr'))
                vMerge_in_tcPr = None
                gridSpan_in_tcPr = None
                if tcPr is not None:
                    vMerge_in_tcPr = tcPr.find(qn('w:vMerge'))
                    gridSpan_in_tcPr = tcPr.find(qn('w:gridSpan'))
                
                # Log full XML structure for problematic cells (Row 10, Col 0-3 and Row 1-9, Col 4)
                should_log_xml = (
                    (row_idx == 10 and col_idx in [0, 1, 2, 3]) or
                    (row_idx in range(1, 10) and col_idx == 4)
                )
                
                if should_log_xml:
                    cell_xml_preview = tc.xml[:800] if hasattr(tc, 'xml') else "N/A"
                    logger.trace(
                        LogModule.EXTRACT,
                        f"[DOCX_MERGED_CELL] Table row {row_idx}, col {col_idx} (XML DEBUG): "
                        f"vMerge={vMerge is not None}, vMerge_val={vMerge_val}, "
                        f"gridSpan={gridSpan is not None}, gridSpan_val={gridSpan_val}, "
                        f"vMerge_in_tcPr={vMerge_in_tcPr is not None}, "
                        f"gridSpan_in_tcPr={gridSpan_in_tcPr is not None}, "
                        f"XML preview: {cell_xml_preview}..."
                    )
                else:
                    logger.trace(
                        LogModule.EXTRACT,
                        f"[DOCX_MERGED_CELL] Table row {row_idx}, col {col_idx}: "
                        f"is_start={is_start}, is_continuation={is_continuation}, "
                        f"vMerge={vMerge is not None}, vMerge_val={vMerge_val}, "
                        f"gridSpan={gridSpan is not None}, gridSpan_val={gridSpan_val}, "
                        f"rowspan={rowspan}, colspan={colspan}, "
                        f"text='{cell_text}...'"
                    )
            except Exception as e:
                logger.warning(LogModule.EXTRACT, f"[DOCX_MERGED_CELL] Error logging cell ({row_idx}, {col_idx}): {e}")
            
            if is_start:
                # Special handling for vertical merge: if cell has vMerge in tcPr with val='restart',
                # but the cell above in the same column also has vMerge in tcPr with val='restart',
                # then this cell is actually a continuation cell, not a start cell.
                # This handles cases where Word marks all cells in a vertical merge as 'restart'.
                if rowspan is not None and row_idx > 0:
                    # Check if this is a vertical merge
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    has_vmerge_restart = False
                    if tcPr is not None:
                        vMerge_in_tcPr = tcPr.find(qn('w:vMerge'))
                        if vMerge_in_tcPr is not None:
                            val = vMerge_in_tcPr.get(qn('w:val'))
                            if val == 'restart':
                                has_vmerge_restart = True
                    
                    # If this cell has vMerge in tcPr with val='restart', check the cell above
                    if has_vmerge_restart and col_idx < len(table.rows[row_idx - 1].cells):
                        above_cell = table.rows[row_idx - 1].cells[col_idx]
                        above_tc = above_cell._tc
                        above_tcPr = above_tc.find(qn('w:tcPr'))
                        if above_tcPr is not None:
                            above_vMerge_in_tcPr = above_tcPr.find(qn('w:vMerge'))
                            if above_vMerge_in_tcPr is not None:
                                above_val = above_vMerge_in_tcPr.get(qn('w:val'))
                                if above_val == 'restart':
                                    # Check if this cell has text content
                                    current_cell_has_text = False
                                    current_cell_text_full = ""
                                    for para in cell.paragraphs:
                                        para_text = para.text.strip() if para.text else ""
                                        if para_text:
                                            current_cell_has_text = True
                                            current_cell_text_full = para_text
                                            break
                                    
                                    # Check if the cell above has text content
                                    above_cell_has_text = False
                                    above_cell_text_full = ""
                                    for para in above_cell.paragraphs:
                                        para_text = para.text.strip() if para.text else ""
                                        if para_text:
                                            above_cell_has_text = True
                                            above_cell_text_full = para_text
                                            break
                                    
                                    # Determine if this is a new merge start:
                                    # 1. If current cell has text and above cell has no text -> new merge start
                                    # 2. If current cell has text and above cell has text but different -> new merge start
                                    is_new_merge_start = False
                                    if current_cell_has_text:
                                        if not above_cell_has_text:
                                            # Case 1: Current has text, above has no text -> new merge start
                                            is_new_merge_start = True
                                        elif current_cell_text_full != above_cell_text_full:
                                            # Case 2: Both have text but different -> new merge start
                                            is_new_merge_start = True
                                            logger.debug(
                                                LogModule.EXTRACT,
                                                f"[DOCX_MERGED_CELL] Cell ({row_idx}, {col_idx}) has different text "
                                                f"('{current_cell_text_full[:50]}') than above cell ('{above_cell_text_full[:50]}'). "
                                                f"This indicates a new merge start."
                                            )
                                    
                                    if is_new_merge_start:
                                        # This is a new merge start, not a continuation
                                        logger.debug(
                                            LogModule.EXTRACT,
                                            f"[DOCX_MERGED_CELL] Cell ({row_idx}, {col_idx}) has vMerge='restart' with text='{current_cell_text_full[:50]}', "
                                            f"and cell above ({row_idx - 1}, {col_idx}) has vMerge='restart' with text='{above_cell_text_full[:50] if above_cell_has_text else '(no text)'}'. "
                                            f"This is a new merge start, not a continuation."
                                        )
                                        # Continue processing this cell as a start cell
                                    else:
                                        # The cell above also has val='restart', and either:
                                        # - current cell has no text (continuation), or
                                        # - above cell also has text and same (both are starts, but current is continuation)
                                        # So this is a continuation cell
                                        logger.debug(
                                            LogModule.EXTRACT,
                                            f"[DOCX_MERGED_CELL] Cell ({row_idx}, {col_idx}) has vMerge in tcPr with val='restart', "
                                            f"but cell above ({row_idx - 1}, {col_idx}) also has val='restart'. "
                                            f"Current has_text={current_cell_has_text}, above has_text={above_cell_has_text}, "
                                            f"texts_match={current_cell_text_full == above_cell_text_full if (current_cell_has_text and above_cell_has_text) else 'N/A'}. "
                                            f"Treating as continuation cell."
                                        )
                                        # Find the actual start cell
                                        start_row, start_col = _find_merged_region_start_docx(table, row_idx, col_idx)
                                        
                                        logger.debug(
                                            LogModule.EXTRACT,
                                            f"[DOCX_MERGED_CELL] _find_merged_region_start_docx returned start=({start_row}, {start_col}) "
                                            f"for continuation cell at ({row_idx}, {col_idx})"
                                        )
                                        
                                        if (start_row, start_col) not in processed_cells:
                                            range_info = get_merged_cell_range_docx(table, start_row, start_col)
                                            merged_regions.append(range_info)
                                            
                                            logger.trace(
                                                LogModule.EXTRACT,
                                                f"[DOCX_MERGED_CELL] Found merged region from continuation at ({row_idx}, {col_idx}): "
                                                f"start=({start_row}, {start_col}), range={range_info}"
                                            )
                                            
                                            # Mark all cells in this region as processed
                                            for r in range(range_info[0], range_info[2] + 1):
                                                for c in range(range_info[1], range_info[3] + 1):
                                                    processed_cells.add((r, c))
                                        continue  # Skip processing this cell as a start cell
                
                # This is a start cell, calculate its range
                range_info = get_merged_cell_range_docx(table, row_idx, col_idx)
                merged_regions.append(range_info)
                
                # Debug: log merged region found
                logger.trace(
                    LogModule.EXTRACT,
                    f"[DOCX_MERGED_CELL] Found merged region start at ({row_idx}, {col_idx}): "
                    f"range={range_info}"
                )
                
                # Mark all cells in this region as processed
                for r in range(range_info[0], range_info[2] + 1):
                    for c in range(range_info[1], range_info[3] + 1):
                        processed_cells.add((r, c))
            elif is_continuation:
                # This is a continuation cell, find the actual start
                start_row, start_col = _find_merged_region_start_docx(table, row_idx, col_idx)
                
                if (start_row, start_col) not in processed_cells:
                    range_info = get_merged_cell_range_docx(table, start_row, start_col)
                    merged_regions.append(range_info)
                    
                    # Debug: log merged region found from continuation
                    logger.trace(
                        LogModule.EXTRACT,
                        f"[DOCX_MERGED_CELL] Found merged region from continuation at ({row_idx}, {col_idx}): "
                        f"start=({start_row}, {start_col}), range={range_info}"
                    )
                    
                    # Mark all cells in this region as processed
                    for r in range(range_info[0], range_info[2] + 1):
                        for c in range(range_info[1], range_info[3] + 1):
                            processed_cells.add((r, c))
    
    logger.trace(LogModule.EXTRACT, f"[DOCX_MERGED_CELL] Total merged regions found: {len(merged_regions)} (XML-based only)")
    return merged_regions


def is_cell_in_merged_region_docx(
    table, 
    row: int, 
    col: int, 
    merged_regions: List[Tuple[int, int, int, int]]
) -> Tuple[bool, Optional[Tuple[int, int, int, int]]]:
    """
    Check if a cell is part of a merged region and return the region info.
    
    Args:
        table: DOCX Table object
        row: Row index (0-based)
        col: Column index (0-based)
        merged_regions: List of merged region tuples (start_row, start_col, end_row, end_col)
        
    Returns:
        Tuple of (is_in_merged_region, merge_range)
        - is_in_merged_region: True if cell is part of a merged region
        - merge_range: The merge range tuple if in merged region, None otherwise
    """
    for merge_range in merged_regions:
        start_row, start_col, end_row, end_col = merge_range
        if start_row <= row <= end_row and start_col <= col <= end_col:
            return (True, merge_range)
    return (False, None)


def is_merged_cell_start_at_position_docx(
    table,
    row: int,
    col: int,
    merged_regions: List[Tuple[int, int, int, int]]
) -> bool:
    """
    Check if a cell at the given position is the start of a merged region.
    
    Args:
        table: DOCX Table object
        row: Row index (0-based)
        col: Column index (0-based)
        merged_regions: List of merged region tuples
        
    Returns:
        True if this cell is the start of a merged region
    """
    for merge_range in merged_regions:
        start_row, start_col, end_row, end_col = merge_range
        if start_row == row and start_col == col:
            return True
    return False

