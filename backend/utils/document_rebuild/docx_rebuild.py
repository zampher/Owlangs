# SPDX-FileCopyrightText: 2026 Zamphersssss
# SPDX-License-Identifier: MPL-2.0

"""DOCX document rebuild from translation segments."""

import ast
import base64
import io
import os
import re
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple

from logger import unified_logger as logger
from logger.logger import LogModule
from utils.translation_segments import get_translation_segments

# Support placeholder IDs with path characters (e.g., "mobi7/Images/image00044.jpeg")
# Characters allowed: letters, numbers, underscore, dot, slash, hyphen
PLACEHOLDER_PATTERN = re.compile(r"<ph-([a-zA-Z0-9_./-]+)>")


def _apply_run_style(run, style, font_name, font_size, is_bold, is_italic):
    if style:
        run.style = style
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if is_bold is not None:
        run.bold = is_bold
    if is_italic is not None:
        run.italic = is_italic


def _copy_paragraph_format(src_para, dst_para) -> None:
    """Copy paragraph-level formatting from src_para to dst_para.

    Preserves alignment (left/center/right/justify), indentation, and spacing
    so the bilingual source paragraph matches the original layout.
    """
    try:
        # Alignment (left, center, right, justify, etc.)
        if src_para.alignment is not None:
            dst_para.alignment = src_para.alignment
    except Exception:
        pass
    try:
        pf = src_para.paragraph_format
        dpf = dst_para.paragraph_format
        # Indentation
        if pf.left_indent is not None:
            dpf.left_indent = pf.left_indent
        if pf.right_indent is not None:
            dpf.right_indent = pf.right_indent
        if pf.first_line_indent is not None:
            dpf.first_line_indent = pf.first_line_indent
        # Spacing
        if pf.space_before is not None:
            dpf.space_before = pf.space_before
        if pf.space_after is not None:
            dpf.space_after = pf.space_after
        if pf.line_spacing is not None:
            dpf.line_spacing = pf.line_spacing
    except Exception:
        pass


def _get_effective_font_size_pt(run, paragraph=None) -> Optional[float]:
    """Return the effective font size of *run* in points.

    Tries the run's explicit size first, then paragraph style, then the
    document Normal style, then the Word default (11 pt).
    """
    try:
        if run.font.size is not None:
            return run.font.size.pt
    except Exception:
        pass
    if paragraph is not None:
        try:
            style = paragraph.style
            if style is not None and style.font.size is not None:
                return style.font.size.pt
        except Exception:
            pass
        try:
            doc = paragraph.part.document
            normal = doc.styles["Normal"]
            if normal.font.size is not None:
                return normal.font.size.pt
        except Exception:
            pass
    return 11.0


def _apply_font_size_delta(paragraph, delta: float) -> None:
    """Adjust font size of every run in *paragraph* by *delta* (in points).

    When a run does not carry an explicit font size the effective size is
    resolved from the paragraph style or document defaults before applying
    the delta.  Minimum size is clamped to 0.5 pt.
    """
    if delta == 0.0:
        return
    from docx.shared import Pt
    for run in paragraph.runs:
        try:
            cur_pt = _get_effective_font_size_pt(run, paragraph) or 11.0
            run.font.size = Pt(max(0.5, cur_pt + delta))
        except Exception:
            pass


def _add_image_to_paragraph(paragraph, image_entry, style, font_name, font_size, is_bold, is_italic) -> bool:
    data_uri = image_entry.get("data")
    if not data_uri or "," not in data_uri or not data_uri.startswith("data:image/"):
        return False

    try:
        base64_data = data_uri.split(",", 1)[1]
        image_bytes = base64.b64decode(base64_data)
    except Exception as e:
        logger.warning(LogModule.RESTOR,f"[DOCX-IMAGE] Failed to decode base64 for placeholder '{image_entry.get('alt')}', error: {e}")
        return False

    alt_text = image_entry.get("alt", "")
    is_formula_or_table = "equation" in alt_text.lower() or "table" in alt_text.lower()
    
    # Determine max width based on image type
    # Formula and table images: smaller max width (4 inches) to match text size
    # Regular images: larger max width (6 inches)
    try:
        from docx.shared import Inches
        from PIL import Image
        
        # Get image dimensions to maintain aspect ratio
        image_stream = io.BytesIO(image_bytes)
        pil_image = Image.open(image_stream)
        original_width, original_height = pil_image.size
        
        # Set max width based on image type
        max_width_inches = 4.0 if is_formula_or_table else 6.0
        max_width_pixels = max_width_inches * 96  # Assume 96 DPI for conversion
        
        # Calculate width and height maintaining aspect ratio
        if original_width > max_width_pixels:
            width_inches = max_width_inches
            height_inches = (original_height / original_width) * max_width_inches
        else:
            # Use original size if smaller than max
            width_inches = original_width / 96.0
            height_inches = original_height / 96.0
        
        # Reset stream for add_picture
        image_stream.seek(0)
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(width_inches))
    except ImportError:
        # Fallback if PIL is not available: use default size
        logger.warning(LogModule.RESTOR,"[DOCX-IMAGE] PIL not available, using default image size")
        image_stream = io.BytesIO(image_bytes)
        run = paragraph.add_run()
        try:
            from docx.shared import Inches
            # Use smaller default size for formula/table images
            max_width = Inches(4.0) if is_formula_or_table else Inches(6.0)
            run.add_picture(image_stream, width=max_width)
        except Exception as e:
            logger.warning(LogModule.RESTOR,f"[DOCX-IMAGE] add_picture failed for placeholder '{alt_text}', error: {e}")
            return False
    except Exception as e:
        logger.warning(LogModule.RESTOR,f"[DOCX-IMAGE] Failed to process image size for placeholder '{alt_text}': {e}, using default size")
        # Fallback: use default size
        image_stream = io.BytesIO(image_bytes)
        run = paragraph.add_run()
        try:
            from docx.shared import Inches
            max_width = Inches(4.0) if is_formula_or_table else Inches(6.0)
            run.add_picture(image_stream, width=max_width)
        except Exception as e2:
            logger.warning(LogModule.RESTOR,f"[DOCX-IMAGE] add_picture failed for placeholder '{alt_text}', error: {e2}")
            return False

    if alt_text:
        caption_run = paragraph.add_run(f" ({alt_text})")
        _apply_run_style(caption_run, style, font_name, font_size, is_bold, is_italic)
    logger.info(LogModule.TRANS, f"[DOCX-IMAGE] Inserted image placeholder='{alt_text}' size={len(image_bytes)} bytes, width={width_inches if 'width_inches' in locals() else 'default'} inches")
    return True


def _apply_translation_to_run_range(
    paragraph,
    target_text: str,
    run_start: int,
    run_end: Optional[int],
) -> bool:
    """
    Write translated text only into runs [run_start, run_end), matching DocxTranslator._after_translate
    (proportional split across multiple runs). Preserves fonts/formatting on runs outside the range.

    Returns False if placeholders are present (caller should fall back to full-paragraph replace)
    or if the slice has no replaceable text runs.
    """
    if PLACEHOLDER_PATTERN.search(target_text or ""):
        return False

    from utils.docx_utils import is_image_run
    from translator.ai_translator.docx_translator import preserve_page_breaks_in_run

    para_runs = list(paragraph.runs)
    if not para_runs:
        return False

    re = len(para_runs) if run_end is None else min(int(run_end), len(para_runs))
    rs = max(0, min(int(run_start), len(para_runs)))
    if rs >= re:
        return False

    text_runs = []
    for ri in range(rs, re):
        r = para_runs[ri]
        if not is_image_run(r):
            text_runs.append(r)

    if not text_runs:
        return False

    ft = target_text or ""

    if len(text_runs) == 1:
        preserve_page_breaks_in_run(text_runs[0], ft)
        return True

    original_lengths = [len(run.text or "") for run in text_runs]
    total_o = sum(original_lengths)
    if total_o == 0:
        preserve_page_breaks_in_run(text_runs[0], ft)
        return True

    cur = 0
    for idx, run in enumerate(text_runs):
        proportion = original_lengths[idx] / total_o
        piece_len = int(len(ft) * proportion)
        if idx == len(text_runs) - 1:
            chunk = ft[cur:]
        else:
            chunk = ft[cur : cur + piece_len]
            cur += piece_len
        preserve_page_breaks_in_run(run, chunk)
    return True


def _insert_text_and_images(paragraph, text: str, image_data_map: Dict[str, Dict[str, str]], style, font_name, font_size, is_bold, is_italic):
    if not PLACEHOLDER_PATTERN.search(text or ""):
        if text:
            run = paragraph.add_run(text)
            _apply_run_style(run, style, font_name, font_size, is_bold, is_italic)
        return

    last_pos = 0
    for match in PLACEHOLDER_PATTERN.finditer(text):
        if match.start() > last_pos:
            preceding = text[last_pos:match.start()]
            if preceding:
                run = paragraph.add_run(preceding)
                _apply_run_style(run, style, font_name, font_size, is_bold, is_italic)

        placeholder_id = match.group(1)
        image_entry = image_data_map.get(placeholder_id)
        if not image_entry or not _add_image_to_paragraph(paragraph, image_entry, style, font_name, font_size, is_bold, is_italic):
            fallback_run = paragraph.add_run(match.group(0))
            _apply_run_style(fallback_run, style, font_name, font_size, is_bold, is_italic)

        last_pos = match.end()

    if last_pos < len(text):
        remaining = text[last_pos:]
        if remaining:
            run = paragraph.add_run(remaining)
            _apply_run_style(run, style, font_name, font_size, is_bold, is_italic)


def _coerce_optional_int(value: Any) -> Optional[int]:
    """Coerce JSON-serialized numeric metadata back to int."""
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _is_expected_untranslated_table_segment(
    segment: Dict[str, Any],
    source_text: str,
    target_text: str,
) -> bool:
    """Return True when source==target is expected (numeric, language_match, identifier)."""
    if not (source_text and target_text and source_text.strip() == target_text.strip()):
        return False

    is_excluded = segment.get("is_excluded", False)
    exclusion_reason = (segment.get("exclusion_reason") or "").lower()
    detected_reason = (segment.get("detected_exclusion_reason") or "").lower()
    combined_reasons = f"{exclusion_reason}|{detected_reason}"
    if is_excluded and any(
        token in combined_reasons
        for token in ("language_match", "identifier", "numeric", "number", "pure")
    ):
        return True

    stripped = source_text.strip()
    try:
        from exclusion.detection.identifier_detector import _is_pure_number

        if _is_pure_number(stripped):
            return True
    except Exception:
        pass

    if stripped.isdigit():
        return True
    return False


def _find_table_cell_paragraph_by_para_index(
    doc,
    table_idx: int,
    row_idx: int,
    cell_idx: int,
    para_index: Optional[int],
    para_index_map: Dict[Tuple, Any],
    calculate_para_index_for_cell,
) -> Optional[Any]:
    """Map global para_index to a top-level cell paragraph (matches extract order)."""
    from utils.docx_utils import paragraph_has_toc_field

    if para_index is None or table_idx >= len(doc.tables):
        return None
    table = doc.tables[table_idx]
    if row_idx >= len(table.rows) or cell_idx >= len(table.rows[row_idx].cells):
        return None
    cell = table.rows[row_idx].cells[cell_idx]

    body_para_count = calculate_para_index_for_cell(table_idx, row_idx, cell_idx)
    cell_local_idx = 0
    for para in cell.paragraphs:
        if not paragraph_has_toc_field(para):
            if body_para_count == para_index:
                para_key = (True, table_idx, row_idx, cell_idx, cell_local_idx)
                return para_index_map.get(para_key)
            body_para_count += 1
            cell_local_idx += 1
    return None


def rebuild_docx_document_from_segments(
    task_state: Dict[str, Any],
    translated_docx_document,
    bilingual_export: bool = False,
    target_first: bool = False,
    source_text_italic: bool = True,
    source_text_color: str = "gray",
    target_text_italic: bool = False,
    target_text_color: Optional[str] = None,
    source_text_font_size_delta: float = 0.0,
    target_text_font_size_delta: float = 0.0,
) -> Optional[Any]:
    """
    Rebuild DOCX Document from revised translation segments.
    
    This function modifies the translated DOCX document by updating paragraph
    texts based on revised segments, preserving styles and structure.
    
    Args:
        task_state: Task state dictionary containing translation_segments
        translated_docx_document: Translated DOCX Document object from workflow
        bilingual_export: If True, insert source text paragraphs alongside translations.
        target_first: If True and bilingual_export is True, place source before target paragraph.
        
    Returns:
        Modified DOCX Document, or None if rebuilding failed
    """
    from ir.document import Document
    
    segments_data = get_translation_segments(None, task_state)
    if not segments_data:
        logger.warning(LogModule.RESTOR,"No translation segments found for rebuilding DOCX document")
        return None
    
    segments = segments_data.get("segments", [])
    if not segments:
        logger.warning(LogModule.RESTOR,"Empty segments list, cannot rebuild DOCX document")
        return None
    
    # Sort segments by index
    segments.sort(key=lambda x: x.get("segment_index", 0))
    
    # Initialize table1_segments for trace logging (will be populated later)
    table1_segments = []
    
    try:
        # Import python-docx for DOCX manipulation
        from docx import Document as DocxDocument
        from docx.text.paragraph import Paragraph
        from docx.text.run import Run
        
        # Get the DOCX content bytes
        docx_bytes = translated_docx_document.content
        logger.info(LogModule.TRANS, f"[DOCX-IMAGE] Rebuild start: doc length={len(docx_bytes)} bytes")
        
        # Load DOCX from bytes
        docx_io = io.BytesIO(docx_bytes)
        doc = DocxDocument(docx_io)
        
        image_data_map = (
            task_state.get("translation_image_data_map")
            or task_state.get("image_data_map")
            or {}
        )

        # Import helper function for TOC detection
        from utils.docx_utils import (
            count_non_toc_paragraphs_in_cell,
            paragraph_has_toc_field,
        )
        
        # Build a mapping from segment index to (target_text, segment_info)
        segment_data_map = {}
        modified_segments = []
        # Note: table1_segments is already initialized at function start for trace logging
        
        for segment in segments:
            segment_index = segment.get("segment_index", -1)

            # Header/footer segments are already applied to the document by the
            # workflow translation phase (apply_headers_footers_flat).  They do
            # not have body segment_info and should not be rebuilt here.
            if segment.get("segment_type") == "header_footer":
                continue

            # CRITICAL: Use the same priority as frontend: modified_text ?? target_text ?? ''
            # Frontend uses: segment['modified_text'] ?? segment['target_text'] ?? ''
            # We must use the same logic to ensure consistency
            modified_text = segment.get("modified_text")
            target_text_raw = segment.get("target_text", "")
            # Use modified_text if available (not None), otherwise use target_text (same as frontend)
            target_text = modified_text if modified_text is not None else target_text_raw
            source_text = segment.get("source_text", "")
            
            # CRITICAL: Log if target_text is same as source_text for table cells (potential issue)
            segment_info = segment.get("segment_info", {})
            is_table_cell = segment_info.get("is_table_cell", False)
            if is_table_cell and target_text and source_text and target_text.strip() == source_text.strip():
                table_idx = segment_info.get("table_index")
                row_idx = segment_info.get("row_index")
                cell_idx = segment_info.get("cell_index")
                same_text_msg = (
                    f"[DOCX-REBUILD] Table cell segment {segment_index} (Table {table_idx}, Row {row_idx}, Cell {cell_idx}): "
                    f"target_text equals source_text. "
                    f"source='{source_text[:100]}...', target='{target_text[:100]}...', "
                    f"modified_text={'present' if modified_text is not None else 'None'}"
                )
                if _is_expected_untranslated_table_segment(segment, source_text, target_text):
                    logger.debug(LogModule.RESTOR, same_text_msg)
                else:
                    logger.warning(
                        LogModule.RESTOR,
                        f"{same_text_msg} This may indicate translation was not saved correctly.",
                    )
            # CRITICAL: Log table cell segments with target_text for debugging (especially for Table 1)
            if is_table_cell:
                table_idx = segment_info.get("table_index")
                row_idx = segment_info.get("row_index")
                cell_idx = segment_info.get("cell_index")
                # Log all table cell segments for Table 1 (table_index=0) to diagnose export issues
                if table_idx == 0:
                    logger.info(
                        LogModule.RESTOR,
                        f"[DOCX-REBUILD-TABLE1] Segment {segment_index} (Table {table_idx}, Row {row_idx}, Cell {cell_idx}): "
                        f"source='{source_text[:80] if source_text else '(empty)'}...', "
                        f"target='{target_text[:80] if target_text else '(empty)'}...', "
                        f"modified_text={'present' if modified_text is not None else 'None'}",
                    )
            is_modified = segment.get("modified", False) or segment.get("retry_count", 0) > 0
            is_failed = segment.get("is_failed", False)
            failure_reason = segment.get("failure_reason")
            is_excluded = segment.get("is_excluded", False)
            exclusion_reason = segment.get("exclusion_reason")
            
            # CRITICAL: Check if segment is cleared (status="cleared" or empty target_text with modified=True)
            # Cleared segments should be exported as empty string
            is_cleared = segment.get("status") == "cleared" or (not target_text and is_modified and segment.get("target_length", -1) == 0)
            
            # CRITICAL: Even if segment is marked as failed or excluded, if target_text exists and differs from source_text,
            # use the target_text. This ensures correct translations are not lost due to incorrect failure/exclusion detection.
            # This is especially important for language_match exclusions where translation might have been done before exclusion.
            # Only use source_text if target_text is empty or same as source_text.
            if (is_failed or is_excluded) and target_text and target_text.strip() != source_text.strip():
                logger.warning(LogModule.RESTOR,
                    f"[DOCX-REBUILD] Segment {segment_index} marked as {'FAILED' if is_failed else 'EXCLUDED'} "
                    f"(reason={failure_reason or exclusion_reason}), "
                    f"but target_text differs from source_text. Using target_text for export: "
                    f"source='{source_text[:50] if source_text else '(empty)'}...', "
                    f"target='{target_text[:50]}...'"
                )
            
            # CRITICAL: For table cells, ensure we always use target_text if it exists and differs from source_text,
            # even if it's empty or marked as failed. This is especially important for table cells where
            # translations might be incorrectly marked as failed but are actually correct.
            segment_info = segment.get("segment_info", {})
            is_table_cell = segment_info.get("is_table_cell", False)
            
            # For table cells, log detailed information about target_text availability
            if is_table_cell:
                table_idx = segment_info.get("table_index")
                row_idx = segment_info.get("row_index")
                cell_idx = segment_info.get("cell_index")
                logger.debug(LogModule.RESTOR,
                    f"[DOCX-REBUILD-TABLE] Processing table cell segment {segment_index}: "
                    f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}, "
                    f"source_text='{source_text[:50] if source_text else '(empty)'}...', "
                    f"target_text='{target_text[:50] if target_text else '(empty)'}...', "
                    f"is_failed={is_failed}, is_excluded={is_excluded}, "
                    f"target_differs_from_source={target_text.strip() != source_text.strip() if target_text and source_text else False}"
                )

            if segment_index >= 0:
                # CRITICAL: Include segment in segment_data_map if:
                # 1. It has target_text (even if same as source_text - frontend shows it, so we should export it)
                # 2. It is cleared (should be exported as empty string)
                # 3. For table cells: always include to ensure we can match and update (frontend shows all table cells)
                # 
                # The key principle: if frontend shows a translation (even if it's the same as source),
                # we should include it in segment_data_map so we can match and update it.
                # The update logic will determine whether to actually update based on current_element_text.
                
                # Determine final_target_text to use
                final_target_text = target_text if not is_cleared else ""
                
                # CRITICAL: Include segment if:
                # - It has target_text (even if same as source - frontend shows it)
                # - It is cleared (should be exported as empty)
                # - It's a table cell (always include for matching)
                should_include = False
                if final_target_text or is_cleared:
                    # Has target_text or is cleared - include it
                    should_include = True
                elif is_table_cell:
                    # For table cells, always include even if target_text is empty
                    # This ensures we can match and update table cells
                    should_include = True
                
                if should_include:
                    segment_data_map[segment_index] = {
                        "target_text": final_target_text,
                        "source_text": source_text,  # Store source_text for logging
                        "segment_info": segment_info,
                        "is_excluded": is_excluded,
                        "is_failed": is_failed,
                        "segment_type": segment.get("segment_type"),
                        "textbox_key": segment.get("textbox_key"),
                    }
                    if is_modified:
                        modified_segments.append(segment_index)
                    
                    
                    # Collect Table1 (table_index=0) segments for trace logging
                    if is_table_cell:
                        table_idx = segment_info.get("table_index")
                        if table_idx == 0:  # Table1
                            table1_segments.append({
                                "segment_index": segment_index,
                                "row_index": segment_info.get("row_index"),
                                "cell_index": segment_info.get("cell_index"),
                                "source_text": source_text,
                                "target_text": final_target_text,
                                "modified_text": modified_text if modified_text is not None else None,
                            })
        
        # CRITICAL: Log Table1 segments with trace level for debugging
        if table1_segments:
            # Sort by row_index and cell_index for better readability
            table1_segments.sort(key=lambda x: (x["row_index"], x["cell_index"]))
            logger.trace(
                LogModule.RESTOR,
                f"[DOCX-REBUILD-TABLE1] Table1 (table_index=0) segments collected for export: {len(table1_segments)} segments",
            )
            for seg in table1_segments:
                logger.trace(
                    LogModule.RESTOR,
                    f"[DOCX-REBUILD-TABLE1] Segment {seg['segment_index']} - "
                    f"Row {seg['row_index']}, Cell {seg['cell_index']}: "
                    f"source='{seg['source_text'][:100] if seg['source_text'] else '(empty)'}...', "
                    f"target='{seg['target_text'][:100] if seg['target_text'] else '(empty)'}...', "
                    f"modified_text={'present' if seg['modified_text'] is not None else 'None'}",
                )
        else:
            logger.trace(
                LogModule.RESTOR,
                "[DOCX-REBUILD-TABLE1] No Table1 segments found in segment_data_map",
            )
        
        hf_bilingual_segments = [
            s for s in segments
            if s.get("segment_type") == "header_footer"
            and (s.get("source_text") or "").strip()
            and not s.get("is_excluded")
            and not s.get("is_failed")
        ]

        if not segment_data_map:
            # All segments are non-body types (e.g. header_footer) that were
            # already applied to the document by the workflow translation phase.
            logger.info(LogModule.RESTOR,
                "No body segments found in segment_data_map — "
                "all segments are header/footer/textbox types already applied by workflow"
            )
            if bilingual_export and hf_bilingual_segments:
                hf_inserted = _insert_bilingual_header_footer_flat_segments(
                    doc,
                    hf_bilingual_segments,
                    target_first=target_first,
                    source_text_italic=source_text_italic,
                    source_text_color=source_text_color,
                    target_text_italic=target_text_italic,
                    target_text_color=target_text_color,
                    source_text_font_size_delta=source_text_font_size_delta,
                    target_text_font_size_delta=target_text_font_size_delta,
                )
                logger.info(
                    LogModule.RESTOR,
                    f"[BILINGUAL-DOCX] HF-only export inserted {hf_inserted} source items "
                    f"from {len(hf_bilingual_segments)} header/footer segments",
                )
                output_io = io.BytesIO()
                doc.save(output_io)
                output_io.seek(0)
                return Document.from_bytes(
                    content=output_io.read(),
                    suffix=translated_docx_document.suffix,
                    stem=translated_docx_document.stem,
                )
            output_io = io.BytesIO()
            doc.save(output_io)
            output_io.seek(0)
            new_doc = Document.from_bytes(
                content=output_io.read(),
                suffix=translated_docx_document.suffix,
                stem=translated_docx_document.stem,
            )
            return new_doc
        
        # Build para_index_map similar to DocxTranslator._pre_translate_with_metadata
        # This maps (is_table_cell, table_idx, row_idx, cell_idx, para_local_idx) -> paragraph
        para_index_map = {}
        para_count = 0
        
        # Map document body paragraphs (non-TOC)
        for para in doc.paragraphs:
            if not paragraph_has_toc_field(para):
                para_key = (False, None, None, None, para_count)
                para_index_map[para_key] = para
                para_count += 1
        
        # OPTIMIZATION: Pre-compute and cache merged regions for all tables
        # This avoids recalculating them multiple times during para_index counting
        from utils.table_utils import (
            get_all_merged_regions_docx,
            is_merged_cell_start_at_position_docx,
            is_cell_in_merged_region_docx,
        )
        
        # Cache merged regions for all tables
        table_merged_regions_cache: Dict[int, List[Tuple[int, int, int, int]]] = {}
        for table_idx, table in enumerate(doc.tables):
            table_merged_regions_cache[table_idx] = get_all_merged_regions_docx(table)
        
        # Helper function to iterate over table cells, skipping merged cell continuation parts
        def iterate_table_cells(table, table_idx: int, merged_regions: List[Tuple[int, int, int, int]]):
            """
            Iterate over table cells, skipping merged cell continuation parts.
            Yields (row_idx, cell_idx, cell) tuples for cells that should be processed.
            """
            processed_cells: set = set()
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Skip if already processed (part of a merged region)
                    if (row_idx, cell_idx) in processed_cells:
                        continue
                    
                    # Check if this cell is part of a merged region
                    is_in_merged, merge_range = is_cell_in_merged_region_docx(
                        table, row_idx, cell_idx, merged_regions
                    )
                    
                    if is_in_merged and merge_range is not None:
                        # Mark all cells in this merged region as processed
                        start_row, start_col, end_row, end_col = merge_range
                        for r in range(start_row, end_row + 1):
                            for c in range(start_col, end_col + 1):
                                processed_cells.add((r, c))
                        
                        # Only process if this is the start of the merged region
                        is_start = is_merged_cell_start_at_position_docx(
                            table, row_idx, cell_idx, merged_regions
                        )
                        if not is_start:
                            # Skip continuation cells (they are part of a merged region but not the start)
                            continue
                    
                    yield (row_idx, cell_idx, cell)
        
        # Map table cell paragraphs using the helper function
        for table_idx, table in enumerate(doc.tables):
            merged_regions = table_merged_regions_cache[table_idx]
            
            for row_idx, cell_idx, cell in iterate_table_cells(table, table_idx, merged_regions):
                # Map paragraphs in this cell (only if it's a start cell or not merged)
                cell_para_count = 0
                for para in cell.paragraphs:
                    if not paragraph_has_toc_field(para):
                        para_key = (True, table_idx, row_idx, cell_idx, cell_para_count)
                        para_index_map[para_key] = para
                        cell_para_count += 1
        
        # OPTIMIZATION: Helper function to calculate para_index for a given table cell position
        # This avoids recalculating merged regions and counting paragraphs multiple times
        def calculate_para_index_for_cell(
            target_table_idx: int,
            target_row_idx: int,
            target_cell_idx: int
        ) -> int:
            """
            Calculate the global para_index for a given table cell position.
            This matches the extraction logic: count all non-TOC paragraphs before this cell,
            skipping merged cell continuation parts.
            
            Returns:
                The global para_index (0-based) for the first paragraph in the target cell.
            """
            para_count = 0
            
            # Count document body paragraphs
            for p in doc.paragraphs:
                if not paragraph_has_toc_field(p):
                    para_count += 1
            
            # Count paragraphs in tables before target table
            for t_idx in range(target_table_idx):
                table_before = doc.tables[t_idx]
                merged_regions_before = table_merged_regions_cache[t_idx]
                
                for row_idx, cell_idx, cell in iterate_table_cells(table_before, t_idx, merged_regions_before):
                    para_count += count_non_toc_paragraphs_in_cell(cell, include_nested=True)
            
            # Count paragraphs in current table before target cell
            if target_table_idx < len(doc.tables):
                table = doc.tables[target_table_idx]
                merged_regions = table_merged_regions_cache[target_table_idx]
                processed_cells: set = set()
                
                # Count rows before target row
                for r_idx in range(target_row_idx):
                    for c_idx in range(len(table.rows[r_idx].cells)):
                        # Skip if already processed (part of a merged region)
                        if (r_idx, c_idx) in processed_cells:
                            continue
                        
                        # Check if this cell is part of a merged region
                        is_in_merged, merge_range = is_cell_in_merged_region_docx(
                            table, r_idx, c_idx, merged_regions
                        )
                        
                        if is_in_merged and merge_range is not None:
                            # Mark all cells in this merged region as processed
                            start_row, start_col, end_row, end_col = merge_range
                            for r_merged in range(start_row, end_row + 1):
                                for c_merged in range(start_col, end_col + 1):
                                    processed_cells.add((r_merged, c_merged))
                            
                            # Only process if this is the start of the merged region
                            is_start = is_merged_cell_start_at_position_docx(
                                table, r_idx, c_idx, merged_regions
                            )
                            if not is_start:
                                continue
                        
                        para_count += count_non_toc_paragraphs_in_cell(
                            table.rows[r_idx].cells[c_idx], include_nested=True
                        )
                
                # Count cells in target row before target cell
                for c_idx in range(target_cell_idx):
                    # Skip if already processed (part of a merged region)
                    if (target_row_idx, c_idx) in processed_cells:
                        continue
                    
                    # Check if this cell is part of a merged region
                    is_in_merged, merge_range = is_cell_in_merged_region_docx(
                        table, target_row_idx, c_idx, merged_regions
                    )
                    
                    if is_in_merged and merge_range is not None:
                        # Mark all cells in this merged region as processed
                        start_row, start_col, end_row, end_col = merge_range
                        for r_merged in range(start_row, end_row + 1):
                            for c_merged in range(start_col, end_col + 1):
                                processed_cells.add((r_merged, c_merged))
                        
                        # Only process if this is the start of the merged region
                        is_start = is_merged_cell_start_at_position_docx(
                            table, target_row_idx, c_idx, merged_regions
                        )
                        if not is_start:
                            continue
                    
                    para_count += count_non_toc_paragraphs_in_cell(
                        table.rows[target_row_idx].cells[c_idx], include_nested=True
                    )
            
            return para_count
        
        # Deep split yields multiple segments per paragraph; process in document order so run-range
        # edits compose predictably (para_index, then run_start_index).
        rebuild_segments_sorted = sorted(
            segment_data_map.items(),
            key=lambda kv: (
                kv[1]["segment_info"].get("para_index") or 0,
                kv[1]["segment_info"].get("run_start_index") or 0,
                kv[0],
            ),
        )

        # Update elements with revised text using segment_info for precise location
        updated_count = 0
        skipped_count = 0
        target_font_delta_applied: set = set()
        
        for segment_index, segment_data in rebuild_segments_sorted:
            target_text = segment_data["target_text"]
            source_text = segment_data.get("source_text", "")
            seg_info = segment_data["segment_info"]
            
            # Get original segment for additional metadata (is_failed, etc.)
            original_segment = None
            for seg in segments:
                if seg.get("segment_index") == segment_index:
                    original_segment = seg
                    break
            
            # CRITICAL: Ensure we use the same logic as frontend for getting target_text
            # Frontend uses: modified_text ?? target_text ?? ''
            # We should do the same to ensure consistency
            # If segment_data_map has empty target_text but original_segment has valid translation, use it
            if original_segment:
                # Get modified_text and target_text from original_segment (same priority as frontend)
                original_modified = original_segment.get("modified_text")
                original_target = original_segment.get("target_text", "")
                # Use modified_text if available, otherwise use target_text (same as frontend)
                original_final_target = original_modified if original_modified is not None else original_target
                original_source = original_segment.get("source_text", "")
                
                # CRITICAL: If target_text from segment_data_map is empty or same as source_text,
                # but original_segment has a valid translation that differs from source_text, use it.
                # This ensures we use the same data that frontend sees.
                if (not target_text or target_text.strip() == source_text.strip()) and original_final_target and original_final_target.strip() != original_source.strip():
                    target_text = original_final_target
                    logger.info(
                        LogModule.RESTOR,
                        f"[DOCX-REBUILD] Segment {segment_index}: "
                        f"target_text was empty or same as source in segment_data_map, but found valid translation in original_segment. "
                        f"Using original_final_target: '{original_final_target[:50]}...' (source: '{original_source[:50]}...')",
                    )
                # Also update source_text if it's different from what we have
                if original_source and original_source.strip() != source_text.strip():
                    source_text = original_source
            
            if not seg_info:
                # Segment info should always be available after refactoring
                # If not available, this is an error condition
                logger.error(LogModule.RESTOR,
                    f"[DOCX-REBUILD] Segment {segment_index} has no segment_info. "
                    "This should not happen after refactoring. Extract phase should always generate segment_info."
                )
                continue
            
            # Use segment_info for precise location
            para_index = seg_info.get('para_index')
            is_table_cell = seg_info.get('is_table_cell', False)
            table_idx = seg_info.get('table_index')
            row_idx = seg_info.get('row_index')
            cell_idx = seg_info.get('cell_index')
            run_start = seg_info.get('run_start_index', 0)
            run_end = seg_info.get('run_end_index')
            
            element = None
            target_stripped = (target_text or "").strip()
            source_stripped = (source_text or "").strip()
            cell_local_idx_hint = _coerce_optional_int(seg_info.get("cell_local_idx"))
            if cell_local_idx_hint is None and original_segment:
                cell_local_idx_hint = _coerce_optional_int(
                    (original_segment.get("segment_info") or {}).get("cell_local_idx")
                )

            if (
                is_table_cell
                and table_idx is not None
                and row_idx is not None
                and cell_idx is not None
            ):
                element = _resolve_table_cell_paragraph(
                    doc,
                    table_idx,
                    row_idx,
                    cell_idx,
                    target_stripped,
                    source_stripped,
                    para_index_map,
                    cell_local_idx_hint,
                )
                if element is not None:
                    logger.debug(
                        LogModule.RESTOR,
                        f"[DOCX-REBUILD] Resolved segment {segment_index} in table cell: "
                        f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}, "
                        f"target='{target_stripped[:60]}...'",
                    )

                if element is None and para_index is not None:
                    element = _find_table_cell_paragraph_by_para_index(
                        doc,
                        table_idx,
                        row_idx,
                        cell_idx,
                        para_index,
                        para_index_map,
                        calculate_para_index_for_cell,
                    )
                    if element is not None:
                        logger.debug(
                            LogModule.RESTOR,
                            f"[DOCX-REBUILD] Resolved segment {segment_index} by para_index={para_index}: "
                            f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}",
                        )
            else:
                # For non-table paragraphs, para_index is the document body paragraph index
                para_key = (False, None, None, None, para_index)
                element = para_index_map.get(para_key)

            if element is None:
                logger.warning(
                    LogModule.RESTOR,
                    f"[DOCX-REBUILD] Could not locate element for segment {segment_index}: "
                    f"para_index={para_index}, is_table_cell={is_table_cell}, "
                    f"table_idx={table_idx}, row_idx={row_idx}, cell_idx={cell_idx}, "
                    f"cell_local_idx={cell_local_idx_hint}, "
                    f"source_text='{source_stripped[:100] if source_stripped else '(empty)'}...', "
                    f"target_text='{target_stripped[:100] if target_stripped else '(empty)'}...'. "
                    f"Segment will be skipped.",
                )
                continue
            
            # CRITICAL: Only update if target_text is different from current element text
            # This prevents unnecessary updates and ensures we apply translations even if marked as failed
            current_element_text = element.text.strip() if element.text else ""
            target_text_stripped = target_text.strip() if target_text else ""
            source_text_stripped = source_text.strip() if source_text else ""
            
            # CRITICAL: For table cells, we need to ensure we apply translations correctly.
            # The key principle: if we have a valid target_text that differs from source_text,
            # we MUST update the element to apply the translation.
            is_table_cell_check = seg_info and seg_info.get('is_table_cell')
            should_update = False
            
            if is_table_cell_check:
                table_idx_log = seg_info.get('table_index')
                row_idx_log = seg_info.get('row_index')
                cell_idx_log = seg_info.get('cell_index')
                
                # CRITICAL: Primary rule - if target_text exists and differs from source_text, we should update
                # unless current_element_text already contains the target_text (already translated correctly)
                if target_text_stripped and target_text_stripped != source_text_stripped:
                    # We have a valid translation that differs from source
                    # Check if current text already contains target_text (already translated correctly)
                    resolve_score = _paragraph_text_match_score(
                        current_element_text,
                        target_text_stripped,
                        source_text_stripped,
                    )
                    if resolve_score < 30:
                        should_update = False
                        logger.warning(
                            LogModule.RESTOR,
                            f"[DOCX-REBUILD-TABLE] Segment {segment_index}: resolved paragraph does not "
                            f"match target/source (score={resolve_score}), skipping update to avoid "
                            f"corrupting Table {table_idx_log}, Row {row_idx_log}, Cell {cell_idx_log}: "
                            f"current='{current_element_text[:80]}...', "
                            f"target='{target_text_stripped[:80]}...'",
                        )
                    elif target_text_stripped in current_element_text:
                        # Current text already contains target_text, no update needed
                        should_update = False
                        logger.debug(LogModule.RESTOR,
                            f"[DOCX-REBUILD-TABLE] Segment {segment_index}: Table cell already contains target text, skipping update: "
                            f"Table {table_idx_log}, Row {row_idx_log}, Cell {cell_idx_log}, "
                            f"current='{current_element_text[:80]}...', target='{target_text_stripped[:80]}...'"
                        )
                    elif current_element_text != target_text_stripped:
                        # Current text is different from target, update needed
                        should_update = True
                        logger.info(
                            LogModule.RESTOR,
                            f"[DOCX-REBUILD-TABLE] Segment {segment_index}: Table cell will be updated with translation: "
                            f"Table {table_idx_log}, Row {row_idx_log}, Cell {cell_idx_log}, "
                            f"current='{current_element_text[:80]}...', source='{source_text_stripped[:80]}...', target='{target_text_stripped[:80]}...'"
                        )
                    else:
                        # Current text already matches target_text exactly, no update needed
                        should_update = False
                        logger.debug(LogModule.RESTOR,
                            f"[DOCX-REBUILD-TABLE] Segment {segment_index}: Table cell already has target text (exact match), skipping update: "
                            f"Table {table_idx_log}, Row {row_idx_log}, Cell {cell_idx_log}"
                        )
                else:
                    # No valid translation (target_text is empty or same as source_text)
                    should_update = False
                    logger.debug(LogModule.RESTOR,
                        f"[DOCX-REBUILD-TABLE] Segment {segment_index}: Table cell has no valid translation, skipping update: "
                        f"Table {table_idx_log}, Row {row_idx_log}, Cell {cell_idx_log}, "
                        f"target_text={'empty' if not target_text_stripped else 'same as source'}"
                    )
            else:
                # For non-table cells, use simple comparison
                should_update = target_text_stripped != current_element_text
            
            # CRITICAL: For table cells, log detailed information for debugging
            if is_table_cell_check:
                logger.info(
                    LogModule.RESTOR,
                    f"[DOCX-REBUILD-TABLE] Segment {segment_index} at Table {seg_info.get('table_index')}, "
                    f"Row {seg_info.get('row_index')}, Cell {seg_info.get('cell_index')}: "
                    f"current_element_text='{current_element_text[:100]}...', "
                    f"target_text='{target_text_stripped[:100]}...', "
                    f"source_text='{source_text_stripped[:100] if source_text_stripped else '(empty)'}...', "
                    f"current==source={current_element_text == source_text_stripped}, "
                    f"target!=source={target_text_stripped != source_text_stripped}, "
                    f"target_in_current={target_text_stripped in current_element_text if target_text_stripped else False}, "
                    f"source_in_current={source_text_stripped in current_element_text if source_text_stripped else False}, "
                    f"will_update={should_update}"
                )
            
            # CRITICAL: Log Table1 updates with trace level for debugging
            if is_table_cell_check and seg_info and seg_info.get('table_index') == 0:
                logger.trace(LogModule.RESTOR,
                    f"[DOCX-REBUILD-TABLE1-UPDATE] Segment {segment_index} - "
                    f"Row {seg_info.get('row_index')}, Cell {seg_info.get('cell_index')}: "
                    f"current_element_text='{current_element_text[:100]}...', "
                    f"target_text='{target_text_stripped[:100] if target_text_stripped else '(empty)'}...', "
                    f"source_text='{source_text_stripped[:100] if source_text_stripped else '(empty)'}...', "
                    f"will_update={should_update}",
                )
            
            # Check if update is needed
            if should_update:
                applied_slice = _apply_translation_to_run_range(
                    element, target_text, run_start, run_end
                )
                if not applied_slice:
                    logger.debug(
                        LogModule.RESTOR,
                        f"[DOCX-REBUILD] Run-range apply failed or skipped (placeholders/empty slice); "
                        f"fallback full-paragraph replace segment={segment_index} "
                        f"run_start={run_start} run_end={run_end}",
                    )
                    runs = element.runs
                    if runs:
                        first_run = runs[0]
                        style = first_run.style
                        font_name = first_run.font.name
                        font_size = first_run.font.size
                        is_bold = first_run.bold
                        is_italic = first_run.italic
                        element.clear()
                        _insert_text_and_images(
                            element,
                            target_text,
                            image_data_map,
                            style,
                            font_name,
                            font_size,
                            is_bold,
                            is_italic,
                        )
                    else:
                        _insert_text_and_images(
                            element,
                            target_text,
                            image_data_map,
                            None,
                            None,
                            None,
                            None,
                            None,
                        )

                updated_count += 1
                
                # Log update for debugging (especially for table cells)
                if seg_info and seg_info.get('is_table_cell'):
                    logger.info(
                        LogModule.RESTOR,
                        f"[DOCX-REBUILD] ✅ Updated segment {segment_index} at Table {seg_info.get('table_index')}, "
                        f"Row {seg_info.get('row_index')}, Cell {seg_info.get('cell_index')}: "
                        f"'{current_element_text[:50]}...' -> '{target_text_stripped[:50]}...'"
                    )
                    # CRITICAL: Log Table1 updates with trace level for debugging
                    if seg_info.get('table_index') == 0:
                        logger.trace(
                            LogModule.RESTOR,
                            f"[DOCX-REBUILD-TABLE1-UPDATED] Segment {segment_index} - "
                            f"Row {seg_info.get('row_index')}, Cell {seg_info.get('cell_index')}: "
                            f"✅ Successfully updated: '{current_element_text[:100]}...' -> '{target_text_stripped[:100]}...'",
                        )
            else:
                # No update needed - target_text matches current element text
                skipped_count += 1
                if seg_info and seg_info.get('is_table_cell'):
                    # For table cells, log more details about why update was skipped
                    logger.info(
                        LogModule.RESTOR,
                        f"[DOCX-REBUILD-TABLE] Skipped segment {segment_index} at Table {seg_info.get('table_index')}, "
                        f"Row {seg_info.get('row_index')}, Cell {seg_info.get('cell_index')}: "
                        f"target_text='{target_text_stripped[:50] if target_text_stripped else '(empty)'}...', "
                        f"current_element_text='{current_element_text[:50]}...', "
                        f"source_text='{source_text_stripped[:50] if source_text_stripped else '(empty)'}...', "
                        f"target==current={target_text_stripped == current_element_text}, "
                        f"target==source={target_text_stripped == source_text_stripped}"
                    )
                else:
                    logger.debug(LogModule.RESTOR,
                        f"[DOCX-REBUILD] Segment {segment_index} skipped (target_text matches current element text): "
                        f"'{target_text_stripped[:50] if target_text_stripped else '(empty)'}...'"
                    )

            # Apply target font-size delta once per resolved paragraph, including
            # segments skipped because translation is already present in the doc.
            if target_text_font_size_delta != 0.0 and element is not None:
                elem_key = id(element._element)
                if elem_key not in target_font_delta_applied:
                    _apply_font_size_delta(element, target_text_font_size_delta)
                    target_font_delta_applied.add(elem_key)
        
        if updated_count == 0 and not bilingual_export:
            logger.info(
                LogModule.RESTOR,
                "No text elements were updated in DOCX document — "
                "translations already present, returning document as-is",
            )
        
        # Bilingual export: insert source text paragraphs alongside translations
        if bilingual_export:
            extras_original = task_state.get("docx_extras_original", {})
            logger.info(
                LogModule.RESTOR,
                f"[BILINGUAL-DOCX] extras_original from task_state: keys={list(extras_original.keys()) if extras_original else 'empty'}"
            )
            if extras_original and "textboxes_sdts" in extras_original:
                tb_items = extras_original["textboxes_sdts"]
                tb_debug = []
                for item in tb_items:
                    if isinstance(item, (list, tuple)) and len(item) >= 2:
                        k = item[0]
                        tb_debug.append(f"type={type(k).__name__},val={k}")
                    else:
                        tb_debug.append(f"bad_item:type={type(item).__name__},val={item!r}")
                logger.info(
                    LogModule.RESTOR,
                    f"[BILINGUAL-DOCX] textboxes_sdts has {len(tb_items)} items: {tb_debug}"
                )
            _insert_bilingual_source_paragraphs(
                doc,
                segment_data_map,
                para_index_map,
                target_first=target_first,
                extras_original=extras_original,
                source_text_italic=source_text_italic,
                source_text_color=source_text_color,
                target_text_italic=target_text_italic,
                target_text_color=target_text_color,
                skip_legacy_header_footer=bool(hf_bilingual_segments),
                source_text_font_size_delta=source_text_font_size_delta,
                target_text_font_size_delta=target_text_font_size_delta,
            )
            if hf_bilingual_segments:
                hf_inserted = _insert_bilingual_header_footer_flat_segments(
                    doc,
                    hf_bilingual_segments,
                    target_first=target_first,
                    source_text_italic=source_text_italic,
                    source_text_color=source_text_color,
                    target_text_italic=target_text_italic,
                    target_text_color=target_text_color,
                    source_text_font_size_delta=source_text_font_size_delta,
                    target_text_font_size_delta=target_text_font_size_delta,
                )
                logger.info(
                    LogModule.RESTOR,
                    f"[BILINGUAL-DOCX] Flat header/footer export inserted {hf_inserted} source items",
                )
        
        # Save modified DOCX back to bytes
        output_io = io.BytesIO()
        doc.save(output_io)
        output_io.seek(0)
        new_bytes = output_io.read()
        
        # Create new Document object with updated content
        new_doc = Document.from_bytes(
            content=new_bytes,
            suffix=translated_docx_document.suffix,
            stem=translated_docx_document.stem
        )
        
        logger.info(
            LogModule.RESTOR,
            f"Rebuilt DOCX Document: updated {updated_count} elements from {len(segments)} segments "
            f"({len(modified_segments)} modified), skipped {skipped_count} (no change needed)"
        )
        return new_doc
        
    except ImportError:
        logger.error(LogModule.RESTOR,"python-docx not available, cannot rebuild DOCX document")
        return None
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        logger.error(LogModule.RESTOR,f"Failed to rebuild DOCX document: {e}\n{tb}")
        return None


def _normalize_key(k) -> tuple:
    """Convert a key to a canonical tuple form.

    Task state data may go through JSON serialization, which converts
    tuples to lists. This helper ensures we can compare reliably.
    """
    if isinstance(k, (list, tuple)):
        return tuple(_normalize_key(item) for item in k)
    return k


def _parse_textbox_sdt_storage_items(raw_items: List[Any]) -> List[Tuple[Any, str]]:
    """Parse textbox/SDT (key, source_text) pairs from task state storage."""
    parsed_items: List[Tuple[Any, str]] = []
    for item in raw_items or []:
        candidate = item
        if isinstance(candidate, str):
            try:
                candidate = ast.literal_eval(candidate)
            except Exception:
                logger.warning(
                    LogModule.RESTOR,
                    f"[BILINGUAL-DOCX] Failed to parse textbox/SDT item: {item[:120]!r}",
                )
                continue
        while isinstance(candidate, str):
            try:
                candidate = ast.literal_eval(candidate)
            except Exception:
                break
        if isinstance(candidate, (list, tuple)) and len(candidate) >= 2:
            key, text = candidate[0], candidate[1]
            if text and str(text).strip():
                parsed_items.append((key, str(text)))
            continue
        logger.warning(
            LogModule.RESTOR,
            f"[BILINGUAL-DOCX] Ignoring invalid textbox/SDT item: {item!r}",
        )
    return parsed_items


def _parse_textbox_key(raw_key: Any) -> Optional[tuple]:
    """Parse a textbox/SDT key from workflow segment metadata."""
    if raw_key is None:
        return None
    if isinstance(raw_key, (list, tuple)):
        return _normalize_key(raw_key)
    if isinstance(raw_key, str):
        try:
            parsed = ast.literal_eval(raw_key)
            if isinstance(parsed, (list, tuple)):
                return _normalize_key(parsed)
        except Exception:
            pass
    return None


def _is_duplicate_legacy_vtextbox_container(element) -> bool:
    """Return True when a legacy v:textbox duplicates a modern txbxContent path."""
    tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
    if tag != "textbox":
        return False
    alt = element.xpath('./ancestor::*[local-name()="AlternateContent"]')
    if alt:
        choice_txbx = alt[0].xpath(
            './/*[local-name()="Choice"]//*[local-name()="txbxContent"]'
        )
        in_fallback = element.xpath('./ancestor::*[local-name()="Fallback"]')
        if choice_txbx and in_fallback:
            return True
    for p in element.xpath('./ancestor::*[local-name()="p"]'):
        if p.xpath('.//*[local-name()="txbxContent"]'):
            return True
        break
    return False


def _build_choice_to_fallback_txbx_map(doc) -> Dict[int, Any]:
    """Map Choice txbxContent elements to their mc:Fallback counterparts."""
    mapping: Dict[int, Any] = {}
    for txbx in doc._element.xpath('.//*[local-name()="txbxContent"]'):
        alt = txbx.xpath('./ancestor::*[local-name()="AlternateContent"]')
        if not alt:
            continue
        in_fallback = txbx.xpath('./ancestor::*[local-name()="Fallback"]')
        choice_txbx = alt[0].xpath(
            './/*[local-name()="Choice"]//*[local-name()="txbxContent"]'
        )
        if in_fallback and choice_txbx:
            mapping[id(choice_txbx[0])] = txbx
    return mapping


def _find_textbox_target_paragraph_element(container, original_text: str):
    """Pick the paragraph that holds translated target text inside a textbox."""
    paras = container.xpath('.//*[local-name()="p"]')
    candidates = []
    for p in paras:
        runs = p.xpath('.//*[local-name()="t"]')
        text = "".join(t.text or "" for t in runs).strip()
        if text:
            candidates.append((p, text))
    if not candidates:
        return None
    normalized_source = _normalize_match_text(original_text)
    for p, text in candidates:
        if _normalize_match_text(text) != normalized_source:
            return p
    return candidates[-1][0]


def _mirror_textbox_source_to_fallback(
    choice_element,
    new_p_elem,
    target_first: bool,
    choice_to_fallback: Dict[int, Any],
) -> None:
    """Mirror an inserted source paragraph to the mc:Fallback txbxContent twin."""
    fallback = choice_to_fallback.get(id(choice_element))
    if fallback is None:
        return
    fb_paras = fallback.xpath('.//*[local-name()="p"]')
    if not fb_paras:
        return
    from copy import deepcopy

    fb_target = fb_paras[0]
    cloned = deepcopy(new_p_elem)
    if target_first:
        fb_target.addnext(cloned)
    else:
        fb_target.addprevious(cloned)


def _normalize_match_text(text: str) -> str:
    """Collapse whitespace for fuzzy paragraph text comparison."""
    return re.sub(r"\s+", " ", (text or "").strip())


def _effective_segment_target_text(segment_data: Dict[str, Any]) -> str:
    """Return target text using the same priority as the frontend."""
    modified_text = segment_data.get("modified_text")
    if modified_text is not None:
        return (modified_text or "").strip()
    return (segment_data.get("target_text") or "").strip()


def _paragraph_text_match_score(
    para_text: str,
    target_text: str,
    source_text: str,
) -> int:
    """Score how well a paragraph matches expected target/source text."""
    normalized_para = _normalize_match_text(para_text)
    normalized_target = _normalize_match_text(target_text)
    normalized_source = _normalize_match_text(source_text)
    if not normalized_para:
        return 0
    if normalized_target and normalized_para == normalized_target:
        return 100
    if normalized_target and normalized_target in normalized_para:
        return 85
    if normalized_target and normalized_para in normalized_target:
        return 75
    if normalized_source and normalized_para == normalized_source:
        return 40
    if normalized_source and normalized_source in normalized_para:
        return 30
    return 0


def _cell_local_idx_hint_accepts(
    hinted_para: Any,
    target_text: str,
    source_text: str,
    min_score: int = 30,
) -> bool:
    """Return True when hinted paragraph content matches segment target/source."""
    if hinted_para is None:
        return False
    if not target_text and not source_text:
        return True
    score = _paragraph_text_match_score(
        hinted_para.text or "", target_text, source_text
    )
    return score >= min_score


def _resolve_table_cell_paragraph(
    doc,
    table_idx: int,
    row_idx: int,
    cell_idx: int,
    target_text: str,
    source_text: str,
    para_index_map: Dict[Tuple, Any],
    cell_local_idx_hint: Optional[int] = None,
) -> Optional[Any]:
    """Locate the translated paragraph inside a table cell by metadata + text."""
    from utils.docx_utils import paragraph_has_toc_field

    if table_idx >= len(doc.tables):
        return None
    table = doc.tables[table_idx]
    if row_idx >= len(table.rows) or cell_idx >= len(table.rows[row_idx].cells):
        return None
    cell = table.rows[row_idx].cells[cell_idx]

    if cell_local_idx_hint is not None:
        hint_key = (True, table_idx, row_idx, cell_idx, cell_local_idx_hint)
        hinted = para_index_map.get(hint_key)
        if hinted is not None and _cell_local_idx_hint_accepts(
            hinted, target_text, source_text
        ):
            return hinted
        if hinted is not None:
            logger.debug(
                LogModule.RESTOR,
                f"[DOCX-REBUILD] Rejected stale cell_local_idx={cell_local_idx_hint} "
                f"for Table {table_idx}, Row {row_idx}, Cell {cell_idx}: "
                f"para='{(hinted.text or '')[:60]}...', target='{target_text[:60]}...'",
            )

    best_para = None
    best_score = 0
    cell_non_toc_idx = 0
    for para in cell.paragraphs:
        if paragraph_has_toc_field(para):
            continue
        score = _paragraph_text_match_score(para.text or "", target_text, source_text)
        mapped = para_index_map.get(
            (True, table_idx, row_idx, cell_idx, cell_non_toc_idx)
        )
        if mapped is para and score > 0:
            score += 5
        if score > best_score:
            best_score = score
            best_para = para
        cell_non_toc_idx += 1

    if best_para is not None and best_score >= 30:
        return best_para
    return None


def _resolve_bilingual_target_paragraph(
    doc,
    segment_data: Dict[str, Any],
    para_index_map: Dict[Tuple, Any],
    segment_index: int,
) -> Optional[Any]:
    """Resolve the translated paragraph for bilingual source insertion."""
    seg_info = segment_data.get("segment_info") or {}
    source_text = (segment_data.get("source_text") or "").strip()
    target_text = _effective_segment_target_text(segment_data)
    if not source_text:
        return None
    if target_text and source_text == target_text:
        return None

    is_table_cell = seg_info.get("is_table_cell", False)
    table_idx = seg_info.get("table_index")
    row_idx = seg_info.get("row_index")
    cell_idx = seg_info.get("cell_index")
    if is_table_cell and table_idx is not None and row_idx is not None and cell_idx is not None:
        para = _resolve_table_cell_paragraph(
            doc,
            table_idx,
            row_idx,
            cell_idx,
            target_text,
            source_text,
            para_index_map,
            seg_info.get("cell_local_idx"),
        )
        if para is not None:
            return para
        logger.info(
            LogModule.RESTOR,
            f"[BILINGUAL-DOCX] Could not match table paragraph for segment {segment_index}: "
            f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}, "
            f"target='{target_text[:60]}...', source='{source_text[:60]}...'",
        )
        return None

    para_index = seg_info.get("para_index")
    if para_index is not None:
        para = para_index_map.get((False, None, None, None, para_index))
        if para is not None:
            return para

    logger.debug(
        LogModule.RESTOR,
        f"[BILINGUAL-DOCX] Could not locate body paragraph for segment {segment_index}, "
        f"para_index={para_index}",
    )
    return None


def _parse_header_footer_key(raw: Any) -> Optional[tuple]:
    """Parse header_footer_key from translation segment metadata."""
    if raw is None:
        return None
    if isinstance(raw, tuple):
        return raw
    if isinstance(raw, list):
        return tuple(raw)
    if isinstance(raw, str):
        try:
            parsed = ast.literal_eval(raw)
            if isinstance(parsed, (list, tuple)):
                return tuple(parsed)
        except (ValueError, SyntaxError):
            pass
    return None


def _lookup_hf_source_text(
    source_map: Dict[tuple, str],
    name: str,
    lookup_idx: int,
    suffix: Tuple,
) -> Optional[str]:
    """Resolve source text for a header/footer flat key (mirrors apply lookup)."""
    from converter.x2md.docx_extras import _PART_NAME_ALIASES

    candidates = [name]
    alias = _PART_NAME_ALIASES.get(name)
    if alias and alias not in candidates:
        candidates.append(alias)
    for candidate in candidates:
        key = _normalize_key((candidate, lookup_idx, *suffix))
        if key in source_map:
            return source_map[key]
    return None


def _insert_bilingual_header_footer_flat_segments(
    doc,
    hf_segments: List[Dict[str, Any]],
    target_first: bool = False,
    source_text_italic: bool = True,
    source_text_color: str = "gray",
    target_text_italic: bool = False,
    target_text_color: Optional[str] = None,
    source_text_font_size_delta: float = 0.0,
    target_text_font_size_delta: float = 0.0,
) -> int:
    """Insert source paragraphs for flat header/footer translation segments."""
    from docx.text.paragraph import Paragraph
    from docx.shared import RGBColor
    from converter.x2md.docx_extras import (
        _cell_should_preserve_pagination,
        _compute_part_fingerprint,
        _is_merged_cell_skip,
        _paragraph_should_preserve_pagination,
        _part_element_id,
    )
    from utils.table_utils import get_all_merged_regions_docx

    source_map: Dict[tuple, str] = {}
    for seg in hf_segments:
        source = (seg.get("source_text") or "").strip()
        if not source:
            continue
        key = _parse_header_footer_key(seg.get("header_footer_key"))
        if key is None:
            continue
        modified_text = seg.get("modified_text")
        target_text = modified_text if modified_text is not None else seg.get("target_text", "")
        if (target_text or "").strip() == source:
            continue
        source_map[_normalize_key(key)] = source

    if not source_map:
        logger.info(
            LogModule.RESTOR,
            "[BILINGUAL-DOCX] No flat header/footer source texts to insert "
            f"(segments={len(hf_segments)})",
        )
        return 0

    _SOURCE_COLOR_MAP = {
        "gray": RGBColor(0x80, 0x80, 0x80),
        "red": RGBColor(0xFF, 0x00, 0x00),
        "blue": RGBColor(0x00, 0x00, 0xFF),
        "green": RGBColor(0x00, 0x80, 0x00),
        "orange": RGBColor(0xFF, 0xA5, 0x00),
        "black": RGBColor(0x00, 0x00, 0x00),
    }
    _resolved_color = _SOURCE_COLOR_MAP.get(source_text_color) if source_text_color else None
    _target_color = (
        _SOURCE_COLOR_MAP.get(target_text_color) if target_text_color else None
    )

    def _apply_target_style(para: Paragraph) -> None:
        if not para:
            return
        for run in para.runs:
            if target_text_italic:
                run.italic = True
            if _target_color:
                try:
                    run.font.color.rgb = _target_color
                except Exception:
                    pass
        if target_text_font_size_delta != 0.0:
            _apply_font_size_delta(para, target_text_font_size_delta)

    def _apply_source_run_style(dst_run) -> None:
        try:
            if source_text_italic:
                dst_run.italic = True
            if _resolved_color:
                dst_run.font.color.rgb = _resolved_color
        except Exception:
            pass

    def _copy_source_format(src_para: Paragraph, dst_para: Paragraph) -> None:
        # Copy run-level formatting
        paired = list(zip(src_para.runs, dst_para.runs))
        for src_run, dst_run in paired:
            try:
                if src_run.font.name:
                    dst_run.font.name = src_run.font.name
                if src_run.bold is not None:
                    dst_run.bold = src_run.bold
                if src_run.font.size:
                    dst_run.font.size = src_run.font.size
                if not source_text_italic and src_run.italic is not None:
                    dst_run.italic = src_run.italic
            except Exception:
                pass
            _apply_source_run_style(dst_run)
        if not dst_para.runs:
            _apply_source_run_style(dst_para.add_run(dst_para.text))
        # Copy paragraph-level formatting (alignment, indentation, spacing)
        # so the bilingual source paragraph matches the original layout.
        _copy_paragraph_format(src_para, dst_para)

    def _insert_adjacent(target_para: Paragraph, source_text: str) -> bool:
        try:
            parent = target_para._element.getparent()
            if parent is None:
                return False

            # Create the new paragraph directly as an XML element at the
            # correct location to avoid the cross-branch move that happens
            # when doc.add_paragraph() (body) is moved into a header/footer
            # or table cell via addprevious/addnext.
            from docx.oxml import OxmlElement
            new_p_elem = OxmlElement('w:p')
            new_r_elem = OxmlElement('w:r')
            new_t_elem = OxmlElement('w:t')
            new_t_elem.text = source_text
            new_r_elem.append(new_t_elem)
            new_p_elem.append(new_r_elem)

            if target_first:
                target_para._element.addnext(new_p_elem)
            else:
                target_para._element.addprevious(new_p_elem)

            new_para = Paragraph(new_p_elem, parent)
            _copy_source_format(target_para, new_para)
            _apply_target_style(target_para)
            if source_text_font_size_delta != 0.0:
                _apply_font_size_delta(new_para, source_text_font_size_delta)
            return True
        except Exception as e:
            logger.warning(
                LogModule.RESTOR,
                f"[BILINGUAL-DOCX] Failed to insert flat header/footer source paragraph: {e}",
            )
            return False

    inserted_count = 0
    content_first_idx: Dict[int, int] = {}
    processed_elements: set = set()

    for idx, section in enumerate(doc.sections):
        for name, part in (
            ("header", section.header),
            ("footer", section.footer),
            ("header_first", section.first_page_header),
            ("footer_first", section.first_page_footer),
        ):
            element_id = _part_element_id(part)
            if element_id in processed_elements:
                continue
            processed_elements.add(element_id)

            fp = _compute_part_fingerprint(part)
            if fp not in content_first_idx:
                content_first_idx[fp] = idx
            lookup_idx = content_first_idx[fp]

            for pi, para in enumerate(part.paragraphs):
                if _paragraph_should_preserve_pagination(para._p, para.text):
                    continue
                source_text = _lookup_hf_source_text(
                    source_map, name, lookup_idx, ("p", pi)
                )
                if source_text:
                    if _insert_adjacent(para, source_text):
                        inserted_count += 1

            for ti, tbl in enumerate(part.tables):
                merged_set: set = set()
                if get_all_merged_regions_docx is not None:
                    merged_set = set(get_all_merged_regions_docx(tbl))
                for ri, row in enumerate(tbl.rows):
                    for ci, cell in enumerate(row.cells):
                        if merged_set and _is_merged_cell_skip(ri, ci, merged_set):
                            continue
                        if _cell_should_preserve_pagination(cell):
                            continue
                        source_text = _lookup_hf_source_text(
                            source_map, name, lookup_idx, ("cell", ti, ri, ci)
                        )
                        if not source_text:
                            continue
                        target_para = None
                        for p in cell.paragraphs:
                            if p.text and p.text.strip():
                                target_para = p
                                break
                        if target_para is None and cell.paragraphs:
                            target_para = cell.paragraphs[0]
                        if target_para is None:
                            continue
                        if _insert_adjacent(target_para, source_text):
                            inserted_count += 1

    logger.info(
        LogModule.RESTOR,
        f"[BILINGUAL-DOCX] Flat header/footer inserted {inserted_count} source items "
        f"(target_first={target_first}, keys={len(source_map)})",
    )
    return inserted_count


def _insert_bilingual_source_paragraphs(
    doc,
    segment_data_map: Dict[int, Dict[str, Any]],
    para_index_map: Dict[Tuple, Any],
    target_first: bool = False,
    extras_original: Optional[Dict[str, Any]] = None,
    source_text_italic: bool = True,
    source_text_color: str = "gray",
    target_text_italic: bool = False,
    target_text_color: Optional[str] = None,
    skip_legacy_header_footer: bool = False,
    source_text_font_size_delta: float = 0.0,
    target_text_font_size_delta: float = 0.0,
) -> None:
    """Insert source-text paragraphs next to translated paragraphs for bilingual export.

    Strategy for body / table cells:
      - Skip excluded or failed segments (emit target only).
      - Group remaining segments by para_index.
      - For each paragraph that has at least one translated segment, create a
        new paragraph containing the concatenated source_text of its segments.
      - Insert BEFORE the translated paragraph when target_first=False
        (source first, target after), otherwise AFTER it.
      - Process paragraphs in *descending* para_index order so insertions do
        not shift the positions of paragraphs yet to be processed.

    Strategy for headers/footers / textboxes:
      - Use extras_original saved during translation to obtain source text.
      - Insert a styled source paragraph after (or before) the translated text.
    """
    from docx.text.paragraph import Paragraph
    from docx.shared import RGBColor

    _SOURCE_COLOR_MAP = {
        "gray": RGBColor(0x80, 0x80, 0x80),
        "red": RGBColor(0xFF, 0x00, 0x00),
        "blue": RGBColor(0x00, 0x00, 0xFF),
        "green": RGBColor(0x00, 0x80, 0x00),
        "orange": RGBColor(0xFF, 0xA5, 0x00),
        "black": RGBColor(0x00, 0x00, 0x00),
    }
    _resolved_color = _SOURCE_COLOR_MAP.get(source_text_color) if source_text_color else None

    def _apply_target_style(para: Paragraph) -> None:
        """Apply target-text style (italic/color) to all runs in a paragraph."""
        if not para:
            return
        _target_color = None
        if target_text_color:
            _target_color = _SOURCE_COLOR_MAP.get(target_text_color)
        for run in para.runs:
            if target_text_italic:
                run.italic = True
            if _target_color:
                try:
                    run.font.color.rgb = _target_color
                except Exception:
                    pass

    def _apply_target_style_with_delta(para: Paragraph) -> None:
        """Apply target style plus font delta (HF/textbox paths not in main rebuild loop)."""
        _apply_target_style(para)
        if target_text_font_size_delta != 0.0:
            _apply_font_size_delta(para, target_text_font_size_delta)

    def _finalize_textbox_source_paragraph(
        new_para: Paragraph, target_para: Paragraph
    ) -> None:
        """Style a textbox source paragraph without copying translated run fonts."""
        _copy_paragraph_format(target_para, new_para)
        for run in new_para.runs:
            _apply_source_run_style(run)
        _apply_target_style_with_delta(target_para)
        if source_text_font_size_delta != 0.0:
            _apply_font_size_delta(new_para, source_text_font_size_delta)

    def _apply_source_run_style(dst_run) -> None:
        """Apply bilingual source italic/color to a single run."""
        try:
            if source_text_italic:
                dst_run.italic = True
            if _resolved_color:
                dst_run.font.color.rgb = _resolved_color
        except Exception:
            pass

    def _copy_source_format(src_para: Paragraph, dst_para: Paragraph) -> None:
        """Copy font/bold/italic/size from src_para runs to dst_para runs.

        Applies source_text_italic and source_text_color when configured.
        """
        paired = list(zip(src_para.runs, dst_para.runs))
        for src_run, dst_run in paired:
            try:
                if src_run.font.name:
                    dst_run.font.name = src_run.font.name
                if src_run.bold is not None:
                    dst_run.bold = src_run.bold
                if src_run.font.size:
                    dst_run.font.size = src_run.font.size
                if source_text_italic:
                    dst_run.italic = True
                elif src_run.italic is not None:
                    dst_run.italic = src_run.italic
                if _resolved_color:
                    dst_run.font.color.rgb = _resolved_color
                elif src_run.font.color and src_run.font.color.rgb:
                    dst_run.font.color.rgb = src_run.font.color.rgb
            except Exception:
                pass

        # Single-run paragraphs from add_paragraph() may not pair with multi-run originals
        for dst_run in dst_para.runs[len(paired):]:
            _apply_source_run_style(dst_run)
        if not paired and dst_para.runs:
            for dst_run in dst_para.runs:
                _apply_source_run_style(dst_run)
        # Copy paragraph-level formatting (alignment, indentation, spacing)
        # so the bilingual source paragraph matches the original layout.
        _copy_paragraph_format(src_para, dst_para)

    # ------------------------------------------------------------------
    # 1. Body paragraphs + table cells (from translation_segments)
    # ------------------------------------------------------------------
    element_id_to_entry: Dict[int, Tuple[Any, List[str], int]] = {}

    for segment_index in sorted(segment_data_map.keys()):
        segment_data = segment_data_map[segment_index]
        is_excluded = segment_data.get("is_excluded", False)
        is_failed = segment_data.get("is_failed", False)
        if is_excluded or is_failed:
            continue

        source_text = (segment_data.get("source_text") or "").strip()
        if not source_text:
            continue

        element = _resolve_bilingual_target_paragraph(
            doc, segment_data, para_index_map, segment_index
        )
        if element is None:
            continue

        elem_id = id(element._element)
        if elem_id not in element_id_to_entry:
            element_id_to_entry[elem_id] = (element, [], segment_index)
        element_id_to_entry[elem_id][1].append(source_text)

    inserted_count = 0

    if element_id_to_entry:
        sorted_entries = sorted(
            element_id_to_entry.values(),
            key=lambda entry: entry[2],
            reverse=True,
        )

        for element, sources, _segment_index in sorted_entries:
            if not sources:
                continue

            combined_source = "\n".join(s.strip() for s in sources if s.strip())
            if not combined_source:
                continue

            try:
                parent = element._element.getparent()
                if parent is None:
                    continue

                # CRITICAL: Create the new paragraph directly as an XML element
                # instead of using doc.add_paragraph().  doc.add_paragraph()
                # appends to the document body and then addprevious/addnext
                # moves the element across XML branches (body → table cell).
                # That cross-branch move corrupts python-docx internal caches
                # and can cause the saved document to place paragraphs in the
                # wrong cells when the target is inside a table.
                from docx.oxml import OxmlElement
                new_p_elem = OxmlElement('w:p')
                new_r_elem = OxmlElement('w:r')
                new_t_elem = OxmlElement('w:t')
                new_t_elem.text = combined_source
                new_r_elem.append(new_t_elem)
                new_p_elem.append(new_r_elem)

                # Insert the new paragraph element next to the target paragraph
                # BEFORE wrapping it in a Paragraph object so format-copy
                # operates on the paragraph in its final location.
                if target_first:
                    element._element.addnext(new_p_elem)
                else:
                    element._element.addprevious(new_p_elem)

                # Wrap the new element in a Paragraph object for format copying.
                new_para = Paragraph(new_p_elem, parent)
                _copy_source_format(element, new_para)
                _apply_target_style(element)
                if source_text_font_size_delta != 0.0:
                    _apply_font_size_delta(new_para, source_text_font_size_delta)
                inserted_count += 1
            except Exception as e:
                logger.warning(
                    LogModule.RESTOR,
                    f"[BILINGUAL-DOCX] Failed to insert source paragraph for segment {_segment_index}: {e}"
                )

    # ------------------------------------------------------------------
    # 2. Headers / footers (legacy joined-text format; skip when flat segments exist)
    # ------------------------------------------------------------------
    if extras_original and not skip_legacy_header_footer:
        hf_items = extras_original.get("headers_footers", [])
        if hf_items:
            for idx, section in enumerate(doc.sections):
                for name, part in (("header", section.header), ("footer", section.footer)):
                    key = (name, idx)
                    # Find the matching original text
                    original_text = None
                    for item in hf_items:
                        if not isinstance(item, (list, tuple)) or len(item) != 2:
                            continue
                        k, text = item
                        if _normalize_key(k) == key and text and text.strip():
                            original_text = text
                            break
                    if not original_text:
                        continue

                    try:
                        if part.paragraphs:
                            target_para = part.paragraphs[0]
                            new_para = part.add_paragraph(original_text)
                            _copy_source_format(target_para, new_para)

                            if target_first:
                                target_para._element.addnext(new_para._element)
                            else:
                                target_para._element.addprevious(new_para._element)

                            _apply_target_style_with_delta(target_para)
                            inserted_count += 1
                    except Exception as e:
                        logger.warning(
                            LogModule.RESTOR,
                            f"[BILINGUAL-DOCX] Failed to insert bilingual header/footer for {key}: {e}"
                        )

    # ------------------------------------------------------------------
    # 3. Textboxes / SDTs
    # ------------------------------------------------------------------
    if extras_original and extras_original.get("textboxes_sdts"):
        raw_tb_items = extras_original.get("textboxes_sdts", [])
        tb_items = _parse_textbox_sdt_storage_items(raw_tb_items)
        logger.info(
            LogModule.RESTOR,
            f"[BILINGUAL-DOCX] Textbox/SDT section: extras_original has textboxes_sdts={bool(raw_tb_items)}, "
            f"count={len(raw_tb_items)}, parsed={len(tb_items)} items"
        )
        if tb_items:
            # Log what keys are available
            keys_found = [_normalize_key(item[0]) if isinstance(item, (list, tuple)) and len(item) >= 2 else f"BAD:{type(item).__name__}:{item!r}" for item in tb_items]
            logger.info(
                LogModule.RESTOR,
                f"[BILINGUAL-DOCX] Textbox/SDT keys in tb_items: {keys_found}"
            )
            # Locate textbox containers in document XML.
            # The iteration order and filtering must match _extract_textboxes
            # in docx_extras.py exactly, otherwise the ("textbox", N) keys
            # will be misaligned and source texts will not be found.
            try:
                from lxml import etree
                from docx.oxml import OxmlElement

                choice_to_fallback = _build_choice_to_fallback_txbx_map(doc)
                all_containers = []

                # Collect txbxContent nodes (skip those in mc:Fallback when
                # mc:Choice already provides the same textbox).
                txbx_nodes = doc._element.xpath('.//*[local-name()="txbxContent"]')
                for txbx in txbx_nodes:
                    _alt_content = txbx.xpath(
                        './ancestor::*[local-name()="AlternateContent"]'
                    )
                    if _alt_content:
                        _in_fallback = txbx.xpath(
                            './ancestor::*[local-name()="Fallback"]'
                        )
                        if _in_fallback:
                            _choice_txbx = _alt_content[0].xpath(
                                './/*[local-name()="Choice"]//*[local-name()="txbxContent"]'
                            )
                            if _choice_txbx:
                                continue  # skip duplicate
                    all_containers.append(txbx)

                # Add legacy v:textbox nodes
                pict_nodes = doc._element.xpath('.//*[local-name()="pict"]//*[local-name()="textbox"]')
                all_containers.extend(pict_nodes)

                # Add drawing elements with text content
                drawing_nodes = doc._element.xpath('.//*[local-name()="drawing"]')
                for drawing in drawing_nodes:
                    text_elements = drawing.xpath('.//*[local-name()="t"]')
                    if text_elements:
                        all_containers.append(drawing)

                logger.info(
                    LogModule.RESTOR,
                    f"[BILINGUAL-DOCX] Textbox containers: {len(txbx_nodes)} txbxContent, "
                    f"{len(pict_nodes)} pict/textbox, {len(drawing_nodes)} drawing "
                    f"-> {len(all_containers)} total"
                )

                _tb_items_by_key: Dict[Tuple, str] = {}
                for item in tb_items:
                    if isinstance(item, (list, tuple)) and len(item) >= 2:
                        _tb_items_by_key[_normalize_key(item[0])] = item[1]
                for segment_data in segment_data_map.values():
                    if segment_data.get("segment_type") != "textbox_sdt":
                        continue
                    source_text = (segment_data.get("source_text") or "").strip()
                    if not source_text:
                        continue
                    key = _parse_textbox_key(segment_data.get("textbox_key"))
                    if key is not None:
                        _tb_items_by_key.setdefault(key, source_text)

                container_idx = 0
                for container in all_containers:
                    key = ("textbox", container_idx)

                    if container.tag.endswith('txbxContent') or container.tag.endswith('textbox'):
                        # Skip textboxes inside SDTs (already handled by SDT
                        # extraction) — must match extraction behaviour.
                        if container.xpath('./ancestor::*[local-name()="sdt"]'):
                            container_idx += 1
                            continue
                        if _is_duplicate_legacy_vtextbox_container(container):
                            logger.debug(
                                LogModule.RESTOR,
                                f"[BILINGUAL-DOCX] Skipping duplicate legacy v:textbox container for key {key}",
                            )
                            container_idx += 1
                            continue
                        original_text = _tb_items_by_key.get(key)
                        container_idx += 1

                        if not original_text:
                            continue

                        target_p_elem = _find_textbox_target_paragraph_element(
                            container, original_text
                        )
                        if target_p_elem is None:
                            continue

                        try:
                            parent = target_p_elem.getparent()
                            target_para = Paragraph(target_p_elem, parent)

                            # Create source paragraph directly in the textbox
                            new_p_elem = OxmlElement('w:p')
                            new_r_elem = OxmlElement('w:r')
                            new_t_elem = OxmlElement('w:t')
                            new_t_elem.text = original_text
                            new_r_elem.append(new_t_elem)
                            new_p_elem.append(new_r_elem)

                            if target_first:
                                target_para._element.addnext(new_p_elem)
                            else:
                                target_para._element.addprevious(new_p_elem)

                            new_para = Paragraph(new_p_elem, parent)
                            _finalize_textbox_source_paragraph(new_para, target_para)
                            if container.tag.endswith('txbxContent'):
                                _mirror_textbox_source_to_fallback(
                                    container,
                                    new_p_elem,
                                    target_first,
                                    choice_to_fallback,
                                )
                            inserted_count += 1
                        except Exception as e:
                            logger.warning(
                                LogModule.RESTOR,
                                f"[BILINGUAL-DOCX] Failed to insert bilingual textbox for {key}: {e}"
                            )
                    elif container.tag.endswith('drawing'):
                        # Drawing elements with text — use same index alignment
                        original_text = _tb_items_by_key.get(key)
                        container_idx += 1

                        if not original_text:
                            continue

                        target_p_elem = _find_textbox_target_paragraph_element(
                            container, original_text
                        )
                        if target_p_elem is None:
                            continue

                        try:
                            parent = target_p_elem.getparent()
                            target_para = Paragraph(target_p_elem, parent)

                            new_p_elem = OxmlElement('w:p')
                            new_r_elem = OxmlElement('w:r')
                            new_t_elem = OxmlElement('w:t')
                            new_t_elem.text = original_text
                            new_r_elem.append(new_t_elem)
                            new_p_elem.append(new_r_elem)

                            if target_first:
                                target_para._element.addnext(new_p_elem)
                            else:
                                target_para._element.addprevious(new_p_elem)

                            new_para = Paragraph(new_p_elem, parent)
                            _finalize_textbox_source_paragraph(new_para, target_para)
                            inserted_count += 1
                        except Exception as e:
                            logger.warning(
                                LogModule.RESTOR,
                                f"[BILINGUAL-DOCX] Failed to insert bilingual drawing text for {key}: {e}"
                            )
                    else:
                        container_idx += 1

                # Handle SDTs
                sdt_elems = doc._element.body.xpath('.//*[local-name()="sdt"]')
                logger.info(
                    LogModule.RESTOR,
                    f"[BILINGUAL-DOCX] SDT section: found {len(sdt_elems)} body SDT elements"
                )
                sdt_index = 0
                for sdt in sdt_elems:
                    parent_sdt = sdt.xpath('./ancestor::*[local-name()="sdt"]')
                    if parent_sdt:
                        continue

                    # Try sdt_content keys
                    for sub_idx in range(10):  # reasonable upper bound
                        key = ("sdt_content", sdt_index, sub_idx)
                        original_text = _tb_items_by_key.get(key)
                        if not original_text:
                            continue

                        sdt_contents = sdt.xpath('.//*[local-name()="sdtContent"]')
                        if sub_idx < len(sdt_contents):
                            para_elems = sdt_contents[sub_idx].xpath('.//*[local-name()="p"]')
                            if para_elems:
                                try:
                                    _parent = para_elems[0].getparent()
                                    target_para = Paragraph(para_elems[0], _parent)

                                    new_p_elem = OxmlElement('w:p')
                                    new_r_elem = OxmlElement('w:r')
                                    new_t_elem = OxmlElement('w:t')
                                    new_t_elem.text = original_text
                                    new_r_elem.append(new_t_elem)
                                    new_p_elem.append(new_r_elem)

                                    if target_first:
                                        target_para._element.addnext(new_p_elem)
                                    else:
                                        target_para._element.addprevious(new_p_elem)

                                    new_para = Paragraph(new_p_elem, _parent)
                                    _finalize_textbox_source_paragraph(new_para, target_para)
                                    inserted_count += 1
                                except Exception as e:
                                    logger.warning(
                                        LogModule.RESTOR,
                                        f"[BILINGUAL-DOCX] Failed to insert bilingual SDT for {key}: {e}"
                                    )

                    # Try direct sdt key
                    key = ("sdt", sdt_index)
                    original_text = _tb_items_by_key.get(key)
                    if original_text:
                        para_elems = sdt.xpath('.//*[local-name()="p"]')
                        if para_elems:
                            try:
                                _parent = para_elems[0].getparent()
                                target_para = Paragraph(para_elems[0], _parent)

                                new_p_elem = OxmlElement('w:p')
                                new_r_elem = OxmlElement('w:r')
                                new_t_elem = OxmlElement('w:t')
                                new_t_elem.text = original_text
                                new_r_elem.append(new_t_elem)
                                new_p_elem.append(new_r_elem)

                                if target_first:
                                    target_para._element.addnext(new_p_elem)
                                else:
                                    target_para._element.addprevious(new_p_elem)

                                new_para = Paragraph(new_p_elem, _parent)
                                _finalize_textbox_source_paragraph(new_para, target_para)
                                inserted_count += 1
                            except Exception as e:
                                logger.warning(
                                    LogModule.RESTOR,
                                    f"[BILINGUAL-DOCX] Failed to insert bilingual SDT for {key}: {e}"
                                )

                    # Try sdt_child keys — iterate child SDTs inside this parent
                    child_sdts = sdt.xpath('.//*[local-name()="sdt"]')
                    for child_idx, child_sdt in enumerate(child_sdts):
                        key = ("sdt_child", sdt_index, child_idx)
                        original_text = _tb_items_by_key.get(key)
                        if not original_text:
                            continue
                        para_elems = child_sdt.xpath('.//*[local-name()="p"]')
                        if para_elems:
                            try:
                                _parent = para_elems[0].getparent()
                                target_para = Paragraph(para_elems[0], _parent)

                                new_p_elem = OxmlElement('w:p')
                                new_r_elem = OxmlElement('w:r')
                                new_t_elem = OxmlElement('w:t')
                                new_t_elem.text = original_text
                                new_r_elem.append(new_t_elem)
                                new_p_elem.append(new_r_elem)

                                if target_first:
                                    target_para._element.addnext(new_p_elem)
                                else:
                                    target_para._element.addprevious(new_p_elem)

                                new_para = Paragraph(new_p_elem, _parent)
                                _finalize_textbox_source_paragraph(new_para, target_para)
                                inserted_count += 1
                            except Exception as e:
                                logger.warning(
                                    LogModule.RESTOR,
                                    f"[BILINGUAL-DOCX] Failed to insert bilingual SDT child for {key}: {e}"
                                )

                    sdt_index += 1
            except Exception as e:
                logger.warning(
                    LogModule.RESTOR,
                    f"[BILINGUAL-DOCX] Failed to process textbox/SDT bilingual insertion: {e}"
                )

    logger.info(
        LogModule.RESTOR,
        f"[BILINGUAL-DOCX] Inserted {inserted_count} source paragraphs (target_first={target_first})"
    )
