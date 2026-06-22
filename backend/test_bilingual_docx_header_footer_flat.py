# SPDX-License-Identifier: AGPL-3.0-or-later
"""Regression tests for bilingual DOCX export with flat header/footer segments."""

from __future__ import annotations

import importlib.util
import sys
from io import BytesIO
from pathlib import Path

import pytest

pytest.importorskip("docx")

from converter.x2md.docx_extras import (  # noqa: E402
    apply_headers_footers_flat,
    extract_headers_footers_flat,
)

_DEBUG_DOCX = Path(
    r"C:\Users\Zampher\AppData\Local\Temp\owlangs_c8aaab84_h973i1c6\AMVR-2104-1.docx"
)
_FALLBACK_DOCX = Path(
    r"C:\Users\Zampher\AppData\Local\Temp\owlangs_a787e0e6_yty9oeyp\AMVR-2104-1.docx"
)


def _resolve_debug_docx() -> Path | None:
    for candidate in (_DEBUG_DOCX, _FALLBACK_DOCX):
        if candidate.is_file():
            return candidate
    return None


def _load_rebuild_module():
    """Load docx_rebuild without pulling optional test-env deps via package __init__."""
    backend = Path(__file__).resolve().parent
    root = backend.parent
    for p in (str(root), str(backend)):
        if p not in sys.path:
            sys.path.insert(0, p)
    spec = importlib.util.spec_from_file_location(
        "docx_rebuild_test",
        backend / "utils" / "document_rebuild" / "docx_rebuild.py",
    )
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


def _build_hf_segments(items, translations):
    segments = []
    for idx, (key, source) in enumerate(items):
        target = translations.get(key, source)
        segments.append(
            {
                "segment_index": idx,
                "source_text": source,
                "target_text": target,
                "modified": False,
                "is_excluded": False,
                "is_failed": False,
                "segment_type": "header_footer",
                "header_footer_key": str(key),
            }
        )
    return segments


@pytest.mark.skipif(_resolve_debug_docx() is None, reason="local debug docx not available")
def test_bilingual_hf_only_export_includes_source_text() -> None:
    """HF-only documents must insert source text alongside translations."""
    pytest.importorskip("httpx")
    from ir.document import Document

    docx_rebuild = _load_rebuild_module()
    docx_path = _resolve_debug_docx()
    assert docx_path is not None
    orig = docx_path.read_bytes()
    items = extract_headers_footers_flat(orig)
    assert items, "expected flat header/footer items"

    translations = {key: f"ZH::{src[:40]}" for key, src in items}
    translated_bytes = apply_headers_footers_flat(orig, translations)

    task_state = {
        "translation_segments": {
            "segments": _build_hf_segments(items, translations),
        }
    }
    translated_doc = Document.from_bytes(
        content=translated_bytes,
        suffix=".docx",
        stem="translated",
    )

    rebuilt = docx_rebuild.rebuild_docx_document_from_segments(
        task_state,
        translated_doc,
        bilingual_export=True,
        target_first=False,
    )
    assert rebuilt is not None

    from docx import Document as PyDocxDocument

    doc = PyDocxDocument(BytesIO(rebuilt.content))
    header_first = doc.sections[0].first_page_header
    label_cell = header_first.tables[0].rows[1].cells[2]
    cell_text = "\n".join(p.text for p in label_cell.paragraphs if p.text.strip())
    assert len(label_cell.paragraphs) >= 2, "bilingual export should add a source paragraph"
    assert "ZH::" in cell_text, "target translation should remain in header cell"
    assert any(source.strip() in cell_text for _key, source in items), (
        "at least one extracted source string should appear in exported header/footer"
    )

    page_cell = header_first.tables[0].rows[1].cells[3].text
    assert page_cell == "1/16", "page-number cell must stay dynamic and untranslated"
