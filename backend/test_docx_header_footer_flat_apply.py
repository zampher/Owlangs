# SPDX-License-Identifier: AGPL-3.0-or-later
"""Regression tests for flat header/footer apply (linked parts across sections)."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest

pytest.importorskip("docx")

from converter.x2md.docx_extras import (  # noqa: E402
    apply_headers_footers_flat,
    extract_headers_footers_flat,
)

_DEBUG_DOCX = Path(
    r"C:\Users\Zampher\AppData\Local\Temp\owlangs_a787e0e6_yty9oeyp\AMVR-2104-1.docx"
)


@pytest.mark.skipif(not _DEBUG_DOCX.is_file(), reason="local debug docx not available")
def test_extract_flat_skips_page_number_cells() -> None:
    """Page-number table cells must not be sent for translation."""
    orig = _DEBUG_DOCX.read_bytes()
    items = extract_headers_footers_flat(orig)
    texts = {text for _key, text in items}
    assert "1/16" not in texts
    assert "16/16" not in texts
    assert "4/16" not in texts


@pytest.mark.skipif(not _DEBUG_DOCX.is_file(), reason="local debug docx not available")
def test_apply_flat_does_not_clear_linked_header_across_sections() -> None:
    """Linked header XML shared by all sections must not be cleared on 2nd pass."""
    orig = _DEBUG_DOCX.read_bytes()
    items = extract_headers_footers_flat(orig)
    translations = {key: f"TRANSLATED::{src}" for key, src in items}

    result = apply_headers_footers_flat(orig, translations)

    from docx import Document

    doc = Document(BytesIO(result))
    header = doc.sections[0].header
    assert header.paragraphs[0].text.startswith("TRANSLATED::")
    page_cell = header.tables[0].rows[1].cells[5].text
    assert page_cell == "16/16"

    header_first = doc.sections[0].first_page_header
    first_page_cell = header_first.tables[0].rows[1].cells[3].text
    assert first_page_cell == "1/16"

    label_cell = header_first.tables[0].rows[1].cells[2].text
    assert label_cell.startswith("TRANSLATED::")

    footer_first = doc.sections[0].first_page_footer
    assert footer_first.paragraphs[1].text.startswith("TRANSLATED::")
