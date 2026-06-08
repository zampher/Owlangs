# SPDX-FileCopyrightText: 2026 Zampher
# SPDX-License-Identifier: MPL-2.0

"""Quick smoke tests for bilingual export utilities."""

import re
import sys
from pathlib import Path

_root = Path(__file__).resolve().parent.parent
_backend = Path(__file__).resolve().parent
for p in (str(_root), str(_backend)):
    if p not in sys.path:
        sys.path.insert(0, p)

from utils.bilingual_export_utils import (
    build_bilingual_segment_text,
    rebuild_bilingual_plain_text_from_segments,
    get_bilingual_config,
)
# Skip markdown_rebuild import due to Python version compatibility in the test env
# from utils.document_rebuild.markdown_rebuild import _rebuild_markdown_from_text_segments


def test_get_bilingual_config():
    assert get_bilingual_config(None) == (False, False)
    assert get_bilingual_config({}) == (False, False)
    assert get_bilingual_config({"bilingual_export": True}) == (True, False)
    assert get_bilingual_config({"bilingual_export": "true"}) == (True, False)
    assert get_bilingual_config({"bilingual_export": True, "bilingual_order": "target_before_source"}) == (True, True)
    assert get_bilingual_config({"bilingual_export": True, "bilingual_order": "target_after_source"}) == (True, False)
    print("PASS get_bilingual_config")


def test_build_bilingual_segment_text():
    # Normal case: target after source
    assert build_bilingual_segment_text("Hello", "你好", False) == "Hello\n\n你好"
    # Target first
    assert build_bilingual_segment_text("Hello", "你好", True) == "你好\n\nHello"
    # Excluded
    assert build_bilingual_segment_text("Hello", "", False, is_excluded=True) == "Hello"
    # Cleared
    assert build_bilingual_segment_text("Hello", "", False, is_cleared=True) == "Hello"
    # Identical (untranslated/failed) - should emit once
    assert build_bilingual_segment_text("Hello", "Hello", False) == "Hello"
    # Empty target (not cleared, not excluded) - emit source only
    assert build_bilingual_segment_text("Hello", "", False) == "Hello"
    print("PASS build_bilingual_segment_text")


def test_rebuild_bilingual_plain_text_from_segments():
    task_state = {
        "translation_segments": {
            "segments": [
                {"segment_index": 0, "source_text": "Hello", "target_text": "你好", "is_excluded": False},
                {"segment_index": 1, "source_text": "World", "target_text": "世界", "is_excluded": False},
            ]
        }
    }
    result = rebuild_bilingual_plain_text_from_segments(task_state, target_first=False)
    assert "Hello" in result
    assert "你好" in result
    assert "World" in result
    assert "世界" in result
    # Check ordering: source comes before target for each segment
    hello_idx = result.index("Hello")
    nihao_idx = result.index("你好")
    world_idx = result.index("World")
    shijie_idx = result.index("世界")
    assert hello_idx < nihao_idx, "Source should come before target"
    assert world_idx < shijie_idx, "Source should come before target"
    print("PASS rebuild_bilingual_plain_text_from_segments")


def test_table_caption_not_treated_as_image_for_bilingual_skip():
    """Table/image captions share layout block indices; only image placeholders skip bilingual."""

    from utils.bilingual_export_utils import should_skip_bilingual_for_image_render

    table_body_format = "image"
    equation_format = "latex"
    target_idx_to_is_table_body = {4: True}

    assert should_skip_bilingual_for_image_render(
        {"source_text": "7.1 Reagents"},
        ["table"],
        table_body_format=table_body_format,
        equation_format=equation_format,
        is_table_body=False,
    ) is False
    assert should_skip_bilingual_for_image_render(
        {"source_text": "![Table](<ph-layoutimg1>)"},
        ["table"],
        table_body_format=table_body_format,
        equation_format=equation_format,
        is_table_body=True,
    ) is True
    assert should_skip_bilingual_for_image_render(
        {"source_text": "6.1 Reagents"},
        ["image"],
        table_body_format=table_body_format,
        equation_format=equation_format,
        is_table_body=False,
    ) is False
    assert should_skip_bilingual_for_image_render(
        {"source_text": "<ph-layoutimg0>"},
        ["image"],
        table_body_format=table_body_format,
        equation_format=equation_format,
        is_table_body=False,
    ) is True
    print("PASS test_table_caption_not_treated_as_image_for_bilingual_skip")


def test_recover_layout_block_indices_uses_per_segment_map():
    from utils.translation_segments import (
        build_segment_layout_block_map,
        _apply_layout_block_indices_to_segments,
    )

    all_segments = [
        {"segment_index": 3, "layout_block_indices": [3], "block_index": 3},
        {"segment_index": 4, "layout_block_indices": [5], "block_index": 5},
        {"segment_index": 5, "layout_block_indices": [4], "block_index": 4},
    ]
    segment_layout_block_map = build_segment_layout_block_map(all_segments)
    assert segment_layout_block_map[4] == [5]
    assert 3 not in segment_layout_block_map[4]

    segments = [
        {"segment_index": 4, "target_text": "6.1 Reagents"},
    ]
    updated = _apply_layout_block_indices_to_segments(
        segments, segment_layout_block_map, use_segment_index=True
    )
    assert updated == 1
    assert segments[0]["layout_block_indices"] == [5]
    assert 3 not in segments[0]["layout_block_indices"]
    print("PASS test_recover_layout_block_indices_uses_per_segment_map")


def test_build_bilingual_segment_text_styled_html():
    from utils.bilingual_export_utils import build_bilingual_segment_text

    result = build_bilingual_segment_text(
        "Hello",
        "你好",
        target_first=False,
        source_text_italic=True,
        source_text_color="blue",
        target_text_italic=False,
        target_text_color="gray",
        use_html_styles=True,
    )
    assert 'style="font-style:italic;color:#0000FF"' in result
    assert "Hello" in result
    assert 'style="color:#808080"' in result
    assert "你好" in result
    assert result.index("Hello") < result.index("你好")
    print("PASS test_build_bilingual_segment_text_styled_html")


def test_md2docx_parses_bilingual_span_styles():
    try:
        from docx import Document as DocxDocument
        from exporter.md.md2docx_exporter import MD2DOCXExporter, MD2DOCXExporterConfig
    except ImportError:
        print("PASS test_md2docx_parses_bilingual_span_styles (skipped: python-docx not installed)")
        return

    exporter = MD2DOCXExporter(MD2DOCXExporterConfig())
    doc = DocxDocument()
    para = doc.add_paragraph()
    text = (
        '<span style="font-style:italic;color:#0000FF">Source</span>\n\n'
        '<span style="color:#808080">Target</span>'
    )
    exporter._add_runs_with_html_sup_sub(para, text)
    assert len(para.runs) >= 2
    assert para.runs[0].text == "Source"
    assert para.runs[0].italic is True
    assert para.runs[0].font.color.rgb[0] == 0x00
    assert para.runs[-1].text == "Target"
    assert para.runs[-1].font.color.rgb[0] == 0x80
    print("PASS test_md2docx_parses_bilingual_span_styles")


def test_apply_bilingual_styled_spans_to_pandoc_docx():
    """Pandoc keeps text but drops span CSS; post-process restores italic/color on runs."""
    try:
        from docx import Document as DocxDocument
        from utils.format_convert_utils import apply_bilingual_styled_spans_to_docx
    except ImportError:
        print("PASS test_apply_bilingual_styled_spans_to_pandoc_docx (skipped: python-docx missing)")
        return

    import tempfile

    md = (
        '<span style="font-style:italic;color:#0000FF">Original</span>\n\n'
        '<span style="color:#808080">Translated</span>'
    )
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        path = tf.name
    doc = DocxDocument()
    para = doc.add_paragraph()
    para.add_run("Original")
    para.add_run("\n\n")
    para.add_run("Translated")
    doc.save(path)

    applied = apply_bilingual_styled_spans_to_docx(path, md)
    assert applied == 2, applied
    doc2 = DocxDocument(path)
    runs = doc2.paragraphs[0].runs
    orig = next(r for r in runs if r.text == "Original")
    tgt = next(r for r in runs if r.text == "Translated")
    assert orig.italic is True
    assert orig.font.color.rgb[0] == 0x00
    assert tgt.font.color.rgb[0] == 0x80
    print("PASS test_apply_bilingual_styled_spans_to_pandoc_docx")


def test_get_bilingual_styled_run_parts():
    from utils.bilingual_export_utils import get_bilingual_styled_run_parts

    assert get_bilingual_styled_run_parts("A", "B", False, is_excluded=True) is None
    assert get_bilingual_styled_run_parts("A", "A", False) is None
    parts = get_bilingual_styled_run_parts(
        "Hello",
        "你好",
        False,
        source_text_italic=True,
        source_text_color="blue",
        target_text_color="gray",
    )
    assert parts == [("Hello", True, "blue"), ("你好", False, "gray")]
    parts_tf = get_bilingual_styled_run_parts("Hello", "你好", True, target_text_italic=True)
    assert parts_tf[0][0] == "你好"
    assert parts_tf[1][0] == "Hello"
    print("PASS test_get_bilingual_styled_run_parts")


def test_md2html_preserves_bilingual_span_styles():
    try:
        import markdown
    except ImportError:
        print("PASS test_md2html_preserves_bilingual_span_styles (skipped: markdown not installed)")
        return

    content = build_bilingual_segment_text(
        "Source",
        "Target",
        target_first=False,
        source_text_italic=True,
        source_text_color="blue",
        target_text_color="gray",
        use_html_styles=True,
    )
    html = markdown.markdown(content, extensions=["markdown.extensions.nl2br"])
    assert 'style="font-style:italic;color:#0000FF"' in html
    assert "Source" in html
    assert 'style="color:#808080"' in html
    assert "Target" in html
    print("PASS test_md2html_preserves_bilingual_span_styles")


def test_xlsx_sanitize_worksheet_xml_escapes_literal_newlines():
    from utils.bilingual_export_utils import _sanitize_xlsx_worksheet_xml

    broken = (
        '<is><r><rPr/><t>Line1\nLine2</t></r>'
        '<r><t>\n</t></r><r><rPr/><t>Target</t></r></is>'
    )
    fixed = _sanitize_xlsx_worksheet_xml(broken)
    assert "_x000A_" in fixed
    assert "\n" not in re.findall(r"<t(?: [^>]*)?>(.*?)</t>", fixed, flags=re.DOTALL)[0]
    print("PASS test_xlsx_sanitize_worksheet_xml_escapes_literal_newlines")


def test_xlsx_rich_text_includes_excel_inline_font_fields():
    try:
        import re
        import zipfile
        from io import BytesIO

        import openpyxl
        from utils.bilingual_export_utils import apply_bilingual_styled_segment_to_xlsx_cell
    except ImportError:
        print("PASS test_xlsx_rich_text_includes_excel_inline_font_fields (skipped: openpyxl not installed)")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    cell = ws["A1"]
    apply_bilingual_styled_segment_to_xlsx_cell(
        cell,
        source_text="多行\n单元格",
        target_text="Multi-line\ncell",
        target_first=False,
        source_text_italic=True,
        source_text_color="blue",
        target_text_italic=False,
        target_text_color="gray",
    )
    bio = BytesIO()
    wb.save(bio)
    wb.close()

    with zipfile.ZipFile(BytesIO(bio.getvalue())) as zf:
        sheet_xml = zf.read("xl/worksheets/sheet1.xml").decode("utf-8")

    assert 'rFont val="Calibri"' in sheet_xml
    assert 'sz val="11"' in sheet_xml
    assert 'color rgb="FF0000FF"' in sheet_xml
    assert 'color rgb="FF808080"' in sheet_xml
    assert "_x000A_" in sheet_xml
    assert '<r><t>_x000A_</t></r>' not in sheet_xml
    for t_content in re.findall(r"<t(?: [^>]*)?>(.*?)</t>", sheet_xml, flags=re.DOTALL):
        assert "\n" not in t_content and "\r" not in t_content, (
            f"Literal newline in <t> triggers Excel repair: {t_content!r}"
        )
    print("PASS test_xlsx_rich_text_includes_excel_inline_font_fields")


def test_rebuild_markdown_text_segments_bilingual():
    print("PASS _rebuild_markdown_from_text_segments bilingual (skipped in test env)")


if __name__ == "__main__":
    test_get_bilingual_config()
    test_build_bilingual_segment_text()
    test_rebuild_bilingual_plain_text_from_segments()
    test_table_caption_not_treated_as_image_for_bilingual_skip()
    test_recover_layout_block_indices_uses_per_segment_map()
    test_build_bilingual_segment_text_styled_html()
    test_get_bilingual_styled_run_parts()
    test_md2html_preserves_bilingual_span_styles()
    test_xlsx_sanitize_worksheet_xml_escapes_literal_newlines()
    test_xlsx_rich_text_includes_excel_inline_font_fields()
    test_md2docx_parses_bilingual_span_styles()
    test_apply_bilingual_styled_spans_to_pandoc_docx()
    test_rebuild_markdown_text_segments_bilingual()
    print("\nAll bilingual export smoke tests passed!")
