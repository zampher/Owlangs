# SPDX-License-Identifier: AGPL-3.0-or-later
"""Tests for bilingual DOCX export in table cells with multiple paragraphs."""

from __future__ import annotations

import importlib.util
import sys
from io import BytesIO
from pathlib import Path

import pytest

pytest.importorskip("docx")


def _load_docx_rebuild():
    backend = Path(__file__).resolve().parent
    root = backend.parent
    for p in (str(root), str(backend)):
        if p not in sys.path:
            sys.path.insert(0, p)
    spec = importlib.util.spec_from_file_location(
        "docx_rebuild_multi_para_test",
        backend / "utils" / "document_rebuild" / "docx_rebuild.py",
    )
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


def test_resolve_table_cell_paragraph_by_target_text() -> None:
    """Each segment in a multi-paragraph cell must map to its own translated paragraph."""
    from docx import Document

    docx_rebuild = _load_docx_rebuild()
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = "评价及建议："
    cell.add_paragraph("通过对确认报告和相关附件的审核...")
    cell.add_paragraph("确认结果的批准")
    cell.add_paragraph("依据确认结果制定多索茶碱片的检验操作规程...")

    cell.paragraphs[0].text = "Evaluation and Recommendation:"
    cell.paragraphs[1].text = "Through the review of the qualification report..."
    cell.paragraphs[2].text = "Approval of Qualification Results"
    cell.paragraphs[3].text = "Based on the qualification results..."

    para_index_map = {}
    for local_idx, para in enumerate(cell.paragraphs):
        para_index_map[(True, 0, 0, 0, local_idx)] = para

    cases = [
        ("评价及建议：", "Evaluation and Recommendation:"),
        ("通过对确认报告和相关附件的审核...", "Through the review of the qualification report..."),
        ("确认结果的批准", "Approval of Qualification Results"),
        ("依据确认结果制定多索茶碱片的检验操作规程...", "Based on the qualification results..."),
    ]
    resolved_targets = []
    for source, target in cases:
        para = docx_rebuild._resolve_table_cell_paragraph(
            doc,
            0,
            0,
            0,
            target,
            source,
            para_index_map,
            None,
        )
        assert para is not None, f"failed to resolve paragraph for target={target!r}"
        assert target in (para.text or ""), (
            f"resolved wrong paragraph: expected target {target!r}, got {para.text!r}"
        )
        resolved_targets.append(para.text)

    assert len(set(resolved_targets)) == len(cases), "each segment must map to a distinct paragraph"


def test_resolve_table_cell_paragraph_by_cell_local_idx_hint() -> None:
    """cell_local_idx hint is used only when paragraph text matches target/source."""
    from docx import Document

    docx_rebuild = _load_docx_rebuild()
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = "First"
    cell.add_paragraph("Second")

    para_index_map = {
        (True, 0, 0, 0, 0): cell.paragraphs[0],
        (True, 0, 0, 0, 1): cell.paragraphs[1],
    }

    para = docx_rebuild._resolve_table_cell_paragraph(
        doc, 0, 0, 0, "Second", "第二", para_index_map, 1
    )
    assert para is not None
    assert para.text == "Second"


def test_resolve_table_cell_paragraph_rejects_stale_cell_local_idx() -> None:
    """Stale cell_local_idx from extract must not override target-text matching."""
    from docx import Document

    docx_rebuild = _load_docx_rebuild()
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = "Approval of Qualification Results"
    cell.add_paragraph("Requalification:")
    cell.add_paragraph("If relevant quality standards change...")

    para_index_map = {
        (True, 0, 0, 0, 0): cell.paragraphs[0],
        (True, 0, 0, 0, 1): cell.paragraphs[1],
        (True, 0, 0, 0, 2): cell.paragraphs[2],
    }

    # Stale hint points at paragraph 1 (Requalification) for segment 49 content.
    para = docx_rebuild._resolve_table_cell_paragraph(
        doc,
        0,
        0,
        0,
        "Approval of Qualification Results",
        "确认结果的批准",
        para_index_map,
        1,
    )
    assert para is not None
    assert "Approval of Qualification Results" in (para.text or "")


def test_nested_table_para_index_count_matches_extractor() -> None:
    """Nested table paragraphs must be included in global para_index counting."""
    from docx import Document
    from utils.docx_utils import count_non_toc_paragraphs_in_cell

    doc = Document()
    doc.add_paragraph("Body")
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = "Outer 1"
    cell.add_paragraph("Outer 2")
    nested = cell.add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "Nested 1"
    cell.add_paragraph("Outer 3")

    outer_only = count_non_toc_paragraphs_in_cell(cell, include_nested=False)
    with_nested = count_non_toc_paragraphs_in_cell(cell, include_nested=True)
    assert with_nested > outer_only
    assert outer_only >= 3


def test_is_expected_untranslated_numeric_segment() -> None:
    docx_rebuild = _load_docx_rebuild()
    segment = {
        "is_excluded": True,
        "exclusion_reason": "identifier",
        "source_text": "3",
        "target_text": "3",
    }
    assert docx_rebuild._is_expected_untranslated_table_segment(segment, "3", "3")


def test_rebuild_returns_doc_when_translations_already_applied() -> None:
    """Rebuild must succeed when every segment already matches the document (updated_count=0)."""
    pytest.importorskip("httpx")
    from docx import Document as PyDocxDocument
    from ir.document import Document

    docx_rebuild = _load_docx_rebuild()
    doc = PyDocxDocument()
    doc.add_paragraph("Hello world")

    buf = BytesIO()
    doc.save(buf)
    translated_doc = Document.from_bytes(content=buf.getvalue(), suffix=".docx", stem="test")

    task_state = {
        "translation_segments": {
            "segments": [
                {
                    "segment_index": 0,
                    "source_text": "Hello world",
                    "target_text": "Hello world",
                    "modified": False,
                    "is_excluded": False,
                    "is_failed": False,
                    "segment_info": {
                        "para_index": 0,
                        "is_table_cell": False,
                        "run_start_index": 0,
                        "run_end_index": 0,
                    },
                }
            ]
        }
    }

    rebuilt = docx_rebuild.rebuild_docx_document_from_segments(
        task_state,
        translated_doc,
        bilingual_export=False,
    )
    assert rebuilt is not None


def test_apply_font_size_delta_uses_default_when_run_has_no_explicit_size() -> None:
    """Font delta must apply even when runs inherit size from document defaults."""
    from docx import Document

    docx_rebuild = _load_docx_rebuild()
    doc = Document()
    para = doc.add_paragraph("Sample text without explicit run size")

    docx_rebuild._apply_font_size_delta(para, 2.0)

    assert para.runs
    size_pt = para.runs[0].font.size.pt
    assert size_pt == pytest.approx(13.0)


def test_bilingual_rebuild_applies_target_delta_when_translation_already_present() -> None:
    """Target font delta must apply on skip path when doc already contains translation."""
    pytest.importorskip("httpx")
    from docx import Document as PyDocxDocument
    from docx.shared import Pt
    from ir.document import Document

    docx_rebuild = _load_docx_rebuild()
    doc = PyDocxDocument()
    para = doc.add_paragraph("Hello world")
    para.runs[0].font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    translated_doc = Document.from_bytes(content=buf.getvalue(), suffix=".docx", stem="test")

    task_state = {
        "translation_segments": {
            "segments": [
                {
                    "segment_index": 0,
                    "source_text": "Hello source",
                    "target_text": "Hello world",
                    "modified": False,
                    "is_excluded": False,
                    "is_failed": False,
                    "segment_info": {
                        "para_index": 0,
                        "is_table_cell": False,
                        "run_start_index": 0,
                        "run_end_index": 0,
                    },
                }
            ]
        }
    }

    rebuilt = docx_rebuild.rebuild_docx_document_from_segments(
        task_state,
        translated_doc,
        bilingual_export=True,
        target_text_font_size_delta=3.0,
        source_text_font_size_delta=-1.0,
    )
    assert rebuilt is not None

    out_doc = PyDocxDocument(BytesIO(rebuilt.content))
    source_para = out_doc.paragraphs[0]
    target_para = out_doc.paragraphs[1]

    assert target_para.runs[0].font.size.pt == pytest.approx(14.0)
    assert source_para.runs[0].font.size.pt == pytest.approx(13.0)


def test_parse_textbox_sdt_storage_items_from_string() -> None:
    docx_rebuild = _load_docx_rebuild()
    raw = [
        "(('textbox', 0), '模板文档，如有雷同纯属巧合！')",
        "(('textbox', 1), 'Second textbox')",
    ]
    parsed = docx_rebuild._parse_textbox_sdt_storage_items(raw)
    assert len(parsed) == 2
    assert docx_rebuild._normalize_key(parsed[0][0]) == ("textbox", 0)
    assert "模板文档" in parsed[0][1]


def test_is_duplicate_legacy_vtextbox_container() -> None:
    docx_rebuild = _load_docx_rebuild()
    from docx.oxml import parse_xml

    alt_xml = """
    <mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml">
      <mc:Choice Requires="wps">
        <w:drawing><w:txbxContent><w:p><w:r><w:t>Target EN</w:t></w:r></w:p></w:txbxContent></w:drawing>
      </mc:Choice>
      <mc:Fallback>
        <w:pict><v:textbox><w:txbxContent><w:p><w:r><w:t>Target EN</w:t></w:r></w:p></w:txbxContent></v:textbox></w:pict>
      </mc:Fallback>
    </mc:AlternateContent>
    """
    alt = parse_xml(alt_xml)
    vtextbox = alt.xpath('.//*[local-name()="textbox"]')[0]
    assert docx_rebuild._is_duplicate_legacy_vtextbox_container(vtextbox) is True


def test_find_textbox_target_paragraph_element() -> None:
    docx_rebuild = _load_docx_rebuild()
    from docx.oxml import parse_xml

    container_xml = """
    <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:p><w:r><w:t>模板文档，如有雷同纯属巧合！</w:t></w:r></w:p>
      <w:p><w:r><w:t>Template document, any resemblance is purely coincidental!</w:t></w:r></w:p>
    </w:txbxContent>
    """
    container = parse_xml(container_xml)
    target = docx_rebuild._find_textbox_target_paragraph_element(
        container, "模板文档，如有雷同纯属巧合！"
    )
    assert target is not None
    runs = target.xpath('.//*[local-name()="t"]')
    text = "".join(t.text or "" for t in runs)
    assert "Template document" in text


def test_bilingual_textbox_insertion_runs_when_skip_legacy_header_footer() -> None:
    """Textbox source insertion must not be gated on legacy header/footer path."""
    from docx import Document

    docx_rebuild = _load_docx_rebuild()
    doc = Document()
    doc.add_paragraph("Body")

    segment_data_map = {
        117: {
            "source_text": "模板文档，如有雷同纯属巧合！",
            "target_text": "Template document, any resemblance is purely coincidental!",
            "segment_info": {},
            "is_excluded": False,
            "is_failed": False,
            "segment_type": "textbox_sdt",
            "textbox_key": "('textbox', 0)",
        }
    }

    extras_original = {
        "textboxes_sdts": [
            "(('textbox', 0), '模板文档，如有雷同纯属巧合！')",
        ]
    }

    docx_rebuild._insert_bilingual_source_paragraphs(
        doc,
        segment_data_map,
        {},
        target_first=False,
        extras_original=extras_original,
        skip_legacy_header_footer=True,
    )
    assert docx_rebuild._parse_textbox_sdt_storage_items(extras_original["textboxes_sdts"])
