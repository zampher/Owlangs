from __future__ import annotations

"""
DOCX utilities
 - extract_headers_footers(docx_bytes) -> list[(location, text)]
 - apply_headers_footers(docx_bytes, translations) -> bytes
 - extract_text_in_textboxes_and_sdts(docx_bytes) -> list[(location, text)]
 - apply_text_in_textboxes_and_sdts(docx_bytes, translations) -> bytes

Notes:
 - Keep field codes (PAGE, NUMPAGES, TOC) intact for headers/footers.
 - Textboxes are located in w:txbxContent (drawing) and sometimes legacy v:textbox; we best-effort handle both.
 - SDTs (Structured Document Tags / Content Controls) are w:sdt nodes.
"""

from io import BytesIO
import re
from typing import List, Tuple, Dict, Set, Optional, Any

try:
    from docx import Document  # type: ignore
except Exception:  # optional dependency
    Document = None  # type: ignore

try:
    from utils.table_utils import get_all_merged_regions_docx, is_cell_in_merged_region_docx
except Exception:
    get_all_merged_regions_docx = None
    is_cell_in_merged_region_docx = None


Location = Tuple[str, int]  # ("header"|"footer", section_index)


def _set_paragraph_text_direct(p_element, text: str):
    """
    Set text for a paragraph lxml element using direct XML manipulation.

    Unlike _set_paragraph_text (which relies on python-docx Paragraph with a
    parent document), this function operates on the raw lxml ``w:p`` element
    directly.  It removes all ``w:r`` children and adds a single new ``w:r`` /
    ``w:t`` pair.  This is necessary for paragraphs inside textboxes (txbxContent)
    and SDT content controls, where python-docx Paragraph with ``parent=None``
    may not properly clear existing runs.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    # 1. Remove all existing w:r children via xpath (local-name avoids namespace issues)
    for r in p_element.xpath('./*[local-name()="r"]'):
        try:
            p_element.remove(r)
        except Exception:
            pass

    # 3. Add a single new w:r / w:t with the text
    if text:
        r_elem = etree.SubElement(p_element, qn('w:r'))
        t_elem = etree.SubElement(r_elem, qn('w:t'))
        t_elem.text = text
        # Ensure xml:space="preserve" so leading/trailing whitespace is kept
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    else:
        # Even when empty, ensure paragraph has at least a run placeholder
        # (some renderers skip empty paragraphs entirely)
        r_elem = etree.SubElement(p_element, qn('w:r'))
        t_elem = etree.SubElement(r_elem, qn('w:t'))
        t_elem.text = ''
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _set_cell_text(cell, text: str):
    """Write text to a table cell using direct lxml manipulation.

    Avoids python-docx Paragraph API which may have issues in header/footer contexts.
    """
    from lxml import etree
    from docx.oxml.ns import qn

    tc = cell._tc  # w:tc element
    p_elements = tc.findall(qn('w:p'))
    for p_elem in p_elements[1:]:
        _set_paragraph_text_direct(p_elem, "")
    if p_elements:
        _set_paragraph_text_direct(p_elements[0], text)
    else:
        # No paragraphs exist — create one
        p_elem = etree.SubElement(tc, qn('w:p'))
        r_elem = etree.SubElement(p_elem, qn('w:r'))
        t_elem = etree.SubElement(r_elem, qn('w:t'))
        t_elem.text = text
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _is_merged_cell_skip(ri: int, ci: int, merged_region_set: Set[Tuple[int, int, int, int]]) -> bool:
    """Check if a cell at (ri, ci) should be skipped because it is a merged cell continuation."""
    for (start_row, start_col, end_row, end_col) in merged_region_set:
        if ri == start_row and ci == start_col:
            continue  # start cell, don't skip
        if start_row <= ri <= end_row and start_col <= ci <= end_col:
            return True  # continuation cell, skip
    return False


# Word field instructions that must stay dynamic (current page / total pages).
_PAGE_FIELD_INSTR_RE = re.compile(
    r"\b(PAGE|NUMPAGES|SECTIONPAGES|SECTION)\b",
    re.IGNORECASE,
)
# Static display snapshots often look like "1/16", "16 / 16", or "3 of 10".
_PAGE_DISPLAY_TEXT_RE = re.compile(
    r"^\s*\d+\s*(?:/\s*\d+|of\s+\d+)\s*$",
    re.IGNORECASE,
)


def _paragraph_has_page_field(p_element) -> bool:
    """True when paragraph XML contains PAGE/NUMPAGES/etc. field instructions."""
    for instr in p_element.xpath('.//*[local-name()="instrText"]'):
        if _PAGE_FIELD_INSTR_RE.search(instr.text or ""):
            return True
    return False


def text_looks_like_page_number_display(text: str) -> bool:
    """True for common page-number display strings (e.g. ``1/16``, ``3 of 10``)."""
    if not text or not text.strip():
        return False
    return bool(_PAGE_DISPLAY_TEXT_RE.match(text.strip()))


def _paragraph_should_preserve_pagination(p_element, display_text: str) -> bool:
    """Do not translate/replace paragraphs that carry dynamic or page-number text."""
    if _paragraph_has_page_field(p_element):
        return True
    return text_looks_like_page_number_display(display_text)


def _cell_should_preserve_pagination(cell) -> bool:
    """Do not translate/replace table cells that contain page fields or page numbers."""
    for p in cell.paragraphs:
        if _paragraph_should_preserve_pagination(p._p, p.text):
            return True
    joined = "\n".join(p.text for p in cell.paragraphs if p.text and p.text.strip())
    return text_looks_like_page_number_display(joined)


def _compute_part_fingerprint(part: Any) -> int:
    """Return a hash of the text content of a header/footer part.

    Used to deduplicate parts that share the same underlying content even when
    python-docx wraps them in different object instances (e.g. different
    sections referencing the same header XML file, or header vs first_page_header
    returning the same part when ``different_first_page_header_footer`` is False).
    """
    texts: List[str] = []
    for p in part.paragraphs:
        t = p.text
        if t:
            texts.append(t)
    for tbl in part.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    if t:
                        texts.append(t)
    # Use a fast hash — we only need to detect equality within one document
    return hash("\n".join(texts))


def _add_distinct_part(processed_parts: Set[int], parts: list,
                        name: str, part: Any) -> None:
    """Add *part* to *parts* if its content fingerprint hasn't been seen yet."""
    fp = _compute_part_fingerprint(part)
    if fp not in processed_parts:
        processed_parts.add(fp)
        parts.append((name, part))


def extract_headers_footers(docx_bytes: bytes) -> List[Tuple[Location, str]]:
    if Document is None:
        return []
    doc = Document(BytesIO(docx_bytes))
    items: List[Tuple[Location, str]] = []
    processed_parts: Set[int] = set()
    for idx, section in enumerate(doc.sections):
        # Collect all distinct header/footer parts for this section.
        # first_page_header/footer may be the same object as header/footer
        # when different_first_page_header_footer is False — dedup by element id.
        parts: List[Tuple[str, Any]] = []
        for name, part in (("header", section.header), ("footer", section.footer)):
            _add_distinct_part(processed_parts, parts, name, part)
        for name, part in (("header_first", section.first_page_header),
                           ("footer_first", section.first_page_footer)):
            _add_distinct_part(processed_parts, parts, name, part)
        for name, part in parts:
            texts: List[str] = []
            # Paragraphs
            for p in part.paragraphs:
                if _paragraph_should_preserve_pagination(p._p, p.text):
                    continue
                texts.append(p.text)
            # Tables — skip merged cell continuations
            for tbl in part.tables:
                merged_set = set()
                if get_all_merged_regions_docx is not None:
                    merged_set = set(get_all_merged_regions_docx(tbl))
                for ri, row in enumerate(tbl.rows):
                    for ci, cell in enumerate(row.cells):
                        if merged_set and _is_merged_cell_skip(ri, ci, merged_set):
                            continue
                        if _cell_should_preserve_pagination(cell):
                            continue
                        for p in cell.paragraphs:
                            if _paragraph_should_preserve_pagination(p._p, p.text):
                                continue
                            texts.append(p.text)
            content = "\n".join(t for t in texts if t)
            if content.strip():
                items.append(((name, idx), content))
    return items


def apply_headers_footers(docx_bytes: bytes, translations: Dict[Location, str]) -> bytes:
    if Document is None:
        return docx_bytes
    doc = Document(BytesIO(docx_bytes))
    processed_parts: Set[int] = set()
    for idx, section in enumerate(doc.sections):
        # Collect all distinct header/footer parts, same order as extraction
        parts: List[Tuple[str, Any]] = []
        for name, part in (("header", section.header), ("footer", section.footer)):
            _add_distinct_part(processed_parts, parts, name, part)
        for name, part in (("header_first", section.first_page_header),
                           ("footer_first", section.first_page_footer)):
            _add_distinct_part(processed_parts, parts, name, part)
        for name, part in parts:
            key: Location = (name, idx)
            if key not in translations:
                continue
            new_text = translations[key]

            # Count translatable structure elements (skip page-number fields/cells).
            para_count = sum(
                1 for p in part.paragraphs
                if p.text and p.text.strip()
                and not _paragraph_should_preserve_pagination(p._p, p.text)
            )
            cell_positions: List[Tuple[object, bool, bool]] = []
            for tbl in part.tables:
                merged_set = set()
                if get_all_merged_regions_docx is not None:
                    merged_set = set(get_all_merged_regions_docx(tbl))
                for ri, row in enumerate(tbl.rows):
                    for ci, cell in enumerate(row.cells):
                        is_merged = bool(merged_set and _is_merged_cell_skip(ri, ci, merged_set))
                        preserve = _cell_should_preserve_pagination(cell)
                        cell_positions.append((cell, is_merged, preserve))

            non_merged_cell_count = sum(
                1 for _, is_merged, preserve in cell_positions
                if not is_merged and not preserve
            )
            total_segments = para_count + non_merged_cell_count

            lines = new_text.split('\n')
            segment_texts = lines[:total_segments]
            if len(segment_texts) < total_segments:
                segment_texts.extend([''] * (total_segments - len(segment_texts)))

            line_idx = 0
            for p in part.paragraphs:
                if not p.text or not p.text.strip():
                    _set_paragraph_text(p, "")
                    continue
                if _paragraph_should_preserve_pagination(p._p, p.text):
                    continue
                if line_idx < len(segment_texts):
                    _set_paragraph_text(p, segment_texts[line_idx])
                    line_idx += 1

            for cell, is_merged, preserve in cell_positions:
                if is_merged or preserve:
                    continue
                if line_idx < len(segment_texts):
                    _set_cell_text(cell, segment_texts[line_idx])
                    line_idx += 1
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# Type alias for flat key: variadic tuple carrying position info
FlatKey = Tuple  # (str, int, str, ...)


def extract_headers_footers_flat(docx_bytes: bytes) -> List[Tuple[FlatKey, str]]:
    """Extract each paragraph and table cell from headers/footers as individual items.

    Unlike :func:`extract_headers_footers` — which joins all texts from one
    header/footer part into a single ``\\n``-delimited string — this function
    returns each paragraph or table cell as a **separate** item.  This makes it
    possible to translate individual cells reliably and to store translation
    segments per cell.

    Returns:
        List of ``(key, text)`` pairs.  The key format is:

        * Paragraph: ``(name, section_idx, "p", paragraph_idx)``
        * Table cell: ``(name, section_idx, "cell", table_idx, row_idx, cell_idx)``

        Only non-empty texts (after ``.strip()``) are included.
        Paragraphs/cells with PAGE/NUMPAGES fields or page-number display
        text (e.g. ``1/16``) are skipped so Word can keep them dynamic.
        Merged-cell continuations are skipped (same logic as
        :func:`extract_headers_footers`).
    """
    if Document is None:
        return []
    doc = Document(BytesIO(docx_bytes))
    items: List[Tuple[FlatKey, str]] = []
    processed_parts: Set[int] = set()
    for idx, section in enumerate(doc.sections):
        # Collect all distinct header/footer parts (same order as old functions)
        parts: List[Tuple[str, Any]] = []
        for name, part in (("header", section.header), ("footer", section.footer)):
            _add_distinct_part(processed_parts, parts, name, part)
        for name, part in (("header_first", section.first_page_header),
                           ("footer_first", section.first_page_footer)):
            _add_distinct_part(processed_parts, parts, name, part)
        for name, part in parts:
            # Paragraphs — one item per non-empty paragraph (skip page-number slots)
            for pi, p in enumerate(part.paragraphs):
                text = p.text
                if not text or not text.strip():
                    continue
                if _paragraph_should_preserve_pagination(p._p, text):
                    continue
                items.append(((name, idx, "p", pi), text))

            # Tables — one item per non-merged cell (combining all cell paragraphs)
            for ti, tbl in enumerate(part.tables):
                merged_set: set = set()
                if get_all_merged_regions_docx is not None:
                    merged_set = set(get_all_merged_regions_docx(tbl))
                for ri, row in enumerate(tbl.rows):
                    for ci, cell in enumerate(row.cells):
                        if merged_set and _is_merged_cell_skip(ri, ci, merged_set):
                            continue
                        if _cell_should_preserve_pagination(cell):
                            continue
                        cell_texts = [
                            p.text for p in cell.paragraphs
                            if p.text and p.text.strip()
                            and not _paragraph_should_preserve_pagination(p._p, p.text)
                        ]
                        if cell_texts:
                            items.append(
                                ((name, idx, "cell", ti, ri, ci), "\n".join(cell_texts))
                            )
    return items


# Maps part names that share content with another part name during extraction.
_PART_NAME_ALIASES: Dict[str, str] = {
    "footer_first": "footer",
}


def _part_element_id(part: Any) -> int:
    """Stable identity for a header/footer part (shared across sections)."""
    return id(part._element)


def _lookup_hf_translation(
    translations: Dict[FlatKey, str],
    name: str,
    section_idx: int,
    suffix: Tuple,
) -> Tuple[Optional[FlatKey], Optional[str]]:
    """Resolve a translation for *name*/*section_idx*, following part-name aliases."""
    candidates = [name]
    alias = _PART_NAME_ALIASES.get(name)
    if alias and alias not in candidates:
        candidates.append(alias)
    for candidate in candidates:
        key: FlatKey = (candidate, section_idx, *suffix)
        if key in translations:
            return key, translations[key]
    return None, None


def apply_headers_footers_flat(
    docx_bytes: bytes, translations: Dict[FlatKey, str]
) -> bytes:
    """Apply translations to individual paragraphs and cells in headers/footers.

    Counterpart to :func:`extract_headers_footers_flat`.  Translates keys
    directly to the matching paragraph or table cell without relying on
    positional ``\\n``-splitting (which is fragile when the LLM alters the
    number of lines).

    Args:
        docx_bytes: Raw DOCX file content.
        translations: Dict mapping flat keys (as produced by
            :func:`extract_headers_footers_flat`) to translated text.

    Returns:
        Modified DOCX bytes.
    """
    if Document is None:
        return docx_bytes
    doc = Document(BytesIO(docx_bytes))
    # Map fingerprint -> first section_idx where this part content was seen.
    # Extraction uses content-fingerprint dedup, so translation keys use the
    # first section's index.  During apply we must process EVERY distinct XML
    # part (even ones whose fingerprint was seen before) because separate XML
    # parts with identical content don't auto-propagate — but we look up
    # translations using the first section's index to match the extraction keys.
    content_first_idx: Dict[int, int] = {}
    processed_elements: Set[int] = set()
    for idx, section in enumerate(doc.sections):
        for name, part in (("header", section.header), ("footer", section.footer),
                           ("header_first", section.first_page_header),
                           ("footer_first", section.first_page_footer)):
            element_id = _part_element_id(part)
            if element_id in processed_elements:
                continue
            processed_elements.add(element_id)

            fp = _compute_part_fingerprint(part)
            if fp not in content_first_idx:
                content_first_idx[fp] = idx
            lookup_idx = content_first_idx[fp]

            # ---- paragraphs -------------------------------------------------
            for pi, p in enumerate(part.paragraphs):
                if _paragraph_should_preserve_pagination(p._p, p.text):
                    continue
                _, translated = _lookup_hf_translation(
                    translations, name, lookup_idx, ("p", pi)
                )
                if translated is not None:
                    _set_paragraph_text(p, translated)

            # ---- table cells ------------------------------------------------
            for ti, tbl in enumerate(part.tables):
                merged_set: set = set()
                if get_all_merged_regions_docx is not None:
                    merged_set = set(get_all_merged_regions_docx(tbl))
                for ri, row in enumerate(tbl.rows):
                    for ci, cell in enumerate(row.cells):
                        if merged_set and _is_merged_cell_skip(ri, ci, merged_set):
                            continue
                        if _cell_should_preserve_pagination(cell):
                            continue
                        _, translated = _lookup_hf_translation(
                            translations, name, lookup_idx, ("cell", ti, ri, ci)
                        )
                        if translated is not None:
                            _set_cell_text(cell, translated)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def has_toc_field(docx_bytes: bytes) -> bool:
    """Detect whether document contains a TOC field ({ TOC ... }).
    We simply scan for 'TOC' in field codes of the main document and headers/footers.
    """
    if Document is None:
        return False
    doc = Document(BytesIO(docx_bytes))
    # check body
    if _document_contains_toc(doc):
        return True
    # check sections' header/footer
    for section in doc.sections:
        for part in (section.header, section.footer):
            if _part_contains_toc(part):
                return True
    return False


def _document_contains_toc(doc) -> bool:
    from docx.oxml import OxmlElement  # type: ignore
    # paragraphs
    for p in doc.paragraphs:
        if _p_has_toc_field(p):
            return True
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _p_has_toc_field(p):
                        return True
    return False


def _part_contains_toc(part) -> bool:
    for p in part.paragraphs:
        if _p_has_toc_field(p):
            return True
    for tbl in part.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _p_has_toc_field(p):
                        return True
    return False


def _p_has_toc_field(paragraph) -> bool:
    try:
        p = paragraph._p  # lxml element
        fldChars = p.xpath('.//*[local-name()="fldChar"]')
        if not fldChars:
            # quick check for instruction text
            instrs = p.xpath('.//*[local-name()="instrText"]')
            for it in instrs:
                if 'TOC' in (it.text or ''):
                    return True
            return False
        instrs = p.xpath('.//*[local-name()="instrText"]')
        for it in instrs:
            if 'TOC' in (it.text or ''):
                return True
    except Exception:
        return False
    return False


def _set_paragraph_text(paragraph, text: str) -> None:
    """Set text for a paragraph, clearing existing content.

    Uses direct XML manipulation (same as _set_paragraph_text_direct) to ensure
    reliability in header/footer contexts where python-docx paragraph.text setter
    may produce inconsistent results.
    """
    _set_paragraph_text_direct(paragraph._p, text)


def refresh_toc_fields(docx_bytes: bytes, method: str = "field_update") -> bytes:
    """
    Refresh TOC (Table of Contents) fields in the DOCX document.
    This function updates all TOC field codes to ensure they display current content.
    
    Args:
        docx_bytes: The DOCX document bytes
        method: The refresh method to use:
            - "field_update": Update field codes to force refresh
            - "field_clear": Clear and recreate TOC fields
            - "full_rebuild": Completely rebuild TOC structure
    """
    if Document is None:
        return docx_bytes
    
    try:
        doc = Document(BytesIO(docx_bytes))
        
        if method == "field_update":
            # Method 1: Update field codes to force refresh
            _refresh_toc_in_document(doc)
            _refresh_toc_in_headers_footers(doc)
            
        elif method == "field_clear":
            # Method 2: Clear and recreate TOC fields
            _clear_and_recreate_toc_fields(doc)
            
        elif method == "full_rebuild":
            # Method 3: Completely rebuild TOC structure
            _rebuild_toc_structure(doc)
            
        else:
            print(f"[DEBUG] Unknown TOC refresh method: {method}, using field_update")
            _refresh_toc_in_document(doc)
            _refresh_toc_in_headers_footers(doc)
        
        # Save the updated document
        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()
        
    except Exception as e:
        print(f"[DEBUG] Error refreshing TOC fields: {e}")
        return docx_bytes


def _refresh_toc_in_headers_footers(doc) -> None:
    """Refresh TOC fields in headers and footers."""
    for section in doc.sections:
        _refresh_toc_in_part(section.header)
        _refresh_toc_in_part(section.footer)


def _clear_and_recreate_toc_fields(doc) -> None:
    """Clear existing TOC fields and recreate them."""
    try:
        # Find and clear all TOC fields
        for para in doc.paragraphs:
            _clear_toc_field_in_paragraph(para)
        
        # Recreate TOC fields
        _recreate_toc_fields(doc)
        
    except Exception as e:
        print(f"[DEBUG] Error clearing and recreating TOC fields: {e}")


def _clear_toc_field_in_paragraph(paragraph) -> None:
    """Clear TOC field in a single paragraph."""
    try:
        p = paragraph._p  # lxml element
        
        # Find TOC field characters
        fldChars = p.xpath('.//*[local-name()="fldChar"]')
        if not fldChars:
            return
        
        # Find instruction text
        instrs = p.xpath('.//*[local-name()="instrText"]')
        for it in instrs:
            if 'TOC' in (it.text or ''):
                # Clear the field
                it.text = ""
                break
                
    except Exception as e:
        print(f"[DEBUG] Error clearing TOC field in paragraph: {e}")


def _recreate_toc_fields(doc) -> None:
    """Recreate TOC fields in the document."""
    try:
        # This is a simplified recreation - in practice, you might want to
        # preserve the original TOC field properties
        for para in doc.paragraphs:
            if _p_has_toc_field(para):
                # Recreate the TOC field
                _add_toc_field_to_paragraph(para)
                
    except Exception as e:
        print(f"[DEBUG] Error recreating TOC fields: {e}")


def _add_toc_field_to_paragraph(paragraph) -> None:
    """Add a TOC field to a paragraph."""
    try:
        # This is a simplified implementation
        # In practice, you would need to create the proper XML structure
        p = paragraph._p  # lxml element
        
        # Add TOC field instruction
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Create field instruction
        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        instr_text.set(qn('w:space'), 'preserve')
        
        # Add to paragraph
        p.append(instr_text)
        
    except Exception as e:
        print(f"[DEBUG] Error adding TOC field to paragraph: {e}")


def _rebuild_toc_structure(doc) -> None:
    """Completely rebuild the TOC structure."""
    try:
        # This is a more complex operation that would involve:
        # 1. Finding all heading paragraphs
        # 2. Creating a new TOC structure
        # 3. Replacing the old TOC with the new one
        
        # For now, we'll use the field update method
        print("[DEBUG] Full TOC rebuild not fully implemented, using field update")
        _refresh_toc_in_document(doc)
        _refresh_toc_in_headers_footers(doc)
        
    except Exception as e:
        print(f"[DEBUG] Error rebuilding TOC structure: {e}")


def force_toc_update(docx_bytes: bytes) -> bytes:
    """
    Force a complete TOC update using multiple methods.
    This function tries different approaches to ensure the TOC is updated.
    """
    if Document is None:
        return docx_bytes
    
    try:
        # Method 1: Field update
        print("[DEBUG] Attempting TOC update with field_update method")
        result = refresh_toc_fields(docx_bytes, "field_update")
        
        # Method 2: If that doesn't work, try field clear
        if result == docx_bytes:
            print("[DEBUG] Field update failed, trying field_clear method")
            result = refresh_toc_fields(docx_bytes, "field_clear")
        
        # Method 3: If that doesn't work, try full rebuild
        if result == docx_bytes:
            print("[DEBUG] Field clear failed, trying full_rebuild method")
            result = refresh_toc_fields(docx_bytes, "full_rebuild")
        
        return result
        
    except Exception as e:
        print(f"[DEBUG] Error in force_toc_update: {e}")
        return docx_bytes


def update_static_toc(docx_bytes: bytes) -> bytes:
    """
    Update static table of contents by regenerating it based on current headings.
    This function creates a new TOC based on the current document structure.
    """
    if Document is None:
        return docx_bytes
    
    try:
        doc = Document(BytesIO(docx_bytes))
        
        # Find all heading paragraphs
        headings = []
        for i, para in enumerate(doc.paragraphs):
            if hasattr(para, 'style') and para.style:
                style_name = para.style.name.lower()
                if 'heading' in style_name:
                    # Extract heading level
                    level = 1
                    if 'heading 1' in style_name:
                        level = 1
                    elif 'heading 2' in style_name:
                        level = 2
                    elif 'heading 3' in style_name:
                        level = 3
                    elif 'heading 4' in style_name:
                        level = 4
                    elif 'heading 5' in style_name:
                        level = 5
                    elif 'heading 6' in style_name:
                        level = 6
                    
                    headings.append({
                        'level': level,
                        'text': para.text.strip(),
                        'paragraph': para
                    })
        
        print(f"[DEBUG] Found {len(headings)} headings in document")
        
        # Generate new TOC content
        toc_content = _generate_toc_content(headings)
        
        # Find and replace existing TOC
        _replace_static_toc(doc, toc_content)
        
        # Save the updated document
        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()
        
    except Exception as e:
        print(f"[DEBUG] Error updating static TOC: {e}")
        return docx_bytes


def _generate_toc_content(headings: List[Dict]) -> str:
    """Generate TOC content from headings."""
    toc_lines = []
    
    for heading in headings:
        level = heading['level']
        text = heading['text']
        
        # Create indentation based on level
        indent = "  " * (level - 1)
        
        # Add bullet point or number
        if level == 1:
            bullet = "•"
        elif level == 2:
            bullet = "◦"
        else:
            bullet = "▪"
        
        toc_lines.append(f"{indent}{bullet} {text}")
    
    return "\n".join(toc_lines)


def _replace_static_toc(doc, new_toc_content: str) -> None:
    """Replace existing static TOC with new content."""
    try:
        # Find the first paragraph that looks like a TOC
        toc_paragraph = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if any(keyword in text for keyword in ["目录", "Table of Contents", "Contents"]):
                toc_paragraph = para
                break
        
        if toc_paragraph:
            # Replace the content
            _set_paragraph_text(toc_paragraph, new_toc_content)
            print(f"[DEBUG] Replaced existing TOC with new content")
        else:
            # Create a new TOC paragraph at the beginning
            if doc.paragraphs:
                # Insert a new paragraph at the beginning
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                
                # Create a new paragraph element
                new_p = OxmlElement('w:p')
                new_r = OxmlElement('w:r')
                new_t = OxmlElement('w:t')
                new_t.text = new_toc_content
                new_r.append(new_t)
                new_p.append(new_r)
                
                # Insert at the beginning of the document
                body = doc._element.body
                body.insert(0, new_p)
                print(f"[DEBUG] Created new TOC paragraph at the beginning")
            else:
                print(f"[DEBUG] No paragraphs found in document")
            
    except Exception as e:
        print(f"[DEBUG] Error replacing static TOC: {e}")


def comprehensive_toc_update(docx_bytes: bytes) -> bytes:
    """
    Comprehensive TOC update that handles both field-based and static TOCs.
    This function tries multiple approaches to ensure the TOC is properly updated.
    """
    if Document is None:
        return docx_bytes
    
    try:
        # First, try to update field-based TOCs
        if has_toc_field(docx_bytes):
            print("[DEBUG] Document contains TOC fields, updating them...")
            result = force_toc_update(docx_bytes)
            if result != docx_bytes:
                return result
        
        # If no field-based TOCs or update failed, try static TOC update
        print("[DEBUG] No TOC fields found or update failed, trying static TOC update...")
        result = update_static_toc(docx_bytes)
        if result != docx_bytes:
            return result
        
        # If all else fails, return original
        print("[DEBUG] All TOC update methods failed, returning original document")
        return docx_bytes
        
    except Exception as e:
        print(f"[DEBUG] Error in comprehensive_toc_update: {e}")
        return docx_bytes


def _refresh_toc_in_document(doc) -> None:
    """Refresh TOC fields in document body."""
    # Refresh TOC in paragraphs
    for p in doc.paragraphs:
        _refresh_toc_in_paragraph(p)
    
    # Refresh TOC in tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _refresh_toc_in_paragraph(p)


def _refresh_toc_in_part(part) -> None:
    """Refresh TOC fields in a document part (header/footer)."""
    # Refresh TOC in paragraphs
    for p in part.paragraphs:
        _refresh_toc_in_paragraph(p)
    
    # Refresh TOC in tables
    for tbl in part.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _refresh_toc_in_paragraph(p)


def _refresh_toc_in_paragraph(paragraph) -> None:
    """Refresh TOC field in a single paragraph."""
    try:
        p = paragraph._p  # lxml element
        
        # Find TOC field characters
        fldChars = p.xpath('.//*[local-name()="fldChar"]')
        if not fldChars:
            return
        
        # Find instruction text
        instrs = p.xpath('.//*[local-name()="instrText"]')
        for it in instrs:
            if 'TOC' in (it.text or ''):
                # This is a TOC field, we need to refresh it
                # The simplest way is to update the field code to force refresh
                original_text = it.text or ''
                if original_text.strip():
                    # Add a timestamp or modify the field to force refresh
                    # This will make Word refresh the TOC when the document is opened
                    it.text = original_text.strip() + ' \\* MERGEFORMAT'
                break
                
    except Exception as e:
        print(f"[DEBUG] Error refreshing TOC in paragraph: {e}")
        pass


# -------------------------------
# Textboxes and SDTs (content controls)
# -------------------------------

def extract_text_in_textboxes_and_sdts(docx_bytes: bytes) -> List[Tuple[Location, str]]:
    """
    Extract text from textboxes (w:txbxContent and legacy v:textbox) and SDTs (w:sdt).
    Returns a list of ((kind, index), text), where kind in {"textbox", "sdt"}.
    """
    if Document is None:
        return []
    from docx.oxml.ns import qn  # type: ignore
    doc = Document(BytesIO(docx_bytes))
    items: List[Tuple[Location, str]] = []

    # Helper function to extract text from a container
    def _extract_text_from_container(container, container_name: str, start_index: int) -> int:
        idx = start_index
        texts: List[str] = []
        # paragraphs
        for p in container.xpath('.//*[local-name()="p"]'):
            runs = p.xpath('.//*[local-name()="t"]')
            if runs:
                txt = ''.join([(t.text or '') for t in runs])
                if txt:
                    texts.append(txt)
        # tables
        for cell in container.xpath('.//*[local-name()="tc"]'):
            for p in cell.xpath('.//*[local-name()="p"]'):
                runs = p.xpath('.//*[local-name()="t"]')
                if runs:
                    txt = ''.join([(t.text or '') for t in runs])
                    if txt:
                        texts.append(txt)
        content = "\n".join(t for t in texts if t)
        if content.strip():
            items.append(((container_name, idx), content))
        idx += 1
        return idx

    # Extract all elements with unified indexing
    try:
        # 1) SDTs in body - only extract top-level SDTs to avoid duplication
        sdt_elems = doc._element.body.xpath('.//*[local-name()="sdt"]')
        sdt_index = 0
        for sdt in sdt_elems:
            # Check if this SDT is nested within another SDT
            parent_sdt = sdt.xpath('./ancestor::*[local-name()="sdt"]')
            if parent_sdt:
                print(f"[DEBUG] Skipping nested SDT: {sdt_index}")
                continue
            
            # Process only direct sdtContent children (not nested within child SDTs)
            sdt_contents = sdt.xpath('./*[local-name()="sdtContent"]')
            print(f"[DEBUG] Found {len(sdt_contents)} sdtContent elements in SDT {sdt_index}")
            
            for i, content in enumerate(sdt_contents):
                texts: List[str] = []
                # paragraphs within sdtContent (skip those inside child SDTs)
                for p in content.xpath('.//*[local-name()="p"]'):
                    if p.xpath('./ancestor::*[local-name()="sdt"][parent::*[local-name()="sdtContent"]]'):
                        continue
                    runs = p.xpath('.//*[local-name()="t"]')
                    if runs:
                        txt = ''.join([(t.text or '') for t in runs])
                        if txt:
                            texts.append(txt)
                # tables within sdtContent (skip cells inside child SDTs)
                for cell in content.xpath('.//*[local-name()="tc"]'):
                    if cell.xpath('./ancestor::*[local-name()="sdt"][parent::*[local-name()="sdtContent"]]'):
                        continue
                    for p in cell.xpath('.//*[local-name()="p"]'):
                        runs = p.xpath('.//*[local-name()="t"]')
                        if runs:
                            txt = ''.join([(t.text or '') for t in runs])
                            if txt:
                                texts.append(txt)
                content_text = "\n".join(t for t in texts if t)
                if content_text.strip():
                    items.append((('sdt_content', sdt_index, i), content_text))
                    print(f"[DEBUG] Found sdtContent {i}: {content_text[:50]}...")
            
            # Process child SDTs within this SDT
            child_sdts = sdt.xpath('.//*[local-name()="sdt"]')
            print(f"[DEBUG] Found {len(child_sdts)} child SDTs in SDT {sdt_index}")
            
            for i, child_sdt in enumerate(child_sdts):
                texts: List[str] = []
                # paragraphs within child SDT
                for p in child_sdt.xpath('.//*[local-name()="p"]'):
                    runs = p.xpath('.//*[local-name()="t"]')
                    if runs:
                        txt = ''.join([(t.text or '') for t in runs])
                        if txt:
                            texts.append(txt)
                # tables within child SDT
                for cell in child_sdt.xpath('.//*[local-name()="tc"]'):
                    for p in cell.xpath('.//*[local-name()="p"]'):
                        runs = p.xpath('.//*[local-name()="t"]')
                        if runs:
                            txt = ''.join([(t.text or '') for t in runs])
                            if txt:
                                texts.append(txt)
                child_text = "\n".join(t for t in texts if t)
                if child_text.strip():
                    items.append((('sdt_child', sdt_index, i), child_text))
                    print(f"[DEBUG] Found child SDT {i}: {child_text[:50]}...")
            
            # Process direct SDT content (if no sdtContent or child SDTs)
            if not sdt_contents and not child_sdts:
                texts: List[str] = []
                # paragraphs within sdt
                for p in sdt.xpath('.//*[local-name()="p"]'):
                    runs = p.xpath('.//*[local-name()="t"]')
                    if runs:
                        txt = ''.join([(t.text or '') for t in runs])
                        if txt:
                            texts.append(txt)
                # tables within sdt
                for cell in sdt.xpath('.//*[local-name()="tc"]'):
                    for p in cell.xpath('.//*[local-name()="p"]'):
                        runs = p.xpath('.//*[local-name()="t"]')
                        if runs:
                            txt = ''.join([(t.text or '') for t in runs])
                            if txt:
                                texts.append(txt)
                content = "\n".join(t for t in texts if t)
                if content.strip():
                    items.append((('sdt', sdt_index), content))
                    print(f"[DEBUG] Found SDT content: {content[:50]}...")
            
            sdt_index += 1
        
        # 2) Textboxes and drawing elements - search in entire document
        all_textbox_elements = []

        # Search for w:txbxContent nodes
        try:
            txbx_nodes = doc._element.xpath('.//*[local-name()="txbxContent"]')
            # Filter: skip txbxContent in mc:Fallback when mc:Choice already has one
            # (this is the same textbox rendered via two paths — DML and VML)
            for txbx in txbx_nodes:
                _alt_content = txbx.xpath('./ancestor::*[local-name()="AlternateContent"]')
                if _alt_content:
                    _in_fallback = txbx.xpath('./ancestor::*[local-name()="Fallback"]')
                    if _in_fallback:
                        _choice_txbx = _alt_content[0].xpath(
                            './/*[local-name()="Choice"]//*[local-name()="txbxContent"]'
                        )
                        if _choice_txbx:
                            print(f"[DEBUG] Skipping txbxContent in mc:Fallback (Choice path already covers it)")
                            continue
                all_textbox_elements.append(txbx)
        except Exception as e:
            print(f"[DEBUG] Error searching w:txbxContent: {e}")

        # Search for legacy v:textbox nodes
        try:
            pict_nodes = doc._element.xpath('.//*[local-name()="pict"]//*[local-name()="textbox"]')
            all_textbox_elements.extend(pict_nodes)
        except Exception as e:
            print(f"[DEBUG] Error searching v:textbox: {e}")
        
        # Search for drawing elements with text
        try:
            drawing_nodes = doc._element.xpath('.//*[local-name()="drawing"]')
            for drawing in drawing_nodes:
                text_elements = drawing.xpath('.//*[local-name()="t"]')
                if text_elements:
                    all_textbox_elements.append(drawing)
        except Exception as e:
            print(f"[DEBUG] Error searching w:drawing: {e}")
        
        print(f"[DEBUG] Found {len(all_textbox_elements)} total textbox/drawing elements in document")
        
        # Process all elements with consistent indexing
        tb_index = 0
        for element in all_textbox_elements:
            if element.tag.endswith('txbxContent') or element.tag.endswith('textbox'):
                # Skip textboxes inside SDTs (already handled by SDT extraction)
                if element.xpath('./ancestor::*[local-name()="sdt"]'):
                    print(f"[DEBUG] Skipping textbox {tb_index} inside SDT")
                    tb_index += 1
                    continue
                # Handle textbox elements
                print(f"[DEBUG] Processing textbox element {tb_index}")
                tb_index = _extract_text_from_container(element, 'textbox', tb_index)
            elif element.tag.endswith('drawing'):
                # Handle drawing elements - always increment index to match apply phase
                text_elements = element.xpath('.//*[local-name()="t"]')
                if text_elements:
                    texts = []
                    for t_elem in text_elements:
                        if t_elem.text:
                            texts.append(t_elem.text)
                    content = "\n".join(t for t in texts if t)
                    if content.strip():
                        print(f"[DEBUG] Found text in drawing: {content[:100]}...")
                        items.append((('textbox', tb_index), content))
                tb_index += 1
            else:
                tb_index += 1
    except Exception as e:
        print(f"[DEBUG] Error extracting elements: {e}")
        pass

    return items


def apply_text_in_textboxes_and_sdts(docx_bytes: bytes, translations: Dict[Location, str]) -> bytes:
    """
    Apply translations back into textboxes (w:txbxContent / legacy v:textbox) and SDTs.
    The matching order follows extract_text_in_textboxes_and_sdts iteration, keyed by (kind, index).
    """
    if Document is None:
        return docx_bytes
    doc = Document(BytesIO(docx_bytes))

    # Helper to set text for a node containing w:p descendants
    def _apply_to_container(container, kind: str, start_index: int) -> int:
        """
        Apply translations to textbox containers using direct XML manipulation.

        Uses ``_set_paragraph_text_direct`` instead of python-docx ``Paragraph``
        because the latter requires a valid parent document.  Textbox paragraphs
        inside ``w:txbxContent`` are not part of the main document body, so
        ``Paragraph(xml_element, None)`` does not reliably clear existing runs.
        """
        idx = start_index
        p_xpath = './/*[local-name()="p"]'
        for node in container:
            key = ('textbox', idx)
            if key in translations:
                new_text = translations[key]
                print(f"[DEBUG] Applying translation to container {idx}: {new_text[:50]}...")
                paragraphs = node.xpath(p_xpath)
                if paragraphs:
                    print(f"[DEBUG] Found {len(paragraphs)} paragraphs in container {idx}")
                    # Set first paragraph text via direct XML (reliable inside textboxes)
                    _set_paragraph_text_direct(paragraphs[0], new_text)
                    # Clear remaining paragraphs
                    for p in paragraphs[1:]:
                        _set_paragraph_text_direct(p, "")
                else:
                    print(f"[DEBUG] No paragraphs found in container {idx}")
            else:
                print(f"[DEBUG] No translation found for key {key}")
            idx += 1
        return idx

    # Apply all elements with unified indexing (same logic as extraction)
    try:
        # 1) Apply SDTs - only apply to top-level SDTs to avoid duplication
        sdt_elems = doc._element.body.xpath('.//*[local-name()="sdt"]')
        sdt_index = 0
        for sdt in sdt_elems:
            # Check if this SDT is nested within another SDT
            parent_sdt = sdt.xpath('./ancestor::*[local-name()="sdt"]')
            if parent_sdt:
                print(f"[DEBUG] Skipping nested SDT in apply: {sdt_index}")
                continue
            
            # Apply sdtContent translations (direct children only, not nested in child SDTs)
            sdt_contents = sdt.xpath('./*[local-name()="sdtContent"]')
            for i, content in enumerate(sdt_contents):
                key = ('sdt_content', sdt_index, i)
                if key in translations:
                    new_text = translations[key]
                    print(f"[DEBUG] Applying sdtContent translation {sdt_index}.{i}: {new_text[:50]}...")
                    # Only modify paragraphs outside child SDTs
                    paragraphs = content.xpath('./*[local-name()="p"]')
                    if not paragraphs:
                        # Fallback: look deeper but skip child SDT paragraphs
                        paragraphs = content.xpath('.//*[local-name()="p"][not(ancestor::*[local-name()="sdt"][parent::*[local-name()="sdtContent"]])]')
                    if paragraphs:
                        _set_paragraph_text_direct(paragraphs[0], new_text)
                        for p in paragraphs[1:]:
                            _set_paragraph_text_direct(p, "")

            # Apply child SDT translations
            child_sdts = sdt.xpath('.//*[local-name()="sdt"]')
            for i, child_sdt in enumerate(child_sdts):
                key = ('sdt_child', sdt_index, i)
                if key in translations:
                    new_text = translations[key]
                    print(f"[DEBUG] Applying child SDT translation {sdt_index}.{i}: {new_text[:50]}...")
                    paragraphs = child_sdt.xpath('.//*[local-name()="p"]')
                    if paragraphs:
                        _set_paragraph_text_direct(paragraphs[0], new_text)
                        for p in paragraphs[1:]:
                            _set_paragraph_text_direct(p, "")

            # Apply direct SDT content (if no sdtContent or child SDTs)
            if not sdt_contents and not child_sdts:
                key = ('sdt', sdt_index)
                if key in translations:
                    new_text = translations[key]
                    print(f"[DEBUG] Applying SDT translation {sdt_index}: {new_text[:50]}...")
                    paragraphs = sdt.xpath('.//*[local-name()="p"]')
                    if paragraphs:
                        _set_paragraph_text_direct(paragraphs[0], new_text)
                        for p in paragraphs[1:]:
                            _set_paragraph_text_direct(p, "")
            
            sdt_index += 1
        
        # 2) Apply textboxes and drawing elements
        all_textbox_elements = []
        # Track mc:AlternateContent Fallback elements that need the Choice translation
        _choice_to_fallback: dict = {}

        # Search for w:txbxContent nodes
        try:
            txbx_nodes = doc._element.xpath('.//*[local-name()="txbxContent"]')
            # Filter: skip txbxContent in mc:Fallback when mc:Choice already has one
            for txbx in txbx_nodes:
                _alt_content = txbx.xpath('./ancestor::*[local-name()="AlternateContent"]')
                if _alt_content:
                    _in_fallback = txbx.xpath('./ancestor::*[local-name()="Fallback"]')
                    _choice_txbx = _alt_content[0].xpath(
                        './/*[local-name()="Choice"]//*[local-name()="txbxContent"]'
                    )
                    if _in_fallback and _choice_txbx:
                        # Map Fallback to its Choice counterpart for later application
                        _choice_to_fallback[id(_choice_txbx[0])] = txbx
                        print(f"[DEBUG] Skipping txbxContent in mc:Fallback in apply (Choice covers it)")
                        continue
                all_textbox_elements.append(txbx)
        except Exception as e:
            print(f"[DEBUG] Error searching w:txbxContent: {e}")
        
        # Search for legacy v:textbox nodes
        try:
            pict_nodes = doc._element.xpath('.//*[local-name()="pict"]//*[local-name()="textbox"]')
            all_textbox_elements.extend(pict_nodes)
        except Exception as e:
            print(f"[DEBUG] Error searching v:textbox: {e}")
        
        # Search for drawing elements with text
        try:
            drawing_nodes = doc._element.xpath('.//*[local-name()="drawing"]')
            for drawing in drawing_nodes:
                text_elements = drawing.xpath('.//*[local-name()="t"]')
                if text_elements:
                    all_textbox_elements.append(drawing)
        except Exception as e:
            print(f"[DEBUG] Error searching w:drawing: {e}")
        
        print(f"[DEBUG] Found {len(all_textbox_elements)} total textbox/drawing elements for applying translations")
        
        # Process all elements with consistent indexing
        tb_index = 0
        for element in all_textbox_elements:
            key = ('textbox', tb_index)

            # Skip textboxes inside SDTs (already handled by SDT apply)
            if element.xpath('./ancestor::*[local-name()="sdt"]'):
                print(f"[DEBUG] Skipping textbox {tb_index} inside SDT in apply")
                tb_index += 1
                continue

            if key in translations:
                new_text = translations[key]
                print(f"[DEBUG] Applying translation to element {tb_index}: {new_text[:50]}...")
                
                if element.tag.endswith('txbxContent') or element.tag.endswith('textbox'):
                    # Handle textbox elements (Choice path)
                    tb_index = _apply_to_container([element], kind='drawing', start_index=tb_index)
                    # Also apply same translation to Fallback counterpart in mc:AlternateContent
                    _fallback_elem = _choice_to_fallback.get(id(element))
                    if _fallback_elem is not None:
                        print(f"[DEBUG] Also applying translation to Fallback counterpart of Choice {tb_index - 1}")
                        _fb_paragraphs = _fallback_elem.xpath('.//*[local-name()="p"]')
                        if _fb_paragraphs:
                            _set_paragraph_text_direct(_fb_paragraphs[0], new_text)
                            for _fb_p in _fb_paragraphs[1:]:
                                _set_paragraph_text_direct(_fb_p, "")
                elif element.tag.endswith('drawing'):
                    # Handle drawing elements - always increment index to match extraction phase
                    text_elements = element.xpath('.//*[local-name()="t"]')
                    if text_elements:
                        print(f"[DEBUG] Found {len(text_elements)} text elements in drawing {tb_index}")
                        # Replace first text element, clear others
                        text_elements[0].text = new_text
                        for t_elem in text_elements[1:]:
                            t_elem.text = ''
                    tb_index += 1
                else:
                    tb_index += 1
            else:
                print(f"[DEBUG] No translation found for key {key}")
                tb_index += 1
    except Exception as e:
        print(f"[DEBUG] Error applying elements: {e}")
        pass

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

