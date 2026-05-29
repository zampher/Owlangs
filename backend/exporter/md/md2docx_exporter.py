# SPDX-FileCopyrightText: 2026 Zampher
# SPDX-License-Identifier: MPL-2.0

"""
Markdown to DOCX exporter.

Converts MarkdownDocument to DOCX format using python-docx library.
"""

import html
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple, TYPE_CHECKING
from xml.etree import ElementTree as ET

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_BREAK

from logger import unified_logger as logger
from logger.logger import LogModule
from utils.math_md_normalize import (
    TEX_LATEX_FENCE_LANGS,
    extract_display_math_inner_from_tex_fence_body,
    parse_opening_markdown_fence_language,
)
from exporter.md.base import MDExporter, MDExporterConfig
from ir.document import Document
from ir.markdown_document import MarkdownDocument

if TYPE_CHECKING:
    from layout.base import LayoutDocument

# Split patterns for _add_formatted_runs_for_text: LaTeX display \[...\], $$...$$, inline \(...\), $...$, then markdown.
_TEX_SPLIT_PARTS_MULTILINE = (
    r"(\\\[[\s\S]*?\\\]|\\\([\s\S]*?\\\)|\$\$[\s\S]*?\$\$|\$[^$]*?\$|"
    r"\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))"
)
_TEX_SPLIT_PARTS_SINGLE_LINE_INLINE = (
    r"(\\\[[\s\S]*?\\\]|\\\([\s\S]*?\\\)|\$\$[\s\S]*?\$\$|\$[^$\n]+?\$|"
    r"\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))"
)
# Inline HTML from translators/rebuild: <sup>n</sup>, <sub>n</sub> inside prose (author affiliations, etc.)
_HTML_SUP_SUB_TAGS = re.compile(
    r"(<sup>\s*[\s\S]*?</sup>|<sub>\s*[\s\S]*?</sub>)",
    re.IGNORECASE,
)


@dataclass
class MD2DOCXExporterConfig(MDExporterConfig):
    """Configuration for Markdown to DOCX exporter."""
    font_name: str = "Calibri"
    font_size: int = 11
    heading_sizes: dict = None  # Will be initialized with default values
    layout_document: 'LayoutDocument | None' = None  # Optional layout document for PDF workflow
    image_data_map: dict = None  # Optional image data map for embedding images in DOCX
    table_body_format: str = "image"  # Table format: "html" or "image"
    equation_format: str = "text"  # Equation format: "text" (LaTeX), "latex", or "image"
    debug_output_dir: Optional[Path] = None  # If set, write MD input to debug_output_dir/docx_export_input.md for debugging

    def __post_init__(self):
        if self.heading_sizes is None:
            self.heading_sizes = {
                1: 24,  # H1
                2: 18,  # H2
                3: 14,  # H3
                4: 12,  # H4
                5: 11,  # H5
                6: 10,  # H6
            }


class MD2DOCXExporter(MDExporter):
    """Exporter to convert MarkdownDocument to DOCX format."""
    
    def __init__(self, config: MD2DOCXExporterConfig = None):
        config = config or MD2DOCXExporterConfig()
        super().__init__(config=config)
        self.font_name = config.font_name
        self.font_size = config.font_size
        self.heading_sizes = config.heading_sizes
        self.layout_document = config.layout_document  # Layout document for PDF workflow
        # Initialize image_data_map from config, default to empty dict
        self.image_data_map = config.image_data_map if config.image_data_map is not None else {}
        # Log image_data_map initialization for debugging
        if self.image_data_map:
            logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] MD2DOCXExporter initialized with {len(self.image_data_map)} images in image_data_map")
            sample_keys = list(self.image_data_map.keys())[:5]
            logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Sample keys: {sample_keys}")
        else:
            logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] MD2DOCXExporter initialized with EMPTY image_data_map!")
        # Table format: "html" (convert to DOCX table) or "image" (keep as image)
        self.table_body_format = config.table_body_format if config.table_body_format else "image"
        # Equation format: "text"/"latex" (render as LaTeX formula) or "image" (render as image)
        self.equation_format = config.equation_format if config.equation_format else "text"
        self.debug_output_dir: Optional[Path] = getattr(config, "debug_output_dir", None) or None

    def _set_run_east_asia_font(self, run, font_name: str | None = None) -> None:
        """Set w:eastAsia so CJK text uses our font. Word uses eastAsia for East Asian chars;
        python-docx run.font.name only sets w:ascii and w:hAnsi, so we set eastAsia explicitly."""
        name = font_name if font_name is not None else self.font_name
        try:
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    rFonts.set(qn('w:eastAsia'), name)
        except Exception:
            pass

    @staticmethod
    def _html_sup_sub_inner_to_plain(fragment: str) -> str:
        """Decode entities, strip nested HTML tags (tag names start with a letter; avoids eating '<2>' literals)."""
        t = html.unescape(fragment.strip())
        t = re.sub(r"</?[A-Za-z][A-Za-z0-9-]*[^>]*>", "", t)
        return t.strip()

    def _apply_default_body_font(self, run) -> None:
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        self._set_run_east_asia_font(run)

    def _add_runs_with_html_sup_sub(self, para, text: str) -> None:
        """Add paragraph runs; interpret <sup>/<sub> as Word superscript/subscript."""
        if not text:
            return
        clean = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]", "", text)
        if not clean:
            return
        lower = clean.lower()
        if "<sup" not in lower and "<sub" not in lower:
            run = para.add_run(clean)
            self._apply_default_body_font(run)
            return
        pos = 0
        for m in _HTML_SUP_SUB_TAGS.finditer(clean):
            if m.start() > pos:
                run = para.add_run(clean[pos : m.start()])
                self._apply_default_body_font(run)
            tag = m.group(0)
            sm = re.search(r"<sup>\s*([\s\S]*?)\s*</sup>", tag, re.IGNORECASE)
            if sm:
                inner = self._html_sup_sub_inner_to_plain(sm.group(1))
                if inner:
                    run = para.add_run(inner)
                    self._apply_default_body_font(run)
                    run.font.superscript = True
            else:
                bm = re.search(r"<sub>\s*([\s\S]*?)\s*</sub>", tag, re.IGNORECASE)
                if bm:
                    inner = self._html_sup_sub_inner_to_plain(bm.group(1))
                    if inner:
                        run = para.add_run(inner)
                        self._apply_default_body_font(run)
                        run.font.subscript = True
                else:
                    run = para.add_run(tag)
                    self._apply_default_body_font(run)
            pos = m.end()
        if pos < len(clean):
            run = para.add_run(clean[pos:])
            self._apply_default_body_font(run)

    def export(self, document: MarkdownDocument) -> Document:
        """Convert MarkdownDocument to DOCX Document."""
        # Get markdown content with robust encoding handling
        if isinstance(document.content, str):
            md_content = document.content
        elif isinstance(document.content, bytes):
            # Try UTF-8 first, then fallback to other encodings
            try:
                md_content = document.content.decode('utf-8')
            except UnicodeDecodeError:
                # Try UTF-8 with BOM
                try:
                    md_content = document.content.decode('utf-8-sig')
                except UnicodeDecodeError:
                    # Try common encodings
                    for encoding in ['gbk', 'gb2312', 'latin-1', 'cp1252']:
                        try:
                            md_content = document.content.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        # Final fallback: UTF-8 with error replacement
                        md_content = document.content.decode('utf-8', errors='replace')
        else:
            raise ValueError(f"Unexpected content type: {type(document.content)}")

        # Write MD to debug directory when configured (for debugging DOCX export input)
        if self.debug_output_dir:
            try:
                debug_dir = Path(self.debug_output_dir)
                debug_dir.mkdir(parents=True, exist_ok=True)
                md_debug_path = debug_dir / "docx_export_input.md"
                md_debug_path.write_text(md_content, encoding="utf-8")
                logger.debug(LogModule.EXPORT, f"[DOCX-DEBUG] Wrote MD input to {md_debug_path}")
            except Exception as e:
                logger.warning(LogModule.EXPORT, f"[DOCX-DEBUG] Failed to write debug MD: {e}")

        # Create a new DOCX document
        docx_doc = DocxDocument()
        
        # Parse markdown and convert to DOCX
        # If layout_document is available (PDF workflow), use layout-based formula detection
        # Otherwise, use code-based detection (for other workflows)
        use_layout_detection = self.layout_document is not None
        
        # For PDF workflow with table_body_format="html", translated content may contain
        # inline HTML <table>...</table> (from HTML rebuild). MD2DOCX only understands
        # markdown table syntax, so we first normalize simple HTML tables into markdown
        # tables to reuse the existing table rendering logic.
        if use_layout_detection and self.table_body_format == "html" and "<table" in md_content:
            md_content = self._html_tables_to_markdown(md_content)
        
        self._markdown_to_docx(md_content, docx_doc, use_layout_detection=use_layout_detection)
        
        # Save DOCX to bytes
        docx_bytes = BytesIO()
        docx_doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return Document.from_bytes(
            content=docx_bytes.getvalue(),
            suffix=".docx",
            stem=document.stem
        )

    def _html_tables_to_markdown(self, md_content: str) -> str:
        """Convert simple HTML <table> blocks into markdown tables.
        
        This is primarily for PDF workflows where table bodies may already be
        rebuilt as HTML tables. We convert them back to markdown so that the
        existing markdown-table parser can render proper DOCX tables.
        """
        try:
            from bs4 import BeautifulSoup  # type: ignore
        except Exception:
            # BeautifulSoup not available; return original content
            return md_content

        soup = BeautifulSoup(md_content, "html.parser")
        changed = False

        # Find all table tags and replace them with markdown text nodes
        for table in soup.find_all("table"):
            # Build markdown from HTML table
            header_cells: list[str] = []
            body_rows: list[list[str]] = []

            # Header
            thead = table.find("thead")
            if thead:
                header_row = thead.find("tr")
            else:
                header_row = table.find("tr")

            if header_row:
                for th in header_row.find_all(["th", "td"]):
                    header_cells.append(th.get_text(strip=True))

            # Body
            tbody = table.find("tbody")
            rows = tbody.find_all("tr") if tbody else table.find_all("tr")[1:]
            for tr in rows:
                row_cells: list[str] = []
                for td in tr.find_all(["td", "th"]):
                    row_cells.append(td.get_text(strip=True))
                if row_cells:
                    body_rows.append(row_cells)

            # Need at least a header and one row to form a markdown table
            if not header_cells or not body_rows:
                continue

            # Normalize body rows column count
            col_count = len(header_cells)
            normalized_rows: list[list[str]] = []
            for row in body_rows:
                if len(row) < col_count:
                    row = row + [""] * (col_count - len(row))
                elif len(row) > col_count:
                    row = row[:col_count]
                normalized_rows.append(row)

            def _escape(cell: str) -> str:
                return cell.replace("|", "\\|")

            lines: list[str] = []
            # Header
            lines.append("| " + " | ".join(_escape(c) for c in header_cells) + " |")
            # Separator
            lines.append("| " + " | ".join("---" for _ in range(col_count)) + " |")
            # Body
            for row in normalized_rows:
                lines.append("| " + " | ".join(_escape(c) for c in row) + " |")

            md_table = "\n".join(lines)

            # Replace table with markdown text (wrapped in newlines to avoid merging)
            table.replace_with(soup.new_string("\n" + md_table + "\n"))
            changed = True

        return str(soup) if changed else md_content
    
    def _markdown_to_docx(self, md_content: str, docx_doc: DocxDocument, use_layout_detection: bool = False):
        """Convert markdown content to DOCX paragraphs.
        
        Args:
            md_content: Markdown content string
            docx_doc: DOCX document object
            use_layout_detection: If True, use layout_document to detect formulas by type.
                                 If False, use code-based detection ($$...$$ format).
        """
        # For PDF workflow: process formulas directly from layout_document
        # Extract formulas from layout blocks and insert them at appropriate positions
        if use_layout_detection and self.layout_document:
            try:
                # Process markdown with layout-aware formula insertion
                self._markdown_to_docx_with_layout(md_content, docx_doc)
                logger.info(LogModule.EXPORT, "[DOCX] Using layout-based formula detection for PDF workflow")
                return
            except Exception as e:
                # If layout-based processing fails, fall back to code detection
                logger.warning(LogModule.EXPORT, f"[DOCX] Layout-based formula detection failed: {e}, falling back to code detection")
                use_layout_detection = False
        
        # For non-PDF workflows (MD files, etc.), use code-based detection
        if not use_layout_detection:
            # Check if markdown contains LaTeX formulas
            has_block_math = '$$' in md_content
            has_inline_math = re.search(r'\$[^$\n]+\$', md_content) is not None
            if has_block_math or has_inline_math:
                logger.info(LogModule.EXPORT, "[DOCX] Using code-based formula detection (LaTeX formulas found in markdown)")
        
        # Split content into lines for processing
        lines = md_content.split('\n')
        
        # Use module-level logger (already imported at top of file)
        
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Skip empty lines (will be handled by paragraph spacing)
            if not line:
                i += 1
                continue
            
            # Handle headings (### Heading)
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                if heading_text:
                    self._add_heading(docx_doc, heading_text, level)
            
            # Handle LaTeX math blocks ($$...$$)
            # For PDF workflow (use_layout_detection=True), skip code detection
            # Formulas will be handled directly from layout_document
            # If equation_format is "image", skip $$...$$ blocks as formulas are already rendered as images
            elif not use_layout_detection and line.strip().startswith('$$') and self.equation_format != "image":
                # Extract LaTeX content (only for non-PDF workflows)
                latex_content = line.strip()[2:].strip()  # Remove leading $$
                if latex_content.endswith('$$'):
                    # Single line formula: $$formula$$
                    latex_content = latex_content[:-2].strip()
                    self._add_math_formula(docx_doc, latex_content, display_mode=True)
                else:
                    # Multi-line formula: $$...$$ on separate lines
                    latex_lines = [latex_content]
                    i += 1
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if next_line.endswith('$$'):
                            latex_lines.append(next_line[:-2].strip())
                            break
                        latex_lines.append(next_line)
                        i += 1
                    latex_content = '\n'.join(latex_lines).strip()
                    self._add_math_formula(docx_doc, latex_content, display_mode=True)
            elif not use_layout_detection and line.strip().startswith('$$') and self.equation_format == "image":
                # Skip $$...$$ blocks when equation_format is "image" - formulas are already rendered as images
                logger.debug(LogModule.EXPORT, f"[DOCX-EQUATION] Skipping $$...$$ block (equation_format=image, formula should be rendered as image)")
                # Skip the $$ line
                i += 1
                # Skip multi-line formula if needed
                if not line.strip().endswith('$$'):
                    while i < len(lines):
                        if lines[i].strip().endswith('$$'):
                            i += 1
                            break
                        i += 1
                continue
            
            # Handle code blocks (```code```); ```tex / ```latex with a single $$...$$ body → OMML not code
            elif line.startswith('```'):
                fence_lang = parse_opening_markdown_fence_language(line)
                code_lines: List[str] = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                code_text = '\n'.join(code_lines)
                if (
                    fence_lang in TEX_LATEX_FENCE_LANGS
                    and self.equation_format != "image"
                ):
                    inner = extract_display_math_inner_from_tex_fence_body(code_text)
                    if inner is not None:
                        logger.info(
                            LogModule.EXPORT,
                            "[DOCX-EQUATION] Fenced tex/latex block rendered as display math (OMML), "
                            "not as code block",
                        )
                        self._add_math_formula(docx_doc, inner, display_mode=True)
                    elif code_text.strip():
                        self._add_code_block(docx_doc, code_text)
                elif code_text.strip():
                    self._add_code_block(docx_doc, code_text)

            # Handle list items (- or *)
            elif line.lstrip().startswith('-') or line.lstrip().startswith('*'):
                list_items = []
                list_indent = len(line) - len(line.lstrip())
                list_items.append(line.lstrip('-* ').strip())
                
                # Collect consecutive list items at same indent level
                i += 1
                while i < len(lines):
                    next_line = lines[i]
                    if not next_line.strip():
                        break
                    next_indent = len(next_line) - len(next_line.lstrip())
                    if (next_line.lstrip().startswith('-') or next_line.lstrip().startswith('*')) and next_indent == list_indent:
                        list_items.append(next_line.lstrip('-* ').strip())
                        i += 1
                    else:
                        break
                
                self._add_list(docx_doc, list_items)
                continue  # Already incremented i
            
            # Handle blockquotes (>)
            elif line.lstrip().startswith('>'):
                quote_lines = []
                quote_lines.append(line.lstrip('> ').strip())
                i += 1
                # Collect consecutive quote lines
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_lines.append(lines[i].lstrip('> ').strip())
                    i += 1
                self._add_quote(docx_doc, '\n'.join(quote_lines))
                continue  # Already incremented i
            
            # Handle horizontal rules (--- or ***)
            elif re.match(r'^[-*_]{3,}$', line.strip()):
                self._add_paragraph(docx_doc, '─' * 50, bold=True)
                i += 1
                continue
            
            # Handle markdown tables (only if table_body_format is "html")
            elif self.table_body_format == "html" and '|' in line and self._is_markdown_table_start(lines, i):
                table_start_idx = i
                table_data = self._parse_markdown_table(lines, i)
                if table_data:
                    self._add_table(docx_doc, table_data)
                    # Calculate how many lines were processed by _parse_markdown_table
                    # It processes: header row + separator row + data rows
                    # Count lines: start from table_start_idx, count until we hit non-table line
                    table_end_idx = table_start_idx
                    while table_end_idx < len(lines):
                        line_check = lines[table_end_idx].strip()
                        if not line_check or '|' not in line_check:
                            break
                        table_end_idx += 1
                    # Skip to after the table
                    i = table_end_idx
                    continue
            
            # Handle images: markdown syntax ![alt](src), HTML <img> tag, or layout placeholders <ph-...>
            # Support placeholder IDs with path characters (e.g., "mobi7/Images/image00044.jpeg")
            elif (
                re.match(r'^!\[.*?\]\(.*?\)$', line.strip())
                or re.match(r'^<img[^>]*>$', line.strip())
                or re.match(r'^<ph-[a-zA-Z0-9_./-]+>$', line.strip())
            ):
                logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Found image line at index {i}: {line.strip()[:100]}")
                self._add_image_from_markdown(docx_doc, line.strip())
            
            # Handle non-standard image formats: !Equation, !Table, or !filename.jpg
            # These may occur when markdown image syntax is partially corrupted
            elif re.match(r'^!(Equation|Table)$', line.strip(), re.IGNORECASE):
                # Try to find image by alt text
                alt_text = line.strip()[1:]  # Remove leading !
                # Search image_data_map for entries with matching alt text
                image_entry = None
                image_ref = None
                for key, value in self.image_data_map.items():
                    if isinstance(value, dict):
                        entry_alt = value.get("alt", "")
                        if alt_text.lower() in entry_alt.lower() or entry_alt.lower() in alt_text.lower():
                            image_entry = value
                            image_ref = key
                            break
                if image_entry and image_ref:
                    self._add_image_from_markdown(docx_doc, f"![{alt_text}]({image_ref})")
                else:
                    logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Could not find image for alt text: {alt_text}")
            elif re.match(r'^![a-zA-Z0-9_\-\.]+\.(jpg|jpeg|png|gif|webp)$', line.strip(), re.IGNORECASE):
                # Try to find image by filename (supports long filenames like hash-based names)
                filename = line.strip()[1:]  # Remove leading !
                # Search image_data_map for entries with matching filename
                image_entry = None
                image_ref = None
                
                # Try direct filename match first
                image_entry = self.image_data_map.get(filename)
                if image_entry:
                    image_ref = filename
                else:
                    # Try case-insensitive filename match
                    filename_lower = filename.lower()
                    for key, value in self.image_data_map.items():
                        if isinstance(key, str):
                            key_filename = key.split('/')[-1].split('\\')[-1]
                            if key_filename.lower() == filename_lower or key.lower() == filename_lower:
                                image_entry = value
                                image_ref = key
                                break
                
                if image_entry and image_ref:
                    self._add_image_from_markdown(docx_doc, f"![Image]({image_ref})")
                else:
                    logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Could not find image for filename: {filename}")
                    logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Available keys in image_data_map (first 20): {list(self.image_data_map.keys())[:20]}")
            
            # Handle regular paragraphs (with inline formatting)
            else:
                # Check if there are more lines for this paragraph
                para_text = line
                i += 1
                # For PDF workflow, preserve line breaks (code snippets, etc.); otherwise join with space
                line_sep = '\n' if use_layout_detection else ' '
                # Collect lines until empty line or special syntax
                while i < len(lines):
                    next_line = lines[i]
                    if not next_line.strip():
                        break
                    # Stop if we hit special syntax, images, or a potential table start
                    if (next_line.strip().startswith('#') or 
                        next_line.strip().startswith('```') or
                        next_line.strip().startswith('-') or
                        next_line.strip().startswith('*') or
                        next_line.strip().startswith('>') or
                        re.match(r'^!\[.*?\]\(.*?\)$', next_line.strip()) or
                        re.match(r'^<img[^>]*>$', next_line.strip()) or
                        (self.table_body_format == "html" and '|' in next_line and self._is_markdown_table_start(lines, i))):
                        break
                    para_text += line_sep + next_line.rstrip()
                    i += 1
                
                self._add_paragraph_with_formatting(docx_doc, para_text)
                continue  # Already incremented i
            
            i += 1
    
    def _add_heading(self, docx_doc: DocxDocument, text: str, level: int):
        """Add a heading paragraph."""
        heading = docx_doc.add_heading(text, level=min(level, 9))
        heading_format = heading.runs[0].font if heading.runs else heading.paragraph_format
        if heading.runs:
            heading.runs[0].font.name = self.font_name
            heading.runs[0].font.size = Pt(self.heading_sizes.get(level, self.font_size))
            self._set_run_east_asia_font(heading.runs[0])
    
    def _add_paragraph(self, docx_doc: DocxDocument, text: str, bold: bool = False, italic: bool = False):
        """Add a simple paragraph."""
        # Clean text to remove NULL bytes and control characters that are not XML compatible
        clean_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        if not clean_text:
            return  # Skip if text becomes empty after cleaning
        
        para = docx_doc.add_paragraph()
        self._add_runs_with_html_sup_sub(para, clean_text)
        for run in para.runs:
            run.bold = bold
            run.italic = italic
    
    def _add_paragraph_with_formatting(self, docx_doc: DocxDocument, text: str, skip_inline_math: bool = False):
        """Add a paragraph with inline markdown formatting.
        
        Args:
            docx_doc: DOCX document
            text: Text with inline formatting
            skip_inline_math: If True, skip inline math detection (for PDF workflow using layout detection)
        """
        # Clean text to remove NULL bytes and control characters that are not XML compatible
        # Keep only printable characters and common whitespace (space, tab, newline, carriage return)
        import re
        # Remove NULL bytes and other control characters except common whitespace
        # Keep: space (0x20), tab (0x09), newline (0x0A), carriage return (0x0D)
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        
        para = docx_doc.add_paragraph()
        
        # Check for images: support multiple markdown images on one line (e.g. side-by-side ![](a) ![](b))
        markdown_image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
        html_image_pattern = re.compile(r'<img[^>]+src=["\']([^"\']+)["\'][^>]*>')
        filename_image_pattern = re.compile(r'!([a-zA-Z0-9_\-\.]+\.(jpg|jpeg|png|gif|webp))', re.IGNORECASE)
        
        def _add_one_image_to_para(image_ref: str, alt_text: str) -> bool:
            """Add a single image to para (data URI or image_data_map). Returns True if processed (added or failed)."""
            if isinstance(image_ref, str) and image_ref.startswith("data:image/"):
                try:
                    import base64
                    import io
                    if "," not in image_ref:
                        return True
                    header, base64_data = image_ref.split(",", 1)
                    image_bytes = io.BytesIO(base64.b64decode(base64_data))
                    from PIL import Image
                    pil_image = Image.open(image_bytes)
                    original_width, original_height = pil_image.size
                    is_formula_or_table = "equation" in (alt_text or "").lower() or "table" in (alt_text or "").lower()
                    max_width_inches = 4.0 if is_formula_or_table else 6.0
                    max_width_pixels = max_width_inches * 96
                    if original_width > max_width_pixels:
                        width_inches = max_width_inches
                    else:
                        width_inches = original_width / 96.0
                    image_bytes.seek(0)
                    run = para.add_run()
                    run.add_picture(image_bytes, width=Inches(width_inches))
                    logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Successfully added inline image from data URI")
                    return True
                except Exception as e:
                    logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Failed to decode data URI: {e}")
                    return True
            filename = image_ref.split('/')[-1].split('\\')[-1]
            image_entry = self.image_data_map.get(image_ref)
            if not image_entry:
                image_entry = self.image_data_map.get(filename)
            if not image_entry:
                for key, value in self.image_data_map.items():
                    if isinstance(key, str):
                        key_filename = key.split('/')[-1].split('\\')[-1]
                        if key_filename.lower() == filename.lower():
                            image_entry = value
                            break
            if not image_entry:
                filename_lower = filename.lower()
                for key, value in self.image_data_map.items():
                    if isinstance(key, str) and key.lower() == filename_lower:
                        image_entry = value
                        break
            if image_entry and isinstance(image_entry, dict):
                data_uri = image_entry.get("data")
                if data_uri and data_uri.startswith("data:image/"):
                    try:
                        import base64
                        import io
                        image_bytes = io.BytesIO(base64.b64decode(data_uri.split(",", 1)[1]))
                        from PIL import Image
                        pil_image = Image.open(image_bytes)
                        original_width, original_height = pil_image.size
                        is_formula_or_table = "equation" in (alt_text or "").lower() or "table" in (alt_text or "").lower()
                        max_width_inches = 4.0 if is_formula_or_table else 6.0
                        max_width_pixels = max_width_inches * 96
                        if original_width > max_width_pixels:
                            width_inches = max_width_inches
                        else:
                            width_inches = original_width / 96.0
                        image_bytes.seek(0)
                        run = para.add_run()
                        run.add_picture(image_bytes, width=Inches(width_inches))
                        logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Successfully added inline image: {image_ref}")
                        return True
                    except Exception as e:
                        logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Failed to add inline image '{image_ref}': {e}")
                        return True
                logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Image data not found for reference: {image_ref}, alt_text: {alt_text}")
            else:
                logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Image data not found for reference: {image_ref}, alt_text: {alt_text}")
            return True
        
        # Process ALL markdown image refs so side-by-side images (e.g. ![](a) ![](b)) all get added
        markdown_matches = list(markdown_image_pattern.finditer(text))
        if markdown_matches:
            for m in markdown_matches:
                alt_text = m.group(1) or ""
                image_ref = m.group(2)
                _add_one_image_to_para(image_ref, alt_text)
            text = markdown_image_pattern.sub('', text).strip()
            if not text:
                return
            if '\n' in text:
                line_list = text.split('\n')
                for idx, single_line in enumerate(line_list):
                    self._add_formatted_runs_for_text(para, single_line, skip_inline_math)
                    if idx < len(line_list) - 1:
                        run = para.add_run()
                        run.add_break(WD_BREAK.LINE)
            else:
                self._add_formatted_runs_for_text(para, text, skip_inline_math)
            return
        
        # Single HTML or filename image (legacy path)
        html_image_match = html_image_pattern.search(text)
        filename_image_match = filename_image_pattern.search(text)
        if html_image_match:
            image_ref = html_image_match.group(1)
            alt_match = re.search(r'alt=["\']([^"\']*)["\']', text)
            alt_text = alt_match.group(1) if alt_match else ""
            _add_one_image_to_para(image_ref, alt_text)
            text = html_image_pattern.sub('', text).strip()
        elif filename_image_match:
            image_ref = filename_image_match.group(1)
            alt_text = ""
            _add_one_image_to_para(image_ref, alt_text)
            text = filename_image_pattern.sub('', text).strip()
        else:
            # No images, fall through to text-only handling below
            pass
        
        if html_image_match or filename_image_match:
            if not text.strip():
                return
            if '\n' in text:
                line_list = text.split('\n')
                for idx, single_line in enumerate(line_list):
                    self._add_formatted_runs_for_text(para, single_line, skip_inline_math)
                    if idx < len(line_list) - 1:
                        run = para.add_run()
                        run.add_break(WD_BREAK.LINE)
            else:
                self._add_formatted_runs_for_text(para, text, skip_inline_math)
            return
        
        # When text contains newlines: if we have inline math ($...$), parse whole paragraph
        # so that $...$ spanning multiple lines is matched; otherwise split by line.
        if '\n' in text:
            if not skip_inline_math and self._paragraph_needs_tex_processing(text):
                self._add_formatted_runs_for_text(para, text, skip_inline_math, allow_multiline_math=True)
            else:
                line_list = text.split('\n')
                for idx, single_line in enumerate(line_list):
                    self._add_formatted_runs_for_text(para, single_line, skip_inline_math)
                    if idx < len(line_list) - 1:
                        run = para.add_run()
                        run.add_break(WD_BREAK.LINE)
            return
        
        # Single line: parse inline formatting and add runs
        self._add_formatted_runs_for_text(para, text, skip_inline_math)
    
    def _add_formatted_runs_for_text(self, para, text: str, skip_inline_math: bool = False, allow_multiline_math: bool = False):
        """Add runs to an existing paragraph. When allow_multiline_math is True, $...$ may span newlines."""
        import re
        # Parse inline formatting: **bold**, *italic*, `code`, [link](url), LaTeX math
        if skip_inline_math:
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
        elif allow_multiline_math:
            parts = re.split(_TEX_SPLIT_PARTS_MULTILINE, text)
        else:
            parts = re.split(_TEX_SPLIT_PARTS_SINGLE_LINE_INLINE, text)
        
        for part in parts:
            if not part:
                continue
            
            # LaTeX display math: \[ ... \]
            if part.startswith("\\[") and part.endswith("\\]") and len(part) > 4:
                latex_content = part[2:-2].strip()
                self._add_inline_math(para, latex_content, display_mode=True)
            # LaTeX inline math: \( ... \)
            elif part.startswith("\\(") and part.endswith("\\)") and len(part) > 4:
                latex_content = part[2:-2].strip()
                self._add_inline_math(para, latex_content, display_mode=False)
            # LaTeX block math: $$formula$$
            elif part.startswith('$$') and part.endswith('$$') and len(part) > 4:
                latex_content = part[2:-2].strip()
                self._add_inline_math(para, latex_content, display_mode=True)
            
            # LaTeX inline math: $formula$ (not $$)
            elif part.startswith('$') and part.endswith('$') and len(part) > 2 and not part.startswith('$$'):
                latex_content = part[1:-1].strip()
                self._add_inline_math(para, latex_content, display_mode=False)
            
            # Bold: **text** (must check before single *)
            elif part.startswith('**') and part.endswith('**') and len(part) > 4:
                clean_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', part[2:-2])
                if clean_text:
                    run = para.add_run(clean_text)
                    run.bold = True
                    run.font.name = self.font_name
                    run.font.size = Pt(self.font_size)
                    self._set_run_east_asia_font(run)
            
            # Italic: *text* (but not **)
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                clean_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', part[1:-1])
                if clean_text:
                    run = para.add_run(clean_text)
                    run.italic = True
                    run.font.name = self.font_name
                    run.font.size = Pt(self.font_size)
                    self._set_run_east_asia_font(run)
            
            # Code: `code`
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                clean_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', part[1:-1])
                if clean_text:
                    run = para.add_run(clean_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(self.font_size - 1)
                    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)  # Dark red
            
            # Link: [text](url)
            elif part.startswith('[') and '](' in part:
                match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if match:
                    link_text = match.group(1)
                    link_url = match.group(2)
                    clean_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', link_text)
                    if clean_text:
                        run = para.add_run(clean_text)
                        run.font.name = self.font_name
                        run.font.size = Pt(self.font_size)
                        self._set_run_east_asia_font(run)
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue
                        run.font.underline = True
            
            # Regular text (may contain HTML <sup>/<sub> from rebuild/translator)
            else:
                clean_part = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', part)
                if clean_part:
                    if allow_multiline_math and '\n' in clean_part:
                        for j, chunk in enumerate(clean_part.split('\n')):
                            if j > 0:
                                run_br = para.add_run()
                                run_br.add_break(WD_BREAK.LINE)
                            if chunk:
                                self._add_runs_with_html_sup_sub(para, chunk)
                    else:
                        self._add_runs_with_html_sup_sub(para, clean_part)
    
    def _add_inline_math(self, para, latex: str, display_mode: bool = False):
        """Add LaTeX math to an existing paragraph as OMML (same normalization as _add_math_formula).

        Repair/snippets often still contain $$...$$ or \\tag {...}; passing raw strings to latex2mathml
        causes failures and LaTeX-as-text OMML fallback that looks like unconverted LaTeX in Word.
        """
        latex_clean, tag_text = self._normalize_formula_latex(latex)
        if not latex_clean:
            return
        latex_preview = latex_clean[:80] + ("..." if len(latex_clean) > 80 else "")
        try:
            logger.info(
                LogModule.EXPORT,
                f"[DOCX-EQUATION] Attempting LaTeX->OMML: display={display_mode}, preview={latex_preview!r}",
            )
            omml_xml = self._latex_to_omml(latex_clean)
            if omml_xml is not None:
                run = para.add_run()
                if display_mode:
                    math_para = OxmlElement("m:oMathPara")
                    math_para.set(qn("m:oMathParaPr"), "")
                    math_para.append(omml_xml)
                    run._element.append(math_para)
                else:
                    run._element.append(omml_xml)
                if tag_text:
                    run_tag = para.add_run(f" ({tag_text})")
                    run_tag.font.size = Pt(self.font_size)
                    run_tag.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                    self._set_run_east_asia_font(run_tag)
            else:
                logger.warning(
                    LogModule.EXPORT,
                    f"[DOCX-EQUATION] LaTeX->OMML returned None, showing as [latex]: preview={latex_preview!r}",
                )
                run = para.add_run(f"[{latex_clean}]")
                run.font.name = "Courier New"
                run.font.size = Pt(self.font_size - 1)
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                run.italic = True
        except Exception as e:
            logger.warning(
                LogModule.EXPORT,
                f"[DOCX-EQUATION] LaTeX->OMML exception, showing as [latex]: preview={latex_preview!r}, error={type(e).__name__}: {e}",
            )
            run = para.add_run(f"[{latex_clean}]")
            run.font.name = "Courier New"
            run.font.size = Pt(self.font_size - 1)
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            run.italic = True
    
    def _add_code_block(self, docx_doc: DocxDocument, code: str):
        """Add a code block paragraph."""
        # Clean text to remove NULL bytes and control characters that are not XML compatible
        clean_code = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', code)
        if not clean_code:
            return  # Skip if code becomes empty after cleaning
        
        para = docx_doc.add_paragraph()
        para.style = 'No Spacing'
        
        run = para.add_run(clean_code)
        run.font.name = 'Courier New'
        run.font.size = Pt(self.font_size - 1)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Dark green
        
        # Add background shading (light gray)
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.right_indent = Inches(0.25)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
    
    def _add_list(self, docx_doc: DocxDocument, items: List[str]):
        """Add a bulleted list."""
        for item in items:
            # Clean text to remove NULL bytes and control characters that are not XML compatible
            clean_item = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', item)
            if not clean_item:
                continue  # Skip empty items after cleaning
            
            para = docx_doc.add_paragraph(style='List Bullet')
            skip_tex = not self._paragraph_needs_tex_processing(clean_item)
            self._add_formatted_runs_for_text(para, clean_item, skip_inline_math=skip_tex)
            for run in para.runs:
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
                self._set_run_east_asia_font(run)
    
    def _add_quote(self, docx_doc: DocxDocument, text: str):
        """Add a blockquote paragraph."""
        # Clean text to remove NULL bytes and control characters that are not XML compatible
        # Keep only printable characters and common whitespace (space, tab, newline, carriage return)
        clean_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        if not clean_text:
            return  # Skip if text becomes empty after cleaning
        
        para = docx_doc.add_paragraph(style='Quote')
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        skip_tex = not self._paragraph_needs_tex_processing(clean_text)
        self._add_formatted_runs_for_text(para, clean_text, skip_inline_math=skip_tex)
        
        for run in para.runs:
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            self._set_run_east_asia_font(run)
            run.italic = True
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)  # Gray
    
    def _normalize_formula_latex(self, latex: str) -> Tuple[str, Optional[str]]:
        """Remove $$ delimiters and \\tag{...} from formula; return (cleaned_latex, tag_content).
        
        - Strips leading/trailing $$ and whitespace.
        - Extracts \\tag{...} / \\tag*{...} content (equation number) and removes the macro.
        - Returns (cleaned_latex, tag_text) so caller can render the number after the formula (OMML has no equation number).
        """
        if not latex or not isinstance(latex, str):
            return ((latex or "").strip(), None)
        s = latex.strip()
        while s.startswith("$$"):
            s = s[2:].strip()
        while s.endswith("$$"):
            s = s[:-2].strip()
        s = s.strip()
        # Extract \tag{...} and \tag*{...} content then remove the macro
        tag_match = re.search(r"\\tag\s*\*?\s*\{([^}]*)\}", s)
        tag_text = tag_match.group(1).strip() if tag_match else None
        s = re.sub(r"\\tag\s*\*?\s*\{[^}]*\}", "", s)

        # For block-level display formulas (used from _add_math_formula, _add_inline_math,
        # and tests), encourage proper lower-limit layout for sum-like operators by
        # inserting \limits when there is only a subscript and no superscript.
        def _maybe_add_limits(expr: str) -> str:
            # Do not touch if user already specified limits behaviour explicitly.
            if "\\limits" in expr or "\\nolimits" in expr:
                return expr
            # Only consider expressions that have a subscript but no superscript.
            if "_" not in expr or "^" in expr:
                return expr
            sum_like_ops = ("\\sum", "\\prod", "\\int", "\\bigcup", "\\bigcap")
            for op in sum_like_ops:
                idx = expr.find(op)
                if idx == -1:
                    continue
                # Require this operator to actually have a subscript following it.
                after = expr[idx + len(op) : idx + len(op) + 8]
                if "_" not in after:
                    continue
                # Insert \limits right after the operator name.
                return expr.replace(op, op + "\\limits", 1)
            return expr

        s = _maybe_add_limits(s)
        return (s.strip(), tag_text if (tag_text and tag_text.strip()) else None)

    # Math font commands whose argument is often split by OCR/repair (e.g. \mathbf{C R}).
    _SANITIZE_SYMBOL_FONT_CMDS: Tuple[str, ...] = (
        "mathbf",
        "mathrm",
        "mathcal",
        "mathbb",
        "mathit",
        "mathsf",
        "mathtt",
        "mathscr",
        "mathfrak",
    )
    @staticmethod
    def _sanitize_latex_for_latex2mathml(s: str) -> str:
        """Fix spacing/OCR patterns that often break latex2mathml (repair outputs, PDF OCR).

        Applied as a second-pass attempt in _latex_to_omml when the original string fails.
        """
        if not s:
            return s
        t = s
        # Merge spaced multi-letter tokens inside font wrappers: \mathbf{C R} -> \mathbf{CR}
        for cmd in MD2DOCXExporter._SANITIZE_SYMBOL_FONT_CMDS:

            def _merge_spaced_symbol_arg(m: re.Match, c: str = cmd) -> str:
                inner = m.group(1)
                return f"\\{c}{{{''.join(inner.split())}}}"

            t = re.sub(
                rf"\\{cmd}\{{\s*([A-Za-z0-9](?:\s+[A-Za-z0-9])+)\s*\}}",
                _merge_spaced_symbol_arg,
                t,
            )
        # Tighten single-token braces: \mathrm { c } -> \mathrm{c}
        for cmd in MD2DOCXExporter._SANITIZE_SYMBOL_FONT_CMDS:

            def _tight_single(m: re.Match, c: str = cmd) -> str:
                return f"\\{c}{{{m.group(1)}}}"

            t = re.sub(rf"\\{cmd}\{{\s*([A-Za-z0-9])\s*\}}", _tight_single, t)
        # OCR: \text { w h e r e } -> \text{where} (spaced Latin letters only)
        t = re.sub(
            r"\\text\s*\{\s*((?:[A-Za-z]\s+)+[A-Za-z])\s*\}",
            lambda m: "\\text{" + "".join(m.group(1).split()) + "}",
            t,
        )
        # Common accents: \bar {x} -> \bar{x}
        for acc in (
            "bar",
            "hat",
            "tilde",
            "acute",
            "grave",
            "vec",
            "dot",
            "ddot",
            "breve",
            "check",
        ):
            t = re.sub(rf"\\{acc}\s*\{{", rf"\\{acc}{{", t)
        # \sum_ { ... } -> \sum_{...} (and prod/int/common operators)
        t = re.sub(
            r"(\\(?:sum|prod|int|oint|iint|iiint|iiiint|bigcup|bigcap|bigoplus|bigotimes|bigodot))"
            r"\s*_\s*\{",
            r"\1_{",
            t,
        )
        # Superscript / subscript open brace: ^{ T} -> ^{T}, _{ i} -> _{i}
        t = re.sub(r"\^\s*\{", "^{", t)
        t = re.sub(r"_\s*\{", "_{", t)
        t = re.sub(r"\\left\s+", r"\\left", t)
        t = re.sub(r"\\right\s+", r"\\right", t)
        if "\n" not in t:
            t = re.sub(r"[ \t]+", " ", t).strip()
        return t

    @staticmethod
    def _latex_fragment_for_log(s: str, max_chars: int = 12000) -> str:
        """Full LaTeX for diagnostics (OMML fallback / latex2mathml failures); avoid megabyte logs."""
        if not s:
            return ""
        if len(s) <= max_chars:
            return s
        return s[:max_chars] + f"\n... [truncated, total_chars={len(s)}]"

    def _add_math_formula(self, docx_doc: DocxDocument, latex: str, display_mode: bool = False):
        """Add a LaTeX formula as Word math equation using OMML.
        
        Args:
            docx_doc: The DOCX document
            latex: LaTeX formula content (without $$ delimiters)
            display_mode: Whether this is a display formula (block) or inline
        """
        latex_clean, tag_text = self._normalize_formula_latex(latex)
        if not latex_clean:
            return
        para = docx_doc.add_paragraph()
        
        # Center align for display mode formulas
        if display_mode:
            para.alignment = 1  # Center alignment
        
        # Try to insert as Word math equation using OMML
        try:
            omml_xml = self._latex_to_omml(latex_clean)
            if omml_xml is not None:
                run = para.add_run()
                if display_mode:
                    math_para = OxmlElement('m:oMathPara')
                    math_para.set(qn('m:oMathParaPr'), '')
                    math_para.append(omml_xml)
                    run._element.append(math_para)
                else:
                    run._element.append(omml_xml)
                # OMML has no equation number; add \tag{...} content as text after the formula
                if tag_text:
                    run_tag = para.add_run(f" ({tag_text})")
                    run_tag.font.size = Pt(self.font_size)
                    run_tag.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            else:
                self._add_latex_fallback(para, latex_clean)
        except Exception:
            self._add_latex_fallback(para, latex_clean)
    
    def _add_latex_fallback(self, para, latex: str):
        """Add LaTeX formula as formatted text (fallback when OMML conversion fails)."""
        run = para.add_run(f'[{latex}]')
        run.font.name = 'Courier New'
        run.font.size = Pt(self.font_size - 1)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Dark green
        run.italic = True

    @staticmethod
    def _paragraph_needs_tex_processing(text: str) -> bool:
        """True if paragraph may contain LaTeX math (must not use skip_inline_math)."""
        if not text:
            return False
        if "$" in text:
            return True
        if "\\[" in text or "\\]" in text:
            return True
        if "\\(" in text or "\\)" in text:
            return True
        if "\\begin{" in text:
            return True
        return False

    def _find_markdown_images_in_line(self, line: str) -> List[Tuple[int, int, str, str]]:
        """Find all markdown image refs in a line. Uses scan for data URI (avoids regex on very long base64).
        Returns list of (start, end, alt_text, image_src) spanning the full ![alt](src) segment."""
        result: List[Tuple[int, int, str, str]] = []
        pos = 0
        data_uri_marker = "](data:image/"
        while True:
            idx = line.find(data_uri_marker, pos)
            if idx == -1:
                break
            alt_start = line.rfind("![", pos, idx)
            if alt_start == -1 or alt_start < pos:
                pos = idx + 1
                continue
            url_open = idx + 1
            url_close = line.find(")", url_open + 1)
            if url_close == -1:
                pos = idx + 1
                continue
            alt_text = line[alt_start + 2:idx]
            image_src = line[url_open + 1:url_close]
            if image_src.startswith("data:image/") and "," in image_src:
                result.append((alt_start, url_close + 1, alt_text, image_src))
            pos = url_close + 1
        if result:
            return result
        pat = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
        for m in pat.finditer(line):
            result.append((m.start(), m.end(), (m.group(1) or ""), m.group(2)))
        return result

    @staticmethod
    def _count_markdown_display_math_blocks(lines: List[str]) -> int:
        """Count $$ display math blocks using the same open/close rules as _markdown_to_docx_with_layout.

        Each block starts at a line whose stripped content begins with $$. Multi-line blocks count once.
        """
        n = 0
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line:
                i += 1
                continue
            if line.strip().startswith("$$"):
                n += 1
                latex_content = line.strip()[2:].strip()
                if not latex_content.endswith("$$"):
                    i += 1
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if next_line.endswith("$$"):
                            break
                        i += 1
            i += 1
        return n

    def _markdown_to_docx_with_layout(self, md_content: str, docx_doc: DocxDocument):
        """Convert markdown to DOCX using layout_document for formula detection.
        
        For PDF workflow: formulas are detected by layout block type (interline_equation)
        rather than markdown code patterns ($$...$$).
        """
        # Extract all interline_equation blocks from layout_document
        equation_blocks: List[Dict[str, Optional[str]]] = []
        if self.layout_document:
            from utils.document_rebuild.table_layout_utils import _extract_equation_from_layout_block
            for block in self.layout_document.iter_blocks():
                if block.type == "interline_equation":
                    equation_content, equation_image_path = _extract_equation_from_layout_block(block)
                    if equation_content or equation_image_path:
                        equation_blocks.append(
                            {"content": equation_content, "image_path": equation_image_path}
                        )
        
        # Process markdown content, replacing $$...$$ with formulas from layout
        # Split content into lines for processing
        lines = md_content.split('\n')

        n_md_blocks = self._count_markdown_display_math_blocks(lines)
        n_layout_blocks = len(equation_blocks)
        if (
            n_layout_blocks > 0
            and n_md_blocks != n_layout_blocks
            and self.equation_format != "image"
        ):
            logger.warning(
                LogModule.EXPORT,
                "[DOCX-EQUATION] Display math block count mismatch: "
                f"markdown_display_math_blocks={n_md_blocks} layout_interline_equation={n_layout_blocks}. "
                "Disabling layout→markdown formula substitution (use LaTeX from each $$ block) "
                "to avoid cumulative drift after segment repairs or extra $$ insertions.",
            )
            equation_blocks = []
        
        i = 0
        equation_index = 0  # Track which equation we're at
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Handle LaTeX math blocks ($$...$$) - replace with layout-based formula
            # If equation_format is "image", skip $$...$$ blocks as formulas are already rendered as images in markdown
            if line.strip().startswith('$$') and self.equation_format != "image":
                # Extract LaTeX content from markdown (for reference)
                latex_content = line.strip()[2:].strip()
                if latex_content.endswith('$$'):
                    # Single line formula: $$formula$$
                    latex_content = latex_content[:-2].strip()
                else:
                    # Multi-line formula: $$...$$ on separate lines
                    latex_lines = [latex_content]
                    i += 1
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if next_line.endswith('$$'):
                            latex_lines.append(next_line[:-2].strip())
                            break
                        latex_lines.append(next_line)
                        i += 1
                    latex_content = '\n'.join(latex_lines).strip()
                
                # Use equation from layout_document if available
                if equation_index < len(equation_blocks):
                    formula_content = equation_blocks[equation_index].get("content") or latex_content
                    equation_index += 1
                else:
                    # Fallback to markdown content if layout doesn't have it
                    formula_content = latex_content
                
                # Render as LaTeX formula
                self._add_math_formula(docx_doc, formula_content, display_mode=True)
            elif line.strip().startswith('$$') and self.equation_format == "image":
                # Render equation as image from layout block (export setting: equation_format=image)
                if equation_index < len(equation_blocks):
                    eq_entry = equation_blocks[equation_index]
                    equation_index += 1
                    image_path = eq_entry.get("image_path")
                    if image_path:
                        filename = image_path.replace("\\", "/").split("/")[-1]
                        inserted = self._add_image_from_markdown(docx_doc, f"![Equation]({filename})")
                        if inserted:
                            logger.info(
                                LogModule.EXPORT,
                                f"[DOCX-EQUATION] Inserted equation image from layout block: {filename}",
                            )
                        else:
                            logger.warning(
                                LogModule.EXPORT,
                                f"[DOCX-EQUATION] Failed to insert equation image from layout block: {filename}",
                            )
                    else:
                        logger.warning(
                            LogModule.EXPORT,
                            "[DOCX-EQUATION] equation_format=image but layout interline_equation has no image_path; skipping formula",
                        )
                # Skip the $$ line
                i += 1
                # Skip multi-line formula if needed
                if not line.strip().endswith('$$'):
                    while i < len(lines):
                        if lines[i].strip().endswith('$$'):
                            i += 1
                            break
                        i += 1
                continue
            
            # Handle headings (### Heading)
            elif line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                if heading_text:
                    self._add_heading(docx_doc, heading_text, level)
                i += 1
                continue
            
            # Handle code blocks (```code```); ```tex / ```latex with a single $$...$$ body → OMML not code
            elif line.startswith('```'):
                fence_lang = parse_opening_markdown_fence_language(line)
                code_lines: List[str] = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                code_text = '\n'.join(code_lines)
                if (
                    fence_lang in TEX_LATEX_FENCE_LANGS
                    and self.equation_format != "image"
                ):
                    inner = extract_display_math_inner_from_tex_fence_body(code_text)
                    if inner is not None:
                        logger.info(
                            LogModule.EXPORT,
                            "[DOCX-EQUATION] Fenced tex/latex block rendered as display math (OMML), "
                            "not as code block",
                        )
                        self._add_math_formula(docx_doc, inner, display_mode=True)
                    elif code_text.strip():
                        self._add_code_block(docx_doc, code_text)
                elif code_text.strip():
                    self._add_code_block(docx_doc, code_text)
                if i < len(lines):
                    i += 1  # skip closing ```
                continue
            
            # Handle list-style line that contains any markdown image (data URI or file path, e.g. "-续 ![](./images/xxx.jpg)")
            # Must run before generic list block so we parse images instead of treating whole line as list text
            _list_style = line.lstrip().startswith('-') or line.lstrip().startswith('*')
            _has_image = '](data:image/' in line or ('![' in line and '](' in line)
            if _list_style and _has_image:
                content = line.lstrip('-* \t').strip()
                img_matches = self._find_markdown_images_in_line(content)
                if img_matches:
                    para = docx_doc.add_paragraph()
                    bullet_run = para.add_run("• ")
                    bullet_run.bold = False
                    self._set_run_east_asia_font(bullet_run)
                    pos = 0
                    for (start, end, _alt, _src) in img_matches:
                        if start > pos:
                            text_segment = content[pos:start].strip()
                            if text_segment:
                                self._add_formatted_runs_for_text(para, text_segment, skip_inline_math=True)
                        self._add_image_from_markdown(docx_doc, f"![{_alt}]({_src})", paragraph=para)
                        pos = end
                    if pos < len(content):
                        text_segment = content[pos:].strip()
                        if text_segment:
                            self._add_formatted_runs_for_text(para, text_segment, skip_inline_math=True)
                    i += 1
                    continue

            # Handle list items (- or *)
            elif line.lstrip().startswith('-') or line.lstrip().startswith('*'):
                list_items = []
                list_indent = len(line) - len(line.lstrip())
                list_items.append(line.lstrip('-* ').strip())
                
                i += 1
                while i < len(lines):
                    next_line = lines[i]
                    if not next_line.strip():
                        break
                    next_indent = len(next_line) - len(next_line.lstrip())
                    if (next_line.lstrip().startswith('-') or next_line.lstrip().startswith('*')) and next_indent == list_indent:
                        list_items.append(next_line.lstrip('-* ').strip())
                        i += 1
                    else:
                        break
                
                self._add_list(docx_doc, list_items)
                continue
            
            # Handle images: one or more ![alt](src) on this line; use scan for data URI (long base64 breaks regex)
            # Also HTML <img> or layout placeholder <ph-...>
            _img_matches = self._find_markdown_images_in_line(line) if "](data:image/" in line else []
            if not _img_matches:
                _img_pat = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
                _img_matches = [(m.start(), m.end(), (m.group(1) or ""), m.group(2)) for m in _img_pat.finditer(line)]
            if _img_matches:
                para = docx_doc.add_paragraph()
                pos = 0
                for (start, end, _alt, _src) in _img_matches:
                    if start > pos:
                        text_segment = line[pos:start].strip()
                        if text_segment:
                            self._add_formatted_runs_for_text(para, text_segment, skip_inline_math=True)
                    self._add_image_from_markdown(docx_doc, f"![{_alt}]({_src})", paragraph=para)
                    pos = end
                if pos < len(line):
                    text_segment = line[pos:].strip()
                    if text_segment:
                        self._add_formatted_runs_for_text(para, text_segment, skip_inline_math=True)
                i += 1
                continue
            if re.match(r'^<img[^>]*>$', line.strip()) or re.match(r'^<ph-[a-zA-Z0-9_./-]+>$', line.strip()):
                logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Found image line at index {i}: {line.strip()[:100]}")
                self._add_image_from_markdown(docx_doc, line.strip())
                i += 1
                continue
            
            # Handle non-standard image formats: !Equation, !Table, or !filename.jpg
            elif re.match(r'^!(Equation|Table)$', line.strip(), re.IGNORECASE):
                # Try to find image by alt text
                alt_text = line.strip()[1:]  # Remove leading !
                # Search image_data_map for entries with matching alt text
                image_entry = None
                image_ref = None
                for key, value in self.image_data_map.items():
                    if isinstance(value, dict):
                        entry_alt = value.get("alt", "")
                        if alt_text.lower() in entry_alt.lower() or entry_alt.lower() in alt_text.lower():
                            image_entry = value
                            image_ref = key
                            break
                if image_entry and image_ref:
                    self._add_image_from_markdown(docx_doc, f"![{alt_text}]({image_ref})")
                else:
                    logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Could not find image for alt text: {alt_text}")
                i += 1
                continue
            
            # Handle blockquotes (>)
            elif line.lstrip().startswith('>'):
                quote_lines = []
                quote_lines.append(line.lstrip('> ').strip())
                i += 1
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_lines.append(lines[i].lstrip('> ').strip())
                    i += 1
                self._add_quote(docx_doc, '\n'.join(quote_lines))
                continue
            
            # Handle horizontal rules (--- or ***)
            elif re.match(r'^[-*_]{3,}$', line.strip()):
                self._add_paragraph(docx_doc, '─' * 50, bold=True)
                i += 1
                continue
            
            # Handle markdown tables (only if table_body_format is "html")
            elif self.table_body_format == "html" and '|' in line and self._is_markdown_table_start(lines, i):
                table_start_idx = i
                table_data = self._parse_markdown_table(lines, i)
                if table_data:
                    self._add_table(docx_doc, table_data)
                    # Calculate how many lines were processed by _parse_markdown_table
                    # It processes: header row + separator row + data rows
                    # Count lines: start from table_start_idx, count until we hit non-table line
                    table_end_idx = table_start_idx
                    while table_end_idx < len(lines):
                        line_check = lines[table_end_idx].strip()
                        if not line_check or '|' not in line_check:
                            break
                        table_end_idx += 1
                    # Skip to after the table
                    i = table_end_idx
                    continue
            
            # Handle regular paragraphs (with inline formatting, but skip $...$ formulas)
            # Layout path is PDF workflow: preserve line breaks in paragraphs (code snippets, etc.)
            else:
                para_text = line
                i += 1
                while i < len(lines):
                    next_line = lines[i]
                    if not next_line.strip():
                        break
                    # Stop if we hit special syntax or a potential table start
                    if (next_line.strip().startswith('#') or 
                        next_line.strip().startswith('```') or
                        next_line.strip().startswith('$$') or
                        next_line.strip().startswith('-') or
                        next_line.strip().startswith('*') or
                        next_line.strip().startswith('>') or
                        (self.table_body_format == "html" and '|' in next_line and self._is_markdown_table_start(lines, i))):
                        break
                    para_text += '\n' + next_line.rstrip()
                    i += 1
                
                # Strip stray $$ delimiters (formula block leftover: next line or end of line)
                para_text = re.sub(r"\n\$\$+\s*$", "", para_text)
                para_text = re.sub(r"\s*\$\$+\s*$", "", para_text)
                para_text = re.sub(r"^\s*\$\$+\s*\n?", "", para_text)
                para_text = para_text.strip()
                if not para_text or para_text == "$$":
                    continue
                # When rebuilt MD has $...$ (mixed formula from markdown_rebuild), render inline math
                skip_inline_math = not self._paragraph_needs_tex_processing(para_text)
                self._add_paragraph_with_formatting(docx_doc, para_text, skip_inline_math=skip_inline_math)
                continue
            
            i += 1
    
    def _latex_to_omml(self, latex: str):
        """Convert LaTeX formula to OMML (Office Math Markup Language) for Word.
        
        Uses latex2mathml (LaTeX -> MathML) and mathml2omml (MathML -> OMML) when
        available so that Word renders native editable equations. Falls back to
        embedding LaTeX as plain text in m:t if conversion fails or libs missing.
        
        Args:
            latex: LaTeX formula string (without $$ delimiters)
            
        Returns:
            OMML XML element (lxml/compatible with docx.oxml), or None if conversion fails
        """
        latex_clean = (latex or "").strip()
        if not latex_clean:
            return None
        latex_preview = latex_clean[:80] + "..." if len(latex_clean) > 80 else latex_clean
        fail_reason = None
        mathml_str = None
        # LaTeX -> MathML: prefer sanitized when it differs (repair/OCR spacing breaks latex2mathml often)
        sanitized = self._sanitize_latex_for_latex2mathml(latex_clean)
        attempts: List[Tuple[str, str]] = []
        if sanitized and sanitized != latex_clean:
            attempts.append((sanitized, "sanitized"))
        attempts.append((latex_clean, "original"))

        import latex2mathml.converter

        for attempt_tex, attempt_label in attempts:
            try:
                m = latex2mathml.converter.convert(attempt_tex)
                if m and m.strip():
                    mathml_str = m
                    if attempt_label == "sanitized" and sanitized != latex_clean:
                        logger.info(
                            LogModule.EXPORT,
                            f"[DOCX-EQUATION] latex2mathml used sanitized input, preview={latex_preview!r}",
                        )
                    break
                fail_reason = "latex2mathml returned empty"
                logger.warning(
                    LogModule.EXPORT,
                    f"[DOCX-EQUATION] {fail_reason}, attempt={attempt_label}, preview={latex_preview!r}",
                )
            except Exception as e:
                fail_reason = f"latex2mathml failed: {type(e).__name__}: {e}"
                logger.warning(
                    LogModule.EXPORT,
                    f"[DOCX-EQUATION] {fail_reason}, attempt={attempt_label}, preview={latex_preview!r}",
                )
        def _parse_omml_to_element(omml_str: str):
            """Parse OMML string to lxml element; return first child or None."""
            from lxml import etree
            OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
            wrapped = f'<w xmlns:m="{OMML_NS}">{omml_str.strip()}</w>'
            root = etree.fromstring(wrapped.encode("utf-8"))
            if root is not None and len(root) > 0:
                return root[0]
            return None

        if mathml_str:
            omml_str = None
            try:
                import mathml2omml
                omml_str = mathml2omml.convert(mathml_str)
                if not omml_str or not omml_str.strip():
                    fail_reason = "mathml2omml returned empty"
                    logger.warning(
                        LogModule.EXPORT,
                        f"[DOCX-EQUATION] {fail_reason}, latex_preview={latex_preview!r}",
                    )
                    omml_str = None
            except Exception as e:
                fail_reason = f"mathml2omml failed: {type(e).__name__}: {e}"
                logger.warning(
                    LogModule.EXPORT,
                    f"[DOCX-EQUATION] {fail_reason}, latex_preview={latex_preview!r}",
                )
                omml_str = None
            if omml_str:
                # Work around known bug in some mathml2omml/mathml2omml-as builds where
                # <m:groupChrPr> is closed with </m:groupChr>, causing XMLSyntaxError:
                # "Opening and ending tag mismatch: groupChrPr ... and groupChr".
                # Fix by rewriting the first closing tag after groupChrPr to </m:groupChrPr>.
                import re as _re
                omml_str = _re.sub(
                    r'(<m:groupChrPr>[\s\S]*?<m:pos[^>]*/>)\s*</m:groupChr>',
                    r'\1</m:groupChrPr>',
                    omml_str,
                )
                try:
                    elt = _parse_omml_to_element(omml_str)
                    if elt is not None:
                        logger.debug(LogModule.EXPORT, "[DOCX-EQUATION] LaTeX -> OMML conversion succeeded")
                        return elt
                    fail_reason = "OMML root empty after parse"
                    logger.warning(
                        LogModule.EXPORT,
                        f"[DOCX-EQUATION] {fail_reason}, latex_preview={latex_preview!r}",
                    )
                except Exception as e:
                    fail_reason = f"OMML parse failed: {type(e).__name__}: {e}"
                    logger.warning(
                        LogModule.EXPORT,
                        f"[DOCX-EQUATION] {fail_reason}, latex_preview={latex_preview!r}",
                    )
            # When mathml2omml output fails to parse (e.g. groupChrPr/groupChr tag mismatch), try mathml2omml-as
            try:
                import mathml2omml_as
                omml_str_as = mathml2omml_as.convert(mathml_str)
                if omml_str_as and omml_str_as.strip():
                    elt = _parse_omml_to_element(omml_str_as)
                    if elt is not None:
                        logger.debug(LogModule.EXPORT, "[DOCX-EQUATION] LaTeX -> OMML succeeded via mathml2omml_as")
                        return elt
            except ImportError:
                pass
            except Exception as e:
                logger.debug(LogModule.EXPORT, f"[DOCX-EQUATION] mathml2omml_as not used: {e}")
        # Fallback: embed LaTeX as plain text in m:t (Word shows formula box with raw LaTeX-like glyphs)
        san_extra = ""
        if sanitized and sanitized != latex_clean:
            san_extra = (
                "\n[DOCX-EQUATION] OMML_FALLBACK sanitized_attempt_was_chars="
                f"{len(sanitized)}\n{self._latex_fragment_for_log(sanitized)}"
            )
        logger.warning(
            LogModule.EXPORT,
            "[DOCX-EQUATION] OMML_FALLBACK_LATEX_AS_TEXT reason=%r latex_chars=%s\nfull_latex=\n%s%s",
            fail_reason or "unknown",
            len(latex_clean),
            self._latex_fragment_for_log(latex_clean),
            san_extra,
        )
        try:
            omath = OxmlElement("m:oMath")
            run = OxmlElement("m:r")
            run_props = OxmlElement("m:rPr")
            run.append(run_props)
            text_el = OxmlElement("m:t")
            text_el.text = " ".join(latex_clean.split())
            run.append(text_el)
            omath.append(run)
            return omath
        except Exception:
            return None
    
    def _is_markdown_table_start(self, lines: List[str], start_idx: int) -> bool:
        """Check if a line is the start of a markdown table.
        
        A markdown table starts with a header row containing | separators,
        followed by a separator row with |---| or similar.
        """
        if start_idx >= len(lines):
            return False
        
        # Check if current line looks like a table row (contains |)
        if '|' not in lines[start_idx]:
            return False
        
        # Check if next line is a separator row (contains |---| or |:---:|)
        if start_idx + 1 >= len(lines):
            return False
        
        next_line = lines[start_idx + 1].strip()
        # Separator row should contain | and - or :
        if '|' not in next_line:
            return False
        
        # Check if separator row contains dashes or colons
        if not re.search(r'[-:]+', next_line):
            return False
        
        return True
    
    def _parse_markdown_table(self, lines: List[str], start_idx: int) -> List[List[str]]:
        """Parse a markdown table into a list of rows, each row is a list of cells.
        
        Returns:
            List of rows, each row is a list of cell strings.
            Returns empty list if parsing fails.
        """
        table_rows = []
        i = start_idx
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Stop if we hit an empty line or a line that doesn't look like a table row
            if not line or '|' not in line:
                break
            
            # Skip separator rows (|---|---| or |:---:|)
            if re.match(r'^\|[\s:|-]+\|$', line):
                i += 1
                continue
            
            # Parse table row: split by | and clean up cells
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty cells at start/end (markdown tables often have leading/trailing |)
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            
            if cells:
                table_rows.append(cells)
            
            i += 1
        
        return table_rows if table_rows else []
    
    def _add_image_from_markdown(self, docx_doc: DocxDocument, image_line: str, paragraph=None) -> bool:
        """Add an image from markdown syntax or HTML img tag.
        
        Args:
            docx_doc: DOCX document
            image_line: Markdown image syntax ![alt](src) or HTML <img> tag
            paragraph: Optional existing paragraph to add the image run to (for side-by-side images).
                       If None, a new paragraph is created.

        Returns:
            True if an image was inserted, False otherwise.
        """
        import re
        import io
        import base64
        from docx.shared import Inches
        from PIL import Image
        
        logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Processing image line: {image_line[:100]}")
        
        # Extract image reference from markdown, HTML, or placeholder
        image_ref = None
        alt_text = ""

        # 1) Placeholder syntax from layout markdown: <ph-layoutimg0> or MOBI: <ph-mobi7/Images/image00044.jpeg>
        # Support placeholder IDs with path characters
        placeholder_match = re.match(r'<ph-([a-zA-Z0-9_./-]+)>', image_line.strip())
        if placeholder_match:
            image_ref = placeholder_match.group(1)
            alt_text = ""  # alt_text will be inferred from image_data_map if needed
        else:
            # 2) Markdown syntax: ![alt](src)
            markdown_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', image_line)
            if markdown_match:
                alt_text = markdown_match.group(1) or ""
                image_ref = markdown_match.group(2)
            else:
                # 3) HTML img tag: <img src="..." alt="..." />
                html_match = re.search(r'<img[^>]+src=["\']([^"\']+)["\']', image_line)
                if html_match:
                    image_ref = html_match.group(1)
                    alt_match = re.search(r'alt=["\']([^"\']*)["\']', image_line)
                    if alt_match:
                        alt_text = alt_match.group(1)

        if not image_ref:
            logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Could not extract image reference from: {image_line}")
            return False

        # Shortcut: if the reference is already a data URI, decode directly and insert
        if isinstance(image_ref, str) and image_ref.startswith("data:image/"):
            try:
                if "," not in image_ref:
                    logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Invalid data URI (no comma): {image_ref[:50]}...")
                    return False
                header, base64_data = image_ref.split(",", 1)
                mime_type = header.split(";")[0].split(":")[1] if ":" in header else "image/png"
                image_bytes = io.BytesIO(base64.b64decode(base64_data))
                pil_image = Image.open(image_bytes)
                original_width, original_height = pil_image.size
                
                # Try to get actual DPI from image metadata, fallback to 96 DPI
                dpi = 200.0  # Default DPI assumption
                if hasattr(pil_image, 'info'):
                    if 'dpi' in pil_image.info and isinstance(pil_image.info['dpi'], tuple):
                        dpi = float(pil_image.info['dpi'][0]) if pil_image.info['dpi'][0] > 0 else 200.0
                    elif 'resolution' in pil_image.info and isinstance(pil_image.info['resolution'], tuple):
                        dpi = float(pil_image.info['resolution'][0]) if pil_image.info['resolution'][0] > 0 else 200.0
                
                if original_width <= 0 or original_height <= 0:
                    logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Invalid image size from data URI: {original_width}x{original_height}")
                    return False
                
                # Convert pixels to inches using actual DPI (or 96 DPI fallback)
                width_inches = original_width / dpi
                height_inches = original_height / dpi
                
                logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] [DATA-URI] Image DPI: {dpi}, size: {original_width}x{original_height} px -> {width_inches:.2f}\" x {height_inches:.2f}\"")
                image_bytes.seek(0)
                para = paragraph if paragraph is not None else docx_doc.add_paragraph()
                run = para.add_run()
                run.add_picture(image_bytes, width=Inches(width_inches), height=Inches(height_inches))
                logger.info(LogModule.EXPORT, f"[DOCX-IMAGE] Successfully inserted image from data URI: alt_text={alt_text}, size={width_inches:.2f}\" x {height_inches:.2f}\"")
                return True
            except Exception as e:
                logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Failed to decode data URI: {e}")
                return False
        
        # Look up image data from image_data_map
        # Try direct match first
        image_entry = self.image_data_map.get(image_ref)
        
        # If not found, try filename matching (for file paths)
        if not image_entry:
            # Extract filename from path (handle data URIs, relative paths, etc.)
            filename = image_ref.split('/')[-1].split('\\')[-1]
            if filename:
                # Try exact filename match
                image_entry = self.image_data_map.get(filename)
                # Try case-insensitive match by filename
                if not image_entry:
                    filename_lower = filename.lower()
                    for key, value in self.image_data_map.items():
                        if isinstance(key, str):
                            key_filename = key.split('/')[-1].split('\\')[-1]
                            if key_filename.lower() == filename_lower:
                                image_entry = value
                                break
                # Try case-insensitive match by full key
                if not image_entry:
                    filename_lower = filename.lower()
                    for key, value in self.image_data_map.items():
                        if isinstance(key, str) and key.lower() == filename_lower:
                            image_entry = value
                            break

        # If still not found, try matching by alt_text (many markdown images keep the original filename in alt)
        if not image_entry and alt_text:
            # direct key
            image_entry = self.image_data_map.get(alt_text)
            if not image_entry:
                alt_lower = alt_text.lower()
                for key, value in self.image_data_map.items():
                    if isinstance(key, str) and key.lower() == alt_lower:
                        image_entry = value
                        logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Found image by alt_text key match: {key}")
                        break
            # match against entries whose "alt" field equals alt_text (case-insensitive)
            if not image_entry:
                alt_lower = alt_text.lower()
                for key, value in self.image_data_map.items():
                    if isinstance(value, dict):
                        entry_alt = value.get("alt", "")
                        if isinstance(entry_alt, str) and entry_alt.lower() == alt_lower:
                            image_entry = value
                            logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Found image by alt_text field match: {key} (alt={entry_alt})")
                            break
            # Also try matching placeholder_id in alt_text (e.g., "image-mobi7/Images/image00044.jpeg")
            if not image_entry and alt_text.startswith("image-"):
                placeholder_id = alt_text[6:]  # Remove "image-" prefix
                image_entry = self.image_data_map.get(placeholder_id)
                if image_entry:
                    logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Found image by placeholder_id from alt_text: {placeholder_id}")

        if not image_entry:
            logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Image data not found for reference: {image_ref}, alt_text: {alt_text}")
            logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Available keys in image_data_map (first 20): {list(self.image_data_map.keys())[:20]}")
            logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Total keys in image_data_map: {len(self.image_data_map)}")
            # Log all keys that contain the filename (for debugging)
            filename = image_ref.split('/')[-1].split('\\')[-1]
            matching_keys = [k for k in self.image_data_map.keys() if isinstance(k, str) and filename.lower() in k.lower()]
            if matching_keys:
                logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Keys containing '{filename}': {matching_keys[:10]}")
            # Also try to find by placeholder_id if alt_text contains it
            if alt_text.startswith("image-"):
                placeholder_id = alt_text[6:]
                logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Trying placeholder_id from alt_text: {placeholder_id}")
                if placeholder_id in self.image_data_map:
                    image_entry = self.image_data_map[placeholder_id]
                    logger.info(LogModule.EXPORT, f"[DOCX-IMAGE] Found image by placeholder_id from alt_text: {placeholder_id}")
            # Try to find by direct filename match (without path)
            if not image_entry and filename:
                logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Trying direct filename match: {filename}")
                image_entry = self.image_data_map.get(filename)
                if image_entry:
                    logger.info(LogModule.EXPORT, f"[DOCX-IMAGE] Found image by direct filename match: {filename}")
            if not image_entry:
                logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Image data not found for reference (may be resolved by fallback): {image_ref}, alt_text: {alt_text}, filename: {filename}")
                return False
        
        # Get image data
        data_uri = image_entry.get("data") if isinstance(image_entry, dict) else None
        if not data_uri or not data_uri.startswith("data:image/"):
            logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Invalid image data URI for reference: {image_ref}")
            return False
        
        try:
            # Decode base64 data
            base64_data = data_uri.split(",", 1)[1]
            image_bytes = io.BytesIO(base64.b64decode(base64_data))
            
            # Get image dimensions and DPI
            pil_image = Image.open(image_bytes)
            original_width, original_height = pil_image.size
            
            # Try to get actual DPI from image metadata, fallback to 96 DPI
            # PIL stores DPI in info dict as 'dpi' tuple or 'resolution' tuple
            dpi = 200.0  # Default DPI assumption
            if hasattr(pil_image, 'info'):
                # Try 'dpi' key first (common for PNG/JPEG)
                if 'dpi' in pil_image.info and isinstance(pil_image.info['dpi'], tuple):
                    dpi = float(pil_image.info['dpi'][0]) if pil_image.info['dpi'][0] > 0 else 200.0
                # Try 'resolution' key (alternative format)
                elif 'resolution' in pil_image.info and isinstance(pil_image.info['resolution'], tuple):
                    dpi = float(pil_image.info['resolution'][0]) if pil_image.info['resolution'][0] > 0 else 200.0
            
            # Use original image size (no down-scaling), using actual DPI if available
            # This keeps formula and table images visually consistent with the source.
            if original_width <= 0 or original_height <= 0:
                logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Invalid image size for reference: {image_ref} ({original_width}x{original_height})")
                return False

            # Convert pixels to inches using actual DPI (or 96 DPI fallback)
            width_inches = original_width / dpi
            height_inches = original_height / dpi
            
            logger.debug(LogModule.EXPORT, f"[DOCX-IMAGE] Image DPI: {dpi}, size: {original_width}x{original_height} px -> {width_inches:.2f}\" x {height_inches:.2f}\"")
            
            # Reset stream for add_picture
            image_bytes.seek(0)
            
            # Add image to given paragraph or a new one (same paragraph for side-by-side images)
            para = paragraph if paragraph is not None else docx_doc.add_paragraph()
            run = para.add_run()
            run.add_picture(image_bytes, width=Inches(width_inches))
            
            logger.info(LogModule.EXPORT, f"[DOCX-IMAGE] Successfully inserted image: reference={image_ref}, alt_text={alt_text}, width={width_inches:.2f} inches")
            return True
        except Exception as e:
            logger.warning(LogModule.EXPORT, f"[DOCX-IMAGE] Failed to add image from {image_ref}: {e}")
            return False

    def _add_table(self, docx_doc: DocxDocument, table_data: List[List[str]]):
        """Add a table to the DOCX document.
        
        Args:
            docx_doc: DOCX document
            table_data: List of rows, each row is a list of cell strings
        """
        if not table_data:
            return
        
        # Determine number of columns (use max columns from all rows)
        num_cols = max(len(row) for row in table_data) if table_data else 0
        if num_cols == 0:
            return
        
        # Create table
        table = docx_doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Light Grid Accent 1'  # Use a built-in table style
        
        # Fill table cells
        for row_idx, row_data in enumerate(table_data):
            for col_idx in range(num_cols):
                cell = table.rows[row_idx].cells[col_idx]
                # Get cell text (or empty string if column doesn't exist in this row)
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                
                # Clear default paragraph and add text
                # Clean text to remove NULL bytes and control characters that are not XML compatible
                clean_cell_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', cell_text)
                cell.paragraphs[0].clear()
                if clean_cell_text:
                    self._add_runs_with_html_sup_sub(cell.paragraphs[0], clean_cell_text)

