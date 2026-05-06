# SPDX-FileCopyrightText: 2026 Zampher
# SPDX-License-Identifier: MPL-2.0

"""Tests for HTML <sup>/<sub> -> Word superscript/subscript in DOCX export."""

import sys
from pathlib import Path

_root = Path(__file__).resolve().parent.parent
_backend = Path(__file__).resolve().parent
for p in (str(_root), str(_backend)):
    if p not in sys.path:
        sys.path.insert(0, p)

from docx import Document

from backend.exporter.md.md2docx_exporter import MD2DOCXExporter, MD2DOCXExporterConfig


def test_sup_sub_inner_plain():
    assert MD2DOCXExporter._html_sup_sub_inner_to_plain(" 1,2 ") == "1,2"
    assert MD2DOCXExporter._html_sup_sub_inner_to_plain("&lt;2&gt;") == "<2>"


def test_add_runs_author_style_line():
    md = "云 李<sup>1</sup> | 志宏 黄<sup>2</sup> | tubs 辰<sup>1,2</sup>"
    doc = Document()
    para = doc.add_paragraph()
    exp = MD2DOCXExporter(config=MD2DOCXExporterConfig())
    exp._add_runs_with_html_sup_sub(para, md)
    texts = [r.text for r in para.runs if r.text]
    assert "云 李" in "".join(texts)
    assert any(r.font.superscript for r in para.runs if r.text in ("1", "2", "1,2"))
    supers = [r.text for r in para.runs if r.font.superscript]
    assert "1" in supers or "2" in supers or "1,2" in supers


def test_subscript_run():
    doc = Document()
    para = doc.add_paragraph()
    exp = MD2DOCXExporter(config=MD2DOCXExporterConfig())
    exp._add_runs_with_html_sup_sub(para, r"H<sub>2</sub>O")
    subs = [r.text for r in para.runs if r.font.subscript]
    assert subs == ["2"]
